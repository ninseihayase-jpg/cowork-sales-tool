"""フェーズ2-1：営業情報DBのブラウザ入力画面（標準ライブラリのみ）。

アカウント・商談・活動、リード・ピッチテーマを入力／一覧し、商談をテーマDBへ同期できる。
入力負荷を抑えるためステージ等はプルダウン。挙動安定を優先し外部依存なし。

起動: python scripts/run_webapp.py  → http://localhost:8787
"""

from __future__ import annotations

import base64
import hashlib
import hmac
import html
import json
import os
import re
import threading
import time
import urllib.parse
import urllib.request
from datetime import date, datetime, timedelta, timezone
from http.cookies import SimpleCookie
from email.message import EmailMessage
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer

from . import sfa_db
from . import leads_csv
from . import deals_csv
from . import csv_utils
from .theme_db import ThemeDBClient
from . import theme_link
from . import dev_project_link

SFA_API_TOKEN = os.environ.get("SFA_API_TOKEN", "")
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
# ブラウザ向け全ページの認証（未設定時はfail-closed=全拒否）。
# フォームログイン(Cookieセッション)＋従来のBasic認証の両方を受け付ける（#54: モバイルのBasic認証
# ダイアログでループする問題への対応。ネイティブダイアログを出さずログイン画面へ誘導する）。
SFA_BASIC_USER = os.environ.get("SFA_BASIC_USER", "")
SFA_BASIC_PASS = os.environ.get("SFA_BASIC_PASS", "")
_SESSION_COOKIE = "sfa_session"
_SESSION_MAX_AGE = 30 * 86400  # 30日


_JST = timezone(timedelta(hours=9))


def _today_jst() -> date:
    """業務上の「今日」＝日本時間の日付。本番(Render)はUTC稼働のため、標準の date.today() だと
    日本の早朝〜午前9時前は前日になってしまう。UIの日付デフォルト/期限比較は必ずこれを使う。"""
    return datetime.now(_JST).date()


def _session_secret() -> bytes:
    # 署名鍵はパスワードから派生（パスワード変更で既存セッションは自動失効）。
    return hashlib.sha256(("sfa-session|" + (SFA_BASIC_PASS or "")).encode("utf-8")).digest()


def _make_session_token() -> str:
    exp = int(time.time()) + _SESSION_MAX_AGE
    sig = hmac.new(_session_secret(), str(exp).encode(), hashlib.sha256).hexdigest()
    return f"{exp}.{sig}"


def _valid_session_token(tok: str) -> bool:
    try:
        exp_s, sig = (tok or "").split(".", 1)
        exp = int(exp_s)
    except (ValueError, AttributeError):
        return False
    if exp < int(time.time()):
        return False
    good = hmac.new(_session_secret(), str(exp).encode(), hashlib.sha256).hexdigest()
    return hmac.compare_digest(sig, good)


def login_page(next_url: str = "/", error: str = "") -> bytes:
    """ネイティブBasic認証ダイアログの代わりに出すログイン画面（モバイル安定）。"""
    nxt = next_url if next_url.startswith("/") else "/"
    err_html = (f'<p style="color:#b91c1c;font-size:13px;margin:0 0 10px">{html.escape(error)}</p>'
                if error else "")
    body = f"""<!doctype html><html lang="ja"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>ログイン ・ Inproc Salesforce</title>
<style>
 body{{font-family:system-ui,'Hiragino Kaku Gothic ProN',sans-serif;background:#f4f6f9;margin:0;
   display:flex;min-height:100vh;align-items:center;justify-content:center;color:#1d2430}}
 .box{{background:#fff;border-radius:14px;box-shadow:0 6px 24px rgba(0,0,0,.10);padding:28px 26px;width:340px;max-width:92vw}}
 h1{{font-size:17px;margin:0 0 4px}} .sub{{color:#8893a8;font-size:12px;margin:0 0 18px}}
 label{{display:block;font-size:12px;color:#6b7689;margin:12px 0 4px}}
 input{{width:100%;box-sizing:border-box;padding:11px 12px;border:1px solid #d4dae4;border-radius:8px;font-size:16px}}
 button{{width:100%;margin-top:18px;background:#2f6fed;color:#fff;border:0;border-radius:9px;padding:12px;font-size:15px;cursor:pointer}}
</style></head><body>
<form class="box" method="post" action="/login">
  <h1>Inproc Salesforce</h1>
  <p class="sub">ログインしてください</p>
  {err_html}
  <input type="hidden" name="next" value="{html.escape(nxt)}">
  <label>ユーザー名</label>
  <input name="username" autocapitalize="off" autocorrect="off" autocomplete="username" spellcheck="false" required>
  <label>パスワード</label>
  <input type="password" name="password" autocapitalize="off" autocomplete="current-password" required>
  <button type="submit">ログイン</button>
</form></body></html>"""
    return body.encode("utf-8")

INPROC_MEMBERS = [
    ("吉江", "takuya.yoshie@inproc.org"),
    ("中島", "yasutaka.nakajima@inproc.org"),
    ("早瀬", "ninsei.hayase@inproc.org"),
    ("岩崎", "eijiro.iwasaki@inproc.org"),
    ("高橋", "masanori.takahashi@inproc.org"),
    ("土屋", "tetsuhiro.tsuchiya@inproc.org"),
    ("戸田", "toda@inproc.org"),
    ("片山", "akito.katayama@inproc.org"),
    ("杉山", "hiroki.sugiyama@inproc.org"),
    ("山端", "rei.yamaberi@inproc.org"),
    ("堀籠", "wataru.horigome@inproc.org"),
]


def _opt(values: list[str], selected: str | None) -> str:
    out = ['<option value=""></option>']
    for v in values:
        sel = " selected" if v == selected else ""
        out.append(f'<option value="{html.escape(v)}"{sel}>{html.escape(v)}</option>')
    return "".join(out)


def _opt_kv(pairs: list[tuple[str, str]], selected: str | None) -> str:
    """(value, label) ペアリストから select options を生成。"""
    out = ['<option value=""></option>']
    for v, label in pairs:
        sel = " selected" if v == selected else ""
        out.append(f'<option value="{html.escape(v)}"{sel}>{html.escape(label)}</option>')
    return "".join(out)


def _esc(v) -> str:
    return "" if v is None else html.escape(str(v))


def _num(v, default=0) -> float:
    """外部入力由来の値を数値に強制する（HTML/SVGへ埋め込む前のXSS・型事故防止）。"""
    if isinstance(v, (int, float)):
        return v
    try:
        return float(v)
    except (TypeError, ValueError):
        return default


def _csv_safe(v):
    """CSVセル値の数式インジェクション中和。=+-@ 等で始まる文字列はExcel/Sheetsで
    数式として実行されうるため、先頭にシングルクォートを付けて無害化する。"""
    if isinstance(v, str) and v and v[0] in ("=", "+", "-", "@", "\t", "\r"):
        return "'" + v
    return v


def _opt_l2(l1: str | None, selected: str | None) -> str:
    """L1に対応するL2選択肢を生成。"""
    opts = ['<option value=""></option>']
    for v in sfa_db.BUSINESS_TYPE_L2_BY_L1.get(l1 or "", []):
        sel = " selected" if v == selected else ""
        opts.append(f'<option value="{html.escape(v)}"{sel}>{html.escape(v)}</option>')
    return "".join(opts)


def _decode_uploaded_csv(file_item) -> str | None:
    """CSV一括取込のファイルアップロード欄（multipart）から取得したバイト列をCSVテキストにする。

    - .xlsx/.xlsは csv_utils.xlsx_to_csv_text() でセル座標ベースに変換する
      （名刺xlsxアップロードは別の固定列レイアウト専用なので、テンプレート形式の
      xlsxをそちらに誤ってアップロードすると列がズレて壊れる。こちらの欄で
      xlsxも直接受け付けることで、その事故を防ぐ）。
    - .csv等のテキストファイルはUTF-8(BOM可)を試し、失敗したらcp932にフォールバック
      （Excel日本語Windowsの「CSV(カンマ区切り)」保存はShift-JISになることが多いため）。
    """
    if not file_item or not isinstance(file_item, tuple):
        return None
    filename, data = file_item
    if not data:
        return None
    if (filename or "").lower().endswith((".xlsx", ".xls")):
        try:
            return csv_utils.xlsx_to_csv_text(data)
        except Exception:
            return None
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError:
        return data.decode("cp932", errors="replace")


def _sticky_th(label: str, width: str | None = None) -> str:
    """縦スクロールしても項目名が見えるよう、テーブル見出しをposition:stickyにする共通ヘルパー。
    親のoverflow:auto付きコンテナ内で使うこと（テーブル単体ではスクロールしないため効果がない）。"""
    w = f' style="width:{width}"' if width else ""
    return f'<th class="sticky"{w}>{label}</th>'


def _call_claude_haiku(prompt: str, *, timeout: int = 20, max_wait: int = 25) -> str:
    """Claude Haikuを呼び出しテキストを返す。APIキー未設定・失敗・タイムアウト時は空文字。"""
    if not ANTHROPIC_API_KEY:
        return ""
    result: list = [None]

    def _do():
        try:
            payload = json.dumps({
                "model": "claude-haiku-4-5-20251001",
                "max_tokens": 512,
                "messages": [{"role": "user", "content": prompt}],
            }).encode()
            req = urllib.request.Request(
                "https://api.anthropic.com/v1/messages", data=payload, headers={
                    "x-api-key": ANTHROPIC_API_KEY,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                })
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                body = json.loads(resp.read())
            result[0] = body["content"][0]["text"].strip()
        except Exception as e:
            print(f"[issue AI summary] Claude API error: {e}", flush=True)

    t = threading.Thread(target=_do, daemon=True)
    t.start()
    t.join(timeout=max_wait)
    return result[0] or ""


def _generate_issue_ai_summary(issue: dict, memos: list[dict]) -> str:
    """論点のメモ全履歴からAIサマリーを生成する。失敗時は空文字（呼び出し側は既存サマリーを保持すること）。"""
    if not memos:
        return ""
    memo_text = "\n".join(
        f"- ({(m.get('created_at') or '')[:16]}) {m.get('body','')}"
        for m in memos
    )
    prompt = (
        f"以下は社内論点「{issue.get('issue','')}」についての議論メモの履歴です。\n"
        "話されているサブトピック・論点ごとに分けて、日本語で箇条書きに要約してください。\n"
        "各行は必ず「・トピック名：内容」の形式で書いてください"
        "（トピック名は名詞句で2〜10文字程度の短い見出し、内容にできれば結論や合意事項が分かるように書く）。\n"
        "まだ結論が出ていない・保留中の論点があれば、そのように分かるように書いてください。\n"
        "箇条書きは最大6行程度、前置きや「要約:」等の言葉は不要で、箇条書き本文だけを出力してください。\n\n"
        f"{memo_text}"
    )
    return _call_claude_haiku(prompt)


def _content_disposition(filename: str) -> str:
    """ダウンロードファイル名に日本語等を含む場合の Content-Disposition ヘッダ値を組み立てる。
    http.server はヘッダをlatin-1でエンコードするため、非ASCII文字を含む filename= だけでは
    UnicodeEncodeErrorでクラッシュする。RFC 6266のfilename*(UTF-8)形式を併記して回避する。"""
    ascii_fallback = filename.encode("ascii", "replace").decode("ascii").replace("?", "_")
    encoded = urllib.parse.quote(filename, safe="")
    return f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{encoded}"


def _tool_link_btn(url, label: str = "🔧 ツール", *, tool_id=None, tool_password=None) -> str:
    """開発案件の「制作したツールのリンク」を、邪魔にならない小さなボタンとして描画する。
    ID/PASSが設定されている場合は、いきなり遷移せずモーダルを開き、ID/PASSをコピー可能な状態で
    表示してから「ツールを開く」導線を出す（`_TOOL_MODAL_HTML`／`openToolModal`）。
    ID/PASSが無い場合は従来どおり直接リンクへ遷移する。一覧・商談詳細など複数画面で共通利用する。"""
    if not url:
        return ""
    _pill = (
        'display:inline-block;background:#ecfdf5;color:#047857;border:1px solid #a7f3d0;'
        'border-radius:6px;padding:2px 8px;font-size:11px;white-space:nowrap'
    )
    if not tool_id and not tool_password:
        return (
            f'<a href="{_esc(url)}" target="_blank" rel="noopener" title="制作したツールを開く" '
            f'style="{_pill};text-decoration:none">{label}</a>'
        )
    # ID/PASSあり → クリックでモーダル（コピペ可能なID/PASS＋「ツールを開く」）
    return (
        f'<button type="button" title="クリックでログイン情報を表示" '
        f'data-url="{_esc(url)}" data-id="{_esc(tool_id or "")}" data-pass="{_esc(tool_password or "")}" '
        f'onclick="openToolModal(this)" style="{_pill};cursor:pointer;font-family:inherit">'
        f'{label} 🔑</button>'
    )


def _seed_chips(seeds: str | None) -> str:
    """開発案件一覧の専用列: 技術シード(カンマ区切り)を小さなチップで表示（複数を見やすく折返し）。"""
    items = [s.strip() for s in (seeds or "").split(",") if s.strip()]
    if not items:
        return '<span class="muted" style="font-size:11px">—</span>'
    return "".join(
        f'<span style="display:inline-block;background:#eef2ff;color:#3730a3;border:1px solid #c7d2fe;'
        f'border-radius:5px;padding:1px 6px;margin:0 3px 3px 0;font-size:10px;white-space:nowrap">'
        f'{_esc(s)}</span>' for s in items)


def _chips(values) -> str:
    """CSV一括取込ページ向け: 選択肢一覧をチップ表示するHTMLを生成。"""
    return "".join(
        f'<span style="display:inline-block;background:#eef1f6;border-radius:4px;'
        f'padding:2px 8px;margin:2px 4px 2px 0;font-size:12px">{html.escape(v)}</span>'
        for v in values
    ) or '<span class="muted">(未設定)</span>'


def _ai_prompt_block(prompt_text: str, download_url: str) -> str:
    """CSV一括取込ページ向け: AI代行入力の指示文セクションを生成。"""
    return f"""
    <details style="margin:12px 0;background:#f3f0ff;border-radius:6px;padding:10px 14px">
      <summary style="cursor:pointer;font-weight:600;color:#5b21b6">
        🤖 AIにCSVを代行入力してもらう場合の指示文（コピーしてAIに貼り付け）
      </summary>
      <p class="muted" style="margin:10px 0 6px">
        調査済みの情報などを渡してAIにCSVを作らせる場合は、下記の指示文をそのままAIへの依頼文の先頭に貼り付けてください
        （選択肢は現在のマスタ設定を反映して自動生成されています）。
      </p>
      <p style="margin:0 0 8px">
        <a class="btn sec" href="{download_url}">📥 指示文をMarkdown(.md)でダウンロード</a>
      </p>
      <textarea readonly rows="14" onclick="this.select()"
        style="font-family:monospace;font-size:11px;background:#fff">{html.escape(prompt_text)}</textarea>
    </details>"""


PAGE = """<!doctype html><html lang="ja"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Inproc Salesforce</title>
<style>
 body{{font-family:system-ui,'Segoe UI','Hiragino Kaku Gothic ProN',sans-serif;margin:0;background:#f4f6f9;color:#1d2430}}
 header{{background:#1f2a44;color:#fff;padding:10px 18px;display:flex;align-items:center;gap:15px;flex-wrap:wrap;position:sticky;top:0;z-index:100}}
 header h1{{font-size:16px;margin:0 4px 0 0;white-space:nowrap}} header a{{color:#cdd7ff;text-decoration:none;font-size:13px;white-space:nowrap}}
 /* 個別編集/入力フォームの上部固定・保存バー（スクロールしても常に保存できる） */
 .save-bar{{position:sticky;top:50px;z-index:40;background:#fff;padding:8px 0;margin:0 0 14px;display:flex;gap:8px;align-items:center;flex-wrap:wrap;border-bottom:1px solid #e6e9f0}}
 .save-bar .sb-title{{font-weight:700;font-size:15px;margin-right:auto}}
 .nav-sep{{width:1px;height:16px;background:rgba(255,255,255,.22);margin:0}}
 .nav-menu{{position:relative}}
 .nav-menu>summary{{list-style:none;cursor:pointer;color:#cdd7ff;font-size:12px;opacity:.75}}
 .nav-menu>summary::-webkit-details-marker{{display:none}}
 .nav-menu[open]>summary{{opacity:1}}
 .nav-menu-panel{{position:absolute;left:0;top:160%;z-index:300;background:#fff;border:1px solid #e6e9f0;border-radius:8px;box-shadow:0 10px 30px rgba(0,0,0,.28);padding:6px;min-width:184px;display:flex;flex-direction:column}}
 .nav-menu-panel a{{color:#1d2430;font-size:13px;padding:7px 10px;border-radius:6px;white-space:nowrap}}
 .nav-menu-panel a:hover{{background:#f1f4f9}}
 .tb-menu{{position:relative;display:inline-block}}
 .tb-menu>summary{{list-style:none;cursor:pointer}}
 .tb-menu>summary::-webkit-details-marker{{display:none}}
 .tb-panel{{position:absolute;right:0;top:112%;z-index:300;background:#fff;border:1px solid #e6e9f0;border-radius:8px;box-shadow:0 10px 30px rgba(0,0,0,.16);padding:6px;min-width:172px;display:flex;flex-direction:column}}
 .tb-panel a{{color:#1d2430;text-decoration:none;font-size:13px;padding:8px 10px;border-radius:6px;white-space:nowrap}}
 .tb-panel a:hover{{background:#f1f4f9}}
 main{{max-width:1440px;margin:20px auto;padding:0 16px}}
 .card{{background:#fff;border-radius:10px;padding:16px 20px;margin-bottom:18px;box-shadow:0 1px 3px rgba(0,0,0,.08)}}
 h2{{font-size:15px;margin:0 0 12px;color:#3a4760}}
 table{{width:100%;border-collapse:collapse;font-size:13px}}
 th,td{{text-align:left;padding:7px 8px;border-bottom:1px solid #eef1f5}}
 th{{color:#8893a8;font-weight:600;font-size:12px}}
 tr:hover td{{background:#fafbfd}}
 /* 縦スクロール時に見出しを固定表示する共通クラス（旧: 各thへの直書きstyleを置換） */
 th.sticky{{position:sticky;top:0;background:#fff;z-index:2}}
 /* 開発案件一覧: 左側カラム(テーマ〜商談)を横スクロールでも固定表示。
    box-sizing:border-boxで幅=leftオフセットを一致させる（paddingでズレるのを防ぐ） */
 #dpTable td.frz,#dpTable th.frz{{box-sizing:border-box}}
 #dpTable td.frz{{position:sticky;background:#fff;z-index:3;box-shadow:1px 0 0 #e6e9f0}}
 #dpTable th.frz{{position:sticky;top:0;background:#fff;z-index:5;box-shadow:1px 0 0 #e6e9f0}}
 /* 開発テーマ詳細は3行までにクランプ（縦伸び防止）。全文はカーソルを合わせるとtitleツールチップで読める */
 #dpTable td.frz{{vertical-align:top}}
 .dp-detail{{display:-webkit-box;-webkit-line-clamp:3;-webkit-box-orient:vertical;overflow:hidden;line-height:1.45;cursor:help}}
 .stage{{display:inline-block;padding:2px 9px;border-radius:12px;font-size:12px;background:#e8edf7;color:#33406b;white-space:nowrap}}
 .tool-link-wrap{{position:relative;display:inline-block}}
 .tool-link-cred{{display:none;position:absolute;top:100%;left:0;z-index:20;background:#fff;
   border:1px solid #d0e4ff;border-radius:6px;padding:6px 10px;font-size:11px;white-space:nowrap;
   box-shadow:0 2px 8px rgba(0,0,0,.12);color:#1d2430;user-select:text;margin-top:4px}}
 .tool-link-wrap:hover .tool-link-cred{{display:block}}
 .btn{{display:inline-block;background:#2f6fed;color:#fff;border:0;border-radius:7px;padding:8px 14px;font-size:13px;cursor:pointer;text-decoration:none}}
 .btn.sec{{background:#e8edf7;color:#33406b}} .btn.sync{{background:#0c9b6a}}
 label{{display:block;font-size:12px;color:#6b7689;margin:10px 0 3px}}
 input,select,textarea{{width:100%;box-sizing:border-box;padding:7px 9px;border:1px solid #d4dae4;border-radius:6px;font-size:13px;font-family:inherit}}
 /* 自由記述textarea共通: フォーカスで拡大・blurで縮小(ヒアリング欄と統一した挙動) #55 */
 .ta-expand{{transition:height .15s ease}}
 .grid{{display:grid;grid-template-columns:1fr 1fr;gap:0 16px}} .full{{grid-column:1/3}}
 .muted{{color:#8893a8;font-size:12px}} .right{{text-align:right}}
 .flash{{background:#e6f7ef;color:#0c6b4a;padding:10px 14px;border-radius:8px;margin-bottom:14px;font-size:13px}}
 .s-new{{background:#f1f5f9;color:#475569}} .s-following{{background:#dbeafe;color:#1e40af}}
 .s-meeting{{background:#fef9c3;color:#92400e}} .s-proposal{{background:#ede9fe;color:#5b21b6}}
 .s-won{{background:#dcfce7;color:#166534}} .s-lost{{background:#fee2e2;color:#991b1b}}
 .theme-dot{{display:inline-block;width:10px;height:10px;border-radius:50%;vertical-align:middle;margin-right:4px}}
 .filter-row{{display:flex;flex-wrap:wrap;gap:8px;margin-bottom:14px;align-items:center}}
 .filter-row select,.filter-row input{{width:auto}}
 pre{{overflow-x:auto;white-space:pre-wrap;font-size:11px;line-height:1.6}}
 .dash-grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(200px,1fr));gap:14px;margin-bottom:20px}}
 .dash-card{{background:#fff;border-radius:10px;padding:20px;box-shadow:0 1px 3px rgba(0,0,0,.08)}}
 .dash-card .icon{{font-size:26px;margin-bottom:6px}}
 .dash-card h3{{font-size:15px;margin:0 0 4px;color:#1d2430}}
 .dash-card .desc{{font-size:12px;color:#8893a8;margin:0 0 10px;line-height:1.5}}
 .dash-card .count{{font-size:28px;font-weight:700;color:#2f6fed;margin-bottom:10px}}
 .dash-card .actions{{display:flex;gap:8px;flex-wrap:wrap}}
 .btn.ext{{background:#f3f0ff;color:#5b21b6}}
 /* スマホ対応（#54）: 折返しヘッダー/保存バーが画面を占有しないよう非固定化・余白圧縮・1カラム化 */
 @media(max-width:640px){{
   .grid{{grid-template-columns:1fr}} .full{{grid-column:1}} .hide-sm{{display:none}}
   table{{display:block;overflow-x:auto}}
   main{{margin:12px auto;padding:0 10px}}
   header{{position:static;padding:8px 12px;gap:8px 10px}}
   header h1{{font-size:15px;width:100%}}
   .save-bar{{position:static}}
   .card{{padding:14px 14px}}
 }}
</style>
<script>
/* 全ページ共通: JSでHTMLを組み立てる際のエスケープ（XSS防止）。動的optionやinnerHTML挿入で必ず使う */
function escH(s) {{
  return String(s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;')
    .replace(/"/g,'&quot;').replace(/'/g,'&#39;');
}}
/* 全ページ共通: 自由記述textareaの挙動。フォーカスで広げ(8行)、離れたら入力ありは4行・空は2行に戻す(#55) */
function taExpand(el) {{ el.rows = 8; }}
function taShrink(el) {{ el.rows = el.value.trim() ? 4 : 2; }}
</script>
</head><body>
<header>
  <h1>Inproc Salesforce</h1>
  <!-- 日常: 商談(ホーム)・開発案件・Delivery(受注後アサイン計画・#75) -->
  <a href="/deals">商談</a>
  <a href="/dev-projects">開発</a>
  <a href="/tasks">✅ タスク</a>
  <a href="/deliveries" style="opacity:.85;font-size:13px">🚚 Delivery</a>
  <span class="nav-sep"></span>
  <!-- 次: ヒアリング・論点 -->
  <a href="/hearings" style="opacity:.85;font-size:13px">ヒアリング</a>
  <a href="/deal-issues" style="opacity:.85;font-size:13px">論点</a>
  <span class="nav-sep"></span>
  <!-- クライアント管理: リード・アカウント -->
  <a href="/leads" style="opacity:.85;font-size:13px">リード</a>
  <a href="/accounts" style="opacity:.85;font-size:13px">アカウント</a>
  <span class="nav-sep"></span>
  <!-- 発信 -->
  <a href="/email-draft" style="opacity:.85;font-size:13px">メール</a>
  <!-- 管理（まとめ） -->
  <details class="nav-menu">
    <summary>⚙ 管理 ▾</summary>
    <div class="nav-menu-panel">
      <a href="/sync-health">🔍 同期チェック</a>
      <a href="/data-tagging">🏷 データ整備</a>
      <a href="/exhibition-tagging">🎪 展示会名タグ付け</a>
      <a href="/slack-memo-backfill">🩹 Slack追記メモ復旧</a>
      <a href="/masters">⚙ マスタ編集</a>
      <a href="/dev-point-master">🎯 開発点数マスタ</a>
      <a href="/tech-seed-master">🌱 技術シードマスタ</a>
      <a href="/tech-seed-tagging">🏷 技術シード一括付け</a>
      <a href="/backups">🗄 バックアップ</a>
      <a href="/logout">🚪 ログアウト</a>
    </div>
  </details>
  <a href="https://hisho-ohxe.onrender.com/dashboard" target="_blank" style="margin-left:auto;background:rgba(255,255,255,.1);border:1px solid rgba(255,255,255,.2);border-radius:6px;padding:5px 11px;font-size:11px;font-weight:600;color:#e0e8ff;text-decoration:none">InProc dashboard ↗</a>
  <a href="/dashboard" style="background:rgba(255,255,255,.08);border:1px solid rgba(255,255,255,.2);border-radius:6px;padding:5px 11px;font-size:11px;font-weight:600;color:#cdd7ff;text-decoration:none">📊 SFA dashboard</a>
  <a href="/reports" style="background:rgba(224,178,122,.16);border:1px solid rgba(224,178,122,.4);border-radius:6px;padding:5px 11px;font-size:11px;font-weight:600;color:#f0d9be;text-decoration:none">📰 週次レポート</a>
</header>
<main>{flash}{body}</main></body></html>"""


# 全ページ共通の「商談クローズ（リードに戻す）」モーダル。各画面のボタンから openCloseModal(id, returnTo)
# で開き、理由(必須)＋詳細(任意)を入力して /deal/{id}/revert_to_lead へPOSTする。運用をこの1本に統一。
_CLOSE_REASON_OPTS = "".join(
    f'<option value="{html.escape(r)}">{html.escape(r)}</option>' for r in sfa_db.CLOSE_REASONS
)
_CLOSE_MODAL_HTML = (
    '<div id="closeModal" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.45);z-index:1000;'
    'align-items:center;justify-content:center" onclick="if(event.target===this)closeCloseModal()">'
    '<div style="background:#fff;border-radius:12px;padding:20px 22px;max-width:460px;width:92%;'
    'box-shadow:0 10px 40px rgba(0,0,0,.25)">'
    '<h3 style="margin:0 0 12px">商談をクローズ（リードに戻す）</h3>'
    '<form id="closeModalForm" method="post">'
    '<input type="hidden" name="return_to" id="closeModalReturn" value="">'
    '<label>終了理由 <span style="color:#c53030">＊必須</span></label>'
    f'<select name="close_reason" id="closeModalReason" required><option value="">選択してください</option>{_CLOSE_REASON_OPTS}</select>'
    '<label>詳細（任意）</label>'
    '<textarea name="memo" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="3" placeholder="補足があれば"></textarea>'
    '<p class="muted" style="font-size:11px;margin:8px 0 12px">この商談はクローズされ、リード（フォロー中）に戻ります。</p>'
    '<div style="display:flex;gap:8px;justify-content:flex-end">'
    '<button type="button" class="btn sec" onclick="closeCloseModal()">キャンセル</button>'
    '<button type="submit" class="btn" style="background:#c53030">クローズする</button>'
    '</div></form></div></div>'
    '<script>'
    'function openCloseModal(id, returnTo){'
    ' var f=document.getElementById("closeModalForm"); f.reset();'
    ' f.action="/deal/"+id+"/revert_to_lead";'
    ' document.getElementById("closeModalReturn").value=returnTo||"";'
    ' document.getElementById("closeModal").style.display="flex";'
    '}'
    'function closeCloseModal(){document.getElementById("closeModal").style.display="none";}'
    'document.addEventListener("keydown",function(e){if(e.key==="Escape")closeCloseModal();});'
    # 全ページ共通のインライン更新（商談一覧の各タブで共有）
    'function updateDealField(id,field,value){'
    ' return fetch("/deal/"+id+"/field",{method:"POST",headers:{"Content-Type":"application/x-www-form-urlencoded"},'
    ' body:"field="+encodeURIComponent(field)+"&value="+encodeURIComponent(value)})'
    ' .then(function(r){return r.json();}).then(function(d){if(!d.ok)alert("更新エラー: "+(d.error||""));return d;})'
    ' .catch(function(){alert("通信エラー");});'
    '}'
    'function updateDevProjectField(id,field,value){'
    ' return fetch("/dev-project/"+id+"/field",{method:"POST",headers:{"Content-Type":"application/x-www-form-urlencoded"},'
    ' body:"field="+encodeURIComponent(field)+"&value="+encodeURIComponent(value)})'
    ' .then(function(r){return r.json();}).then(function(d){if(!d.ok)alert("更新エラー");return d;})'
    ' .catch(function(){alert("通信エラー");});'
    '}'
    # 事業種別L1変更時はL2の選択肢が変わるため、保存後にリロードして再描画（シンプル・確実）
    'function updateDealL1(id,l1){updateDealField(id,"business_type_l1",l1).then(function(){location.reload();});}'
    # 商談一覧の絞り込み（アカウント名テキスト＋ステージ複数選択チェック。全タブ共通）。
    'function filterDealsByAccount(){'
    ' var i=document.getElementById("accSearchInput"); var q=i?(i.value||"").toLowerCase():"";'
    ' var st=[]; document.querySelectorAll(".stg-cb:checked").forEach(function(b){st.push(b.value);});'
    ' var useSt=st.length>0;'
    ' document.querySelectorAll("tr[data-account]").forEach(function(tr){'
    '  var okA=tr.getAttribute("data-account").indexOf(q)>=0;'
    '  var okS=!useSt||st.indexOf(tr.getAttribute("data-stage")||"")>=0;'
    '  tr.style.display=(okA&&okS)?"":"none";});'
    ' var lbl=document.getElementById("stgFilterLbl");'
    ' if(lbl)lbl.textContent=useSt?("ステージ:"+st.length+"選択"):"ステージ:全て";'
    '}'
    # 初期表示時、URL由来で初期チェックされたステージがあれば絞り込みを適用（日付変更で保持）。
    'document.addEventListener("DOMContentLoaded",function(){'
    ' try{if(document.querySelector(".stg-cb:checked"))filterDealsByAccount();}catch(e){}});'
    '</script>'
)


_TOOL_MODAL_HTML = (
    '<div id="toolModal" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.45);z-index:1000;'
    'align-items:center;justify-content:center" onclick="if(event.target===this)closeToolModal()">'
    '<div style="background:#fff;border-radius:12px;padding:20px 22px;max-width:440px;width:92%;'
    'box-shadow:0 10px 40px rgba(0,0,0,.25)">'
    '<h3 style="margin:0 0 14px">制作ツールのログイン情報</h3>'
    '<div id="toolModalIdRow" style="display:flex;align-items:center;gap:8px;margin:0 0 10px">'
    '<span style="width:44px;color:#7a7266;font-size:12px;flex:none">ID</span>'
    '<code id="toolModalIdVal" style="flex:1;background:#f4f6f9;padding:6px 8px;border-radius:6px;'
    'font-size:13px;word-break:break-all"></code>'
    '<button type="button" id="toolModalIdCopy" class="btn sec" style="font-size:12px;padding:4px 10px;flex:none">コピー</button>'
    '</div>'
    '<div id="toolModalPassRow" style="display:flex;align-items:center;gap:8px;margin:0 0 10px">'
    '<span style="width:44px;color:#7a7266;font-size:12px;flex:none">PASS</span>'
    '<code id="toolModalPassVal" style="flex:1;background:#f4f6f9;padding:6px 8px;border-radius:6px;'
    'font-size:13px;word-break:break-all"></code>'
    '<button type="button" id="toolModalPassCopy" class="btn sec" style="font-size:12px;padding:4px 10px;flex:none">コピー</button>'
    '</div>'
    '<div style="display:flex;gap:8px;justify-content:flex-end;margin-top:16px">'
    '<button type="button" class="btn sec" onclick="closeToolModal()">閉じる</button>'
    '<a id="toolModalOpen" href="#" target="_blank" rel="noopener" class="btn" '
    'style="text-decoration:none">ツールを開く ↗</a>'
    '</div></div></div>'
    '<script>'
    'function _tmCopyFallback(txt){var t=document.createElement("textarea");t.value=txt;'
    't.style.position="fixed";t.style.opacity="0";document.body.appendChild(t);t.focus();t.select();'
    'try{document.execCommand("copy");}catch(e){}document.body.removeChild(t);}'
    'function _tmCopy(txt,btn){var done=function(){var o=btn.textContent;btn.textContent="✓ コピー済";'
    'setTimeout(function(){btn.textContent=o;},1200);};'
    'if(navigator.clipboard&&navigator.clipboard.writeText){'
    'navigator.clipboard.writeText(txt).then(done,function(){_tmCopyFallback(txt);done();});}'
    'else{_tmCopyFallback(txt);done();}}'
    'function openToolModal(btn){'
    ' var url=btn.getAttribute("data-url")||"",id=btn.getAttribute("data-id")||"",pass=btn.getAttribute("data-pass")||"";'
    ' var idRow=document.getElementById("toolModalIdRow"),passRow=document.getElementById("toolModalPassRow");'
    ' if(id){document.getElementById("toolModalIdVal").textContent=id;idRow.style.display="flex";'
    '  document.getElementById("toolModalIdCopy").onclick=function(){_tmCopy(id,this);};}else{idRow.style.display="none";}'
    ' if(pass){document.getElementById("toolModalPassVal").textContent=pass;passRow.style.display="flex";'
    '  document.getElementById("toolModalPassCopy").onclick=function(){_tmCopy(pass,this);};}else{passRow.style.display="none";}'
    ' var open=document.getElementById("toolModalOpen");'
    ' if(url){open.href=url;open.style.display="";}else{open.style.display="none";}'
    ' document.getElementById("toolModal").style.display="flex";'
    '}'
    'function closeToolModal(){document.getElementById("toolModal").style.display="none";}'
    'document.addEventListener("keydown",function(e){if(e.key==="Escape")closeToolModal();});'
    '</script>'
)


def render(body: str, flash: str = "") -> bytes:
    flash_html = f'<div class="flash">{html.escape(flash)}</div>' if flash else ""
    _sid = os.environ.get("SALES_SHEET_ID", "")
    delivery_url = f"https://docs.google.com/spreadsheets/d/{_sid}/edit" if _sid else "#"
    return PAGE.format(
        body=body + _CLOSE_MODAL_HTML + _TOOL_MODAL_HTML,
        flash=flash_html,
        delivery_url=delivery_url,
    ).encode("utf-8")


# ── 週次レポートのアーカイブ（読み物サイト） ─────────────────────────────────
# 本文（社外秘: クライアント名・金額）はGit(public)には置かず、永続ディスク上のDB
# (weekly_reports)にのみ格納する。閲覧・編集は既存Basic認証越し（社内限定）。
# 一覧・記事はCRM本体のガワを使わず、この専用の読み物デザイン(_reports_doc)で配信する。
_SLUG_RE = re.compile(r"[0-9A-Za-z_-]+")

_REPORTS_CSS = """
:root{
  --paper:#f7f4ec; --card:#fffdf8; --ink:#2b2620; --muted:#7c7264; --faint:#a99f8d;
  --accent:#a9743f; --accent-soft:#c39a68; --rule:#e7dfce; --tint:#efe9db; --shadow:rgba(60,45,25,.10);
  --serif:"Hiragino Mincho ProN","Yu Mincho","YuMincho","Noto Serif JP",serif;
  --sans:system-ui,-apple-system,"Hiragino Kaku Gothic ProN","Yu Gothic","Meiryo",sans-serif;
}
@media (prefers-color-scheme:dark){
  :root{--paper:#1c1916; --card:#252019; --ink:#ece4d5; --muted:#a99d8a; --faint:#7d7261;
    --accent:#cba06a; --accent-soft:#b3854f; --rule:#38301f; --tint:#282216; --shadow:rgba(0,0,0,.42);}
}
:root[data-theme="light"]{--paper:#f7f4ec; --card:#fffdf8; --ink:#2b2620; --muted:#7c7264; --faint:#a99f8d;
  --accent:#a9743f; --accent-soft:#c39a68; --rule:#e7dfce; --tint:#efe9db; --shadow:rgba(60,45,25,.10);}
:root[data-theme="dark"]{--paper:#1c1916; --card:#252019; --ink:#ece4d5; --muted:#a99d8a; --faint:#7d7261;
  --accent:#cba06a; --accent-soft:#b3854f; --rule:#38301f; --tint:#282216; --shadow:rgba(0,0,0,.42);}
*{box-sizing:border-box}
html{-webkit-text-size-adjust:100%}
@media (prefers-reduced-motion:no-preference){html{scroll-behavior:smooth}}
body{margin:0;background:var(--paper);color:var(--ink);font-family:var(--serif);
  font-size:clamp(1.02rem,.96rem + .35vw,1.14rem);line-height:2.02;
  -webkit-font-smoothing:antialiased;text-rendering:optimizeLegibility;
  font-feature-settings:"palt" 1;letter-spacing:.02em;}
a{color:var(--accent)}
img{max-width:100%;height:auto;display:block}
.rwrap{max-width:64rem;margin:0 auto;padding:0 clamp(1.2rem,5vw,2.6rem);}
.rmast{display:flex;align-items:flex-end;justify-content:space-between;gap:1rem;flex-wrap:wrap;
  padding:clamp(1.8rem,4vw,2.8rem) 0 1.3rem;border-bottom:1px solid var(--rule);margin-bottom:clamp(1.8rem,4vw,3rem);}
.rmark{font-family:var(--serif);font-weight:700;font-size:clamp(1.15rem,2.6vw,1.5rem);
  letter-spacing:.06em;color:var(--ink);text-decoration:none;display:flex;align-items:center;gap:.5rem;}
.rmark .dot{width:.62rem;height:.62rem;border-radius:50%;background:var(--accent);flex:none;}
.rtag{font-family:var(--sans);font-size:.74rem;color:var(--faint);letter-spacing:.06em;margin-top:.35rem;}
.rnav{display:flex;gap:1.1rem;font-family:var(--sans);font-size:.82rem;}
.rnav a{color:var(--muted);text-decoration:none;letter-spacing:.04em;padding-bottom:.15rem;
  border-bottom:1.5px solid transparent;transition:color .15s,border-color .15s;}
.rnav a:hover{color:var(--ink);border-color:var(--accent);}
.rfoot{margin:clamp(3rem,6vw,5rem) 0 3rem;padding-top:1.4rem;border-top:1px solid var(--rule);
  font-family:var(--sans);font-size:.76rem;color:var(--faint);display:flex;justify-content:space-between;gap:1rem;flex-wrap:wrap;}
.rfoot a{color:var(--muted);text-decoration:none;}
.kicker{font-family:var(--sans);font-size:.74rem;letter-spacing:.2em;color:var(--accent);
  font-variant-numeric:tabular-nums;text-transform:uppercase;}
/* ── 一覧 ── */
.feat{display:grid;grid-template-columns:1.15fr 1fr;gap:clamp(1.4rem,3vw,2.6rem);align-items:stretch;
  margin-bottom:clamp(2.4rem,5vw,3.6rem);}
@media (max-width:44rem){.feat{grid-template-columns:1fr}}
.feat-cover{border-radius:14px;min-height:15rem;background-size:cover;background-position:center;
  box-shadow:0 10px 30px var(--shadow);}
.feat-body{align-self:center;}
.feat-body h2{font-size:clamp(1.5rem,3.6vw,2.15rem);line-height:1.48;margin:.5rem 0 .7rem;
  text-wrap:balance;letter-spacing:.02em;}
.feat-body .lead{color:var(--muted);margin:0 0 1.1rem;line-height:1.95;text-wrap:pretty;}
.readmore{font-family:var(--sans);font-size:.86rem;font-weight:700;color:var(--accent);text-decoration:none;letter-spacing:.04em;}
.readmore:hover{text-decoration:underline;}
.bn-h{font-family:var(--serif);font-size:1.05rem;font-weight:700;letter-spacing:.1em;color:var(--ink);
  display:flex;align-items:center;gap:.7rem;margin:0 0 1.3rem;}
.bn-h::before{content:"";width:1.4rem;height:2px;background:var(--accent);border-radius:2px;flex:none;}
.grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(15rem,1fr));gap:clamp(1.1rem,2.5vw,1.8rem);}
.rcard{display:block;text-decoration:none;color:inherit;background:var(--card);border:1px solid var(--rule);
  border-radius:12px;overflow:hidden;transition:transform .16s ease,box-shadow .16s ease;}
.rcard:hover{transform:translateY(-3px);box-shadow:0 12px 26px var(--shadow);}
.rcard-cover{height:8.5rem;background-size:cover;background-position:center;}
.rcard-body{padding:.9rem 1.1rem 1.2rem;}
.rcard-body h3{font-size:1.08rem;font-weight:700;line-height:1.5;margin:.3rem 0 .4rem;letter-spacing:.02em;}
.rcard-body .lead{font-family:var(--sans);font-size:.8rem;color:var(--muted);line-height:1.7;
  display:-webkit-box;-webkit-line-clamp:2;-webkit-box-orient:vertical;overflow:hidden;}
.empty{color:var(--muted);text-align:center;padding:4rem 0;font-family:var(--sans);}
/* ── 記事 ── */
.hero{border-radius:0 0 18px 18px;margin-bottom:clamp(1.8rem,4vw,2.8rem);position:relative;overflow:hidden;}
.hero.has-cover{min-height:clamp(15rem,40vw,24rem);display:flex;align-items:flex-end;
  background-size:cover;background-position:center;color:#fff;}
.hero.has-cover::after{content:"";position:absolute;inset:0;
  background:linear-gradient(180deg,rgba(0,0,0,.05) 0%,rgba(0,0,0,.15) 45%,rgba(0,0,0,.72) 100%);}
.hero.has-cover .hero-inner{position:relative;z-index:1;padding:clamp(1.6rem,4vw,2.8rem);}
.hero.has-cover .kicker{color:#f2d9bd;}
.hero.has-cover h1{color:#fff;}
.hero.plain{background:linear-gradient(135deg,var(--tint),var(--card));border:1px solid var(--rule);
  border-radius:16px;padding:clamp(1.8rem,4vw,3rem) clamp(1.4rem,4vw,2.6rem);}
h1.art-title{font-size:clamp(1.7rem,4.6vw,2.5rem);font-weight:700;line-height:1.42;margin:.6rem 0 0;
  text-wrap:balance;letter-spacing:.02em;}
.art{max-width:46rem;margin:0 auto;}
.art .lead{font-size:clamp(1.16rem,2.1vw,1.36rem);line-height:2.1;margin:0 0 2.6rem;text-wrap:pretty;color:var(--ink);}
.art p{margin:0 0 1.4rem;text-wrap:pretty;}
.art h2{font-family:var(--serif);font-size:1.28rem;font-weight:700;letter-spacing:.06em;
  margin:3rem 0 1.3rem;display:flex;align-items:center;gap:.7rem;scroll-margin-top:1.5rem;}
.art h2::before{content:"";width:1.5rem;height:2px;background:var(--accent);border-radius:2px;flex:none;}
.art h3{font-size:1.16rem;font-weight:700;margin:2rem 0 .5rem;letter-spacing:.02em;}
.art .who{font-family:var(--sans);font-size:.86rem;color:var(--muted);letter-spacing:.02em;margin:0 0 .4rem;}
.art figure{margin:1.8rem 0;}
.art figure img{width:100%;border-radius:12px;box-shadow:0 8px 24px var(--shadow);}
.art figcaption{font-family:var(--sans);font-size:.78rem;color:var(--faint);margin-top:.55rem;text-align:center;letter-spacing:.02em;}
.art blockquote,.art .pull{margin:2rem 0;padding:0 0 0 1.2rem;border-left:3px solid var(--accent-soft);
  font-size:1.16rem;line-height:1.9;color:var(--ink);font-style:normal;}
.art .note{background:var(--tint);border-radius:10px;padding:1rem 1.2rem;color:var(--muted);
  font-family:var(--sans);font-size:.92rem;line-height:1.9;margin:1.6rem 0;}
.art hr{border:none;border-top:1px solid var(--rule);margin:2.6rem 0;}
.sign{margin-top:2.6rem;text-align:right;color:var(--muted);letter-spacing:.16em;}
/* 数字パック（本文中に置ける） */
.numbers{display:grid;grid-template-columns:repeat(auto-fit,minmax(9rem,1fr));gap:.7rem;margin:1.8rem 0 2.2rem;}
.numbers .n{background:var(--card);border:1px solid var(--rule);border-radius:10px;padding:.85rem 1rem;}
.numbers .k{font-family:var(--sans);font-size:.72rem;color:var(--muted);letter-spacing:.02em;line-height:1.5;}
.numbers .v{font-family:var(--sans);font-weight:700;font-size:1.5rem;color:var(--ink);
  font-variant-numeric:tabular-nums;line-height:1.3;margin-top:.15rem;}
.numbers .v .u{font-size:.72rem;font-weight:500;color:var(--muted);margin-left:.2em;}
.funnel{display:flex;flex-wrap:wrap;gap:.5rem 1.4rem;font-family:var(--sans);margin:0 0 1.4rem;
  padding:.9rem 1.1rem;background:var(--tint);border-radius:10px;}
.funnel span{font-size:.82rem;color:var(--muted);font-variant-numeric:tabular-nums;}
.funnel b{color:var(--ink);font-weight:700;margin-left:.35em;}
@media (prefers-reduced-motion:no-preference){
  .art>*,.feat,.grid>*{animation:rrise .7s cubic-bezier(.2,.6,.2,1) both}
  @keyframes rrise{from{opacity:0;transform:translateY(10px)}to{opacity:1;transform:none}}
}
"""


def _cover_bg(data_uri: str) -> str:
    """cover_image(data: URI)を background-image スタイル値に。安全な data:image/ のみ許可。"""
    v = (data_uri or "").strip()
    if v.startswith("data:image/") and '"' not in v and ")" not in v.split(",", 1)[0]:
        return f"background-image:url('{v}')"
    return ""


def _reports_doc(inner: str, *, page_title: str) -> str:
    """読み物サイト共通のガワ（standalone HTML）。一覧・記事の両方をこれで包む。"""
    return f"""<!doctype html><html lang="ja"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{_esc(page_title)}</title>
<style>{_REPORTS_CSS}</style>
</head><body>
<div class="rwrap">
  <header class="rmast">
    <div>
      <a class="rmark" href="/reports"><span class="dot"></span>InProc 営業レポート</a>
      <div class="rtag">現場で、いま何が起きているか。</div>
    </div>
    <nav class="rnav">
      <a href="/reports">一覧</a>
      <a href="/reports/manage">＋ 号を書く</a>
      <a href="/">SFA ↗</a>
    </nav>
  </header>
  {inner}
  <footer class="rfoot">
    <span>InProc 営業週次レポート（社外秘・社内限定）</span>
    <a href="/reports">レポート一覧へ</a>
  </footer>
</div>
</body></html>"""


def reports_index_page(con) -> str:
    """週次レポートの号一覧（読み物サイトのトップ）。最新号をフィーチャー＋バックナンバー。"""
    entries = sfa_db.list_weekly_reports(con)
    if not entries:
        inner = ('<div class="empty">まだレポートがありません。'
                 '<br><br><a class="readmore" href="/reports/manage">＋ 最初の号を書く</a></div>')
        return _reports_doc(inner, page_title="InProc 営業レポート")

    def _card(r, featured=False):
        slug = _esc(r.get("slug"))
        date = _esc(r.get("report_date") or "")
        title = _esc(r.get("title") or r.get("slug"))
        lead = _esc(r.get("lead") or "")
        bg = _cover_bg(r.get("cover_image") or "")
        if featured:
            cover = (f'<a href="/reports/{slug}" class="feat-cover" style="{bg}"></a>' if bg
                     else f'<a href="/reports/{slug}" class="feat-cover" style="background:linear-gradient(135deg,var(--accent-soft),var(--tint))"></a>')
            return f"""
            <div class="feat">
              {cover}
              <div class="feat-body">
                <div class="kicker">{date}</div>
                <h2><a href="/reports/{slug}" style="color:inherit;text-decoration:none">{title}</a></h2>
                <p class="lead">{lead}</p>
                <a class="readmore" href="/reports/{slug}">この号を読む →</a>
              </div>
            </div>"""
        cover = (f'<div class="rcard-cover" style="{bg}"></div>' if bg
                 else '<div class="rcard-cover" style="background:linear-gradient(135deg,var(--tint),var(--accent-soft))"></div>')
        return f"""
        <a class="rcard" href="/reports/{slug}">
          {cover}
          <div class="rcard-body">
            <div class="kicker">{date}</div>
            <h3>{title}</h3>
            <div class="lead">{lead}</div>
          </div>
        </a>"""

    featured = _card(entries[0], featured=True)
    rest = entries[1:]
    back = ""
    if rest:
        cards = "".join(_card(r) for r in rest)
        back = f'<h2 class="bn-h">バックナンバー</h2><div class="grid">{cards}</div>'
    return _reports_doc(featured + back, page_title="InProc 営業レポート")


# 記事(号)の読み物デザイン。artifactの2カラム・マガジン設計をアプリ側が保持し、
# 本文は「中身のfragment(.cols)」だけをDBに置く。フォント等の調整はここ(コード)で完結し、
# 貼り直し不要にする。フォントはフル幅ブラウザでも読みやすいサイズに調整済み。
_REPORT_ARTICLE_CSS = """
:root{
  --paper:#f5f2ea; --ink:#2a2622; --muted:#7a7266; --faint:#a99f8f;
  --ai:#9c7b4f; --rule:#e6dfd1; --tint:#efeade;
  --serif:"Hiragino Mincho ProN","Yu Mincho","YuMincho","Noto Serif JP",serif;
  --sans:system-ui,-apple-system,"Hiragino Kaku Gothic ProN","Yu Gothic","Meiryo",sans-serif;
}
@media (prefers-color-scheme:dark){
  :root{--paper:#211d18; --ink:#ece5d8; --muted:#aba08e; --faint:#867c6c;
    --ai:#c6a578; --rule:#372f26; --tint:#282219;}
}
:root[data-theme="light"]{--paper:#f5f2ea; --ink:#2a2622; --muted:#7a7266; --faint:#a99f8f;
  --ai:#9c7b4f; --rule:#e6dfd1; --tint:#efeade;}
:root[data-theme="dark"]{--paper:#211d18; --ink:#ece5d8; --muted:#aba08e; --faint:#867c6c;
  --ai:#c6a578; --rule:#372f26; --tint:#282219;}
*{box-sizing:border-box}
html{-webkit-text-size-adjust:100%}
@media (prefers-reduced-motion:no-preference){html{scroll-behavior:smooth}}
body{margin:0;background:var(--paper);color:var(--ink);font-family:var(--serif);
  font-size:clamp(.9rem,.87rem + .18vw,1.0rem);line-height:1.95;
  -webkit-font-smoothing:antialiased;text-rendering:optimizeLegibility;
  font-feature-settings:"palt" 1;letter-spacing:.02em;}
.wrap{max-width:60rem;margin:0 auto;padding:clamp(1.8rem,4vw,3.2rem) clamp(1.2rem,4vw,2.4rem) 4rem;}
.backlink{font-family:var(--sans);font-size:.76rem;color:var(--faint);text-decoration:none;letter-spacing:.05em;}
.backlink:hover{color:var(--ai);}
.mast{border-bottom:1px solid var(--rule);padding-bottom:1.3rem;margin:1rem 0 2.4rem;}
.date{font-family:var(--sans);font-size:.72rem;letter-spacing:.2em;color:var(--faint);font-variant-numeric:tabular-nums;}
h1{font-size:clamp(1.35rem,2.6vw,1.75rem);font-weight:700;line-height:1.5;margin:.7rem 0 0;text-wrap:balance;letter-spacing:.03em;}
.cols{display:grid;grid-template-columns:13rem 1fr;gap:clamp(1.8rem,3.5vw,3rem);}
.rail{align-self:start;position:sticky;top:2rem;}
@media (max-width:52rem){
  .cols{grid-template-columns:1fr;gap:0}
  .rail{position:static;top:auto;margin-bottom:2.4rem;padding-bottom:2rem;border-bottom:1px solid var(--rule);}
  .toc{display:none}
}
.toc{margin-bottom:1.8rem;}
.toc a{display:block;font-family:var(--serif);font-size:.86rem;color:var(--muted);text-decoration:none;
  padding:.3rem 0;letter-spacing:.03em;border-left:2px solid transparent;padding-left:.8rem;transition:color .15s,border-color .15s;}
.toc a:hover,.toc a:focus-visible{color:var(--ink);border-left-color:var(--ai);outline:none;}
.rail-h{display:flex;align-items:center;gap:.6rem;font-family:var(--serif);font-size:.9rem;font-weight:700;letter-spacing:.06em;color:var(--ink);margin:0 0 .9rem;}
.rail-h::before{content:"";width:1.2rem;height:2px;background:var(--ai);border-radius:2px;flex:none;}
.rail-h.sub{margin-top:1.6rem;}
.stat{padding:.5rem 0;border-bottom:1px solid var(--rule);}
.stat:first-of-type{border-top:1px solid var(--rule);}
.stat .k{font-family:var(--sans);font-size:.72rem;color:var(--muted);letter-spacing:.03em;}
.stat .v{font-family:var(--sans);font-weight:700;font-size:1.08rem;color:var(--ink);font-variant-numeric:tabular-nums;line-height:1.4;margin-top:.1rem;}
.stat .v .u{font-size:.68rem;font-weight:500;color:var(--muted);margin-left:.15em;}
.stat .v .ar{color:var(--muted);margin:0 .22em;}
.fn{display:flex;justify-content:space-between;font-family:var(--sans);font-size:.78rem;color:var(--muted);padding:.3rem 0;font-variant-numeric:tabular-nums;}
.fn b{color:var(--ink);font-weight:700;}
.article{min-width:0;max-width:40rem;}
.lead{font-size:clamp(1.02rem,1.4vw,1.16rem);line-height:2.0;margin:0 0 2.6rem;text-wrap:pretty;}
.lead .q{color:var(--ink);font-weight:600;}
.article p{margin:0 0 1.3rem;text-wrap:pretty;}
.label{display:flex;align-items:center;gap:.7rem;font-family:var(--serif);font-size:.98rem;font-weight:700;letter-spacing:.08em;color:var(--ink);margin:3rem 0 1.3rem;scroll-margin-top:1.5rem;}
.label::before{content:"";width:1.5rem;height:2px;background:var(--ai);border-radius:2px;flex:none;}
#lead{scroll-margin-top:1.5rem;}
.story{margin:0 0 1.8rem;}
.story .who{font-family:var(--serif);font-size:.88rem;color:var(--muted);letter-spacing:.02em;margin-bottom:.5rem;}
.story h3{font-size:1.05rem;font-weight:700;margin:0 0 .5rem;letter-spacing:.02em;}
.note{background:var(--tint);border-radius:10px;padding:.9rem 1.1rem;color:var(--muted);font-family:var(--sans);font-size:.86rem;line-height:1.9;margin:0 0 1.3rem;}
.sign{margin-top:2.6rem;padding-top:1.4rem;border-top:1px solid var(--rule);font-size:.98rem;color:var(--muted);text-align:right;letter-spacing:.14em;}
@media (prefers-reduced-motion:no-preference){
  .article>*{animation:arise .7s cubic-bezier(.2,.6,.2,1) both}
  .article>*:nth-child(2){animation-delay:.05s}
  .article>*:nth-child(3){animation-delay:.1s}
  @keyframes arise{from{opacity:0;transform:translateY(10px)}to{opacity:1;transform:none}}
}
"""


def report_article_html(rep: dict, rail_html: str = "") -> str:
    """1号の読み物ページ。旧・単体HTML本文は後方互換でそのまま返し、
    新形式（本文fragment=.cols の中身）はアプリ側の2カラム・マガジン設計に包んで返す。
    デザイン・フォント調整は _REPORT_ARTICLE_CSS 側で完結する（本文の貼り直し不要）。
    本文中の <!--NUMBERS--> は rail_html（自動集計した数字レール）に差し替える（#39）。"""
    body = rep.get("html_body") or ""
    # 数字レール自動注入: プレースホルダを実データのレールに差し替え（人は数字を手打ちしない）
    body = body.replace("<!--NUMBERS-->", rail_html or "")
    low = body.lstrip().lower()
    if low.startswith("<!doctype") or low.startswith("<html"):
        return body  # 旧形式（単体HTML）はそのまま配信
    date = _esc(rep.get("report_date") or "")
    title = _esc(rep.get("title") or rep.get("slug"))
    return (
        '<!doctype html><html lang="ja"><head><meta charset="utf-8">'
        '<meta name="viewport" content="width=device-width,initial-scale=1">'
        f'<title>{title}｜InProc 営業レポート</title>'
        f'<style>{_REPORT_ARTICLE_CSS}</style></head><body>'
        '<div class="wrap">'
        '<a class="backlink" href="/reports">← レポート一覧</a>'
        f'<header class="mast"><p class="date">営業週次レポート ・ {date}</p><h1>{title}</h1></header>'
        f'{body}'
        '</div></body></html>'
    )


def reports_manage_page(con, slug: str = "") -> str:
    """レポートの新規作成／編集フォーム（CRM側のガワ内・管理用）。
    本文は記事の中身HTML（fragment）を書く。カバー画像・本文中画像は
    ブラウザ側でリサイズしdata URI化して埋め込む（サーバはテキスト保存のみ）。"""
    import cowork.weekly_report as weekly_report
    rep = sfa_db.get_weekly_report(con, slug) if slug else None
    v_slug = _esc(rep["slug"]) if rep else ""
    v_date = _esc(rep["report_date"] or "") if rep else ""
    v_title = _esc(rep["title"] or "") if rep else ""
    v_lead = _esc(rep["lead"] or "") if rep else ""
    v_body = _esc(rep["html_body"] or "") if rep else ""
    v_cover = _esc(rep["cover_image"] or "") if rep else ""
    # 対象週(week_start=月曜)。新規は「今週の月曜(JST)」を既定に。数字レール自動注入の基準（#39）。
    _dflt_ws, _, _ = weekly_report._week_bounds(_today_jst())
    v_ws = _esc((rep.get("week_start") or "") if rep else "") or _dflt_ws
    cover_prev = (f'<img id="coverPreview" src="{v_cover}" style="max-height:120px;border-radius:8px;margin-top:8px">'
                  if v_cover else '<img id="coverPreview" style="display:none;max-height:120px;border-radius:8px;margin-top:8px">')
    slug_readonly = "readonly" if rep else ""
    heading = f"レポートを編集: {v_slug}" if rep else "レポートを新規作成"
    del_btn = ""
    if rep:
        del_btn = (
            f'<form method="post" action="/reports/manage/delete" style="display:inline" '
            f'onsubmit="return confirm(\'この号を削除します。よろしいですか？\')">'
            f'<input type="hidden" name="slug" value="{v_slug}">'
            f'<button type="submit" class="btn sec" style="color:#c53030">この号を削除</button></form>'
        )
    others = "".join(
        f'<a href="/reports/manage?slug={_esc(r["slug"])}" class="btn sec" '
        f'style="font-size:12px;margin:0 6px 6px 0">{_esc(r["slug"])}</a>'
        for r in sfa_db.list_weekly_reports(con)
    )
    others_block = f'<p class="muted" style="margin:0 0 6px">既存号を編集:</p>{others}' if others else ""
    return f"""
    <div class="card">
      <p style="margin:0 0 10px"><a href="/reports">← レポート一覧を見る</a></p>
      <h2 style="margin:0 0 4px">{heading}</h2>
      <p class="muted" style="margin:0 0 6px">本文は「記事の単体HTML」をそのまま貼り付けます（前回のデザインの読み物HTMLをそのまま配信します）。
        本文中に写真を入れる場合は、その単体HTMLの中に画像を埋め込んでおいてください。</p>
      <p class="muted" style="margin:0 0 14px">カバー画像は<b>一覧ページのサムネイル用</b>です（記事本体には影響しません）。
        選ぶと自動で縮小し、Gitには置かずこのDBに埋め込みます（社外秘・社内限定）。</p>
      <form method="post" action="/reports/manage/save" id="repForm">
        <label>スラッグ（URL・英数と-_のみ）＊必須</label>
        <input type="text" name="slug" value="{v_slug}" required pattern="[0-9A-Za-z_-]+"
               placeholder="2026-07-12" {slug_readonly}>
        <label>期間（一覧カードの日付表示）</label>
        <input type="text" name="report_date" value="{v_date}" placeholder="2026.7.6 – 7.12">
        <label>対象週（週の月曜・数字レールの基準週）＊自動集計の元になります</label>
        <input type="date" name="week_start" id="repWeekStart" value="{v_ws}"
               onchange="loadRailPreview()">
        <p class="muted" style="margin:4px 0 0">この週の数字が本文の <code>&lt;!--NUMBERS--&gt;</code> に自動で入ります（人は数字を手打ちしません）。</p>
        <label>表題（一覧カードの見出し）</label>
        <input type="text" name="title" value="{v_title}" placeholder="展示会が終わって、最初の一週間">
        <label>一言（一覧カードのリード）</label>
        <input type="text" name="lead" value="{v_lead}" placeholder="「そんなことができるんですか」。今週、その一言を何度も聞きました。">

        <label>カバー画像（一覧カードのサムネ用・任意）</label>
        <input type="file" accept="image/*" id="coverFile" onchange="onCoverPick(this)">
        <button type="button" class="btn sec" style="font-size:12px;margin-top:6px" onclick="clearCover()">カバーを外す</button>
        <input type="hidden" name="cover_image" id="coverData" value="{v_cover}">
        {cover_prev}

        <label style="margin-top:14px">数字ブロック（自動集計）</label>
        <div style="border:1px solid #e6e9f0;border-radius:8px;padding:10px 12px;background:#fbfaf7">
          <div style="display:flex;gap:8px;align-items:center;flex-wrap:wrap;margin-bottom:8px">
            <button type="button" class="btn sec" style="font-size:12px" onclick="insertNumbersPlaceholder()">
              本文に数字ブロックを挿入</button>
            <span class="muted" style="font-size:12px">本文の入れたい位置にカーソルを置いて押すと <code>&lt;!--NUMBERS--&gt;</code> を挿入します。</span>
          </div>
          <p class="muted" style="margin:0 0 6px;font-size:12px">選択中の対象週のプレビュー（保存後、記事ではこの数字が入ります）:</p>
          <div id="railPreview" style="max-height:320px;overflow:auto;background:#fff;border:1px solid #eee;border-radius:6px;padding:8px">
            <span class="muted">読み込み中…</span></div>
        </div>
        <label style="margin-top:14px">本文HTML（記事の中身fragment。デザインはアプリ側で固定。数字は上のブロックで自動注入）＊必須</label>
        <textarea name="html_body" id="repBody" rows="18" required
                  style="font-family:monospace;font-size:11px;background:#fff;white-space:pre">{v_body}</textarea>
        <div style="display:flex;gap:8px;justify-content:flex-end;margin-top:12px">
          {del_btn}
          <button type="submit" class="btn">保存</button>
        </div>
      </form>
      <div style="margin-top:18px;padding-top:14px;border-top:1px solid #e6e9f0">{others_block}</div>
    </div>
    <script>
    // 画像をブラウザ側で最大1600pxに縮小しJPEG(data URI)化。Gitに置かずDBにテキスト保存する。
    function _resizeToDataURL(file, maxDim, cb){{
      var reader=new FileReader();
      reader.onload=function(e){{
        var img=new Image();
        img.onload=function(){{
          var w=img.width,h=img.height;
          if(w>maxDim||h>maxDim){{var s=maxDim/Math.max(w,h);w=Math.round(w*s);h=Math.round(h*s);}}
          var c=document.createElement('canvas');c.width=w;c.height=h;
          c.getContext('2d').drawImage(img,0,0,w,h);
          cb(c.toDataURL('image/jpeg',0.85));
        }};
        img.src=e.target.result;
      }};
      reader.readAsDataURL(file);
    }}
    function onCoverPick(input){{
      if(!input.files||!input.files[0])return;
      _resizeToDataURL(input.files[0],1600,function(d){{
        document.getElementById('coverData').value=d;
        var p=document.getElementById('coverPreview');p.src=d;p.style.display='';
      }});
    }}
    function clearCover(){{
      document.getElementById('coverData').value='';
      var p=document.getElementById('coverPreview');p.src='';p.style.display='none';
      document.getElementById('coverFile').value='';
    }}
    // #39: 対象週の数字レールをプレビュー表示（保存後、記事の<!--NUMBERS-->にこの数字が入る）。
    function loadRailPreview(){{
      var ws=(document.getElementById('repWeekStart')||{{}}).value||'';
      var box=document.getElementById('railPreview');
      if(!box) return;
      box.innerHTML='<span class="muted">読み込み中…</span>';
      fetch('/reports/rail-preview?ws='+encodeURIComponent(ws)).then(function(r){{return r.text();}})
        .then(function(html){{ box.innerHTML=html||'<span class="muted">数字がありません。</span>'; }})
        .catch(function(){{ box.innerHTML='<span class="muted" style="color:#b91c1c">プレビュー取得に失敗しました。</span>'; }});
    }}
    // 本文テキストエリアのカーソル位置に <!--NUMBERS--> を挿入（既にあれば警告）。
    function insertNumbersPlaceholder(){{
      var ta=document.getElementById('repBody'); if(!ta) return;
      if(ta.value.indexOf('<!--NUMBERS-->')>=0){{ alert('本文に既に数字ブロック（<!--NUMBERS-->）があります。'); return; }}
      var s=ta.selectionStart||0, e=ta.selectionEnd||0;
      ta.value=ta.value.slice(0,s)+'\\n<!--NUMBERS-->\\n'+ta.value.slice(e);
      ta.focus(); ta.selectionStart=ta.selectionEnd=s+13;
    }}
    document.addEventListener('DOMContentLoaded', loadRailPreview);
    </script>"""


# ── Delivery（受注後・納品）アサイン計画（#75）─────────────────────────────

def _fmt_week(mon: str) -> str:
    """YYYY-MM-DD(月曜) → 'M/D' 週ラベル。"""
    try:
        d = date.fromisoformat(mon)
        return f"{d.month}/{d.day}"
    except Exception:
        return mon or ""


def _num_pct(v) -> str:
    """FTE%を整数なら整数、端数あれば小数1桁で表示（50.0→50, 12.5→12.5）。"""
    try:
        f = float(v or 0)
    except (TypeError, ValueError):
        return "0"
    return str(int(f)) if f == int(f) else f"{f:.1f}"


def _snap_monday(s) -> str:
    """任意の日付文字列(YYYY-MM-DD)をその週の月曜に丸める。空/不正は空文字。"""
    s = (s or "").strip()
    if not s:
        return ""
    try:
        return sfa_db._monday_of(date.fromisoformat(s))
    except ValueError:
        return ""


def _delivery_confidence(deal_stage: str, deal_status: str) -> tuple[str, str]:
    """(ラベル, 色) を返す。受注=確定/提案・クロージング=見込み/クローズ非受注=無効。"""
    if deal_status == "closed" and deal_stage != "受注":
        return ("無効(終了)", "#9ca3af")
    if deal_stage == "受注":
        return ("確定", "#047857")
    return ("見込み", "#b45309")


def _heat_style(pct: float) -> str:
    """総合負荷率%→背景色（100/150閾値。#75）。"""
    t = sfa_db.DELIVERY_HEAT_THRESHOLDS
    if pct <= 0:
        return "background:transparent;color:#cbd5e1"
    if pct < t["ok"]:
        return "background:#f0fdf4;color:#166534"
    if pct < t["full"]:
        return "background:#dcfce7;color:#166534"
    if pct < t["over"]:
        return "background:#fef3c7;color:#92400e"
    return "background:#dc2626;color:#fff;font-weight:700"


def deliveries_page(con) -> str:
    """Delivery案件一覧（旧・外部スプシリンクの置換）。#75。"""
    rows = ""
    for dv in sfa_db.list_deliveries(con):
        lbl, col = _delivery_confidence(dv.get("deal_stage") or "", dv.get("deal_status") or "open")
        blocks = sfa_db.list_delivery_assignments(con, dv["id"])
        who = "、".join(sorted({b["owner"] for b in blocks})) or "—"
        period = (f'{_fmt_week(dv.get("start_week"))}〜{_fmt_week(dv.get("end_week"))}'
                  if dv.get("start_week") or dv.get("end_week") else "—")
        rows += (
            f'<tr>'
            f'<td class="muted" style="font-size:.8em">#{dv["id"]}</td>'
            f'<td><a href="/delivery/{dv["id"]}"><b>{_esc(dv.get("title") or "(無題)")}</b></a></td>'
            f'<td>{_esc(dv.get("account_name") or "—")}</td>'
            f'<td><span style="background:{col};color:#fff;border-radius:5px;padding:1px 7px;font-size:11px">{lbl}</span></td>'
            f'<td>{_esc(dv.get("status") or "")}</td>'
            f'<td class="muted">{period}</td>'
            f'<td>{_esc(who)}</td>'
            f'<td><a class="btn sec" style="font-size:11px" href="/delivery/{dv["id"]}">編集</a></td>'
            f'</tr>')
    if not rows:
        rows = '<tr><td colspan=8 class=muted>Deliveryはまだありません。商談が「提案」に至ると自動で起票されます。</td></tr>'
    # 手動起票（提案以降のopen商談から選択）
    _cands = [d for d in sfa_db.list_deals(con, status="open")
              if (d.get("stage") or "") in sfa_db.DELIVERY_TRIGGER_STAGES]
    _cand_opts = "".join(
        f'<option value="{d["id"]}">{_esc(d.get("account_name") or "")}：{_esc(d.get("deal_name") or "")}（{_esc(d.get("stage") or "")}）</option>'
        for d in _cands)
    return f"""
    <div class="card">
      <h2 style="margin:0 0 4px">🚚 Delivery（受注後・アサイン計画）</h2>
      <p class="muted" style="margin:0 0 12px">商談が「提案」に至ると自動でDelivery案件が起票されます。ここでアサイン（誰が・いつ・何%）を入力すると、
        Hisho経営ダッシュボードの「稼働予定」に反映されます。単位＝FTE割合(%)。
        <a href="/base-workload">▶ ベース工数（営業・管理など恒常稼働）を設定</a></p>
      <div style="overflow:auto"><table style="min-width:900px">
        <tr><th>#</th><th>案件</th><th>クライアント</th><th>確度</th><th>状態</th><th>期間(週)</th><th>アサイン</th><th></th></tr>
        {rows}
      </table></div>
      <form method="post" action="/deliveries/new" style="margin-top:14px;display:flex;gap:8px;align-items:center;flex-wrap:wrap">
        <span class="muted" style="font-size:12px">手動で追加:</span>
        <select name="deal_id" required style="font-size:12px;max-width:420px"><option value="">商談を選択（提案以降）</option>{_cand_opts}</select>
        <button class="btn sec" style="font-size:12px">＋Delivery追加</button>
      </form>
    </div>"""


def delivery_form(con, delivery_id: int) -> str:
    """1Deliveryの編集：ヘッダ＋アサインブロック入力＋プレビューグリッド。#75。"""
    dv = sfa_db.get_delivery(con, delivery_id)
    if not dv:
        return '<div class="card"><p>Deliveryが見つかりません。<a href="/deliveries">← 一覧</a></p></div>'
    lbl, col = _delivery_confidence(dv.get("deal_stage") or "", dv.get("deal_status") or "open")
    status_opts = _opt(sfa_db.DELIVERY_STATUSES, dv.get("status") or "進行中")
    owner_opts = _opt(sfa_db.OWNERS, None)
    # 既存ブロック
    blocks = sfa_db.list_delivery_assignments(con, delivery_id)
    brows = ""
    for b in blocks:
        brows += (
            f'<tr>'
            f'<td>{_esc(b["owner"])}</td>'
            f'<td>{_fmt_week(b["from_week"])}〜{_fmt_week(b["to_week"])}</td>'
            f'<td style="text-align:right">{_num_pct(b["fte_pct"])}%</td>'
            f'<td class="muted">{_esc(b.get("note") or "")}</td>'
            f'<td><form method="post" action="/delivery/{delivery_id}/assignment/{b["id"]}/delete" style="margin:0" '
            f'onsubmit="return confirm(\'このブロックを削除しますか？\')">'
            f'<button class="btn sec" style="font-size:11px;color:#c53030">×</button></form></td>'
            f'</tr>')
    if not brows:
        brows = '<tr><td colspan=5 class=muted>まだアサインがありません。下で追加してください。</td></tr>'
    # プレビューグリッド
    grid = sfa_db.delivery_grid(con, delivery_id)
    if grid["weeks"]:
        head = "".join(f'<th style="font-size:11px;white-space:nowrap">{_fmt_week(w)}</th>' for w in grid["weeks"])
        grows = ""
        for ow in grid["owners"]:
            cells = ""
            for w in grid["weeks"]:
                v = grid["cells"].get(ow, {}).get(w, 0)
                cells += f'<td style="text-align:center;{_heat_style(v)}">{_num_pct(v)+"%" if v else "·"}</td>'
            grows += f'<tr><th style="text-align:left;white-space:nowrap">{_esc(ow)}</th>{cells}</tr>'
        grid_html = (f'<div style="overflow:auto"><table style="border-collapse:collapse">'
                     f'<tr><th></th>{head}</tr>{grows}</table></div>'
                     '<p class="muted" style="font-size:11px;margin:6px 0 0">※このグリッドはこのDelivery分のみ。'
                     '全社の総工数（デモ開発＋Delivery＋ベース）と負荷色はHishoダッシュボードで見ます。</p>')
    else:
        grid_html = '<p class="muted">アサインを追加するとここに週別グリッドが表示されます。</p>'

    return f"""
    <div class="card">
      <p style="margin:0 0 8px"><a href="/deliveries">← Delivery一覧</a></p>
      <h2 style="margin:0 0 2px">{_esc(dv.get("title") or "(無題)")}
        <span style="background:{col};color:#fff;border-radius:5px;padding:1px 8px;font-size:12px;vertical-align:middle">{lbl}</span></h2>
      <p class="muted" style="margin:0 0 12px">
        <a href="/deal/{dv["deal_id"]}">{_esc(dv.get("account_name") or "")}：{_esc(dv.get("deal_name") or "")}</a>
        （ステージ: {_esc(dv.get("deal_stage") or "")}）</p>

      <form method="post" action="/delivery/{delivery_id}/save" style="border:1px solid #e6e9f0;border-radius:8px;padding:12px;margin-bottom:14px">
        <div style="display:flex;gap:12px;flex-wrap:wrap;align-items:flex-end">
          <label style="font-size:12px">案件名<br><input type="text" name="title" value="{_esc(dv.get("title") or "")}" style="width:240px"></label>
          <label style="font-size:12px">開始週<br><input type="date" name="start_week" value="{_esc(dv.get("start_week") or "")}"></label>
          <label style="font-size:12px">終了週<br><input type="date" name="end_week" value="{_esc(dv.get("end_week") or "")}"></label>
          <label style="font-size:12px">状態<br><select name="status">{status_opts}</select></label>
          <button class="btn" style="font-size:12px">保存</button>
        </div>
        <label style="font-size:12px;display:block;margin-top:8px">概要・納品方針<br>
          <textarea name="overview" rows="2" style="width:100%">{_esc(dv.get("overview") or "")}</textarea></label>
        <p class="muted" style="font-size:11px;margin:4px 0 0">※週は入力後、自動でその週の月曜に丸めます。</p>
      </form>

      <h3 style="margin:0 0 6px;font-size:14px">アサイン（メンバー × 期間 × FTE%）</h3>
      <div style="overflow:auto"><table style="min-width:520px">
        <tr><th>メンバー</th><th>期間</th><th style="text-align:right">FTE%</th><th>メモ</th><th></th></tr>
        {brows}
      </table></div>

      <form method="post" action="/delivery/{delivery_id}/assignment/add" id="blkForm"
            style="display:flex;gap:8px;flex-wrap:wrap;align-items:flex-end;background:#f8fafc;border-radius:8px;padding:10px;margin-top:8px">
        <label style="font-size:12px">メンバー<br><select name="owner" required style="font-size:12px">{owner_opts}</select></label>
        <label style="font-size:12px">開始週<br><input type="date" name="from_week" required></label>
        <label style="font-size:12px">終了週<br><input type="date" name="to_week" required></label>
        <label style="font-size:12px">FTE%<br>
          <input type="number" name="fte_pct" id="blkFte" min="0" max="300" step="5" value="50" style="width:70px">
          <span style="white-space:nowrap">
            <button type="button" class="btn sec" style="font-size:10px;padding:2px 5px" onclick="document.getElementById('blkFte').value=25">25</button>
            <button type="button" class="btn sec" style="font-size:10px;padding:2px 5px" onclick="document.getElementById('blkFte').value=50">50</button>
            <button type="button" class="btn sec" style="font-size:10px;padding:2px 5px" onclick="document.getElementById('blkFte').value=100">100</button>
          </span></label>
        <label style="font-size:12px">メモ<br><input type="text" name="note" placeholder="担当領域など" style="width:160px"></label>
        <button class="btn" style="font-size:12px">＋アサイン追加</button>
      </form>

      <h3 style="margin:16px 0 6px;font-size:14px">プレビュー（週別・このDelivery分）</h3>
      {grid_html}

      <div style="margin-top:18px;border-top:1px solid #eee;padding-top:10px">
        <form method="post" action="/delivery/{delivery_id}/delete" style="display:inline"
              onsubmit="return confirm('このDelivery案件を削除します。アサインも消えます。よろしいですか？')">
          <button class="btn sec" style="font-size:12px;color:#c53030">このDeliveryを削除</button></form>
      </div>
    </div>"""


def base_workload_page(con) -> str:
    """ベース工数（人×機能×%）の編集。#75。SFA正本。"""
    rows = ""
    for r in sfa_db.list_base_workload(con):
        rows += (
            f'<tr>'
            f'<td>{_esc(r["owner"])}</td>'
            f'<td>{_esc(r["function"])}</td>'
            f'<td style="text-align:right">{_num_pct(r["pct"])}%</td>'
            f'<td><form method="post" action="/base-workload/{r["id"]}/delete" style="margin:0" '
            f'onsubmit="return confirm(\'削除しますか？\')"><button class="btn sec" style="font-size:11px;color:#c53030">×</button></form></td>'
            f'</tr>')
    if not rows:
        rows = '<tr><td colspan=4 class=muted>まだありません。</td></tr>'
    # 担当ごとの合算
    by_owner = sfa_db.base_workload_by_owner(con)
    sums = "、".join(f"{_esc(o)} {_num_pct(p)}%" for o, p in sorted(by_owner.items())) or "—"
    owner_opts = _opt(sfa_db.OWNERS, None)
    return f"""
    <div class="card">
      <p style="margin:0 0 8px"><a href="/deliveries">← Delivery一覧</a></p>
      <h2 style="margin:0 0 4px">ベース工数（恒常稼働：人 × 機能 × %）</h2>
      <p class="muted" style="margin:0 0 12px">案件に紐づかない恒常的な稼働（営業・管理・採用など）を人ごとに%で登録します。
        Hishoの総工数（デモ開発＋Delivery＋<b>ベース</b>）に加算されます。例: 早瀬 営業 30%。</p>
      <div style="overflow:auto"><table style="min-width:420px">
        <tr><th>メンバー</th><th>機能</th><th style="text-align:right">%</th><th></th></tr>
        {rows}
      </table></div>
      <p class="muted" style="font-size:12px;margin:8px 0 0">合算: {sums}</p>
      <form method="post" action="/base-workload/save" style="display:flex;gap:8px;flex-wrap:wrap;align-items:flex-end;background:#f8fafc;border-radius:8px;padding:10px;margin-top:10px">
        <label style="font-size:12px">メンバー<br><select name="owner" required style="font-size:12px">{owner_opts}</select></label>
        <label style="font-size:12px">機能（自由入力）<br><input type="text" name="function" required placeholder="営業 / 管理 / 採用 …" style="width:160px"></label>
        <label style="font-size:12px">%<br><input type="number" name="pct" min="0" max="100" step="5" value="20" style="width:70px"></label>
        <button class="btn" style="font-size:12px">保存（同一メンバー×機能は上書き）</button>
      </form>
    </div>"""


# ── メールパターン管理 ───────────────────────────────────────────────────────────

def email_patterns_page(con) -> str:
    patterns = sfa_db.list_email_patterns(con)
    rows = ""
    for p in patterns:
        cc = _esc(p.get("cc_addresses") or "")
        rows += (
            f'<tr>'
            f'<td><a href="/email-patterns/{p["id"]}/edit"><strong>{_esc(p["name"])}</strong></a></td>'
            f'<td>{_esc(p.get("from_address") or "—")}</td>'
            f'<td class="muted">{cc or "—"}</td>'
            f'<td>{_esc(p.get("subject") or "")}</td>'
            f'<td><form method="post" action="/email-patterns/{p["id"]}/delete" style="display:inline">'
            f'<button class="btn sec" style="font-size:11px;padding:4px 8px" '
            f'onclick="return confirm(\'削除しますか？\')">削除</button></form></td>'
            f'</tr>'
        )
    count = sfa_db.list_leads(con, status=None)
    assigned = sum(1 for l in count if l.get("email_pattern_id"))
    return f"""
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>メールパターン管理</span>
        <span style="display:flex;gap:8px">
          <a class="btn sec" href="/email-draft">ドラフト生成 ({assigned}件選択中)</a>
          <a class="btn" href="/email-patterns/new">＋パターン追加</a>
        </span>
      </h2>
      <p class="muted" style="margin-bottom:14px">テンプレート変数: <code>{{company}}</code> 社名　<code>{{name}}</code> 氏名　<code>{{title}}</code> 役職</p>
      <table>
        <tr><th>パターン名</th><th>From</th><th>CC</th><th>件名テンプレート</th><th></th></tr>
        {rows or '<tr><td colspan=5 class="muted">パターンがありません。</td></tr>'}
      </table>
    </div>"""


def email_pattern_form(con, pattern=None) -> str:
    pid = pattern["id"] if pattern else None
    action = f"/email-patterns/{pid}/save" if pid else "/email-patterns/save"
    title = "パターン編集" if pid else "パターン追加"
    from_opts = '<option value=""></option>' + "".join(
        f'<option value="{email}"{" selected" if pattern and pattern.get("from_address") == email else ""}>'
        f'{name} &lt;{email}&gt;</option>'
        for name, email in INPROC_MEMBERS
    )
    cc_existing = set((pattern.get("cc_addresses") or "").split(",")) if pattern else set()
    cc_checks = "".join(
        f'<div style="display:flex;align-items:center;gap:8px;margin:3px 0">'
        f'<input type="checkbox" name="cc" value="{email}" id="cc_{email}"'
        f'{" checked" if email in cc_existing else ""} style="flex-shrink:0;width:14px;height:14px">'
        f'<label for="cc_{email}" style="display:inline;font-size:13px;color:#2a3245;margin:0;cursor:pointer">'
        f'{name} &lt;{email}&gt;</label></div>'
        for name, email in INPROC_MEMBERS
    )
    return f"""
    <div class="card" style="max-width:900px">
      <h2>{title}</h2>
      <form method="post" action="{action}">
        <label>パターン名</label>
        <input name="name" required value="{_esc(pattern.get('name') if pattern else '')}">
        <label>From（送信元）</label>
        <select name="from_address">{from_opts}</select>
        <label>CC</label>
        <div style="background:#f4f6f9;border-radius:6px;padding:10px 14px;border:1px solid #d4dae4">{cc_checks}</div>
        <label>件名テンプレート <span class="muted">（{{company}} 等使用可）</span></label>
        <input name="subject" required value="{_esc(pattern.get('subject') if pattern else '')}">
        <label>本文テンプレート <span class="muted">（{{company}} / {{name}} / {{title}} 使用可）</span></label>
        <textarea name="body" rows="12" style="min-height:220px">{_esc(pattern.get('body') if pattern else '')}</textarea>
        <div style="margin-top:14px;display:flex;gap:8px">
          <button class="btn" type="submit">保存</button>
          <a class="btn sec" href="/email-patterns">キャンセル</a>
        </div>
      </form>
    </div>"""


def _render_tmpl(tmpl, lead) -> str:
    return (tmpl or "").replace("{company}", lead.get("company") or "").replace(
        "{name}", lead.get("name") or "").replace("{title}", lead.get("title") or "")


def build_eml_bytes(p, lead) -> bytes:
    """メールパターン + リードからEMLファイルのバイト列を生成する。"""
    subj = _render_tmpl(p.get("subject", ""), lead)
    body_raw = _render_tmpl(p.get("body", ""), lead)
    escaped = html.escape(body_raw)
    escaped = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', escaped)
    escaped = re.sub(
        r'\[([^\]]+)\]',
        r'<span style="background-color:yellow">[\1]</span>',
        escaped,
    )
    body_content = escaped.replace('\n', '<br>')
    full_html = (
        '<html><head><meta charset="UTF-8"></head>'
        '<body style="font-family:Meiryo UI,Meiryo,sans-serif;font-size:13px;line-height:1.7;color:#222">'
        f'{body_content}'
        '</body></html>'
    )
    msg = EmailMessage()
    to_addr = lead.get("email") or ""
    msg['To'] = to_addr
    if p.get("cc_addresses"):
        msg['CC'] = p["cc_addresses"]
    if p.get("from_address"):
        msg['From'] = p["from_address"]
    msg['Subject'] = subj
    msg.set_content(full_html, subtype='html')
    return msg.as_bytes()


def email_draft_page(con, *, status_filter=None, q=None) -> str:
    """メール送信ワークスペース。
    上段: リードごとにパターンを選択（前回値プリセット、変更はAJAX保存）
    下段: 選択済みリードのドラフトを常時表示
    """
    patterns_list = sfa_db.list_email_patterns(con)
    patterns = {p["id"]: p for p in patterns_list}

    # リード取得（デフォルト: converted/lost 除外）
    all_leads = sfa_db.list_leads(con, q=q)
    if status_filter:
        leads = [l for l in all_leads if l.get("lead_status") == status_filter]
    else:
        leads = [l for l in all_leads if l.get("lead_status") not in ("converted", "lost")]

    def _render(tmpl, lead):
        return _render_tmpl(tmpl, lead)

    def _cc_param(p):
        """CC アドレスをOutlook互換の '; ' 区切りでURLエンコードして返す。"""
        cc_raw = p.get("cc_addresses") or ""
        if not cc_raw:
            return ""
        cc_str = "; ".join(a.strip() for a in cc_raw.split(",") if a.strip())
        return "&cc=" + urllib.parse.quote(cc_str, safe="@;, ")

    def _mailto(p, lead):
        to_addr = lead.get("email") or ""
        if not to_addr:
            return ""
        subj = _render(p.get("subject", ""), lead)
        body_plain = _render(p.get("body", ""), lead).replace("**", "")
        qs = "subject=" + urllib.parse.quote(subj) + "&body=" + urllib.parse.quote(body_plain)
        qs += _cc_param(p)
        return "mailto:" + urllib.parse.quote(to_addr, safe="@") + "?" + qs

    def _mailto_noBody(p, lead):
        """To/CC/Subject のみ（bodyなし）のmailtoリンク。クリップボード貼り付け用。"""
        to_addr = lead.get("email") or ""
        if not to_addr:
            return ""
        subj = _render(p.get("subject", ""), lead)
        qs = "subject=" + urllib.parse.quote(subj)
        qs += _cc_param(p)
        return "mailto:" + urllib.parse.quote(to_addr, safe="@") + "?" + qs

    def _clipboard_html(text):
        """クリップボード用HTMLボディ: ** → <strong>、[括弧] → 黄色ハイライト（Outlook互換）。"""
        escaped = html.escape(text)
        escaped = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', escaped)
        escaped = re.sub(
            r'\[([^\]]+)\]',
            r'<span style="background:yellow;mso-highlight:yellow">[\1]</span>',
            escaped,
        )
        body_content = escaped.replace('\n', '<br>')
        return (
            '<html><head><meta charset="UTF-8"></head>'
            '<body style="font-family:Meiryo UI,Meiryo,sans-serif;font-size:13px;line-height:1.7;color:#222">'
            f'{body_content}</body></html>'
        )

    def _body_html(text):
        escaped = html.escape(text)
        escaped = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', escaped)
        escaped = re.sub(r'\[([^\]]+)\]',
                         r'<mark style="background:#fef08a;border-radius:2px;padding:1px 3px">\1</mark>',
                         escaped)
        return escaped.replace('\n', '<br>')

    # ── 上段: リスト + パターン選択 ──
    pattern_opts_base = '<option value="">— なし —</option>' + "".join(
        f'<option value="{p["id"]}">{_esc(p["name"])}</option>'
        for p in patterns_list
    )

    status_filter_opts = (
        '<option value="">アクティブ（default）</option>'
        + "".join(
            f'<option value="{s}"{" selected" if s == status_filter else ""}>'
            f'{sfa_db.LEAD_STATUS_LABELS[s]}</option>'
            for s in sfa_db.LEAD_STATUSES
        )
    )

    sel_rows = ""
    for ld in leads:
        cur_pid = ld.get("email_pattern_id")
        opts = '<option value="">— なし —</option>' + "".join(
            f'<option value="{p["id"]}"{" selected" if p["id"] == cur_pid else ""}>'
            f'{_esc(p["name"])}</option>'
            for p in patterns_list
        )
        has_email = bool(ld.get("email"))
        email_badge = (f'<span style="color:#16a34a;font-size:11px">✓</span>'
                       if has_email else '<span class="muted" style="font-size:11px">—</span>')
        sel_rows += (
            f'<tr>'
            f'<td><strong>{_esc(ld.get("company"))}</strong>'
            f'<span class="muted" style="font-size:11px;margin-left:6px">{_esc(ld.get("name"))}</span></td>'
            f'<td style="text-align:center">{email_badge}</td>'
            f'<td><select onchange="setLeadPattern({ld["id"]}, this.value)"'
            f' style="font-size:12px;padding:2px 4px;width:100%;">'
            f'{opts}</select></td>'
            f'</tr>'
        )

    if not leads:
        sel_rows = '<tr><td colspan=3 class="muted">リードがありません。</td></tr>'

    # ── 下段: ドラフト ──
    assigned = [l for l in sfa_db.list_leads(con) if l.get("email_pattern_id")]
    by_pattern: dict = {}
    for lead in assigned:
        pid = lead["email_pattern_id"]
        by_pattern.setdefault(pid, []).append(lead)

    preview_data = []
    draft_sections = []
    all_mailto_all = []
    for pid, p_leads in by_pattern.items():
        p = patterns.get(pid)
        if not p:
            continue
        cc_str = _esc(p.get("cc_addresses") or "")
        fr_str = _esc(p.get("from_address") or "")
        mailto_list = [_mailto(p, l) for l in p_leads if l.get("email")]
        all_mailto_all.extend(mailto_list)
        js_links = json.dumps(mailto_list, ensure_ascii=False)
        open_js = f"var lnks={js_links};lnks.forEach(function(u){{window.open(u)}});"
        rows = ""
        for lead in p_leads:
            mailto = _mailto(p, lead)
            subj = _render(p.get("subject", ""), lead)
            body_raw = _render(p.get("body", ""), lead)
            pidx = len(preview_data)
            preview_data.append({
                "label": f'{lead.get("company") or ""} / {lead.get("name") or ""}',
                "subject": subj,
                "body_html": _body_html(body_raw),
                "mailto": mailto,
                "mailto_noBody": _mailto_noBody(p, lead),
                "clipboard_html": _clipboard_html(body_raw),
                "eml_url": f"/email-draft/eml?lead_id={lead['id']}&pattern_id={pid}",
            })
            btn = (
                f'<button class="btn" onclick="showEmailPreview({pidx})" style="font-size:12px;padding:4px 10px">プレビュー</button>'
                if mailto else '<span class="muted" style="font-size:11px">アドレス未登録</span>'
            )
            rows += (
                f'<tr>'
                f'<td><strong>{_esc(lead.get("company"))}</strong>'
                f'<span class="muted" style="font-size:11px;margin-left:6px">{_esc(lead.get("name"))}</span></td>'
                f'<td class="muted">{_esc(lead.get("email") or "—")}</td>'
                f'<td>{_esc(subj)}</td>'
                f'<td>{btn}</td>'
                f'</tr>'
            )
        draft_sections.append(f"""
        <div style="margin-bottom:16px">
          <div style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px;margin-bottom:8px">
            <strong>{_esc(p["name"])}</strong>
            <span style="display:flex;gap:8px;align-items:center">
              <span class="muted" style="font-size:11px">From: {fr_str}{"　CC: "+cc_str if cc_str else ""}</span>
              <button class="btn sec" style="font-size:11px;padding:4px 10px"
                onclick="{html.escape(open_js)}">このパターン全件開く</button>
            </span>
          </div>
          <table><tr><th>会社 / 氏名</th><th>メールアドレス</th><th>件名プレビュー</th><th></th></tr>
          {rows}</table>
        </div>""")

    draft_count = len(assigned)
    draft_body = "".join(draft_sections) if draft_sections else '<p class="muted">パターンが選択されているリードがありません。</p>'
    all_js = json.dumps(all_mailto_all, ensure_ascii=False)
    all_open_js = f"var all={all_js};all.forEach(function(u){{window.open(u)}});"

    no_patterns_note = (
        f'<p class="muted" style="margin-bottom:10px">'
        f'パターンがまだありません。<a href="/email-patterns">パターン管理</a> から作成してください。</p>'
        if not patterns_list else ""
    )

    return f"""
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>メール送信ワークスペース</span>
        <a class="btn sec" href="/email-patterns" style="font-size:12px">パターン管理</a>
      </h2>
      {no_patterns_note}
      <p class="muted" style="margin-bottom:12px;font-size:12px">
        各リードにパターンを選択してください。選択内容は自動保存され、次回も引き継がれます。
      </p>
      <form class="filter-row" style="margin-bottom:10px">
        <select name="status" onchange="this.form.submit()" style="width:auto">{status_filter_opts}</select>
        <input name="q" placeholder="会社・氏名検索" value="{_esc(q or '')}" style="min-width:140px">
        <button class="btn sec" type="submit">絞り込み</button>
        <a class="btn sec" href="/email-draft">リセット</a>
      </form>
      <div style="overflow-x:auto">
      <table>
        <tr><th>会社 / 氏名</th><th style="text-align:center;width:40px">メール</th><th>パターン選択</th></tr>
        {sel_rows}
      </table>
      </div>
    </div>

    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>ドラフト <span class="muted" style="font-weight:normal">({draft_count}件選択中)</span></span>
        {'<button class="btn" style="font-size:12px" onclick="' + html.escape(all_open_js) + '">全件まとめて開く</button>' if all_mailto_all else ''}
      </h2>
      {draft_body}
    </div>

    <!-- メールプレビューモーダル -->
    <div id="emailPreviewModal" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.5);z-index:9999;align-items:flex-start;justify-content:center;padding-top:60px">
      <div style="background:#fff;border-radius:10px;max-width:620px;width:92%;max-height:80vh;overflow-y:auto;padding:24px 28px;box-shadow:0 12px 40px rgba(0,0,0,.3)">
        <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:18px">
          <strong id="epLabel" style="font-size:14px;color:#2a3245"></strong>
          <button onclick="closeEmailPreview()" style="background:none;border:none;font-size:20px;cursor:pointer;color:#999;line-height:1">×</button>
        </div>
        <div style="font-size:11px;color:#8a98b4;margin-bottom:3px;letter-spacing:.04em">件名</div>
        <div id="epSubject" style="font-size:13px;font-weight:600;padding:8px 12px;background:#f5f7fa;border-radius:5px;margin-bottom:16px;color:#2a3245"></div>
        <div style="font-size:11px;color:#8a98b4;margin-bottom:3px;letter-spacing:.04em">本文</div>
        <div id="epBody" style="font-size:13px;line-height:1.75;padding:14px 16px;background:#f5f7fa;border-radius:5px;color:#2a3245"></div>
        <div style="margin-top:10px;font-size:11px;color:#aab">
          <mark style="background:#fef08a;padding:1px 4px;border-radius:2px">黄色</mark>＝送信前に書き換えが必要な箇所
        </div>
        <div style="margin-top:20px;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
          <a id="epEml" href="#" style="font-size:11px;color:#8a98b4">EMLダウンロード（書式あり）</a>
          <div style="display:flex;gap:8px">
            <button onclick="closeEmailPreview()" class="btn sec" style="font-size:13px">閉じる</button>
            <button id="epOpen" onclick="openWithOutlook()" class="btn" style="font-size:13px">Outlookで開く</button>
          </div>
        </div>
      </div>
    </div>

    <script>
    var _epData = {json.dumps(preview_data)};
    var _epCurrent = null;
    function showEmailPreview(idx) {{
      _epCurrent = _epData[idx];
      document.getElementById('epLabel').textContent = _epCurrent.label;
      document.getElementById('epSubject').textContent = _epCurrent.subject;
      document.getElementById('epBody').innerHTML = _epCurrent.body_html;
      document.getElementById('epEml').href = _epCurrent.eml_url || '#';
      var m = document.getElementById('emailPreviewModal');
      m.style.display = 'flex';
    }}
    function closeEmailPreview() {{
      document.getElementById('emailPreviewModal').style.display = 'none';
    }}
    document.getElementById('emailPreviewModal').addEventListener('click', function(e) {{
      if (e.target === this) closeEmailPreview();
    }});
    function openWithOutlook() {{
      if (!_epCurrent) return;
      var mailto = _epCurrent.mailto_noBody || _epCurrent.mailto;
      // execCommand('copy') でレンダリング済みDOMをコピー → background-color が Outlook に伝わる
      try {{
        var parser = new DOMParser();
        var doc = parser.parseFromString(_epCurrent.clipboard_html, 'text/html');
        var temp = document.createElement('div');
        var bodyStyle = doc.body.getAttribute('style') || '';
        temp.style.cssText = bodyStyle + ';position:fixed;left:-9999px;top:0;width:600px;pointer-events:none';
        temp.innerHTML = doc.body.innerHTML;
        document.body.appendChild(temp);
        var range = document.createRange();
        range.selectNodeContents(temp);
        var sel = window.getSelection();
        sel.removeAllRanges();
        sel.addRange(range);
        document.execCommand('copy');
        sel.removeAllRanges();
        document.body.removeChild(temp);
        window.location.href = mailto;
        setTimeout(function() {{ showToast('本文をコピーしました。Outlookの本文欄に Ctrl+V で貼り付けてください'); }}, 600);
      }} catch(e) {{
        window.location.href = _epCurrent.mailto;
        setTimeout(function() {{ showToast('Outlookを開きます（書式なし）'); }}, 600);
      }}
    }}
    function showToast(msg) {{
      var t = document.createElement('div');
      t.textContent = msg;
      t.style.cssText = 'position:fixed;bottom:24px;left:50%;transform:translateX(-50%);background:#2a3245;color:#fff;padding:10px 20px;border-radius:6px;font-size:13px;z-index:99999;box-shadow:0 4px 12px rgba(0,0,0,.3);white-space:nowrap';
      document.body.appendChild(t);
      setTimeout(function(){{ t.style.opacity='0'; t.style.transition='opacity .4s'; setTimeout(function(){{t.remove()}},400); }}, 3500);
    }}
    function setLeadPattern(id, patternId) {{
      fetch('/leads/' + id + '/set_pattern', {{
        method: 'POST',
        headers: {{'Content-Type': 'application/x-www-form-urlencoded'}},
        body: 'pattern_id=' + encodeURIComponent(patternId)
      }}).then(r => r.json()).then(d => {{
        if (d.ok) {{ location.reload(); }} else {{ alert('エラー: ' + (d.error||'')); }}
      }}).catch(() => alert('通信エラー'));
    }}
    </script>"""


# ── ダッシュボード ──────────────────────────────────────────────────────────────

def dashboard_page(con) -> str:
    """SFA内部の俯瞰ダッシュボード。数値サマリ（アカウント/リード/商談）＋直近5日間の次回MS。"""
    deals = sfa_db.list_deals(con, status="open")
    accounts = sfa_db.list_accounts(con)
    leads = sfa_db.list_leads(con)
    hisho_url = os.environ.get("THEME_API_URL", "https://hisho-ohxe.onrender.com").rstrip("/") + "/dashboard"

    # 当日〜5日以内に次回MSがある商談
    today_str = _today_jst().isoformat()
    horizon_str = (_today_jst() + timedelta(days=5)).isoformat()
    recent_deals = sorted(
        [d for d in deals
         if d.get("next_milestone_date")
         and today_str <= d["next_milestone_date"] <= horizon_str],
        key=lambda d: d["next_milestone_date"],
    )
    recent_rows = ""
    for d in recent_deals:
        ms_raw = d.get("next_milestone_date", "")
        if ms_raw == today_str:
            ms = f'<span style="color:#dc2626;font-weight:700">今日 {_esc(ms_raw)}</span>'
        else:
            ms = _esc(ms_raw)
        if d.get("next_milestone_label"):
            ms += f'<br><span class="muted" style="font-size:.85em">{_esc(d["next_milestone_label"])}</span>'
        recent_rows += (
            f'<tr><td><a href="/deal/{d["id"]}">{_esc(d.get("account_name"))}</a></td>'
            f'<td>{_esc(d.get("deal_name"))}</td>'
            f'<td><span class="stage">{_esc(d.get("stage"))}</span></td>'
            f'<td>{ms}</td></tr>'
        )

    def _stat(icon, label, n, href):
        return (f'<div class="dash-card"><div class="icon">{icon}</div><h3>{label}</h3>'
                f'<div class="count">{n}</div>'
                f'<div class="actions"><a class="btn sec" href="{href}">一覧</a></div></div>')

    return f"""
    <div class="card">
      <h2 style="margin-top:0">ダッシュボード</h2>
      <div class="dash-grid">
        {_stat("🏢", "アカウント", len(accounts), "/accounts")}
        {_stat("🎯", "リード", len(leads), "/leads")}
        {_stat("💼", "商談（進行中）", len(deals), "/deals")}
      </div>
    </div>
    <div class="card">
      <h2>直近5日間の次回MS</h2>
      <table>
        <tr><th>アカウント</th><th>案件名</th><th>ステージ</th><th>次回MS</th></tr>
        {recent_rows or '<tr><td colspan=4 class=muted>直近5日以内に次回MSがある商談はありません</td></tr>'}
      </table>
      <p style="margin-top:10px">
        <a class="btn sec" href="/deals">すべての商談を見る</a>
        <a class="btn sec" href="/weekly-numbers" style="margin-left:8px">📊 週次レポート数字</a>
        <a class="btn ext" href="{hisho_url}" target="_blank" style="margin-left:8px">Inproc Dashboard ↗</a>
      </p>
    </div>"""


def _wow_note(delta) -> str:
    """前週比の注記（Noneや未確定は空文字）。"""
    if delta is None:
        return ""
    d = round(delta)
    sign = "±" if d == 0 else ("+" if d > 0 else "−")
    return f'<span class="muted" style="font-size:.85em">（前週比 {sign}{abs(d):g}）</span>'


def weekly_numbers_page(con) -> str:
    """週次レポート②の数字パック。開くと今週分スナップショットを自動記録し、貼れるテキストも生成。"""
    import cowork.weekly_report as weekly_report
    # 今週分スナップショットを自動記録（前週比の「時計」を回す。同週再訪は上書き）
    try:
        weekly_report.record_snapshot(con)
    except Exception as exc:  # noqa: BLE001 — 記録失敗でもレポート表示は続ける
        print(f"[weekly_numbers] snapshot record failed: {exc}", flush=True)
    r = weekly_report.compute_weekly_numbers(con)
    flow, stock, wow = r["flow"], r["stock"], r["wow"]
    exh = r["cohort"]["exhibition"]
    wk = f'{r["week_start"]} 〜 {r["week_end"]}'

    def dlt(key):
        return _wow_note(wow.get(key)) if wow.get("available") else ""

    # フロー前週比（フローは日付から都度計算するのでスナップショット無しでも出せる）
    mtg_d = flow["meetings"] - flow["prev"]["meetings"]
    lead_d = flow["new_leads"] - flow["prev"]["new_leads"]

    exh_src = flow["new_leads_by_source"].get("exhibition", 0)

    funnel_rows = "".join(
        f'<tr><td>{_esc(f["stage"])}</td><td style="text-align:right">{f["count"]}</td>'
        f'<td style="text-align:right">{f["lump"]:,.0f}</td>'
        f'<td style="text-align:right">{f["recurring"]:,.0f}</td>'
        f'<td>{_wow_note((wow.get("funnel") or {}).get(f["stage"])) if wow.get("available") else ""}</td></tr>'
        for f in stock["funnel"])
    closing_rows = "".join(
        f'<tr><td>{_esc(c.get("account"))}</td><td>{_esc(c.get("deal_name"))}</td>'
        f'<td style="text-align:right">{(c.get("lump") or 0):,.0f}</td>'
        f'<td>{_esc(c.get("ms_date") or "—")}</td><td>{_esc(c.get("owner") or "—")}</td></tr>'
        for c in stock["closing_deals"]) or '<tr><td colspan=5 class=muted>クロージング商談なし</td></tr>'

    wow_banner = ("" if wow.get("available")
                  else '<p class="muted">※ 前週比は来週号から表示されます（今週が最初のスナップショット記録）。</p>')

    # ②に貼れるプレーンテキスト（温度ゼロ・事実のみ）
    def sgn(x):
        x = round(x); return "±0" if x == 0 else (f"+{x:g}" if x > 0 else f"−{abs(x):g}")
    wow_pipe = f"（前週比 {sgn(wow['pipeline_lump'])}万）" if wow.get("available") else ""
    fn = {f["stage"]: f["count"] for f in stock["funnel"]}
    paste = (
        f"② 今週の数字（{wk}）\n"
        f"・面談：{flow['meetings']}件（相手 {flow['meeting_companies']}社）"
        f"{'（前週比 ' + sgn(mtg_d) + '）' if True else ''}\n"
        f"・新規リード：{flow['new_leads']}件（うち展示会 {exh_src}件）\n"
        f"・新規商談：{flow['new_deals']}件\n"
        f"・現ファネル：初回アポ {fn.get('初回アポ実施',0)}／要件詰め {fn.get('要件詰め',0)}"
        f"／提案 {fn.get('提案',0)}／クロージング {fn.get('クロージング',0)}\n"
        f"・パイプライン総額：{stock['pipeline_lump']:,.0f}万円{wow_pipe}\n"
        f"・展示会ファネル：商談化 {exh['total']}（有効母数 {exh['valid_total']}／ニーズなし {exh['no_need']}）"
        f" → 初回面談 {exh['first_meeting']} → 次商談 {exh['second_meeting']} → 受注 {exh['won']}\n"
        f"※キャンセル率は手元集計とマージ"
    )

    return f"""
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>週次レポート 数字パック <span class="muted" style="font-size:.6em">{_esc(wk)}</span></span>
        <a class="btn sec" href="/weekly-numbers/audit">🔍 集計を検証（元データ）</a>
      </h2>
      {wow_banner}
      <h3>今週の動き（フロー）</h3>
      <ul>
        <li>面談：<b>{flow['meetings']}件</b>（相手 {flow['meeting_companies']}社） {_wow_note(mtg_d)}</li>
        <li>新規リード：<b>{flow['new_leads']}件</b>（うち展示会 {exh_src}件） {_wow_note(lead_d)}</li>
        <li>新規商談：<b>{flow['new_deals']}件</b></li>
      </ul>
      <h3>現ファネル（ストック）　パイプライン単発総額 <b>{stock['pipeline_lump']:,.0f}万円</b> {dlt('pipeline_lump')}</h3>
      <table>
        <tr><th>ステージ</th><th>件数</th><th>単発計(万)</th><th>継続月額計(万)</th><th>前週比</th></tr>
        {funnel_rows}
      </table>
      <h3>クロージング商談</h3>
      <table>
        <tr><th>アカウント</th><th>案件名</th><th>単発(万)</th><th>次回MS</th><th>担当</th></tr>
        {closing_rows}
      </table>
      <h3>展示会ファネル（コホート）</h3>
      <p>商談化 <b>{exh['total']}</b>
         <span class="muted" style="font-size:.85em">（有効母数 {exh['valid_total']}／ニーズなし {exh['no_need']}・キャンセル {exh['canceled']}）</span>
         → 初回面談到達 <b>{exh['first_meeting']}</b>
         → 次の商談に進んだ <b>{exh['second_meeting']}</b> → 受注 <b>{exh['won']}</b></p>
      <p class="muted" style="font-size:.85em">有効母数＝総数−ニーズなし（最初から見込みゼロを分母から外す）。キャンセル率の手元集計とマージ想定。</p>
      <h3>レポート②に貼れるテキスト</h3>
      <textarea rows="12" style="width:100%;font-family:monospace;font-size:13px"
        onclick="this.select()">{_esc(paste)}</textarea>
    </div>"""


def _audit_table(headers: list, rows: list) -> str:
    """監査用の素朴なテーブル（ヘッダ＋行）。rowsは各要素がセル文字列のlist。"""
    if not rows:
        return '<p class="muted" style="font-size:.85em;margin:.3rem 0">該当行なし</p>'
    thead = "".join(f"<th>{_esc(h)}</th>" for h in headers)
    trs = "".join("<tr>" + "".join(f"<td>{c}</td>" for c in r) + "</tr>" for r in rows)
    return (f'<table style="font-size:12px;width:100%;border-collapse:collapse">'
            f'<tr>{thead}</tr>{trs}</table>')


def _audit_section(title: str, definition: str, value_html: str, detail_html: str,
                   warn: str = "") -> str:
    """1指標の監査ブロック: 定義・値・元データ(details)・注意。"""
    warn_html = (f'<div style="background:#fef2f2;border-left:3px solid #dc2626;padding:6px 10px;'
                 f'margin:6px 0;font-size:12px;color:#991b1b">⚠ {warn}</div>' if warn else "")
    return (f'<div class="card" style="margin-bottom:10px">'
            f'<h3 style="margin:0 0 4px">{_esc(title)} — {value_html}</h3>'
            f'<div class="muted" style="font-size:12px;margin-bottom:4px">定義: {definition}</div>'
            f'{warn_html}'
            f'<details><summary style="cursor:pointer;font-size:12px;color:#2563eb">元データを表示（数えている行）</summary>'
            f'<div style="overflow-x:auto;margin-top:6px">{detail_html}</div></details></div>')


def weekly_numbers_audit_page(con, as_of=None, exh_filter=None) -> str:
    """数字パックの「集計監査」ページ（#27）。各指標を 定義＋値＋元データ に展開し、
    集計の確からしさを人が元データで検証できるようにする。スナップショットは記録しない（読み取り専用）。"""
    import cowork.weekly_report as weekly_report
    ws, we, _prev = weekly_report._week_bounds(as_of)
    open_cond = "(d.status='open' OR d.status IS NULL)"

    # 1) 面談（フロー）。同一商談(SFA#)×同一日は1面談に重複排除して数える。
    mtg_rows = con.execute(
        "SELECT a.id, a.deal_id, a.occurred_on, acc.name acc, d.deal_name, a.contact_name, a.type "
        "FROM activities a LEFT JOIN deals d ON d.id=a.deal_id LEFT JOIN accounts acc ON acc.id=d.account_id "
        "WHERE a.type='面談' AND a.occurred_on BETWEEN ? AND ? ORDER BY a.occurred_on", (ws, we)).fetchall()
    mtg_companies = len({r["acc"] for r in mtg_rows if r["acc"]})
    # 重複排除キー: deal_idがあれば (deal_id, 日付)、無ければ活動id単位
    _mtg_keys = {}
    for r in mtg_rows:
        _k = (f"d{r['deal_id']}|{r['occurred_on']}" if r["deal_id"] is not None else f"a{r['id']}")
        _mtg_keys.setdefault(_k, 0)
        _mtg_keys[_k] += 1
    mtg_dedup = len(_mtg_keys)
    _dup_activities = len(mtg_rows) - mtg_dedup  # 同一商談・同日で潰れた活動数
    mtg_tbl = _audit_table(
        ["面談日", "アカウント", "案件", "相手", "重複"],
        [[_esc(r["occurred_on"]), _esc(r["acc"] or "—"), _esc(r["deal_name"] or "—"),
          _esc(r["contact_name"] or "—"),
          ("↩ 同一商談・同日" if (r["deal_id"] is not None
            and _mtg_keys.get(f"d{r['deal_id']}|{r['occurred_on']}", 0) > 1) else "")]
         for r in mtg_rows])
    sec_mtg = _audit_section(
        "面談数", f"activities.type='面談' かつ occurred_on が {ws}〜{we}。"
        "同一商談(SFA#)×同一日は1面談に重複排除。相手社数=案件経由のaccount重複排除。",
        f'<b>{mtg_dedup}</b>件（相手 {mtg_companies}社／元activities {len(mtg_rows)}件・'
        f'重複排除 {_dup_activities}件）', mtg_tbl,
        warn="occurred_on 未入力の面談はここに出ません。面談を activities に『面談』種別で登録しているか要確認。"
             "『重複』列が付いた行は同一商談・同日のため件数上は1件に集約しています。")

    # 2) 新規リードは集計不要（#27・ユーザー確定）のため監査ページから除外。

    # 3) 新規商談（フロー）＝「その週に初めて“面談”した商談」（#27で定義変更）
    deal_rows = con.execute(
        "SELECT fa.first_act, acc.name acc, d.deal_name, d.stage, d.owner FROM ("
        "  SELECT a.deal_id, MIN(a.occurred_on) first_act FROM activities a"
        "  WHERE a.type='面談' AND a.occurred_on IS NOT NULL AND a.occurred_on != ''"
        "  GROUP BY a.deal_id HAVING first_act BETWEEN ? AND ?"
        ") fa JOIN deals d ON d.id=fa.deal_id LEFT JOIN accounts acc ON acc.id=d.account_id "
        "ORDER BY fa.first_act", (ws, we)).fetchall()
    deal_tbl = _audit_table(
        ["初回面談日", "アカウント", "案件名", "ステージ", "担当"],
        [[_esc(r["first_act"]), _esc(r["acc"] or "—"), _esc(r["deal_name"] or "—"),
          _esc(r["stage"] or "—"), _esc(r["owner"] or "—")] for r in deal_rows])
    sec_deal = _audit_section(
        "新規商談", f"その商談の最初の“面談”(activities.type='面談' の occurred_on 最小)が {ws}〜{we} にある商談。"
        "（旧: 登録日created_at基準→初回活動基準→初回面談基準 と定義変更 #27）",
        f'<b>{len(deal_rows)}</b>件', deal_tbl,
        warn="面談(type='面談')が1件も登録されていない商談は『新規』に数えられない。メモ/メール等の"
             "非面談活動が初回でも新規には数えない（継続案件のメモ起点の誤カウントを防ぐ）。occurred_on が空の面談は無視。")

    # 4) パイプライン（ストック＝open かつ「要件詰め以降」の商談。#27でユーザー確定）
    _pipe_ph = ", ".join("?" for _ in sfa_db.PIPELINE_STAGES)
    open_rows = con.execute(
        f"SELECT acc.name acc, d.deal_name, d.stage, d.value_lumpsum lump, d.value_recurring rec "
        f"FROM deals d LEFT JOIN accounts acc ON acc.id=d.account_id "
        f"WHERE {open_cond} AND d.stage IN ({_pipe_ph}) "
        f"ORDER BY d.value_lumpsum DESC", list(sfa_db.PIPELINE_STAGES)).fetchall()
    sum_lump = sum((r["lump"] or 0) for r in open_rows)
    sum_rec = sum((r["rec"] or 0) for r in open_rows)
    open_tbl = _audit_table(
        ["アカウント", "案件名", "ステージ", "単発(万)", "継続月(万)"],
        [[_esc(r["acc"] or "—"), _esc(r["deal_name"] or "—"), _esc(r["stage"] or "—"),
          f'{(r["lump"] or 0):,.0f}', f'{(r["rec"] or 0):,.0f}'] for r in open_rows])
    sec_pipe = _audit_section(
        "パイプライン（要件詰め以降）",
        f"status='open'(またはNULL) かつ ステージが {' / '.join(sfa_db.PIPELINE_STAGES)} の商談。"
        "初回アポ実施・受注・保留中は除外。金額は提案金額（value_lumpsum 単発 / value_recurring 継続月）の合計＝顧客予算(client_budget)ではない。",
        f'<b>{len(open_rows)}</b>件・単発計 <b>{sum_lump:,.0f}</b>万 / 継続 <b>{sum_rec:,.0f}</b>万', open_tbl,
        warn="金額未入力(NULL)は0扱い。桁違いの仮入力・重複商談があると総額が跳ねる。単位が『万円』で統一されているか要確認。")

    # 5) 展示会ファネル（クライアント側フィルタ・分類別 #27）。展示会/分類の絞り込みは遷移なしで即時。
    import cowork.weekly_report as _wr
    _today = _today_jst().isoformat()
    exh_all = _wr.exhibition_deal_rows(con)
    for _r in exh_all:
        _r["_bucket"] = _wr.classify_exhibition_deal(_r, _today)
    _bucket_label = dict(_wr.EXH_BUCKETS)
    _bucket_order = {k: i for i, (k, _) in enumerate(_wr.EXH_BUCKETS)}
    exh_all.sort(key=lambda r: (_bucket_order.get(r["_bucket"], 99),
                                -(r.get("mtg") or 0), r.get("id") or 0))
    _exh_names = sfa_db.list_exhibition_names(con)
    _total_all = len(exh_all)
    # フィルタ選択肢（すべてクライアント側）
    _exh_sel = ('<option value="__all__">全展示会</option><option value="__none__">（未設定）</option>'
                + "".join(f'<option value="{_esc(n)}">{_esc(n)}</option>' for n in _exh_names))
    _bkt_sel = ('<option value="__all__">全分類</option>'
                + "".join(f'<option value="{k}">{_esc(lbl)}</option>' for k, lbl in _wr.EXH_BUCKETS))
    _cnt_html = "".join(
        f'<span style="display:inline-block;margin:0 14px 4px 0;font-size:12px">'
        f'{_esc(lbl)}: <b id="exc_{k}">0</b></span>' for k, lbl in _wr.EXH_BUCKETS)
    _js_cnt = "".join(
        f"var _c=document.getElementById('exc_{k}');if(_c)_c.textContent=(counts['{k}']||0);"
        for k, _l in _wr.EXH_BUCKETS)
    _thead = "".join(_sticky_th(h) for h in
                     ["区分", "展示会", "案件名", "アカウント", "面談", "ステージ", "状態", "終了理由", "次回MS", "開発"])
    _exrows = "".join(
        f'<tr class="exrow" data-exh="{_esc(r.get("exhibition_name") or "")}" data-bucket="{r["_bucket"]}">'
        f'<td>{_esc(_bucket_label.get(r["_bucket"], r["_bucket"]))}</td>'
        f'<td>{_esc(r.get("exhibition_name") or "—")}</td>'
        f'<td>{_esc(r["deal_name"] or "—")}</td><td>{_esc(r["acc"] or "—")}</td>'
        f'<td>{r["mtg"] or 0}</td><td>{_esc(r["stage"] or "—")}</td>'
        f'<td>{"クローズ" if (r.get("status") == "closed") else "open"}</td>'
        f'<td>{_esc(r["close_reason"] or "—")}</td><td>{_esc(r["next_milestone_date"] or "—")}</td>'
        f'<td>{"✓" if r.get("has_dev") else "—"}</td></tr>'
        for r in exh_all) or '<tr><td colspan=10 class=muted>展示会由来(lead_pattern=Exh.)の商談はありません。</td></tr>'
    _exh_detail = f"""
      <div class="filter-row" style="margin:0 0 8px">
        <label style="font-size:13px">展示会: <select id="exhSel" onchange="exhFilter()">{_exh_sel}</select></label>
        <label style="font-size:13px">分類: <select id="bktSel" onchange="exhFilter()">{_bkt_sel}</select></label>
        <button class="btn sec" type="button"
          onclick="document.getElementById('exhSel').value='__all__';document.getElementById('bktSel').value='__all__';exhFilter()">クリア</button>
        <a class="btn sec" href="/exhibition-tagging">🎪 展示会名をタグ付け</a>
        <span id="exhVis" class="muted" style="font-size:12px;align-self:center"></span>
      </div>
      <div style="margin:0 0 10px;padding:8px 10px;background:#f8fafc;border-radius:6px">{_cnt_html}</div>
      <div style="overflow:auto;max-height:64vh"><table id="exhTbl" style="min-width:980px">
      <tr>{_thead}</tr>{_exrows}</table></div>
      <script>
      function exhFilter() {{
        var ex = document.getElementById('exhSel').value;
        var bk = document.getElementById('bktSel').value;
        var counts = {{}}; var cohort = 0; var vis = 0;
        document.querySelectorAll('#exhTbl tr.exrow').forEach(function(tr) {{
          var e = tr.getAttribute('data-exh') || ''; var b = tr.getAttribute('data-bucket') || '';
          var inC = (ex === '__all__') || (ex === '__none__' ? e === '' : e === ex);
          if (inC) {{ cohort++; counts[b] = (counts[b] || 0) + 1; }}
          var show = inC && (bk === '__all__' || b === bk);
          tr.style.display = show ? '' : 'none'; if (show) vis++;
        }});
        {_js_cnt}
        var v = document.getElementById('exhVis');
        if (v) v.textContent = '対象 ' + cohort + '件 / 表示 ' + vis + '件';
      }}
      exhFilter();
      </script>
    """
    sec_exh = _audit_section(
        "展示会ファネル（分類・展示会でフィルタ／画面遷移なし）",
        "lead_pattern='Exh.' が母集団。面談は『同一日=1面談』(日付なしは除外)。面談0回で次回MS当日以降=初回面談待ち。"
        "1次実施でも次回MSが『アポ』(2次商談)なら『2次面談アポ済』＝進捗扱い。展示会/分類の選択は即時フィルタ(遷移なし)。",
        f'総数 <b>{_total_all}</b>件（下の分類別カウントは選択中の展示会に連動）',
        _exh_detail,
        warn="母集団は『lead_pattern=Exh.』タグに全依存。『不成立(要検証)』＝面談0回だが初回面談待ちでない件"
             "(closed/次回MS無し等)。分類を選ぶとその区分の行だけ表示。展示会名は🎪タグ付けで付与。")

    # 6) ステージ別（open）
    stage_rows = con.execute(
        f"SELECT COALESCE(d.stage,'未設定') s, COUNT(*) n FROM deals d WHERE {open_cond} "
        f"GROUP BY d.stage ORDER BY n DESC").fetchall()
    stage_tbl = _audit_table(["ステージ", "件数"], [[_esc(r["s"]), str(r["n"])] for r in stage_rows])
    sec_stage = _audit_section(
        "ステージ別（進行中）", f"status='open'(またはNULL)を stage で集計。",
        f'<b>{sum(r["n"] for r in stage_rows)}</b>件', stage_tbl)

    week_input = (as_of.isoformat() if hasattr(as_of, "isoformat") else "")
    return f"""
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>🔍 数字パック 集計監査 <span class="muted" style="font-size:.6em">{ws} 〜 {we}</span></span>
        <a class="btn sec" href="/weekly-numbers">← 数字パックに戻る</a>
      </h2>
      <p class="muted" style="font-size:13px">各指標を「定義・値・元データ（実際に数えている行）」に展開しました。
      自動注入を有効化する前に、ここで集計の確からしさを確認します。定義に違和感があれば、その定義自体を直します（#27）。</p>
      <form method="get" action="/weekly-numbers/audit" class="filter-row">
        <label style="font-size:13px">対象週（この日を含む月〜日）:
          <input type="date" name="as_of" value="{week_input}"></label>
        <button class="btn sec" type="submit">この週で見る</button>
        <a class="btn sec" href="/weekly-numbers/audit">今週</a>
      </form>
    </div>
    {sec_mtg}{sec_deal}{sec_pipe}{sec_exh}{sec_stage}"""


def sync_health_page(con, theme_client) -> str:
    """SFA↔Hisho同期の整合性診断ページ（オンデマンドでHishoへ照会）。"""
    if theme_client is None:
        return """
        <div class="card">
          <h2>🔍 同期チェック</h2>
          <p class="muted" style="margin:0">テーマDB連携が無効です（THEME_API_TOKEN未設定）。</p>
        </div>"""
    diag = dev_project_link.diagnose_sync(theme_client, con)
    if diag.get("error"):
        return f"""
        <div class="card">
          <h2>🔍 同期チェック</h2>
          <p class="muted" style="margin:0">Hishoへの照会に失敗しました: {_esc(diag['error'])}</p>
        </div>"""

    def _section(title, items, render_row, hint):
        if not items:
            return (f'<div class="card"><h2>{title} <span class="stage" '
                    f'style="background:#dcfce7;color:#166534">0件</span></h2>'
                    f'<p class="muted" style="margin:0">問題なし。</p></div>')
        rows = "".join(render_row(x) for x in items)
        return (f'<div class="card"><h2>{title} <span class="stage" '
                f'style="background:#fee2e2;color:#991b1b">{len(items)}件</span></h2>'
                f'<p class="muted" style="margin:0 0 8px">{hint}</p>'
                f'<table>{rows}</table></div>')

    unsynced = _section(
        "開発案件: 未同期（Hishoへ未連携）", diag["dev_unsynced"],
        lambda p: f'<tr><td><a href="/dev-project/{p["id"]}/edit">{_esc(p.get("theme"))}</a></td>'
                  f'<td class="muted">{_esc(p.get("deal_name") or "—")}</td></tr>',
        "保存し直すか「同期失敗の再同期」で解消できます。")
    broken = _section(
        "開発案件: リンク切れ（Hisho側に対応行が無い）", diag["dev_broken_link"],
        lambda p: f'<tr><td><a href="/dev-project/{p["id"]}/edit">{_esc(p.get("theme"))}</a></td>'
                  f'<td class="muted">hisho_id={p.get("hisho_id")}</td></tr>',
        "Hisho側で削除された可能性。保存し直すと再作成されます。")
    orphan = _section(
        "Hisho側の孤児（SFAに存在しない開発案件）", diag["dev_orphan_hisho"],
        lambda r: f'<tr><td>Hisho id={r.get("id")}</td><td class="muted">sfa_id={r.get("sfa_id")}</td></tr>',
        "SFAで削除済みだがHisho側に残っている行。Hishoダッシュボードで確認・削除してください。")
    deal_unsynced = _section(
        "商談: テーマDB未同期", diag["deal_unsynced"],
        lambda d: f'<tr><td><a href="/deal/{d["id"]}">{_esc(d.get("deal_name"))}</a></td><td></td></tr>',
        "商談一覧の「テーマDB未同期 N件を同期」で解消できます。")

    # 記録済みの同期失敗（エラー文つき）。再同期しても消えない＝原因が要る失敗の中身を見える化。
    _kind_label = {"deal": "商談", "dev_project": "開発案件", "dev_project_delete": "開発案件(削除)"}

    def _fail_row(frec):
        kind = frec.get("kind", "")
        ref = frec.get("ref_id")
        label = _kind_label.get(kind, kind)
        if kind == "deal":
            ref_html = f'<a href="/deal/{ref}">商談 #{ref}</a>'
        elif kind == "dev_project":
            ref_html = f'<a href="/dev-project/{ref}/edit">開発案件 #{ref}</a>'
        else:
            ref_html = f'{_esc(label)} #{ref}'
        return (f'<tr><td style="white-space:nowrap">{_esc(label)}<br>{ref_html}</td>'
                f'<td style="font-size:.85em;color:#991b1b;word-break:break-all">{_esc(frec.get("error") or "—")}</td>'
                f'<td class="muted" style="white-space:nowrap;font-size:.8em">{_esc((frec.get("created_at") or "")[:16])}</td></tr>')

    failures = sfa_db.list_sync_failures(con)
    if failures:
        fail_rows = "".join(_fail_row(fr) for fr in failures)
        failures_html = (
            '<div class="card"><h2>同期失敗の詳細 '
            f'<span class="stage" style="background:#fee2e2;color:#991b1b">{len(failures)}件</span></h2>'
            '<p class="muted" style="margin:0 0 8px">再同期しても解消しない失敗の中身です。エラー文から原因を特定してください。</p>'
            '<table><tr><th>対象</th><th>エラー</th><th>記録日時</th></tr>' + fail_rows + '</table></div>'
        )
    else:
        failures_html = ('<div class="card"><h2>同期失敗の詳細 '
                         '<span class="stage" style="background:#dcfce7;color:#166534">0件</span></h2>'
                         '<p class="muted" style="margin:0">記録された同期失敗はありません。</p></div>')

    # #67: ステージ='失注' の残存商談。失注運用は「クローズ＋理由=失注」へ一本化済みのため、
    #      旧ステージ='失注'（特にopenのまま滞留中）をクローズへ移行する片付けボタンを出す。
    lost_all = sfa_db.list_lost_stage_deals(con)
    lost_open = [d for d in lost_all if (d.get("status") or "open") != "closed"]
    if lost_all:
        _lost_rows = "".join(
            f'<tr><td><a href="/deal/{d["id"]}">{_esc(d.get("deal_name") or "—")}</a>'
            f'<span class="muted" style="font-size:.85em"> / {_esc(d.get("account_name") or "")}</span></td>'
            f'<td class="muted" style="white-space:nowrap">{"クローズ済" if (d.get("status") or "open")=="closed" else "🟠 未クローズ(滞留)"}</td>'
            f'<td class="muted">{_esc(d.get("close_reason") or "—")}</td></tr>'
            for d in lost_all)
        _mig_btn = (
            f'<form method="post" action="/deals/migrate_lost_stage" style="margin:8px 0 0">'
            f'<button class="btn" style="background:#c53030" '
            f'onclick="return confirm(\'ステージ=失注の商談 {len(lost_all)}件を『クローズ＋終了理由=失注』にし、'
            f'各社を『フォロー中リード』として作成/再活性化します（Hisho再同期込み）。同名リードは再利用・冪等。よろしいですか？\')">'
            f'🧹 失注 {len(lost_all)}件をクローズ＋フォロー中リード化</button></form>')
        lost_stage_html = (
            '<div class="card"><h2>ステージ「失注」の残存商談 '
            f'<span class="stage" style="background:{"#fee2e2;color:#991b1b" if lost_open else "#fef9c3;color:#854d0e"}">'
            f'{len(lost_all)}件（未クローズ {len(lost_open)}）</span></h2>'
            '<p class="muted" style="margin:0 0 8px">失注は「クローズ＋終了理由=失注」に一本化されました（ステージ選択肢からは撤廃済み）。'
            '下のボタンで旧ステージ=失注の商談を一括クローズ＋終了理由=失注にし、各社を「フォロー中リード」として'
            '作成/再活性化（再接触できるように）＋Hisho再同期します。同名リードは再利用するため冪等（何度押しても安全）。</p>'
            f'<table><tr><th>商談</th><th>状態</th><th>終了理由</th></tr>{_lost_rows}</table>{_mig_btn}</div>')
    else:
        lost_stage_html = ""

    total = (len(diag["dev_unsynced"]) + len(diag["dev_broken_link"])
             + len(diag["dev_orphan_hisho"]) + len(diag["deal_unsynced"]))
    banner_bg = "#dcfce7" if total == 0 else "#fff7ed"
    banner_border = "#86efac" if total == 0 else "#fed7aa"
    summary = "同期は健全です（ズレなし）。" if total == 0 else f"合計 {total} 件のズレを検出しました。"
    return f"""
    <div class="card" style="background:{banner_bg};border:1.5px solid {banner_border}">
      <p style="margin:0 0 10px"><a class="btn sec" href="/deals">← 商談一覧へ</a></p>
      <h2>🔍 SFA ↔ Hisho 同期チェック</h2>
      <p class="muted" style="margin:0">{summary} このページを開くたびにHishoへ最新状態を照会します。</p>
    </div>
    {lost_stage_html}{unsynced}{broken}{orphan}{deal_unsynced}{failures_html}"""


def backups_page(db_path: str) -> str:
    """DBバックアップの一覧・復元・ダウンロード画面。"""
    backups = sfa_db.list_backups(db_path)

    def _fmt_size(n):
        return f"{n/1024:.0f} KB" if n < 1024 * 1024 else f"{n/1024/1024:.1f} MB"

    def _fmt_time(ts):
        import time as _t
        return _t.strftime("%Y-%m-%d %H:%M", _t.localtime(ts))

    rows = "".join(
        f'<tr>'
        f'<td>{_esc(b["name"])}</td>'
        f'<td class="muted" style="white-space:nowrap">{_fmt_time(b["mtime"])}</td>'
        f'<td class="muted">{_fmt_size(b["size"])}</td>'
        f'<td><a class="btn sec" style="font-size:11px;padding:4px 8px" '
        f'href="/backups/download?name={urllib.parse.quote(b["name"])}">⬇ DL</a></td>'
        f'<td><form method="post" action="/backups/restore" style="display:inline" '
        f'onsubmit="return confirm(\'このバックアップで現在のDBを置き換えます。\\n復元前の現DBは自動退避されますが、実行後は元に戻すのに再度復元操作が必要です。\\n本当に復元しますか？\')">'
        f'<input type="hidden" name="name" value="{_esc(b["name"])}">'
        f'<button class="btn" style="font-size:11px;padding:4px 8px;background:#c53030">↩ 復元</button></form></td>'
        f'</tr>'
        for b in backups
    ) or '<tr><td colspan=5 class=muted>バックアップがまだありません</td></tr>'

    return f"""
    <div class="card" style="background:#f0f4f8;border:1.5px solid #d4dae4">
      <h2>🗄 DBバックアップ</h2>
      <p class="muted" style="margin:0">サーバー起動時に日次で自動取得されます。ここから手動バックアップ・
      復元・ダウンロードができます。復元は現在のDBを丸ごと置き換える操作です（復元前に現DBは自動退避されます）。</p>
    </div>
    <div class="card">
      <div style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px;margin-bottom:10px">
        <span class="muted">保存されているバックアップ（{len(backups)}件）</span>
        <form method="post" action="/backups/create" style="display:inline">
          <button class="btn sec" type="submit">＋ 今すぐバックアップを作成</button>
        </form>
      </div>
      <div style="overflow:auto;max-height:70vh">
      <table>
        <tr>{_sticky_th('ファイル名')}{_sticky_th('作成日時')}{_sticky_th('サイズ')}{_sticky_th('')}{_sticky_th('')}</tr>
        {rows}
      </table>
      </div>
    </div>"""


def _tech_seed_block(l1: str, leaves: list) -> str:
    """技術シードマスタの1カテゴリ(L1)ブロック: カテゴリ名＋L2シード(1行ずつ)のtextarea。"""
    lines = "\n".join(leaves)
    return (
        '<div class="ts-block" style="border:1px solid #e6e9f0;border-radius:8px;padding:12px;'
        'margin-bottom:10px;background:#fafbfc">'
        '<div style="display:flex;gap:8px;align-items:center;margin-bottom:6px">'
        f'<input name="l1_name[]" value="{_esc(l1)}" placeholder="カテゴリ名（例: 研究テーマ(SCM)）" '
        'style="flex:1;font-weight:700">'
        '<button type="button" class="btn sec" style="font-size:11px;white-space:nowrap" '
        "onclick=\"this.closest('.ts-block').remove()\">このカテゴリを削除</button>"
        '</div>'
        f'<textarea name="l2_lines[]" rows="6" placeholder="シードを1行ずつ入力" '
        f'style="font-size:13px">{_esc(lines)}</textarea>'
        '</div>')


def tech_seed_master_page(con) -> str:
    """技術シード マスタ（ツリー: L1カテゴリ→L2シード）の編集ページ（#60）。"""
    tree = sfa_db.get_tech_seed_tree(con)
    blocks = "".join(_tech_seed_block(l1, leaves) for l1, leaves in tree.items()) \
        or _tech_seed_block("", [])
    return f"""
    <div class="card" style="max-width:920px">
      <p style="margin:0 0 10px"><a class="btn sec" href="/dev-projects">← 開発案件一覧へ</a></p>
      <h2>技術シード マスタ（カテゴリ=L1 ／ シード=L2）</h2>
      <p class="muted" style="margin:0 0 12px;font-size:12px">カテゴリ(L1)ごとに、シード(L2)を1行ずつ入力します。
        事業が増えたら「＋カテゴリを追加」で <b>研究テーマ(◯◯)</b> を足せます。カテゴリ名が空の枠は保存時に削除されます。</p>
      <form method="post" action="/tech-seed-master/save">
        <div id="tsBlocks">{blocks}</div>
        <button type="button" class="btn sec" onclick="tsAddBlock()">＋ カテゴリを追加</button>
        <template id="tsBlockTpl">{_tech_seed_block("", [])}</template>
        <div style="margin-top:16px">
          <button class="btn" type="submit">保存</button>
          <a class="btn sec" href="/dev-projects">キャンセル</a>
        </div>
      </form>
    </div>
    <script>
    function tsAddBlock() {{
      var t = document.getElementById('tsBlockTpl');
      document.getElementById('tsBlocks').insertAdjacentHTML('beforeend', t.innerHTML);
    }}
    </script>"""


def _seed_checkboxes_grouped(con, field_name: str, selected: set, tree: dict | None = None) -> str:
    """技術シードのチェックボックス群をL1カテゴリでグループ表示（一括タグ付け用・コンパクト）。"""
    tree = tree if tree is not None else sfa_db.get_tech_seed_tree(con)
    parts = []
    for l1, leaves in tree.items():
        if not leaves:
            continue
        boxes = "".join(
            f'<label style="display:inline-flex;align-items:center;gap:3px;margin:0 10px 4px 0;'
            f'font-weight:400;font-size:12px;white-space:nowrap">'
            f'<input type="checkbox" name="{field_name}" value="{_esc(s)}"{" checked" if s in selected else ""}'
            f' style="width:auto">{_esc(s)}</label>' for s in leaves)
        parts.append(
            f'<div style="margin-bottom:4px"><span style="font-size:10px;font-weight:700;color:#4338ca;'
            f'margin-right:6px">{_esc(l1)}</span>{boxes}</div>')
    return "".join(parts) or '<span class="muted" style="font-size:11px">シード未登録</span>'


def tech_seed_tagging_page(con) -> str:
    """技術シード 一括タグ付け（全開発案件を1画面で、必要な技術シードをまとめてチェック→保存）（#58）。"""
    projects = sfa_db.list_dev_projects(con)
    tree = sfa_db.get_tech_seed_tree(con)
    has_seeds = any(v for v in tree.values())
    _rows = []
    for p in projects:
        pid = p["id"]
        sel = {s for s in (p.get("tech_seeds") or "").split(",") if s}
        _rows.append(
            f'<tr><td style="min-width:230px;vertical-align:top">'
            f'<a href="/dev-project/{pid}/edit"><b>{_esc(p.get("theme") or "(無題)")}</b></a>'
            f'<div class="muted" style="font-size:11px">{_esc(p.get("account_name") or "-")}／'
            f'{_esc(p.get("deal_name") or "-")}・{_esc(p.get("status") or "")}</div>'
            f'<input type="hidden" name="pids[]" value="{pid}"></td>'
            f'<td style="vertical-align:top">{_seed_checkboxes_grouped(con, f"seeds__{pid}", sel, tree)}</td></tr>')
    body = "".join(_rows) or '<tr><td colspan=2 class=muted>開発案件がありません。</td></tr>'
    warn = ("" if has_seeds else
            '<p class="flash" style="background:#fef3c7;color:#92400e">技術シードのマスタが空です。'
            '先に「管理 ▾ → 🌱 技術シードマスタ」でシードを登録してください。</p>')
    return f"""
    <form method="post" action="/tech-seed-tagging/save" id="tsTagForm">
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>技術シード 一括付け（{len(projects)}件）</span>
        <a class="btn sec" href="/dev-projects">← 開発案件一覧へ</a></h2>
      {_save_bar('tsTagForm', cancel_url='/dev-projects', label='💾 まとめて保存')}
      {warn}
      <p class="muted" style="font-size:12px">各開発案件に必要な技術シードをチェックして「まとめて保存」。チェックを外して保存すると解除されます（表示中の全案件が対象）。</p>
      <div style="overflow:auto;max-height:72vh"><table style="min-width:900px">
        <tr>{_sticky_th('開発案件', '240px')}{_sticky_th('必要な技術シード')}</tr>
        {body}
      </table></div>
      <div style="margin-top:12px"><button class="btn" type="submit">💾 まとめて保存</button>
        <a class="btn sec" href="/dev-projects">キャンセル</a></div>
    </div>
    </form>"""


_TASK_CAT_PALETTE = ["#4f46e5", "#0891b2", "#059669", "#b45309", "#be185d", "#7c3aed",
                     "#0d9488", "#c2410c", "#4338ca", "#15803d", "#9333ea", "#6b7280"]


def _task_category_color(name: str) -> str:
    """タスク種類名から決定的に色を割り当てる（プロセス間で安定）。"""
    if not name:
        return "#6b7280"
    return _TASK_CAT_PALETTE[sum(ord(c) for c in name) % len(_TASK_CAT_PALETTE)]


_TASKS_JS = """
<style>
/* 画面幅いっぱいに5列を伸縮配置。広い画面では横スクロールなし、狭い画面(=合計が入り切らない時)だけ横スクロール */
#taskBoard{display:flex;gap:10px;overflow-x:auto;padding-bottom:8px;align-items:flex-start}
/* 各列を内部スクロールにし、ヘッダ(h3)は列上部に常駐＝縦に増えてもタイトルが見える（ナビと干渉しない） */
.task-col{flex:1 1 0;min-width:185px;background:#f1f4f9;border-radius:10px;padding:8px 7px 10px;display:flex;flex-direction:column;max-height:calc(100vh - 200px)}
.task-col h3{font-size:13px;margin:0 0 6px;flex:none}
.tc-col-body{overflow-y:auto;flex:1 1 auto;min-height:18px}
.task-card{background:#fff;border:1px solid #e6e9f0;border-radius:8px;padding:5px 8px;margin-bottom:6px;box-shadow:0 1px 2px rgba(0,0,0,.05)}
.task-card.saved{outline:2px solid #10b981;transition:outline .15s}
.task-card.pinned{border-color:#f59e0b}
.tc-head{display:flex;align-items:center;gap:6px;cursor:pointer}
.tc-dot{width:9px;height:9px;border-radius:50%;flex:none}
.tc-ttl{flex:1;font-weight:600;font-size:12px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
.tc-duem{font-size:10px;white-space:nowrap}
.tc-car{font-size:10px;color:#94a3b8;transition:transform .15s}
.task-card.open .tc-car{transform:rotate(90deg)}
.tc-pin{border:none;background:transparent;color:#d1d5db;cursor:pointer;font-size:12px;padding:0;line-height:1}
.tc-pin.on{color:#f59e0b}
.tc-body{display:none;margin-top:5px}
.task-card.open .tc-body{display:block}
.tc-q{font-size:9px;padding:1px 5px;border:1px solid #e2e8f0;border-radius:4px;background:#fff;cursor:pointer}
.tc-q.rec{border-color:#93c5fd;color:#2563eb;font-weight:600}
.tc-ai{font-size:10px;padding:1px 5px;border:1px solid #e2e8f0;border-radius:4px;background:#fff;cursor:pointer}
.tc-lbl{font-size:9px;color:#94a3b8}
.tc-mini{display:flex;flex-wrap:wrap;align-items:center;gap:5px;margin-top:3px}
.tc-mini .tc-actions{margin:0}
.m-pj{font-size:9px;background:#eef2ff;color:#4338ca;border-radius:4px;padding:1px 5px;max-width:120px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
.m-asg{font-size:10px;color:#475569}
.m-due{font-size:10px;font-weight:600;white-space:nowrap}
.tc-foot{display:flex;gap:10px;align-items:center;margin-top:7px;padding-top:6px;border-top:1px dashed #eef1f5}
.tc-foot .del{border:none;background:transparent;color:#c53030;cursor:pointer;font-size:11px;padding:0}
.tc-disc{border:1px solid #e2e8f0;background:#fff;border-radius:5px;cursor:pointer;font-size:11px;padding:2px 8px}
.tc-sum{font-size:10px;color:#3730a3;background:#eef2ff;border-radius:4px;padding:3px 6px;margin:3px 0;cursor:pointer;max-height:44px;overflow:hidden}
.np-sum{background:#eef2ff;border-radius:6px;padding:6px 8px;font-size:12px;color:#1e3a8a;margin:6px 0;white-space:pre-wrap;max-height:200px;overflow:auto}
.np-sum-h{font-weight:700;font-size:11px;margin-bottom:3px}
.pj-strip{display:flex;gap:6px;overflow-x:auto;padding:2px 0 8px;margin-bottom:6px}
.pj-chip{flex:none;display:flex;align-items:center;gap:4px;font-size:11px;text-decoration:none;color:#334155;background:#f1f5f9;border:1px solid #e2e8f0;border-radius:20px;padding:3px 10px;white-space:nowrap}
.pj-chip.active{background:#dbeafe;border-color:#93c5fd;color:#1e40af}
.pj-chip.mng{background:#fff;color:#64748b}
.pj-dot{width:8px;height:8px;border-radius:50%}
.pj-cnt{color:#64748b}
.tc-title{width:100%;border:none;font-weight:600;font-size:13px;padding:2px;background:transparent}
.tc-title:hover,.tc-title:focus{background:#f1f5f9;border-radius:4px;outline:none}
.tc-na{display:flex;align-items:center;gap:3px;margin:3px 0}
.tc-na>span{color:#0369a1;font-size:11px}
.tc-na input{flex:1;border:none;font-size:11px;color:#0369a1;padding:2px;background:transparent}
.tc-na input:hover,.tc-na input:focus{background:#f0f9ff;border-radius:4px;outline:none}
.tc-na.empty input{color:#c53030}
.tc-na.empty input::placeholder{color:#c53030}
.tc-meta{display:flex;flex-wrap:wrap;gap:4px;align-items:center;margin:4px 0}
.tc-sel,.tc-due{font-size:10px;padding:1px 2px;border:1px solid #e2e8f0;border-radius:4px;max-width:100%;background:#fff}
.tc-notes{font-size:10px;color:#64748b;background:#f8fafc;border-radius:4px;padding:3px 6px;margin:3px 0;cursor:pointer;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
.tc-notes:hover{background:#eef2f7}
.tc-actions{display:flex;gap:4px;flex-wrap:wrap;align-items:center;margin-top:5px;min-height:24px}
.tc-actions button{font-size:10px;padding:2px 8px;border-radius:5px;border:1px solid #cbd5e1;background:#fff;cursor:pointer}
.tc-actions .go{background:#2563eb;color:#fff;border-color:#2563eb}
.tc-actions .done{background:#10b981;color:#fff;border-color:#10b981}
.tc-actions .del{color:#c53030;border:none;background:transparent;margin-left:auto}
.tc-edit{font-size:10px;color:#94a3b8;text-decoration:none}
#notesBackdrop{position:fixed;inset:0;z-index:9998;display:none;background:rgba(15,23,42,.15)}
#notesPop{position:fixed;z-index:9999;display:none;background:#fff;border:1px solid #cbd5e1;border-radius:10px;box-shadow:0 8px 30px rgba(0,0,0,.18);width:340px;max-width:92vw;max-height:70vh;overflow:auto;padding:12px}
</style>
<script>
var _TC_STATUSES=["受信箱","未着手","対応中","保留","完了"];
function _tcEsc(s){return (s==null?'':String(s)).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');}
function _tcFlash(id){var c=document.getElementById('tc-'+id); if(!c)return; c.classList.add('saved'); setTimeout(function(){c.classList.remove('saved');},600);}
function tcActions(id,status){
  if(status==='受信箱') return '<span style="font-size:9px;color:#94a3b8">担当＋期限で自動整理</span>';
  if(status==='未着手') return '<button class="go" onclick="taskField('+id+',\\'status\\',\\'対応中\\')">▶ 開始</button>';
  if(status==='対応中') return '<button onclick="taskField('+id+',\\'status\\',\\'保留\\')">⏸ 保留</button><button class="done" onclick="taskField('+id+',\\'status\\',\\'完了\\')">✓ 完了</button>';
  if(status==='保留') return '<button class="go" onclick="taskField('+id+',\\'status\\',\\'対応中\\')">▶ 再開</button>';
  if(status==='完了') return '<button onclick="taskField('+id+',\\'status\\',\\'対応中\\')">↩ 戻す</button>';
  return '';
}
function tcRenderActions(card){var id=card.id.slice(3); var st=card.getAttribute('data-status'); var a=card.querySelector('.tc-actions'); if(a)a.innerHTML=tcActions(id,st);}
function tcCounts(){_TC_STATUSES.forEach(function(s){var body=document.querySelector('[data-col="'+s+'"]'); var cnt=document.querySelector('[data-count="'+s+'"]'); if(body&&cnt)cnt.textContent=body.querySelectorAll('.task-card').length;});}
function tcMove(id,status){var card=document.getElementById('tc-'+id); if(!card)return; var body=document.querySelector('[data-col="'+status+'"]'); if(!body)return; card.setAttribute('data-status',status); body.appendChild(card); tcRenderActions(card); tcCounts();}
function taskField(id,field,value){
  return fetch('/task/'+id+'/field',{method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},
    body:'field='+encodeURIComponent(field)+'&value='+encodeURIComponent(value)})
   .then(function(r){return r.json();}).then(function(d){
     if(!d.ok){ alert('更新エラー: '+(d.error||'')); return; }
     _tcFlash(id);
     var card=document.getElementById('tc-'+id);
     if(field==='next_action'&&card){ var na=card.querySelector('.tc-na'); if(na)na.classList.toggle('empty',!(value||'').trim()); }
     if(d.status&&card&&card.getAttribute('data-status')!==d.status){ tcMove(id,d.status); }
   }).catch(function(){ alert('通信エラー'); });
}
function taskDelete(id){ if(!confirm('このタスクを削除しますか？')) return;
  fetch('/task/'+id+'/delete',{method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},body:'ajax=1'})
   .then(function(r){return r.json();}).then(function(d){ if(d.ok){ var c=document.getElementById('tc-'+id); if(c)c.remove(); tcCounts(); } });
}
function openNotes(id,kind){
  kind=kind||'progress'; var isD=(kind==='discussion');
  var pop=document.getElementById('notesPop'), bd=document.getElementById('notesBackdrop');
  pop.setAttribute('data-tid',id); pop.setAttribute('data-kind',kind);
  var title=isD?'💬 議論メモ':'📝 進捗ログ';
  var ph=isD?'議論・検討メモを追記（たっぷり書けます）':'進捗を追記';
  var inp=isD
    ? '<textarea id="noteInput" rows="12" placeholder="'+ph+'" style="width:100%;min-height:240px;font-size:13px;padding:8px;box-sizing:border-box;line-height:1.6" onkeydown="if((event.ctrlKey||event.metaKey)&&event.key===&#39;Enter&#39;){event.preventDefault();addNote();}"></textarea><div style="display:flex;justify-content:space-between;align-items:center;margin-top:4px"><span class="muted" style="font-size:10px">Enter=改行 ／ Ctrl+Enter=確定 ／ Esc=閉じる</span><button class="btn np-add" style="font-size:11px;padding:3px 10px" onclick="addNote()">追記＋サマリ更新</button></div>'
    : '<div style="display:flex;gap:6px"><input id="noteInput" placeholder="'+ph+'" style="flex:1;font-size:12px;padding:4px 6px" onkeydown="if(event.key===&#39;Enter&#39;)addNote()"><button class="btn np-add" style="font-size:11px;padding:3px 8px" onclick="addNote()">追記</button></div>';
  pop.innerHTML='<div style="display:flex;justify-content:space-between;align-items:center"><b style="font-size:13px">'+title+'</b><span onclick="closeNotes()" style="cursor:pointer;color:#94a3b8">✕</span></div>'
    +(isD?'<div id="notesSummary" class="np-sum"></div>':'')
    +'<div id="notesBody" class="muted" style="font-size:12px">読み込み中…</div>'
    +'<div style="margin-top:8px">'+inp+'</div>';
  bd.style.display='block'; pop.style.display='block'; pop.style.width=(isD?'min(600px,94vw)':'340px');
  pop.style.left=Math.max(8,(window.innerWidth-pop.offsetWidth)/2)+'px';
  pop.style.top=Math.max(8,(window.innerHeight-pop.offsetHeight)/2)+'px';
  fetch('/task/'+id+'/notes?kind='+kind).then(function(r){return r.json();}).then(function(d){ _tcRenderNotes(d.notes||[]); if(isD)_npSummary(d.summary); });
}
function _npSummary(s){ var el=document.getElementById('notesSummary'); if(!el)return;
  el.innerHTML=s?('<div class="np-sum-h">🧠 AIサマリ</div>'+_tcEsc(s).replace(/\\n/g,'<br>')):'<span class="muted" style="font-size:11px">メモを追記するとAIサマリが自動生成されます。</span>'; }
function _tcRenderNotes(notes){
  var body=document.getElementById('notesBody'); if(!body)return;
  if(!notes.length){ body.innerHTML='<i>まだメモはありません。下の欄から追記できます。</i>'; return; }
  body.innerHTML=notes.map(function(n){ return '<div style="border-left:3px solid #cbd5e1;padding:2px 0 2px 8px;margin:6px 0"><div style="font-size:10px;color:#94a3b8">'+(n.created_at||'').slice(0,16)+(n.author?(' ・'+_tcEsc(n.author)):'')+'</div>'+_tcEsc(n.body).replace(/\\n/g,'<br>')+'</div>'; }).join('');
}
function addNote(){
  var pop=document.getElementById('notesPop'), id=pop.getAttribute('data-tid'), kind=pop.getAttribute('data-kind')||'progress';
  var inp=document.getElementById('noteInput'), v=(inp.value||'').trim(); if(!v)return;
  var btn=pop.querySelector('.np-add'); if(btn){btn.disabled=true;btn.textContent='保存中…';}
  fetch('/task/'+id+'/note',{method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},body:'ajax=1&kind='+kind+'&body='+encodeURIComponent(v)})
   .then(function(r){return r.json();}).then(function(d){ if(btn){btn.disabled=false;btn.textContent=(kind==='discussion'?'追記＋サマリ更新':'追記');}
     if(!d.ok)return; inp.value=''; _tcRenderNotes(d.notes||[]); var card=document.getElementById('tc-'+id);
     if(kind==='discussion'){ _npSummary(d.summary);
       if(card&&d.summary){ var sc=card.querySelector('.tc-sum'); if(sc){sc.innerHTML='🧠 '+_tcEsc(d.summary).slice(0,130);} } }
     else { if(card){ var chip=card.querySelector('.tc-notes'); if(chip)chip.textContent='📝 '+v.slice(0,40); if(d.status&&card.getAttribute('data-status')!==d.status)tcMove(id,d.status); } }
   });
}
function closeNotes(){ var i=document.getElementById('noteInput');
  if(i&&(i.value||'').trim()&&!confirm('入力中のメモを破棄して閉じますか？')) return;
  document.getElementById('notesPop').style.display='none'; document.getElementById('notesBackdrop').style.display='none'; }
document.addEventListener('keydown',function(e){ if(e.key==='Escape'){ var p=document.getElementById('notesPop'); if(p&&p.style.display==='block'){ e.preventDefault(); closeNotes(); } } });
function taskFilter(){ var q=(document.getElementById('taskSearch').value||'').toLowerCase().trim();
  document.querySelectorAll('.task-card').forEach(function(c){ c.style.display=(!q||(c.getAttribute('data-search')||'').indexOf(q)>=0)?'':'none'; }); }
var _TC_IDLE_MS=45000;
function _tcTouch(card){ if(card)card.setAttribute('data-touch',Date.now()); }
// カード内のどこをクリックしても展開（操作系＝ボタン/リンク/入力は除く）。開いているときは何もしない。
function tcCardClick(ev,card){ if(ev.target.closest('button,a,select,input,textarea,label')) return;
  if(!card.classList.contains('open')){ card.classList.add('open'); _tcTouch(card); } }
function tcSyncTitle(id,v){ var c=document.getElementById('tc-'+id); if(c){ var t=c.querySelector('.tc-ttl'); if(t)t.textContent=v; } }
function _tcUrg(due){ var D=window._TC||{}; if(!due)return '#cbd5e1'; if(due<D.today)return '#dc2626'; if(due<=D.d3)return '#f59e0b'; if(due<=D.weekend)return '#eab308'; return '#94a3b8'; }
function tcDue(id,ds){ var c=document.getElementById('tc-'+id); if(c){ var col=_tcUrg(ds); var inp=c.querySelector('.tc-due'); if(inp){inp.value=ds;inp.style.color=col;} var m=c.querySelector('.tc-duem'); if(m){m.textContent=ds||'期限なし';m.style.color=col;} var d=c.querySelector('.tc-dot'); if(d)d.style.background=col; } taskField(id,'due_date',ds); }
function tcPin(id){ var c=document.getElementById('tc-'+id); var cur=c.getAttribute('data-pinned')==='1'; var nv=cur?'0':'1';
  fetch('/task/'+id+'/field',{method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},body:'field=pinned&value='+nv})
   .then(function(r){return r.json();}).then(function(d){ if(!d.ok)return; c.setAttribute('data-pinned',nv); c.classList.toggle('pinned',nv==='1'); var b=c.querySelector('.tc-pin'); if(b)b.classList.toggle('on',nv==='1'); _tcFlash(id); }); }
function tcAiCat(id){ var c=document.getElementById('tc-'+id);
  fetch('/task/'+id+'/ai-category',{method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},body:'x=1'})
   .then(function(r){return r.json();}).then(function(d){ if(!d.ok||!d.category){ if(d&&!d.category)alert('AIが種類を判定できませんでした'); return; }
     var sel=c.querySelector('select[data-field="category"]'); if(sel){ sel.value=d.category; } _tcFlash(id); }); }
document.addEventListener('DOMContentLoaded',function(){ document.querySelectorAll('.task-card').forEach(tcRenderActions);
  var bd=document.getElementById('notesBackdrop'); if(bd)bd.addEventListener('click',closeNotes);
  // 「当該カードの外」をクリックしたら、そのカードを閉じる（他カードをクリックした時も閉じる）
  document.addEventListener('click',function(e){ if(e.target.closest('#notesPop'))return;
    document.querySelectorAll('.task-card.open').forEach(function(c){ if(!c.contains(e.target)) c.classList.remove('open'); }); });
  // 開いているカード内の操作で無操作タイマーをリセット
  var board=document.getElementById('taskBoard');
  if(board){ ['input','change','click','keydown'].forEach(function(evt){ board.addEventListener(evt,function(e){ var c=e.target.closest('.task-card.open'); if(c)_tcTouch(c); }); }); }
  // 一定時間(45秒)操作の無い開いたカードは自動でコンパクトに戻す
  setInterval(function(){ var now=Date.now(); document.querySelectorAll('.task-card.open').forEach(function(c){ var t=parseInt(c.getAttribute('data-touch')||'0'); if(t&&now-t>_TC_IDLE_MS)c.classList.remove('open'); }); },5000);
  // Slack等からの直リンク(#tc-<id>)で該当タスクを開いてスクロール
  if(location.hash&&location.hash.indexOf('#tc-')===0){ var fc=document.getElementById(location.hash.slice(1));
    if(fc){ fc.classList.add('open'); _tcTouch(fc); fc.style.outline='2px solid #2563eb';
      setTimeout(function(){fc.scrollIntoView({behavior:'smooth',block:'center'});},150);
      setTimeout(function(){fc.style.outline='';},2500); } }
});
</script>
<div id="notesBackdrop"></div><div id="notesPop"></div>
"""


def _tc_sel(tid: int, field: str, values: list, cur, placeholder: str, blank: bool = True) -> str:
    """カード上のインライン編集<select>（変更で即保存・遷移なし）。
    現値がマスタ未登録でも選択肢に含めて必ず表示する（値が"消える"のを防ぐ）。"""
    cur_s = str(cur or "")
    vals = [str(v) for v in values]
    if cur_s and cur_s not in vals:
        vals = [cur_s] + vals  # マスタ外の現値を先頭に補って表示
    opts = (f'<option value="">{html.escape(placeholder)}</option>' if blank else "")
    for v in vals:
        opts += (f'<option value="{html.escape(v)}"'
                 f'{" selected" if v == cur_s else ""}>{html.escape(v)}</option>')
    return (f'<select class="tc-sel" data-field="{field}" title="{html.escape(placeholder)}" '
            f'onchange="taskField({tid},&#39;{field}&#39;,this.value)">{opts}</select>')


def _task_cat_optgroups(cats: list, selected) -> str:
    """種類<select>用の<optgroup>群（L1見出し＋L2オプション）。ツリー外はまとめて『未分類』へ。"""
    sel = str(selected or "")
    used = set()
    out = ""
    for l1, leaves in sfa_db.TASK_CATEGORY_TREE.items():
        opts = ""
        for l2 in leaves:
            if l2 in cats:
                used.add(l2)
                opts += (f'<option value="{html.escape(l2)}"'
                         f'{" selected" if l2 == sel else ""}>{html.escape(l2)}</option>')
        if opts:
            out += f'<optgroup label="{html.escape(l1)}">{opts}</optgroup>'
    extra = [c for c in cats if c not in used]
    if extra:
        opts = "".join(f'<option value="{html.escape(c)}"{" selected" if c == sel else ""}>'
                       f'{html.escape(c)}</option>' for c in extra)
        out += f'<optgroup label="未分類">{opts}</optgroup>'
    return out


def _task_auto_triage(con, tid: int) -> str | None:
    """受信箱のタスクに担当＋期限が揃ったら自動で「未着手」へ整理する。現在のstatusを返す。"""
    t = sfa_db.get_task(con, tid)
    if not t:
        return None
    if (t.get("status") == "受信箱" and (t.get("assignee") or "").strip()
            and (t.get("due_date") or "").strip()):
        sfa_db.set_task_status(con, tid, "未着手")
        return "未着手"
    return t.get("status")


def _task_auto_start(con, tid: int) -> str | None:
    """進捗を追記＝着手の合図。受信箱/未着手/保留なら自動で「対応中」へ。現在のstatusを返す。"""
    t = sfa_db.get_task(con, tid)
    if not t:
        return None
    if t.get("status") in ("受信箱", "未着手", "保留"):
        sfa_db.set_task_status(con, tid, "対応中")
        return "対応中"
    return t.get("status")


def _task_urgency(due: str, today: str, d3: str, weekend: str) -> tuple:
    """期限から緊急度を算出（色, ラベル）。主観の優先度を廃し客観の期限ベースに（#30 ⑤）。"""
    due = (due or "").strip()
    if not due:
        return ("#cbd5e1", "期限未設定")
    if due < today:
        return ("#dc2626", "超過")        # 🔴
    if due <= d3:
        return ("#f59e0b", "まもなく")     # 🟠 〜3営業日
    if due <= weekend:
        return ("#eab308", "今週")        # 🟡 〜今週末
    return ("#94a3b8", "先")             # ⚪


_WD_JP = "月火水木金土日"


def _due_compact(due: str) -> str:
    """期限をコンパクト表記に（例: 2026-07-31 → 7/31金）。"""
    try:
        d = date.fromisoformat(due)
        return f"{d.month}/{d.day}{_WD_JP[d.weekday()]}"
    except (ValueError, TypeError):
        return due or ""


def _ai_guess_task_category(title: str, detail: str = "") -> str | None:
    """タイトル/詳細からタスク種類をAI(Haiku)で1つ推定。失敗時はNone（＝未設定のまま）。"""
    title = (title or "").strip()
    if not title:
        return None
    cats = "／".join(sfa_db.TASK_CATEGORIES)
    prompt = (
        "あなたは社内タスクの分類器です。次のタスクを、以下のカテゴリのちょうど1つに分類し、"
        "カテゴリ名だけを余計な語なしで出力してください。\n"
        f"カテゴリ: {cats}\n"
        f"タイトル: {title}\n" + (f"詳細: {detail.strip()}\n" if detail else ""))
    try:
        ans = (_call_claude_haiku(prompt, timeout=15, max_wait=18) or "").strip()
    except Exception:
        return None
    if not ans:
        return None
    for c in sfa_db.TASK_CATEGORIES:      # 完全一致優先
        if ans == c:
            return c
    for c in sfa_db.TASK_CATEGORIES:      # 部分一致（前後に語が付いた場合）
        if c in ans:
            return c
    return None


def _ai_summarize_task(title: str, memos: list) -> str | None:
    """タスクの議論メモ(新しい順のbody list)をHaikuで要約。失敗時None。"""
    memos = [m for m in memos if (m or "").strip()]
    if not memos:
        return None
    joined = "\n---\n".join(reversed(memos))  # 古い順に並べて文脈を作る
    prompt = (
        f"次はタスク「{title}」についての議論メモの履歴（古い順）です。\n"
        "全体を踏まえ、①現状/論点 ②決まったこと ③未決・懸念 ④次の一手 を、"
        "各1〜2行の簡潔な箇条書き（日本語）でまとめてください。前置き不要。\n\n"
        f"{joined}")
    try:
        ans = (_call_claude_haiku(prompt, timeout=25, max_wait=30) or "").strip()
    except Exception:
        return None
    return ans or None


def _ai_summarize_project(pname: str, meta: str, task_lines: list) -> str | None:
    """PJ配下タスクの状態＋各サマリ/次アクションを俯瞰要約（Haiku）。失敗時None。"""
    if not task_lines:
        return None
    body = "\n".join(task_lines)
    prompt = (
        f"プロジェクト「{pname}」（{meta}）の配下タスクの状況一覧です。\n"
        "全体として ①進捗サマリ ②主要な論点/意思決定 ③リスク・停滞 ④次にやるべきこと を、"
        "各1〜3行の簡潔な箇条書き（日本語）で俯瞰的にまとめてください。前置き不要。\n\n"
        f"{body}")
    try:
        ans = (_call_claude_haiku(prompt, timeout=25, max_wait=30) or "").strip()
    except Exception:
        return None
    return ans or None


def task_form(con, task=None) -> str:
    """タスクの新規/編集フォーム（純粋な入力フォーム, #30）。"""
    task = task or {}
    is_edit = bool(task.get("id"))
    owners = sfa_db.get_master_list(con, "owners")
    cats = sfa_db.get_master_list(con, "task_categories")
    projects = [p["name"] for p in sfa_db.list_task_projects(con)]
    _cur_pj = (task.get("project") or "").strip()
    if _cur_pj and _cur_pj not in projects:
        projects = [_cur_pj] + projects
    dev_opts = '<option value="">（なし）</option>'
    for dp in sfa_db.list_dev_projects(con):
        sel = " selected" if (task.get("link_type") == "dev_project" and task.get("link_id") == dp["id"]) else ""
        dev_opts += (f'<option value="dev_project:{dp["id"]}"{sel}>🛠 {_esc(dp.get("theme") or "開発案件")}'
                     f'（{_esc(dp.get("account_name") or "-")}）</option>')
    _blank = '<option value=""></option>'
    # 進捗ログ（追記式・履歴）は編集時のみ表示
    notes_block = ""
    if is_edit:
        notes = sfa_db.list_task_notes(con, task["id"])
        rows = "".join(
            f'<div style="border-left:3px solid #cbd5e1;padding:2px 0 2px 8px;margin:4px 0;font-size:12px">'
            f'<span class="muted" style="font-size:10px">{_esc((n.get("created_at") or "")[:16])}'
            f'{" ・" + _esc(n.get("author")) if n.get("author") else ""}</span><br>'
            f'{_esc(n.get("body")).replace(chr(10), "<br>")}</div>'
            for n in notes) or '<div class="muted" style="font-size:11px">まだ進捗ログはありません。</div>'
        notes_block = f"""
      <div class="card" style="max-width:720px;margin-top:12px">
        <h3 style="margin:0 0 6px">進捗ログ（追記式）</h3>
        <form method="post" action="/task/{task['id']}/note" style="margin-bottom:8px">
          <textarea name="body" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="2" placeholder="今日やったこと・状況を追記（履歴が残ります）" required></textarea>
          <div style="margin-top:6px"><button class="btn" type="submit">＋進捗を追記</button></div>
        </form>
        {rows}
      </div>"""
    return f"""
    <div class="card" style="max-width:720px">
      <h2>{'タスクを編集' if is_edit else '新規タスク'}</h2>
      {_save_bar('taskForm', cancel_url='/tasks')}
      <form id="taskForm" method="post" action="/tasks/save">
        <input type="hidden" name="id" value="{_esc(task.get('id'))}">
        <label>タイトル（タスク名・中項目） *</label>
        <input name="title" required value="{_esc(task.get('title'))}" placeholder="例: 〇〇のデモ環境を用意する">
        <label>次アクション（次にやる具体的な一手。空だと「止まっている」合図）</label>
        <input name="next_action" value="{_esc(task.get('next_action'))}" placeholder="例: △△さんにサンプルデータを依頼する">
        <label>詳細・文脈</label>
        <textarea name="detail" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="2">{_esc(task.get('detail'))}</textarea>
        <div class="grid">
          <div><label>プロジェクト（大項目）</label><select name="project">{_blank}{_opt(projects, task.get('project'))}</select></div>
          <div><label>担当</label><select name="assignee">{_blank}{_opt(owners, task.get('assignee'))}</select></div>
          <div><label>期限</label><input type="date" name="due_date" value="{_esc(task.get('due_date'))}"></div>
          <div><label>種類（空ならAIが自動判定）</label><select name="category"><option value=""></option>{_task_cat_optgroups(cats, task.get('category'))}</select></div>
          <div><label>状態</label><select name="status">{_opt(sfa_db.TASK_STATUSES, task.get('status') or ('未着手' if is_edit else '受信箱'))}</select></div>
          <div class="full"><label>関連（開発案件）</label><select name="link">{dev_opts}</select></div>
        </div>
        <div style="margin-top:14px"><button class="btn" type="submit">保存</button>
          <a class="btn sec" href="/tasks">キャンセル</a></div>
      </form>
    </div>{notes_block}"""


def tasks_page(con, *, assignee: str | None = None, category: str | None = None,
               project: str | None = None, urgency: str | None = None) -> str:
    """タスクボード（状態別カンバン）。コンパクト折りたたみカード＋その場編集＋緊急度自動＋
    プロジェクト一覧（期限・状態別内訳）＋期限クイック/逆算推奨（#30）。"""
    owners = sfa_db.get_master_list(con, "owners")
    cats = sfa_db.get_master_list(con, "task_categories")
    proj_objs = sfa_db.list_task_projects(con)
    projects = [p["name"] for p in proj_objs]
    proj_deadline = {p["name"]: (p.get("deadline") or "") for p in proj_objs if p.get("deadline")}
    dev_map = {dp["id"]: dp for dp in sfa_db.list_dev_projects(con)}
    # プロジェクトはカンマ区切りで複数選択可（チップのトグルで増減）
    sel_projects = [p for p in (project or "").split(",") if p.strip()]
    tasks = sfa_db.list_tasks(con, assignee=assignee or None, category=category or None)
    if sel_projects:
        _selset = set(sel_projects)
        tasks = [t for t in tasks if (t.get("project") or "") in _selset]
    latest_notes = {}
    for r in con.execute(
        "SELECT n.task_id, n.body, n.created_at FROM task_notes n WHERE n.kind='progress' AND n.id=("
        "SELECT id FROM task_notes WHERE task_id=n.task_id AND kind='progress' "
        "ORDER BY created_at DESC, id DESC LIMIT 1)"):
        latest_notes[r["task_id"]] = dict(r)
    _td = _today_jst()
    today = _td.isoformat()
    d3 = sfa_db.add_business_days(_td, 3).isoformat()
    weekend = (_td + timedelta(days=6 - _td.weekday())).isoformat()
    # 期限クイック候補（今日＋N営業日）を先に計算してJSへ
    qdates = {n: sfa_db.add_business_days(_td, n).isoformat() for n in (1, 3, 5, 8)}
    # 緊急度フィルタ（期限ベース）: overdue=超過 / week=今週まで / nodue=期限なし
    if urgency:
        def _uok(t):
            due = (t.get("due_date") or "").strip()
            if urgency == "overdue":
                return bool(due) and due < today
            if urgency == "week":
                return bool(due) and today <= due <= weekend
            if urgency == "nodue":
                return not due
            return True
        tasks = [t for t in tasks if _uok(t)]
    cols = {s: [] for s in sfa_db.TASK_STATUSES}
    for t in tasks:
        cols.setdefault(t.get("status") or "受信箱", []).append(t)

    def card(t):
        tid = t["id"]
        status = t.get("status") or "受信箱"
        due = (t.get("due_date") or "").strip()
        ucolor, ulabel = _task_urgency(due, today, d3, weekend)
        pinned = bool(t.get("pinned"))
        na = (t.get("next_action") or "").strip()
        proj = (t.get("project") or "").strip()
        proj_sel = (_tc_sel(tid, "project", projects, proj, "📁PJ")
                    if (projects or proj) else "")
        asg_sel = _tc_sel(tid, "assignee", owners, t.get("assignee"), "👤担当")
        cat_sel = (f'<select class="tc-sel" data-field="category" title="種類" '
                   f'onchange="taskField({tid},&#39;category&#39;,this.value)">'
                   f'<option value="">種類</option>{_task_cat_optgroups(cats, t.get("category"))}</select>')
        ai_btn = (f'<button type="button" class="tc-ai" title="AIで種類を判定" '
                  f'onclick="tcAiCat({tid})">🤖種類</button>')
        due_input = (f'<input type="date" class="tc-due" style="color:{ucolor}" value="{_esc(due)}" '
                     f'title="期限" onchange="tcDue({tid},this.value)">')
        quick = "".join(
            f'<button type="button" class="tc-q" onclick="tcDue({tid},&#39;{qdates[n]}&#39;)">+{n}営</button>'
            for n in (1, 3, 5, 8))
        rec = ""
        dl = proj_deadline.get(proj)
        if dl:
            try:
                rec_d = sfa_db.add_business_days(date.fromisoformat(dl), -5)
                if rec_d < _td:
                    rec_d = _td
                rec = (f'<button type="button" class="tc-q rec" title="PJ期限{_esc(dl)}から逆算" '
                       f'onclick="tcDue({tid},&#39;{rec_d.isoformat()}&#39;)">推奨 {rec_d.isoformat()[5:]}</button>')
            except ValueError:
                rec = ""
        link_html = ""
        if t.get("link_type") == "dev_project" and t.get("link_id") in dev_map:
            dp = dev_map[t["link_id"]]
            link_html = (f'<a href="/dev-project/{dp["id"]}/edit" style="font-size:10px" '
                         f'title="{_esc(dp.get("account_name") or "")}">🛠{_esc(dp.get("theme") or "開発案件")}</a>')
        search = _esc(" ".join(str(t.get(k) or "") for k in
                               ("title", "detail", "assignee", "category", "project", "next_action")).lower())
        note = latest_notes.get(tid)
        note_snip = _esc((note.get("body") or "")[:40]) if note else "進捗を追記…"
        summary_html = (f'<div class="tc-sum" title="議論メモのAIサマリ" '
                        f'onclick="openNotes({tid},&#39;discussion&#39;)">🧠 '
                        f'{_esc((t.get("summary") or "").strip()[:130])}</div>'
                        if (t.get("summary") or "").strip() else "")
        pin_btn = (f'<button type="button" class="tc-pin{" on" if pinned else ""}" '
                   f'onclick="tcPin({tid})" title="★最優先ピン">★</button>')
        asg = (t.get("assignee") or "").strip()
        # コンパクト表示（常時）: PJ・担当・期限(7/31金)・遷移ボタン
        m_pj = f'<span class="m-pj" title="{_esc(proj)}">📁{_esc(proj)}</span>' if proj else ""
        m_asg = f'<span class="m-asg">👤{_esc(asg)}</span>' if asg else ""
        m_due = (f'<span class="m-due" style="color:{ucolor}" title="{ulabel}">📅{_esc(_due_compact(due))}</span>'
                 if due else '<span class="m-due" style="color:#cbd5e1">📅—</span>')
        return (
            f'<div class="task-card{" pinned" if pinned else ""}" id="tc-{tid}" '
            f'data-status="{_esc(status)}" data-pinned="{1 if pinned else 0}" data-search="{search}" '
            f'onclick="tcCardClick(event,this)">'
            f'<div class="tc-head">'
            f'<span class="tc-dot" style="background:{ucolor}" title="{ulabel}"></span>'
            f'<span class="tc-ttl">{_esc(t.get("title"))}</span>'
            f'{pin_btn}<span class="tc-car">▸</span></div>'
            f'<div class="tc-mini">{m_pj}{m_asg}{m_due}</div>'
            f'<div class="tc-actions"></div>'
            f'<div class="tc-body">'
            f'<input class="tc-title" value="{_esc(t.get("title"))}" title="タイトル" '
            f'onchange="taskField({tid},&#39;title&#39;,this.value);tcSyncTitle({tid},this.value)">'
            f'<div class="tc-na{" empty" if not na else ""}"><span>▶</span>'
            f'<input value="{_esc(na)}" placeholder="次アクション未設定" title="次アクション" '
            f'onchange="taskField({tid},&#39;next_action&#39;,this.value)"></div>'
            f'<div class="tc-meta">{proj_sel}{asg_sel}{cat_sel}{ai_btn}{link_html}</div>'
            f'<div class="tc-meta"><span class="tc-lbl">期限</span>{due_input}{quick}{rec}</div>'
            f'<div class="tc-notes" onclick="openNotes({tid},&#39;progress&#39;)" title="進捗ログを見る・追記">📝 {note_snip}</div>'
            f'{summary_html}'
            f'<div class="tc-foot">'
            f'<button type="button" class="tc-disc" onclick="openNotes({tid},&#39;discussion&#39;)" title="議論メモ＋AIサマリ">💬 議論メモ</button>'
            f'<a class="tc-edit" href="/tasks/{tid}/edit">編集</a>'
            f'<button type="button" class="del" onclick="taskDelete({tid})">🗑 削除</button></div>'
            f'</div></div>')

    columns = ""
    for s in sfa_db.TASK_STATUSES:
        ts = cols.get(s, [])
        inner = "".join(card(t) for t in ts)
        columns += (f'<div class="task-col"><h3>{s}（<span class="tc-count" data-count="{s}">{len(ts)}</span>）</h3>'
                    f'<div class="tc-col-body" data-col="{s}">{inner}</div></div>')

    # /tasks のURLを組み立て（PJ複数＋担当/種類フィルタを保持）
    def _tasks_url(projs_list):
        params = []
        if projs_list:
            params.append("project=" + urllib.parse.quote(",".join(projs_list)))
        if assignee:
            params.append("assignee=" + urllib.parse.quote(assignee))
        if category:
            params.append("category=" + urllib.parse.quote(category))
        return "/tasks" + ("?" + "&".join(params) if params else "")

    # プロジェクト一覧ストリップ（状態別内訳＋期限。クリックで絞り込み・再クリックで解除・複数選択）
    pcounts = sfa_db.task_counts_by_project_status(con)
    strip = ""
    if proj_objs:
        chips = ""
        for p in proj_objs:
            name = p["name"]
            c = pcounts.get(name, {})
            open_n = sum(c.get(s, 0) for s in ("受信箱", "未着手", "対応中", "保留"))
            done_n = c.get("完了", 0)
            dl = p.get("deadline") or ""
            dcol = _task_urgency(dl, today, d3, weekend)[0] if dl else "#cbd5e1"
            in_sel = name in sel_projects
            # トグル: 選択中なら外す、未選択なら足す
            new_sel = [x for x in sel_projects if x != name] if in_sel else sel_projects + [name]
            done_txt = f'・完{done_n}' if done_n else ""
            dl_txt = f'・〆{_esc(dl[5:])}' if dl else ""
            chips += (f'<a class="pj-chip{" active" if in_sel else ""}" href="{_tasks_url(new_sel)}">'
                      f'<span class="pj-dot" style="background:{dcol}"></span>'
                      f'{"✓" if in_sel else ""}📁{_esc(name)}'
                      f'<span class="pj-cnt">{open_n}件{done_txt}{dl_txt}</span></a>')
        clear = (f'<a class="pj-chip mng" href="{_tasks_url([])}">絞り込み解除</a>' if sel_projects else "")
        strip = (f'<div class="pj-strip">{chips}{clear}'
                 f'<a class="pj-chip mng" href="/task-projects">⚙ PJ管理</a></div>')

    def _fopt(values, cur, alllabel):
        return f'<option value="">{alllabel}</option>' + "".join(
            f'<option value="{html.escape(v)}"{" selected" if v == cur else ""}>{html.escape(v)}</option>'
            for v in values)
    _urg_opts = "".join(
        f'<option value="{v}"{" selected" if urgency == v else ""}>{lbl}</option>'
        for v, lbl in (("", "緊急度:全て"), ("overdue", "🔴 超過"), ("week", "🟡 今週まで"),
                       ("nodue", "⚪ 期限なし")))
    filter_row = f"""<form method="get" action="/tasks" class="filter-row">
      <input type="hidden" name="project" value="{_esc(','.join(sel_projects))}">
      <select name="assignee" onchange="this.form.submit()">{_fopt(owners, assignee, '担当:全て')}</select>
      <select name="category" onchange="this.form.submit()"><option value="">種類:全て</option>{_task_cat_optgroups(cats, category)}</select>
      <select name="urgency" onchange="this.form.submit()">{_urg_opts}</select>
      <input type="text" id="taskSearch" placeholder="🔍 タイトル・詳細・次アクションで検索…" oninput="taskFilter()" style="max-width:260px">
      <a class="btn sec" href="/tasks">リセット</a>
    </form>"""
    has_test = any((t.get("source") == "test") for t in tasks)
    test_bar = ""
    if has_test:
        test_bar = ('<div class="card" style="background:#fef3c7;border-color:#fcd34d;'
                    'display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">'
                    '<span>🧪 <b>【テスト】</b>で始まるタスクは検証用のサンプルです（本番データではありません）。</span>'
                    '<form method="post" action="/tasks/delete-test">'
                    '<button class="btn sec" type="submit">🧪 テストデータを削除</button></form></div>')
    seed_btn = ('<form method="post" action="/tasks/seed-test" style="display:inline">'
                '<button class="btn sec" type="submit" title="検証用の【テスト】サンプルを投入">🧪 テストデータ投入</button></form>')
    filt_note = (f'<span class="muted" style="font-size:12px">📁 {_esc("・".join(sel_projects))} で絞り込み中 '
                 f'<a href="{_tasks_url([])}">解除</a></span>' if sel_projects else "")
    quick_js = (f'<script>window._TC={{today:"{today}",d3:"{d3}",weekend:"{weekend}"}};</script>')
    return f"""
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>タスク（{len(tasks)}） {filt_note}</span>
        <span style="display:flex;gap:8px;flex-wrap:wrap">
          {seed_btn}
          <a class="btn sec" href="/tasks/digest">🔔 朝ダイジェスト</a>
          <a class="btn" href="/tasks/new">＋新規タスク</a>
        </span>
      </h2>
      {test_bar}
      {strip}
      {filter_row}
      <div id="taskBoard">{columns}</div>
    </div>{quick_js}{_TASKS_JS}"""


def task_projects_page(con) -> str:
    """タスクのプロジェクト（大項目）管理: 期限・状態の編集＋追加＋削除（#30 ⑦）。"""
    projs = sfa_db.list_task_projects(con)
    counts = sfa_db.task_counts_by_project_status(con)
    blocks = ""
    for p in projs:
        c = counts.get(p["name"], {})
        open_n = sum(c.get(s, 0) for s in ("受信箱", "未着手", "対応中", "保留"))
        done_n = c.get("完了", 0)
        _sum = (p.get("summary") or "").strip()
        sum_disp = (f'<div id="pjsum-{p["id"]}" style="margin-top:8px;background:#eef2ff;border-radius:6px;'
                    f'padding:8px 10px;font-size:12px;color:#1e3a8a;white-space:pre-wrap">🧠 {_esc(_sum)}</div>'
                    if _sum else f'<div id="pjsum-{p["id"]}"></div>')
        blocks += f"""
      <div class="card" style="margin-bottom:8px">
        <form method="post" action="/task-projects/save" class="filter-row" style="margin:0">
          <input type="hidden" name="id" value="{p['id']}">
          <input name="name" value="{_esc(p['name'])}" style="min-width:220px" title="プロジェクト名">
          <label style="font-size:12px">期限 <input type="date" name="deadline" value="{_esc(p.get('deadline') or '')}"></label>
          <select name="status">{_opt(sfa_db.TASK_PROJECT_STATUSES, p.get('status') or '進行中')}</select>
          <span class="muted" style="font-size:12px">進行 {open_n}件・完了 {done_n}件</span>
          <button class="btn sec" type="submit">保存</button>
          <a class="btn sec" href="/tasks?project={urllib.parse.quote(p['name'])}">タスクを見る</a>
          <button class="btn sec" type="button" onclick="pjSummary({p['id']})">🧠 PJサマリ生成</button>
          <button class="btn sec" type="submit" formaction="/task-project/{p['id']}/delete"
            style="color:#c53030" formnovalidate
            onclick="return confirm('プロジェクト定義を削除しますか？（紐づくタスク自体は残ります）')">削除</button>
        </form>
        {sum_disp}
      </div>"""
    if not projs:
        blocks = '<div class="card muted">プロジェクトはまだありません。下から追加してください。</div>'
    return f"""
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>📁 プロジェクト管理（大項目）</span>
        <a class="btn sec" href="/tasks">← タスク看板へ</a>
      </h2>
      <p class="muted" style="font-size:13px">プロジェクトに期限と状態を持たせると、看板上部の一覧に出て、
      タスクの期日はプロジェクト期限から逆算して推奨されます。「🧠PJサマリ生成」で配下タスクの
      議論・進捗をAIが俯瞰要約します。</p>
      <form method="post" action="/task-projects/save" class="filter-row">
        <input name="name" placeholder="新しいプロジェクト名（例: セキュリティISO取得）" required style="min-width:260px">
        <label style="font-size:12px">期限 <input type="date" name="deadline"></label>
        <select name="status">{_opt(sfa_db.TASK_PROJECT_STATUSES, '進行中')}</select>
        <button class="btn" type="submit">＋ 追加</button>
      </form>
    </div>
    {blocks}
    <script>
    function pjSummary(id){{
      var box=document.getElementById('pjsum-'+id);
      if(box){{ box.style.display='block'; box.style.marginTop='8px'; box.style.background='#eef2ff';
        box.style.borderRadius='6px'; box.style.padding='8px 10px'; box.style.fontSize='12px';
        box.style.color='#1e3a8a'; box.style.whiteSpace='pre-wrap'; box.textContent='🧠 生成中…（配下タスクを要約しています）'; }}
      fetch('/task-project/'+id+'/summary',{{method:'POST',headers:{{'Content-Type':'application/x-www-form-urlencoded'}},body:'x=1'}})
       .then(function(r){{return r.json();}}).then(function(d){{
         if(box) box.textContent = d.ok&&d.summary ? ('🧠 '+d.summary) : '⚠ サマリを生成できませんでした（配下タスクが無い/AI応答なし）。';
       }}).catch(function(){{ if(box)box.textContent='⚠ 通信エラー'; }});
    }}
    </script>"""


# ---- 朝のタスクダイジェスト（Slack DM）(#30) ----

def _load_owner_slack_map() -> dict:
    """config/owner_slack_map.json（担当者名→email）を読む。_始まりのキーは除外。"""
    p = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                     "config", "owner_slack_map.json")
    try:
        with open(p, encoding="utf-8") as fp:
            d = json.load(fp)
        return {k: v for k, v in d.items() if not k.startswith("_")}
    except Exception:
        return {}


def _task_digest_groups(con) -> dict:
    """未完了タスクを担当ごとにまとめる（list_tasks の順＝期限昇順→優先度）。"""
    groups: dict = {}
    for t in sfa_db.list_tasks(con, exclude_done=True):
        groups.setdefault(t.get("assignee") or "（担当未設定）", []).append(t)
    return groups


def build_task_digest_text(owner: str, tasks: list, tool_url: str) -> str:
    """1担当ぶんの朝ダイジェストDM本文を組み立てる（超過/今週/その他に区分）。"""
    import urllib.parse as _up
    today = _today_jst().isoformat()
    week = (_today_jst() + timedelta(days=7)).isoformat()
    overdue, thisweek, rest = [], [], []
    for t in tasks:
        due = (t.get("due_date") or "").strip()
        if due and due < today:
            overdue.append(t)
        elif due and today <= due <= week:
            thisweek.append(t)
        else:
            rest.append(t)

    def _mmdd(due):
        try:
            d = date.fromisoformat(due)
            return f"{d.month}/{d.day}"
        except (ValueError, TypeError):
            return "期限なし"

    def line(t):
        due = (t.get("due_date") or "").strip()
        dd = _mmdd(due) if due else "期限なし"
        na = (t.get("next_action") or "").strip()
        na_s = f"　→ {na}" if na else "　→ 次アクション未設定"
        # タイトルをその場のカードへの直リンクに（Slack→該当タスクが1クリック）
        title = str(t.get("title", "")).replace("|", "｜").replace(">", "＞").replace("<", "＜")
        link = f"{tool_url}/tasks?assignee={_up.quote(owner)}#tc-{t.get('id')}"
        return f"  • {dd}｜<{link}|{title}>{na_s}"

    lines = [f"【朝のタスクダイジェスト {today}】",
             f"{owner}さん、未完了タスク {len(tasks)}件です。", ""]
    if overdue:
        lines.append(f"🔴 超過 {len(overdue)}件")
        lines += [line(t) for t in overdue[:10]]
        lines.append("")
    if thisweek:
        lines.append(f"🟡 今週まで {len(thisweek)}件")
        lines += [line(t) for t in thisweek[:10]]
        lines.append("")
    if rest:
        lines.append(f"📋 その他 {len(rest)}件（先の期限・期限なし）")
        lines.append("")
    lines.append(f"🔗 ボード: {tool_url}/tasks?assignee={_up.quote(owner)}")
    lines.append("今日はどれを前に進めますか？ 一手だけ決めましょう。")
    return "\n".join(lines)


def _slack_api_post(method: str, **kwargs) -> dict:
    import urllib.request
    token = os.environ.get("SLACK_BOT_TOKEN", "")
    req = urllib.request.Request(
        f"https://slack.com/api/{method}", data=json.dumps(kwargs).encode("utf-8"),
        headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"})
    with urllib.request.urlopen(req, timeout=10) as r:
        return json.loads(r.read())


def _slack_api_get(method: str, params: dict) -> dict:
    import urllib.request
    import urllib.parse as _up
    token = os.environ.get("SLACK_BOT_TOKEN", "")
    req = urllib.request.Request(
        f"https://slack.com/api/{method}?{_up.urlencode(params)}",
        headers={"Authorization": f"Bearer {token}"})
    with urllib.request.urlopen(req, timeout=10) as r:
        return json.loads(r.read())


def send_task_digests(con, only: str | None = None) -> str:
    """朝ダイジェストをSlack DMで送る。only指定時はその担当のみ（テスト送信用）。結果文字列を返す。"""
    if not os.environ.get("SLACK_BOT_TOKEN"):
        return "SLACK_BOT_TOKEN が未設定のため送信できません（本番のみ有効）。"
    owner_map = _load_owner_slack_map()
    tool_url = os.environ.get("SFA_TOOL_URL", "") or "https://sfa-crm.onrender.com"
    results = []
    for owner, tasks in _task_digest_groups(con).items():
        if only and owner != only:
            continue
        if owner == "（担当未設定）":
            continue
        email = owner_map.get(owner)
        if not email:
            results.append(f"{owner}: メール未設定でスキップ")
            continue
        try:
            lk = _slack_api_get("users.lookupByEmail", {"email": email})
            if not lk.get("ok"):
                results.append(f"{owner}: Slackユーザー不明")
                continue
            dm = _slack_api_post("conversations.open", users=lk["user"]["id"])
            if not dm.get("ok"):
                results.append(f"{owner}: DMを開けず")
                continue
            from cowork import slack_tasks as _st
            res = _slack_api_post("chat.postMessage", channel=dm["channel"]["id"],
                                  text=build_task_digest_text(owner, tasks, tool_url),
                                  blocks=_st.build_digest_blocks(owner, tasks, tool_url))
            results.append(f"{owner}: {'送信OK' if res.get('ok') else 'エラー ' + str(res.get('error'))}"
                           f"（{len(tasks)}件）")
        except Exception as e:  # noqa: BLE001
            results.append(f"{owner}: 例外 {e}")
    return " ／ ".join(results) or "送信対象がありませんでした。"


def tasks_digest_page(con, result: str | None = None) -> str:
    """朝ダイジェストのプレビュー＋Slackテスト送信画面（#30）。"""
    tool_url = os.environ.get("SFA_TOOL_URL", "") or "https://sfa-crm.onrender.com"
    owner_map = _load_owner_slack_map()
    groups = _task_digest_groups(con)
    result_html = (f'<div class="card" style="background:#ecfdf5;border-color:#a7f3d0">'
                   f'<b>送信結果:</b> {_esc(result)}</div>' if result else "")
    blocks = ""
    for owner, tasks in groups.items():
        if owner == "（担当未設定）":
            note = '<span style="color:#c53030">※担当未設定＝DM送信対象外</span>'
            btn = ""
        elif owner not in owner_map:
            note = '<span style="color:#c53030">※owner_slack_mapにメール未設定＝送信不可</span>'
            btn = ""
        else:
            note = ""
            btn = (f'<form method="post" action="/tasks/digest/send" style="display:inline">'
                   f'<input type="hidden" name="only" value="{_esc(owner)}">'
                   f'<button class="btn sec" type="submit">この人に送信</button></form>')
        preview = _esc(build_task_digest_text(owner, tasks, tool_url))
        blocks += (f'<div class="card" style="max-width:720px">'
                   f'<h3 style="margin:0 0 4px;display:flex;justify-content:space-between;align-items:center">'
                   f'<span>👤 {_esc(owner)}（{len(tasks)}件）{note}</span>{btn}</h3>'
                   f'<pre style="white-space:pre-wrap;font-size:12px;background:#f8fafc;'
                   f'border-radius:6px;padding:8px;margin:0">{preview}</pre></div>')
    if not blocks:
        blocks = '<div class="card">未完了タスクがありません。</div>'
    return f"""
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>🔔 朝のタスクダイジェスト（プレビュー）</span>
        <a class="btn sec" href="/tasks">← ボードに戻る</a>
      </h2>
      <p class="muted" style="font-size:13px">担当ごとに、未完了タスクを「超過／今週まで／その他」に整理してSlack DMで送ります。
      まずは<b>自分に送って手触りを確認</b>してください。将来は毎朝1回、cronで自動送信します。</p>
      <div style="display:flex;gap:8px;flex-wrap:wrap">
        <form method="post" action="/tasks/digest/send">
          <input type="hidden" name="only" value="早瀬">
          <button class="btn" type="submit">自分（早瀬）にテスト送信</button>
        </form>
        <form method="post" action="/tasks/digest/send"
              onsubmit="return confirm('担当が設定された全員にSlack DMを送信します。よろしいですか？')">
          <button class="btn sec" type="submit">全員に送信</button>
        </form>
      </div>
    </div>
    {result_html}
    {blocks}"""


def masters_page(con) -> str:
    """入力マスタ編集ページ。各リストの選択肢を追加・削除・並び替えできる。"""
    cards = []
    for key, label in sfa_db.MASTER_LABELS.items():
        values = sfa_db.get_master_list(con, key)
        items_html = "".join(
            f'<div class="master-item" draggable="true" data-key="{html.escape(key)}" data-idx="{i}">'
            f'<span class="drag-handle" title="ドラッグで並び替え">⠿</span>'
            f'<span class="item-label">{html.escape(v)}</span>'
            f'<button type="button" onclick="delItem(\'{html.escape(key)}\',{i})" '
            f'style="background:none;border:none;color:#ef4444;cursor:pointer;font-size:14px;padding:0 4px">✕</button>'
            f'</div>'
            for i, v in enumerate(values)
        )
        hidden_inputs = "".join(
            f'<input type="hidden" name="{html.escape(key)}[]" value="{html.escape(v)}">'
            for v in values
        )
        cards.append(f"""
        <div class="card" id="master_{key}">
          <h2>{label}</h2>
          <div id="items_{key}" style="display:flex;flex-wrap:wrap;gap:6px;margin-bottom:10px">
            {items_html}
          </div>
          <div style="display:flex;gap:8px;align-items:center">
            <input id="new_{key}" placeholder="新しい選択肢を追加" style="max-width:200px">
            <button type="button" class="btn sec" onclick="addItem('{html.escape(key)}')">追加</button>
          </div>
          <div id="hidden_{key}">{hidden_inputs}</div>
        </div>""")

    return f"""
    <div class="card" style="background:#f0f4f8;border:1.5px solid #d4dae4">
      <h2>⚙ 入力マスタの編集</h2>
      <p class="muted">各項目の選択肢を編集できます。変更は「すべて保存」ボタンで反映されます。</p>
    </div>
    <form method="post" action="/masters/save" id="master_form">
      {''.join(cards)}
      <p><button class="btn">すべて保存</button>
         <a class="btn sec" href="/">キャンセル</a></p>
    </form>
    <style>
      .master-item{{display:inline-flex;align-items:center;background:#e8edf7;border-radius:20px;padding:3px 10px;font-size:13px;gap:4px;cursor:default;user-select:none}}
      .master-item.drag-over{{outline:2px dashed #2f6fed;background:#dbeafe}}
      .drag-handle{{cursor:grab;color:#aab;font-size:15px;line-height:1}}
    </style>
    <script>
    function rebuildHidden(key) {{
      const container = document.getElementById('items_' + key);
      const hidden = document.getElementById('hidden_' + key);
      const items = Array.from(container.querySelectorAll('.master-item'));
      hidden.innerHTML = items.map(el =>
        `<input type="hidden" name="${{key}}[]" value="${{escH(el.querySelector('.item-label').textContent)}}">`
      ).join('');
      items.forEach((el, i) => {{
        el.dataset.idx = i;
        el.querySelector('button').setAttribute('onclick', `delItem('${{key}}',${{i}})`);
      }});
    }}
    function delItem(key, idx) {{
      const container = document.getElementById('items_' + key);
      Array.from(container.querySelectorAll('.master-item'))[idx].remove();
      rebuildHidden(key);
    }}
    function addItem(key) {{
      const input = document.getElementById('new_' + key);
      const val = input.value.trim();
      if (!val) return;
      const container = document.getElementById('items_' + key);
      const idx = container.querySelectorAll('.master-item').length;
      container.insertAdjacentHTML('beforeend',
        `<div class="master-item" draggable="true" data-key="${{key}}" data-idx="${{idx}}">` +
        `<span class="drag-handle" title="ドラッグで並び替え">⠿</span>` +
        `<span class="item-label">${{escH(val)}}</span>` +
        `<button type="button" onclick="delItem('${{key}}',${{idx}})" style="background:none;border:none;color:#ef4444;cursor:pointer;font-size:14px;padding:0 4px">✕</button>` +
        `</div>`
      );
      rebuildHidden(key);
      input.value = '';
      initDrag(key);
    }}
    function initDrag(key) {{
      const container = document.getElementById('items_' + key);
      container.querySelectorAll('.master-item[draggable]').forEach(item => {{
        item.ondragstart = e => {{
          e.dataTransfer.effectAllowed = 'move';
          container._dragging = item;
        }};
        item.ondragover = e => {{
          e.preventDefault();
          const dragging = container._dragging;
          if (!dragging || dragging === item) return;
          const rect = item.getBoundingClientRect();
          if (e.clientY < rect.top + rect.height / 2) container.insertBefore(dragging, item);
          else container.insertBefore(dragging, item.nextSibling);
        }};
        item.ondragend = () => {{ rebuildHidden(key); container._dragging = null; }};
      }});
    }}
    document.addEventListener('DOMContentLoaded', () => {{
      {'; '.join(f"initDrag('{html.escape(key)}')" for key in sfa_db.MASTER_LABELS)}
    }});
    </script>"""


def activity_deal_picker(con) -> str:
    deals = sfa_db.list_deals(con, status="open")
    rows = "".join(
        f'<tr style="cursor:pointer" onclick="location.href=\'/deal/{d["id"]}#activity\'">'
        f'<td><a href="/deal/{d["id"]}">{_esc(d.get("account_name"))}</a></td>'
        f'<td>{_esc(d.get("deal_name"))}</td>'
        f'<td><span class="stage">{_esc(d.get("stage"))}</span></td>'
        f'<td>{_esc(d.get("next_milestone_date") or "—")}</td></tr>'
        for d in deals
    ) or '<tr><td colspan=4 class=muted>進行中の商談がありません</td></tr>'
    return f"""
    <div class="card">
      <h2>活動を追加する商談を選択</h2>
      <p class="muted">行をクリックすると商談ページへ移動し、活動履歴を追加できます。</p>
      <table>
        <tr><th>アカウント</th><th>案件名</th><th>ステージ</th><th>次回MS</th></tr>
        {rows}
      </table>
    </div>"""


# ── 既存ページ（商談・アカウント）─────────────────────────────────────────────

def _udeal_sel(deal_id, field, values, current, *, sel_id=None, cascade_l1=False, disabled=False):
    """商談一覧の共通インライン編集セレクト。cascade_l1=Trueなら事業種別L1(保存後リロード)。
    disabled=Trueで編集不可（クローズ済み商談などの閲覧専用行に使う）。"""
    opts = "".join(
        f'<option value="{html.escape(v)}"{" selected" if v == current else ""}>{html.escape(v)}</option>'
        for v in values)
    id_attr = f' id="{sel_id}"' if sel_id else ""
    if disabled:
        return (f'<select{id_attr} disabled '
                f'style="font-size:11px;padding:1px 2px;min-width:84px;max-width:120px;'
                f'background:#f1f5f9;color:#64748b">'
                f'<option value=""></option>{opts}</select>')
    if cascade_l1:
        onchange = f'updateDealL1({deal_id}, this.value)'
    else:
        onchange = f"updateDealField({deal_id}, '{field}', this.value)"
    return (f'<select{id_attr} onchange="{onchange}" '
            f'style="font-size:11px;padding:1px 2px;min-width:84px;max-width:120px">'
            f'<option value=""></option>{opts}</select>')


def unified_deal_table(con, deals: list, *, return_to_url: str, bulk: bool = False) -> str:
    """商談一覧の共通テーブル（全タブ同一の14列・全インライン編集・全行クローズ）を返す。
    bulk=Trueで先頭に一括選択チェックボックス列を出す（呼び出し側が<form>で包む）。"""
    stages = sfa_db.get_master_list(con, "deal_stages")
    owners = sfa_db.get_master_list(con, "owners")
    l1_list = sfa_db.get_master_list(con, "business_type_l1")
    # ツール表示用に、開発案件を商談ごとにまとめる（1クエリ）
    dev_by_deal: dict = {}
    for dp in sfa_db.list_dev_projects(con):
        dev_by_deal.setdefault(dp.get("deal_id"), []).append(dp)
    # 次回MSの本数（#48）: 一覧の「ほかN件」バッジ用（未完了MSが2件以上のとき表示）
    ms_counts = sfa_db.count_open_milestones(con, [d["id"] for d in deals])

    cb_th = ('<th class="sticky" style="width:28px"><input type="checkbox" id="deal_chk_all" title="全選択"'
             ' onchange="var v=this.checked;document.querySelectorAll(\'[name=ids]\').forEach(function(c){c.checked=v;});">'
             '</th>') if bulk else ""
    _th_total = _sticky_th("提案総額<br><span style='font-size:10px;color:#8893a8'>(万円)</span>")
    header = (
        f'<tr>{cb_th}{_sticky_th("#")}{_sticky_th("アカウント")}{_sticky_th("案件名")}{_sticky_th("ステージ")}'
        f'{_sticky_th("主担当")}{_sticky_th("サブ担当")}{_sticky_th("種別L1")}{_sticky_th("種別L2")}'
        f'{_sticky_th("予算")}{_th_total}'
        f'{_sticky_th("次回MS日")}{_sticky_th("次回MS")}{_sticky_th("ツール")}{_sticky_th("クローズ")}</tr>'
    )

    _ro_input_style = "background:#f1f5f9;color:#64748b"  # 編集不可の見た目（クローズ済み行）
    rows = []
    for d in deals:
        did = d["id"]
        _ro = (d.get("status") == "closed")  # クローズ済みは閲覧専用（編集不可）
        _dis = " disabled" if _ro else ""
        _ro_sty = f";{_ro_input_style}" if _ro else ""
        l2_values = sfa_db.BUSINESS_TYPE_L2_BY_L1.get(d.get("business_type_l1") or "", [])
        inp_budget = (f'<input type="text" value="{_esc(d.get("client_budget") or "")}"{_dis}'
                      f' onchange="updateDealField({did}, \'client_budget\', this.value)"'
                      f' style="font-size:11px;padding:1px 2px;width:72px{_ro_sty}">')
        inp_total = (f'<input type="number" step="0.1" value="{_esc(d.get("value_lumpsum") or "")}"{_dis}'
                     f' onchange="updateDealField({did}, \'value_lumpsum\', this.value)"'
                     f' style="font-size:11px;padding:1px 2px;width:72px{_ro_sty}">')
        inp_ms_date = (f'<input type="date" id="msdate_{did}" value="{_esc(d.get("next_milestone_date"))}"{_dis}'
                       f' onchange="updateDealField({did}, \'next_milestone_date\', this.value)"'
                       f' style="font-size:11px;padding:1px 2px{_ro_sty}">')
        _ms_open = ms_counts.get(did, 0)
        # 一覧から全MSを確認/追加/編集できるパネルのトリガー（未完了2件以上は強調）。
        # クローズ済みは編集不可のため、パネルを開かない静的表示にする。
        if _ro:
            _ms_trigger = (f'<span class="ms-trigger" style="cursor:default;opacity:.6">'
                           f'🗂 MS·{_ms_open}</span>')
        else:
            _ms_trigger = (
                f'<button type="button" id="mstrg_{did}" class="ms-trigger{" has-multi" if _ms_open >= 2 else ""}" '
                f'onclick="openMsPanel({did}, this)" title="この商談の次回MSを一覧・追加・編集">'
                f'🗂 MS·<span id="mscount_{did}">{_ms_open}</span></button>')
        inp_ms_label = (f'<input type="text" id="mslabel_{did}" value="{_esc(d.get("next_milestone_label"))}"{_dis}'
                        f' onchange="updateDealField({did}, \'next_milestone_label\', this.value)"'
                        f' style="font-size:11px;padding:1px 2px;width:130px{_ro_sty}"><br>'
                        + f'<span id="mstype_{did}">'
                        + _udeal_sel(did, "next_milestone_type", sfa_db.NEXT_MS_TYPES,
                                     d.get("next_milestone_type") or "", disabled=_ro)
                        + '</span> ' + _ms_trigger)
        tool_btns = " ".join(
            _tool_link_btn(dp.get("tool_url"), tool_id=dp.get("tool_login_id"), tool_password=dp.get("tool_login_pass"))
            for dp in dev_by_deal.get(did, [])
            if dp.get("tool_url") and dp.get("status") != "中止"
        ) or "—"
        # 開発案件への導線: 既存の案件リンクを並べ、その後に必ず「＋開発案件」を出す
        # （1商談に複数の開発案件がつくため、既存があっても追加できるようにする）。
        # いずれも return_to に現在の一覧URLを渡し、遷移後に一覧へ戻れるようにする。
        _dev_pill = ('display:inline-block;background:#eff6ff;color:#1d4ed8;border:1px solid #bfdbfe;'
                     'border-radius:6px;padding:2px 8px;font-size:11px;white-space:nowrap;text-decoration:none')
        _add_pill = ('display:inline-block;background:#f5f3ff;color:#6d28d9;border:1px dashed #c4b5fd;'
                     'border-radius:6px;padding:2px 8px;font-size:11px;white-space:nowrap;text-decoration:none')
        _ret_q = urllib.parse.quote(return_to_url, safe="")
        _devs = dev_by_deal.get(did, [])
        dev_links = "".join(
            f'<a href="/dev-project/{dp["id"]}/edit?return_to={_ret_q}" title="開発案件を開く" '
            f'style="{_dev_pill}">🛠 {_esc(dp.get("theme") or "開発案件")}</a>'
            for dp in _devs
        )
        dev_links += (f'<a href="/dev-projects/new?deal_id={did}&return_to={_ret_q}" '
                      f'title="この商談に開発案件を新規追加" style="{_add_pill}">＋開発案件</a>')
        # 開発案件のツールリンク（主/追加）をこの場で貼れるフローティングを開くトリガー（#2）
        _link_trg = (f'<button type="button" class="toollink-trigger" '
                     f'onclick="openToolLinkPanel({did}, this)" '
                     f'title="制作したツール/追加リンクを貼る">🔗 リンク貼付</button>')
        tool_cell = (f'<div style="min-width:150px;display:flex;flex-direction:column;gap:4px;align-items:flex-start">'
                     f'<div>{tool_btns}</div>'
                     f'<div style="display:flex;flex-wrap:wrap;gap:4px;align-items:center">{dev_links}{_link_trg}</div></div>')
        if d.get("status") != "closed":
            close_btn = (f'<button type="button" class="btn sec" style="font-size:11px;padding:4px 8px;'
                         f'background:#c53030;color:#fff;border-color:#c53030"'
                         f' onclick="openCloseModal({did}, \'{return_to_url}\')">クローズ</button>')
        else:
            close_btn = '<span class="muted">クローズ済</span>'
        cb_td = (f'<td style="width:28px"><input type="checkbox" name="ids" value="{did}"'
                 f'{" disabled" if _ro else ""}></td>') if bulk else ""
        _closed_badge = ('<br><span style="display:inline-block;background:#c53030;color:#fff;'
                         'border-radius:4px;padding:1px 5px;font-size:10px;margin-top:3px;'
                         'white-space:nowrap">Close(編集不可)</span>') if _ro else ""
        _name_input = (f'<input type="text" value="{_esc(d.get("deal_name"))}"{_dis} '
                       f'onchange="updateDealField({did}, \'deal_name\', this.value)" '
                       f'style="font-size:12px;padding:2px 4px;width:150px{_ro_sty}">')
        rows.append(
            f'<tr class="deal-row{" deal-row-closed" if _ro else ""}" data-account="{_esc((d.get("account_name") or "").lower())}" data-stage="{_esc(d.get("stage") or "")}">'
            f'{cb_td}'
            f'<td class="muted" style="font-size:.8em;color:#888;white-space:nowrap">#{did}{_closed_badge}</td>'
            f'<td><a href="/deal/{did}?return_to={urllib.parse.quote(return_to_url, safe="")}">{_esc(d.get("account_name"))}</a></td>'
            f'<td>{_name_input}</td>'
            f'<td>{_udeal_sel(did, "stage", stages, d.get("stage") or "", disabled=_ro)}</td>'
            f'<td>{_udeal_sel(did, "owner", owners, d.get("owner") or "", disabled=_ro)}</td>'
            f'<td>{_udeal_sel(did, "sub_owner", owners, d.get("sub_owner") or "", disabled=_ro)}</td>'
            f'<td>{_udeal_sel(did, "business_type_l1", l1_list, d.get("business_type_l1") or "", cascade_l1=True, disabled=_ro)}</td>'
            f'<td>{_udeal_sel(did, "business_type_l2", l2_values, d.get("business_type_l2") or "", sel_id=f"l2_{did}", disabled=_ro)}</td>'
            f'<td>{inp_budget}</td><td>{inp_total}</td>'
            f'<td>{inp_ms_date}</td><td>{inp_ms_label}</td>'
            f'<td>{tool_cell}</td><td>{close_btn}</td></tr>'
        )
    body = "".join(rows) or f'<tr><td colspan={16 if bulk else 15} class=muted>商談がありません。</td></tr>'
    return (f'<div style="overflow:auto;max-height:70vh"><table style="min-width:1560px">'
            f'{header}{body}</table></div>' + _MS_PANEL_BLOCK + _TOOL_LINK_PANEL_BLOCK)


# 一覧の「次回MS」欄から全MSを確認/追加/編集/削除するポップオーバー（#48）。全タブ共通・1回だけ出力。
_MS_PANEL_BLOCK = """
<style>
.ms-trigger{font-size:10px;background:#f1f5f9;color:#334155;border:1px solid #cbd5e1;border-radius:5px;
  padding:1px 6px;cursor:pointer;white-space:nowrap}
.ms-trigger:hover{background:#e2e8f0}
.ms-trigger.has-multi{background:#eef2ff;color:#4338ca;border-color:#c7d2fe;font-weight:700}
#dealMsPanel{position:fixed;z-index:1200;display:none;background:#fff;border:1px solid #cbd5e1;
  border-radius:10px;box-shadow:0 12px 40px rgba(0,0,0,.22);padding:12px 14px;width:430px;max-width:94vw;
  max-height:70vh;overflow:auto;font-size:12px}
#dealMsPanel h4{margin:0 0 8px;font-size:13px;display:flex;justify-content:space-between;align-items:center}
#dealMsPanel .msp-row,#dealMsPanel .msp-add{display:flex;gap:5px;align-items:center;margin-bottom:6px;flex-wrap:wrap}
#dealMsPanel input[type=date]{font-size:12px;padding:2px 4px}
#dealMsPanel input[type=text]{font-size:12px;padding:2px 4px;flex:1;min-width:110px}
#dealMsPanel select{font-size:12px;padding:2px 4px;width:auto}
#dealMsPanel .msp-add{border-bottom:1px dashed #cbd5e1;padding-bottom:10px;margin-bottom:10px;background:#f8fafc;border-radius:6px;padding:8px}
#dealMsPanel .msp-del{background:#fef2f2;color:#b91c1c;border:1px solid #fecaca;border-radius:5px;
  cursor:pointer;padding:1px 7px;font-size:12px}
#dealMsPanel .msp-add-btn{background:#4f46e5;color:#fff;border:0;border-radius:6px;cursor:pointer;padding:3px 10px;font-size:12px}
#dealMsPanel .msp-hint{color:#64748b;font-size:10px;margin:0 0 8px}
</style>
<div id="dealMsPanel"></div>
<script>
(function(){
  var TYPES=['アポ','タスク'];
  function esc(s){return (s==null?'':String(s)).replace(/[&<>\\"]/g,function(c){
    return {'&':'&amp;','<':'&lt;','>':'&gt;','\\"':'&quot;'}[c];});}
  function typeOpts(v){var o='<option value=""></option>';TYPES.forEach(function(t){
    o+='<option value="'+t+'"'+(t===v?' selected':'')+'>'+t+'</option>';});return o;}
  function panelEl(){var p=document.getElementById('dealMsPanel');return p;}
  function post(url,body){return fetch(url,{method:'POST',
    headers:{'Content-Type':'application/x-www-form-urlencoded'},body:body}).then(function(r){return r.json();});}
  function syncInline(d){
    var did=d.deal_id,c=d.cache||{};
    var dt=document.getElementById('msdate_'+did); if(dt) dt.value=c.date||'';
    var lb=document.getElementById('mslabel_'+did); if(lb) lb.value=c.label||'';
    var tw=document.getElementById('mstype_'+did); if(tw){var s=tw.querySelector('select'); if(s) s.value=c.type||'';}
    var cnt=document.getElementById('mscount_'+did); if(cnt) cnt.textContent=d.open_count;
    var trg=document.getElementById('mstrg_'+did); if(trg){ if(d.open_count>=2) trg.classList.add('has-multi'); else trg.classList.remove('has-multi'); }
  }
  function render(d, keepPos){
    var p=panelEl(); window.__msDid=d.deal_id;
    var mss=(d.milestones||[]).slice().sort(function(a,b){
      var da=a.date||'',db=b.date||''; if(!da&&!db)return 0; if(!da)return -1; if(!db)return 1;
      return db.localeCompare(da);});  // 新しい日付→古い日付（最古＝次回MSが一番下・空欄は上）
    var rows=mss.map(function(m){
      return '<div class="msp-row">'
        +'<input type="date" data-mid="'+m.id+'" data-mf="date" value="'+esc(m.date)+'">'
        +'<input type="text" data-mid="'+m.id+'" data-mf="label" value="'+esc(m.label)+'" placeholder="ラベル">'
        +'<select data-mid="'+m.id+'" data-mf="type">'+typeOpts(m.type)+'</select>'
        +'<label style="font-size:11px"><input type="checkbox" data-mid="'+m.id+'" data-mf="done"'+(m.done?' checked':'')+'>完了</label>'
        +'<button type="button" class="msp-del" data-del="'+m.id+'">×</button>'
        +'</div>';
    }).join('');
    if(!rows) rows='<p class="msp-hint">MSはまだありません。上の欄で追加してください。</p>';
    var add='<div class="msp-add">'
      +'<input type="date" id="mspNewDate">'
      +'<input type="text" id="mspNewLabel" placeholder="ラベル（例：初回アポ）">'
      +'<select id="mspNewType">'+typeOpts('')+'</select>'
      +'<button type="button" class="msp-add-btn" data-add="1">＋追加</button></div>';
    p.innerHTML='<h4>次回マイルストーン <span data-close="1" style="cursor:pointer;color:#94a3b8">✕</span></h4>'
      +'<p class="msp-hint">新規入力は一番上。下ほど古い日付で、最下段の最も早い日付が「次回MS」として集計されます。変更は即保存。</p>'
      +add+rows;
    p.querySelectorAll('[data-mf]').forEach(function(el){
      el.addEventListener('change',function(){
        var val=(el.type==='checkbox')?(el.checked?'1':'0'):el.value;
        post('/milestone/'+el.dataset.mid+'/field','field='+encodeURIComponent(el.dataset.mf)+'&value='+encodeURIComponent(val))
          .then(function(r){ if(r&&r.ok!==false){ render(r,true); syncInline(r);} });
      });
    });
    p.querySelectorAll('[data-del]').forEach(function(el){
      el.addEventListener('click',function(){ if(!confirm('このMSを削除しますか？')) return;
        post('/milestone/'+el.dataset.del+'/delete','').then(function(r){ if(r&&r.ok!==false){ render(r,true); syncInline(r);} });
      });
    });
    var addBtn=p.querySelector('[data-add]');
    if(addBtn) addBtn.addEventListener('click',function(){
      var dd=document.getElementById('mspNewDate').value;
      var ll=document.getElementById('mspNewLabel').value;
      var tt=document.getElementById('mspNewType').value;
      if(!dd&&!ll){ alert('日付かラベルを入力してください。'); return; }
      post('/deal/'+window.__msDid+'/milestones/add','ms_date='+encodeURIComponent(dd)+'&ms_label='+encodeURIComponent(ll)+'&ms_type='+encodeURIComponent(tt))
        .then(function(r){ if(r&&r.ok!==false){ render(r,true); syncInline(r);} });
    });
    var cl=p.querySelector('[data-close]'); if(cl) cl.addEventListener('click',closePanel);
  }
  function closePanel(){var p=panelEl(); if(p) p.style.display='none';}
  window.openMsPanel=function(did, btn){
    fetch('/deal/'+did+'/milestones').then(function(r){return r.json();}).then(function(d){
      var p=panelEl(); render(d,false); p.style.display='block';
      var rc=btn.getBoundingClientRect();
      var left=Math.min(rc.left, window.innerWidth-p.offsetWidth-10);
      var top=rc.bottom+6;
      if(top+p.offsetHeight>window.innerHeight) top=Math.max(8, window.innerHeight-p.offsetHeight-8);
      p.style.left=Math.max(8,left)+'px'; p.style.top=top+'px';
    });
  };
  document.addEventListener('click',function(e){
    var p=panelEl(); if(!p||p.style.display!=='block') return;
    if(p.contains(e.target)) return;
    if(e.target.closest('.ms-trigger')) return;
    closePanel();
  });
})();
</script>
"""


# 一覧のツール欄から、開発案件の「制作したツール(主)」「追加リンク」をその場で貼れるポップオーバー（#2）。
# 全タブ共通・1回だけ出力。1商談に開発案件が複数ある場合はセレクトで選択、0件なら作成を促す。
_TOOL_LINK_PANEL_BLOCK = """
<style>
.toollink-trigger{font-size:11px;background:#ecfeff;color:#0e7490;border:1px solid #a5f3fc;border-radius:6px;
  padding:2px 8px;cursor:pointer;white-space:nowrap}
.toollink-trigger:hover{background:#cffafe}
#dealToolPanel{position:fixed;z-index:1200;display:none;background:#fff;border:1px solid #cbd5e1;
  border-radius:10px;box-shadow:0 12px 40px rgba(0,0,0,.22);padding:12px 14px;width:420px;max-width:94vw;
  max-height:74vh;overflow:auto;font-size:12px}
#dealToolPanel h4{margin:0 0 8px;font-size:13px;display:flex;justify-content:space-between;align-items:center}
#dealToolPanel .tlp-sec{border:1px solid #e5e7eb;border-radius:8px;padding:8px 10px;margin-bottom:10px}
#dealToolPanel .tlp-sec h5{margin:0 0 6px;font-size:12px;color:#334155}
#dealToolPanel label.tlp-f{display:block;margin-bottom:5px;color:#475569;font-size:11px}
#dealToolPanel input[type=text]{width:100%;box-sizing:border-box;font-size:12px;padding:3px 5px;margin-top:1px}
#dealToolPanel select{font-size:12px;padding:3px 5px;max-width:100%}
#dealToolPanel .tlp-btn{background:#0891b2;color:#fff;border:0;border-radius:6px;cursor:pointer;padding:4px 12px;font-size:12px}
#dealToolPanel .tlp-btn.add{background:#4f46e5}
#dealToolPanel .tlp-hint{color:#64748b;font-size:10px;margin:0 0 8px}
#dealToolPanel .tlp-links a{display:inline-block;background:#eff6ff;color:#1d4ed8;border:1px solid #bfdbfe;
  border-radius:6px;padding:1px 7px;font-size:11px;margin:0 4px 4px 0;text-decoration:none;max-width:100%;
  overflow:hidden;text-overflow:ellipsis;white-space:nowrap;vertical-align:bottom}
#dealToolPanel .tlp-msg{font-size:11px;margin:4px 0}
</style>
<div id="dealToolPanel"></div>
<script>
(function(){
  function esc(s){return (s==null?'':String(s)).replace(/[&<>\\"]/g,function(c){
    return {'&':'&amp;','<':'&lt;','>':'&gt;','\\"':'&quot;'}[c];});}
  function panelEl(){return document.getElementById('dealToolPanel');}
  function post(url,body){return fetch(url,{method:'POST',
    headers:{'Content-Type':'application/x-www-form-urlencoded'},body:body}).then(function(r){return r.json();});}
  function enc(v){return encodeURIComponent(v||'');}
  function proj(){ if(!window.__tlData) return null;
    return window.__tlData.projects.filter(function(p){return p.id===window.__tlSel;})[0]||null; }
  function render(){
    var p=panelEl(), d=window.__tlData;
    if(!d){ p.innerHTML=''; return; }
    var head='<h4>ツールリンクを貼る <span data-close="1" style="cursor:pointer;color:#94a3b8">✕</span></h4>';
    if(!d.projects.length){
      p.innerHTML=head+'<p class="tlp-hint">この商談には開発案件がまだありません。先に開発案件を作成してください。</p>'
        +'<a class="tlp-btn add" style="text-decoration:none" href="/dev-projects/new?deal_id='+d.deal_id+'">＋開発案件を作成</a>';
      bindClose(); return;
    }
    var selHtml='';
    if(d.projects.length>1){
      selHtml='<label class="tlp-f">対象の開発案件<select id="tlpProj">'
        +d.projects.map(function(x){return '<option value="'+x.id+'"'+(x.id===window.__tlSel?' selected':'')+'>'+esc(x.theme)+'</option>';}).join('')
        +'</select></label>';
    } else {
      selHtml='<p class="tlp-hint">開発案件: <b>'+esc(d.projects[0].theme)+'</b></p>';
    }
    var pr=proj();
    var mainSec='<div class="tlp-sec"><h5>制作したツール（主リンク）</h5>'
      +'<label class="tlp-f">URL<input type="text" id="tlpMainUrl" value="'+esc(pr.tool_url)+'" placeholder="https://..."></label>'
      +'<label class="tlp-f">ログインID<input type="text" id="tlpMainId" value="'+esc(pr.tool_login_id)+'"></label>'
      +'<label class="tlp-f">ログインPASS<input type="text" id="tlpMainPass" value="'+esc(pr.tool_login_pass)+'"></label>'
      +'<button type="button" class="tlp-btn" data-savemain="1">主リンクを保存</button>'
      +'<span class="tlp-msg" id="tlpMainMsg"></span></div>';
    var links=(pr.extra||[]).map(function(e){
      return '<a href="'+esc(e.url)+'" target="_blank" rel="noopener" title="'+esc(e.url)+'">'+esc(e.label||e.url)+'</a>';}).join('');
    var addSec='<div class="tlp-sec"><h5>追加リンク</h5>'
      +'<div class="tlp-links">'+(links||'<span class="tlp-hint">まだありません</span>')+'</div>'
      +'<label class="tlp-f">URL<input type="text" id="tlpAddUrl" placeholder="https://..."></label>'
      +'<label class="tlp-f">ラベル<input type="text" id="tlpAddLabel" placeholder="表示名（任意）"></label>'
      +'<label class="tlp-f">ログインID<input type="text" id="tlpAddId"></label>'
      +'<label class="tlp-f">ログインPASS<input type="text" id="tlpAddPass"></label>'
      +'<button type="button" class="tlp-btn add" data-addlink="1">追加リンクを保存</button>'
      +'<span class="tlp-msg" id="tlpAddMsg"></span></div>';
    p.innerHTML=head+selHtml+mainSec+addSec
      +'<p class="tlp-hint">保存内容は個別開発案件のリンクに直接連携されます。一覧のボタン表示はページ再読込後に反映されます。</p>';
    var sp=document.getElementById('tlpProj');
    if(sp) sp.addEventListener('change',function(){ window.__tlSel=parseInt(sp.value,10); render(); });
    var sm=p.querySelector('[data-savemain]');
    if(sm) sm.addEventListener('click',function(){
      var msg=document.getElementById('tlpMainMsg'); msg.textContent='保存中…'; msg.style.color='#64748b';
      post('/dev-project/'+window.__tlSel+'/tool-main',
        'url='+enc(document.getElementById('tlpMainUrl').value)
        +'&login_id='+enc(document.getElementById('tlpMainId').value)
        +'&login_pass='+enc(document.getElementById('tlpMainPass').value))
      .then(function(r){ if(r&&r.ok){ msg.textContent='✓ 保存しました'; msg.style.color='#059669';
          var pp=proj(); if(pp){pp.tool_url=document.getElementById('tlpMainUrl').value;
            pp.tool_login_id=document.getElementById('tlpMainId').value;
            pp.tool_login_pass=document.getElementById('tlpMainPass').value;} }
        else { msg.textContent='✕ '+((r&&r.error)||'保存に失敗しました'); msg.style.color='#b91c1c'; } });
    });
    var al=p.querySelector('[data-addlink]');
    if(al) al.addEventListener('click',function(){
      var msg=document.getElementById('tlpAddMsg'); var url=document.getElementById('tlpAddUrl').value.trim();
      if(!/^https?:\\/\\//i.test(url)){ msg.textContent='✕ URLはhttp(s)で入力してください'; msg.style.color='#b91c1c'; return; }
      msg.textContent='保存中…'; msg.style.color='#64748b';
      post('/dev-project/'+window.__tlSel+'/tool-add',
        'url='+enc(url)+'&label='+enc(document.getElementById('tlpAddLabel').value)
        +'&login_id='+enc(document.getElementById('tlpAddId').value)
        +'&login_pass='+enc(document.getElementById('tlpAddPass').value))
      .then(function(r){ if(r&&r.ok){ reload(); }
        else { msg.textContent='✕ '+((r&&r.error)||'保存に失敗しました'); msg.style.color='#b91c1c'; } });
    });
    bindClose();
  }
  function bindClose(){var cl=panelEl().querySelector('[data-close]'); if(cl) cl.addEventListener('click',closePanel);}
  function closePanel(){var p=panelEl(); if(p) p.style.display='none';}
  function reload(){
    fetch('/deal/'+window.__tlDid+'/dev-tools').then(function(r){return r.json();}).then(function(d){
      window.__tlData=d;
      if(d.projects.length && !d.projects.some(function(x){return x.id===window.__tlSel;})) window.__tlSel=d.projects[0].id;
      render();
    });
  }
  window.openToolLinkPanel=function(did, btn){
    window.__tlDid=did; window.__tlSel=null;
    fetch('/deal/'+did+'/dev-tools').then(function(r){return r.json();}).then(function(d){
      window.__tlData=d; if(d.projects.length) window.__tlSel=d.projects[0].id;
      var p=panelEl(); render(); p.style.display='block';
      var rc=btn.getBoundingClientRect();
      var left=Math.min(rc.left, window.innerWidth-p.offsetWidth-10);
      var top=rc.bottom+6;
      if(top+p.offsetHeight>window.innerHeight) top=Math.max(8, window.innerHeight-p.offsetHeight-8);
      p.style.left=Math.max(8,left)+'px'; p.style.top=top+'px';
    });
  };
  document.addEventListener('click',function(e){
    var p=panelEl(); if(!p||p.style.display!=='block') return;
    if(p.contains(e.target)) return;
    if(e.target.closest('.toollink-trigger')) return;
    closePanel();
  });
})();
</script>
"""


def _ms_type_opts(ms_type: str | None) -> str:
    """次回MS種別の絞り込み<select>用<option>群（全MS種別／各種別／未設定）。"""
    return (
        '<option value="">全MS種別</option>'
        + "".join(f'<option value="{html.escape(t)}"{" selected" if t == ms_type else ""}>{html.escape(t)}</option>'
                  for t in sfa_db.NEXT_MS_TYPES)
        + f'<option value="none"{" selected" if ms_type == "none" else ""}>未設定</option>'
    )


def _stage_multi_filter(stages: list[str], selected=()) -> str:
    """ステージの複数選択フィルタ（クライアント側・チェックボックスのドロップダウン）。
    チェック変更で filterDealsByAccount() を呼び、data-stage を持つ行を絞り込む。全タブ共通。
    name="stg" を付け、囲みフォーム送信(日付変更等)でも選択を持ち越せる。selectedは初期チェック。"""
    sel = set(selected or ())
    boxes = "".join(
        f'<label style="display:block;padding:3px 8px;white-space:nowrap;font-size:12px;cursor:pointer">'
        f'<input type="checkbox" class="stg-cb" name="stg" value="{html.escape(s)}"'
        f'{" checked" if s in sel else ""} onchange="filterDealsByAccount()"> '
        f'{html.escape(s)}</label>'
        for s in stages)
    _lbl = f"ステージ:{len(sel)}選択" if sel else "ステージ:全て"
    return (
        '<details class="tb-menu" style="display:inline-block">'
        f'<summary class="btn sec" style="font-size:12px"><span id="stgFilterLbl">{_lbl}</span> ▾</summary>'
        f'<div class="tb-panel" style="padding:4px 0">{boxes}'
        '<div style="border-top:1px solid #e5e7eb;margin-top:4px;padding-top:4px">'
        '<a href="#" style="font-size:11px;padding:3px 8px;display:block" '
        'onclick="document.querySelectorAll(\'.stg-cb\').forEach(function(c){c.checked=false;});'
        'filterDealsByAccount();return false;">クリア</a></div></div></details>')


def _filter_deals_by_ms_type(deals: list, ms_type: str | None) -> list:
    """次回MS種別で商談リストを後段フィルタ（"none"=未設定・未指定はそのまま）。"""
    if ms_type == "none":
        return [d for d in deals if not (d.get("next_milestone_type"))]
    if ms_type in sfa_db.NEXT_MS_TYPES:
        return [d for d in deals if (d.get("next_milestone_type") or "") == ms_type]
    return deals


def _ms_type_options(selected: str = "") -> str:
    """次回MS種別<select>の<option>群（先頭に空選択）。"""
    return '<option value=""></option>' + "".join(
        f'<option value="{_esc(t)}"{" selected" if t == selected else ""}>{_esc(t)}</option>'
        for t in sfa_db.NEXT_MS_TYPES)


def _ms_row_html(ms: dict) -> str:
    """個別商談フォームの次回MS 1行分（日付/ラベル/種別/完了/削除）。"""
    d = _esc(ms.get("ms_date") or "")
    lb = _esc(ms.get("ms_label") or "")
    tp = ms.get("ms_type") or ""
    done = 1 if ms.get("done") else 0
    return (
        '<div class="ms-row" style="display:flex;gap:6px;align-items:center;margin-bottom:6px;flex-wrap:wrap">'
        f'<input type="date" name="ms_date[]" value="{d}" style="font-size:13px">'
        f'<input type="text" name="ms_label[]" value="{lb}" placeholder="ラベル（例：初回アポ）" '
        'style="flex:1;min-width:150px;font-size:13px">'
        f'<select name="ms_type[]" style="font-size:13px;width:auto">{_ms_type_options(tp)}</select>'
        '<span style="font-size:12px;display:inline-flex;align-items:center;gap:3px">'
        f'<input type="checkbox" class="ms-done-chk"{" checked" if done else ""} '
        "onchange=\"this.parentElement.querySelector('input[type=hidden]').value=this.checked?'1':'0'\">"
        f'<input type="hidden" name="ms_done[]" value="{done}">完了</span>'
        '<button type="button" class="btn sec" style="font-size:11px;padding:3px 8px" '
        "onclick=\"this.closest('.ms-row').remove()\">削除</button>"
        '</div>'
    )


def _ms_editor_html(milestones: list[dict]) -> str:
    """個別商談フォームの次回MS複数行エディタ（行群＋追加ボタン＋テンプレ＋JS）。"""
    rows = "".join(_ms_row_html(m) for m in milestones)
    tpl = _ms_row_html({})  # 追加用の空行テンプレ
    js = ("<script>function addMsRow(){"
          "var w=document.getElementById('msRows');"
          "var t=document.getElementById('msRowTpl');"
          "w.insertAdjacentHTML('beforeend', t.innerHTML);}</script>")
    return (
        '<label>次回マイルストーン（複数設定可・未完了で最も早い日付が「次回MS」として集計されます）</label>'
        f'<div id="msRows" style="margin:2px 0 6px">{rows}</div>'
        '<button type="button" class="btn sec" style="font-size:12px" onclick="addMsRow()">＋ MSを追加</button>'
        f'<template id="msRowTpl">{tpl}</template>' + js
    )


def _ms_panel_json(con, deal_id: int) -> dict:
    """一覧のMS管理パネル用JSON（全MS＋キャッシュ＋未完了件数）。"""
    ms = sfa_db.list_deal_milestones(con, deal_id)
    d = sfa_db.get_deal(con, deal_id) or {}
    return {
        "ok": True, "deal_id": deal_id,
        "milestones": [{"id": m["id"], "date": m.get("ms_date") or "", "label": m.get("ms_label") or "",
                        "type": m.get("ms_type") or "", "done": bool(m.get("done"))} for m in ms],
        "cache": {"date": d.get("next_milestone_date") or "", "label": d.get("next_milestone_label") or "",
                  "type": d.get("next_milestone_type") or ""},
        "open_count": sum(1 for m in ms if not m.get("done")),
    }


def _save_bar(form_id: str, title: str = "", cancel_url: str | None = None, label: str = "💾 保存",
              extra: str = "") -> str:
    """個別編集/入力フォーム上部の固定・保存バー。ボタンは form="<id>" で対象フォームを送信する
    （HTML5のform属性。バー自体は<form>の外にあってよい）。
    extra: タイトルとボタンの間に差し込む任意HTML（ステータス表示などの固定表示用・rawで挿入）。"""
    t = f'<span class="sb-title">{_esc(title)}</span>' if title else ""
    ex = f'<span class="sb-extra" style="margin-right:12px;display:flex;align-items:center;gap:6px">{extra}</span>' if extra else ""
    cancel = f'<a class="btn sec" href="{_esc(cancel_url)}">キャンセル</a>' if cancel_url else ""
    # 二重送信防止: フォームが妥当なら送信後にボタンを無効化（連打で多重登録されるのを防ぐ）。
    # 未入力等でバリデーションに落ちる場合はロックしない。
    guard = ("var f=this.form;if(f&&!f.checkValidity())return true;"
             "if(this.dataset.busy)return false;this.dataset.busy=1;var b=this;"
             "setTimeout(function(){b.disabled=true;b.textContent='保存中…';},50);")
    return (f'<div class="save-bar">{t}{ex}'
            f'<button class="btn" type="submit" form="{form_id}" onclick="{guard}">{label}</button>{cancel}</div>')


def home_page(con, owner: str | None = None, status_filter: str | None = None,
              stage_filter: str | None = None, ms_type: str | None = None, stages_sel=None) -> str:
    # デフォルトでclosedを除外（NULLもopenとして扱う）。"all"は全件表示
    effective_status = None if status_filter == "all" else (status_filter or "open")
    # ステージ絞り込みはクライアント側の複数選択(_stage_multi_filter)に統一したためサーバ側では絞らない。
    deals = sfa_db.list_deals(con, status=effective_status, owner=owner)
    deals = _filter_deals_by_ms_type(deals, ms_type)  # 次回MS種別で後段フィルタ
    pending_sync = con.execute("SELECT COUNT(*) c FROM deals WHERE theme_id IS NULL").fetchone()["c"]
    owners = sfa_db.get_master_list(con, "owners")
    stages = sfa_db.get_master_list(con, "deal_stages")
    biz_l1_list = sfa_db.get_master_list(con, "business_type_l1")
    owner_opts = '<option value="">全担当</option>' + "".join(
        f'<option value="{html.escape(o)}"{" selected" if o == owner else ""}>{html.escape(o)}</option>'
        for o in owners
    )
    status_opts = (
        f'<option value="all"{"  selected" if status_filter=="all" else ""}>全て（クローズ含む）</option>'
        + f'<option value="open"{"  selected" if status_filter is None or status_filter=="open" else ""}>進行中のみ</option>'
        + f'<option value="closed"{" selected" if status_filter=="closed" else ""}>クローズ済のみ</option>'
    )
    # 選択するだけで自動絞り込み（onchangeで即送信・「絞り込み」ボタンは廃止）
    # フィルタUIは全タブで「サーバ側select群 → 次回MS種別 → 🔍アカウント検索 → リセット」の順で統一。
    filter_row = f"""<form method="get" action="/deals" class="filter-row">
      <input type="hidden" name="tab" value="active">
      <select name="owner" onchange="this.form.submit()">{owner_opts}</select>
      <select name="status" onchange="this.form.submit()">{status_opts}</select>
      {_stage_multi_filter(stages, selected=stages_sel)}
      <select name="ms_type" onchange="this.form.submit()" title="次回MSの種別で絞り込み">{_ms_type_opts(ms_type)}</select>
      <input type="text" id="accSearchInput" placeholder="🔍 アカウント名で検索..."
        oninput="filterDealsByAccount()" style="max-width:220px">
      <a class="btn sec" href="/deals?tab=active">リセット</a>
    </form>"""
    # バルク編集用JSオブジェクト構築
    deal_bulk_options = {
        "stage": [["", "（変更なし）"]] + [[s, s] for s in stages],
        "owner": [["", "（変更なし）"]] + [[o, o] for o in owners],
        "sub_owner": [["", "（変更なし）"]] + [[o, o] for o in owners],
        "business_type_l1": [["", "（変更なし）"]] + [[v, v] for v in biz_l1_list],
    }
    deal_bulk_options_json = json.dumps(deal_bulk_options, ensure_ascii=False)

    accounts = sfa_db.list_accounts(con)
    acc_rows = "".join(
        f'<tr><td><a href="/account/{a["id"]}">{_esc(a["name"])}</a></td>'
        f'<td>{_esc(a.get("industry"))}</td><td>{_esc(a.get("company_size"))}</td></tr>'
        for a in accounts
    )
    sync_btn = (
        f'<form method="post" action="/deals/sync_pending" style="display:inline">'
        f'<button class="btn sec" style="font-size:12px" '
        f'onclick="return confirm(\'テーマDB未同期の商談{pending_sync}件をバックグラウンドで同期します。よろしいですか？\')">'
        f'⏳ テーマDB未同期 {pending_sync}件を同期</button></form>'
        if pending_sync else ""
    )
    # Hisho同期に失敗して記録が残っている件数（あれば再同期ボタンを出す）
    _sync_fail_n = sfa_db.count_sync_failures(con)
    sync_fail_btn = (
        f'<form method="post" action="/sync-failures/retry" style="display:inline">'
        f'<button class="btn" style="font-size:12px;background:#c53030" '
        f'onclick="return confirm(\'Hisho同期に失敗した{_sync_fail_n}件の再同期を試みます。よろしいですか？\')">'
        f'⚠ 同期失敗 {_sync_fail_n}件を再同期</button></form>'
        if _sync_fail_n else ""
    )
    return f"""
    <div class="card"><h2 style="display:flex;justify-content:space-between;align-items:center">
      <span>商談 ({len(deals)})</span>
      <span style="display:flex;gap:8px;flex-wrap:wrap;align-items:center">
        {sync_fail_btn}
        {sync_btn}
        <details class="tb-menu">
          <summary class="btn sec">＋ 新規 ▾</summary>
          <div class="tb-panel">
            <a href="/dev-projects/new">＋ 開発案件</a>
            <a href="/hearing/new">＋ ヒアリング</a>
            <a href="/deal-issue/new">＋ 論点</a>
          </div>
        </details>
        <details class="tb-menu">
          <summary class="btn sec">取込・整備 ▾</summary>
          <div class="tb-panel">
            <a href="/deals/import">📥 CSV取込</a>
            <a href="/data-tagging">🏷 データ整備</a>
            <a href="/sync-health">🔍 同期チェック</a>
          </div>
        </details>
        <a class="btn" href="/deal/new">＋商談追加</a>
      </span>
    </h2>
    {filter_row}
    <form id="deal_bulk_form" method="post" action="/deals/bulk_edit">
    {unified_deal_table(con, deals, return_to_url="/deals", bulk=True)}
    <div style="display:flex;align-items:center;gap:8px;margin-top:10px;flex-wrap:wrap">
      <select id="deal_bulk_field" name="field" style="width:auto">
        <option value="stage">ステージ</option>
        <option value="owner">主担当</option>
        <option value="sub_owner">サブ担当</option>
        <option value="business_type_l1">事業種別L1</option>
      </select>
      <select id="deal_bulk_value" name="value" style="width:auto"></select>
      <button class="btn sec" type="submit">選択した件を一括変更</button>
      <span style="width:1px;height:20px;background:#e5e7eb;margin:0 2px"></span>
      <button class="btn sec" type="submit" formaction="/deliveries/bulk_new" formnovalidate
              title="チェックした商談にDelivery（受注後アサイン計画）を登録します"
              style="background:#ecfeff;color:#0e7490;border-color:#a5f3fc">🚚 選択商談にDeliveryを登録</button>
    </div>
    </form>
    </div>
    <div class="card"><h2>アカウント ({len(accounts)})</h2>
    <table><tr><th>企業名</th><th>業界</th><th>企業規模</th></tr>
    {acc_rows or '<tr><td colspan=3 class=muted>まだアカウントがありません。</td></tr>'}
    </table></div>
    <script>
    const DEAL_BULK_OPTIONS = {deal_bulk_options_json};
    const DEAL_L2_MAP = {json.dumps(sfa_db.BUSINESS_TYPE_L2_BY_L1, ensure_ascii=False)};
    function updateDealL1(id, l1_value) {{
      updateDealField(id, 'business_type_l1', l1_value);
      var l2sel = document.getElementById('l2_' + id);
      if (l2sel) {{
        var opts = DEAL_L2_MAP[l1_value] || [];
        l2sel.innerHTML = '<option value=""></option>' +
          opts.map(function(v) {{ return '<option value="' + escH(v) + '">' + escH(v) + '</option>'; }}).join('');
        l2sel.value = '';
        updateDealField(id, 'business_type_l2', '');
      }}
    }}
    function updateDealField(id, field, value) {{
      fetch('/deal/' + id + '/field', {{
        method: 'POST',
        headers: {{'Content-Type': 'application/x-www-form-urlencoded'}},
        body: 'field=' + encodeURIComponent(field) + '&value=' + encodeURIComponent(value)
      }}).then(r => r.json()).then(d => {{
        if (!d.ok) alert('更新エラー');
      }}).catch(() => alert('通信エラー'));
    }}
    function filterDealsByAccount() {{
      var q = document.getElementById('accSearchInput').value.trim().toLowerCase();
      document.querySelectorAll('#deal_bulk_form tr.deal-row').forEach(function(row) {{
        row.style.display = (!q || row.dataset.account.includes(q)) ? '' : 'none';
      }});
    }}
    function repopulateDealBulkValue() {{
      var field = document.getElementById('deal_bulk_field').value;
      var opts = DEAL_BULK_OPTIONS[field] || [];
      var sel = document.getElementById('deal_bulk_value');
      sel.innerHTML = opts.map(function(pair) {{
        return '<option value="' + escH(pair[0]) + '">' + escH(pair[1]) + '</option>';
      }}).join('');
    }}
    document.getElementById('deal_bulk_field').addEventListener('change', repopulateDealBulkValue);
    repopulateDealBulkValue();
    </script>
    """


def resolve_default_deals_tab(con, explicit_tab: str | None, owner: str | None = None) -> str:
    """商談一覧のデフォルトタブを決める。ユーザーがtabを明示していればそれを尊重。
    未指定のときは「MS超過(overdue)」を既定とするが、MS超過は既定で「当日MSを除く」ため、
    当日ちょうどのMSを除いた要フォローが0件なら「進行中(active)」へフォールバックする
    （＝当日MSしか無い/滞留も無いなら空のMS超過タブを見せない）。"""
    if explicit_tab:
        return explicit_tab
    try:
        return "overdue" if sfa_db.list_overdue_deals(
            con, owner=owner, today=_today_jst().isoformat(), exclude_today=True) else "active"
    except Exception:  # noqa: BLE001 — 集計失敗時は従来どおりoverdue
        return "overdue"


def deals_page(con, *, tab: str = "active", owner: str | None = None, status_filter: str | None = None,
               stage_filter: str | None = None, date: str | None = None, ms_type: str | None = None,
               week: str | None = None, exclude_today: bool = False, stages_sel=None) -> str:
    """商談一覧をタブ化: 「進行中の商談」(既存home_page)・「特定日の商談」・「MS超過の商談」。"""
    is_by_date = tab == "byDate"
    is_overdue = tab == "overdue"
    tab_nav = f"""
    <div style="display:flex;gap:8px;margin-bottom:14px">
      <a class="{'btn' if is_overdue else 'btn sec'}" href="/deals?tab=overdue">MS超過の商談</a>
      <a class="{'btn' if is_by_date else 'btn sec'}" href="/deals?tab=byDate">特定日の商談</a>
      <a class="{'btn' if not (is_by_date or is_overdue) else 'btn sec'}" href="/deals?tab=active">進行中の商談</a>
    </div>"""
    if is_overdue:
        body = overdue_deals_page(con, owner=owner, ms_type=ms_type, exclude_today=exclude_today)
    elif is_by_date:
        body = deals_by_date_page(con, target_date=date, owner=owner, ms_type=ms_type, week=week,
                                  stages_sel=stages_sel)
    else:
        body = home_page(con, owner=owner, status_filter=status_filter, stage_filter=stage_filter,
                         ms_type=ms_type, stages_sel=stages_sel)
    return tab_nav + body


def deals_by_date_page(con, *, target_date: str | None = None, owner: str | None = None,
                       ms_type: str | None = None, week: str | None = None, stages_sel=None) -> str:
    """特定日 または 特定週 に次回MS日付または活動がある商談の一覧・その場編集画面。
    week(YYYY-Www)が指定されればその週(月〜日)、無ければ日付(既定は当日)で表示する。"""
    owners = sfa_db.get_master_list(con, "owners")
    stages = sfa_db.get_master_list(con, "deal_stages")
    owner_opts = '<option value="">全担当</option>' + "".join(
        f'<option value="{html.escape(o)}"{" selected" if o == owner else ""}>{html.escape(o)}</option>'
        for o in owners
    )
    # 週モード判定（YYYY-Www）。妥当なら月曜〜日曜を算出。
    week_mode = False
    wk_start = wk_end = None
    if week and "-W" in week:
        try:
            _y, _w = week.split("-W")
            wk_start = date.fromisocalendar(int(_y), int(_w), 1)
            wk_end = date.fromisocalendar(int(_y), int(_w), 7)
            week_mode = True
        except (ValueError, TypeError):
            week_mode = False
    target_date = target_date or _today_jst().isoformat()
    try:
        _cur = date.fromisoformat(target_date)
    except ValueError:
        _cur = _today_jst()
        target_date = _cur.isoformat()

    _owner_qs = f"&owner={urllib.parse.quote(owner)}" if owner else ""
    _mstype_qs = f"&ms_type={urllib.parse.quote(ms_type)}" if ms_type else ""
    # 日付/週を変えてもステージ絞り込みを保持（前後ナビのURLにstgを引き回す）。
    _stg_qs = "".join(f"&stg={urllib.parse.quote(s)}" for s in (stages_sel or []))
    _keep_qs = f"{_owner_qs}{_mstype_qs}{_stg_qs}"
    if week_mode:
        _prev_href = f"/deals?tab=byDate&week={(wk_start - timedelta(days=7)).strftime('%G-W%V')}{_keep_qs}"
        _next_href = f"/deals?tab=byDate&week={(wk_start + timedelta(days=7)).strftime('%G-W%V')}{_keep_qs}"
        _prev_lbl, _next_lbl = "◀ 前週", "翌週 ▶"
    else:
        _prev_href = f"/deals?tab=byDate&date={(_cur - timedelta(days=1)).isoformat()}{_keep_qs}"
        _next_href = f"/deals?tab=byDate&date={(_cur + timedelta(days=1)).isoformat()}{_keep_qs}"
        _prev_lbl, _next_lbl = "◀ 前日", "翌日 ▶"
    _date_val = "" if week_mode else target_date
    _week_val = week if week_mode else ""
    # 日入力・週入力は片方を選ぶともう片方をクリアして自動送信（選択即反映・「表示」ボタン廃止）
    form = f"""<form method="get" action="/deals" class="filter-row">
      <input type="hidden" name="tab" value="byDate">
      <a class="btn sec" href="{_prev_href}">{_prev_lbl}</a>
      <input type="date" name="date" value="{_esc(_date_val)}" title="日で表示"
        onchange="this.form.week.value='';this.form.submit()">
      <input type="week" name="week" value="{_esc(_week_val)}" title="週で表示"
        onchange="this.form.date.value='';this.form.submit()">
      <a class="btn sec" href="{_next_href}">{_next_lbl}</a>
      <select name="owner" onchange="this.form.submit()">{owner_opts}</select>
      <select name="ms_type" onchange="this.form.submit()" title="次回MSの種別で絞り込み">{_ms_type_opts(ms_type)}</select>
      {_stage_multi_filter(stages, selected=stages_sel)}
      <input type="text" id="accSearchInput" placeholder="🔍 アカウント名で検索..."
        oninput="filterDealsByAccount()" style="max-width:220px">
      <a class="btn sec" href="/deals?tab=byDate">リセット</a>
    </form>"""

    if week_mode:
        deals = [dict(d) for d in sfa_db.list_deals_by_week(
            con, wk_start.isoformat(), wk_end.isoformat(), owner=owner)]
        _title = (f"特定週の商談（{wk_start.month}/{wk_start.day}〜{wk_end.month}/{wk_end.day}"
                  f"・{_esc(week)}）")
        _ret = {"tab": "byDate", "week": week or "", "owner": owner or "", "ms_type": ms_type or ""}
    else:
        deals = [dict(d) for d in sfa_db.list_deals_by_date(con, target_date, owner=owner)]
        _title = f"特定日の商談（{_esc(target_date)}）"
        _ret = {"tab": "byDate", "date": target_date or "", "owner": owner or "", "ms_type": ms_type or ""}
    deals = _filter_deals_by_ms_type(deals, ms_type)
    return_to_url = f"/deals?{urllib.parse.urlencode(_ret)}"

    return f"""
    <div class="card"><h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
      <span>{_title} {len(deals)}件</span><a class="btn" href="/deal/new">＋商談追加</a></h2>
    {form}
    {unified_deal_table(con, deals, return_to_url=return_to_url, bulk=False)}
    </div>
    """


def overdue_deals_page(con, *, owner: str | None = None, ms_type: str | None = None,
                       exclude_today: bool = False) -> str:
    """次回MSが超過した／未設定の進行中商談の一覧・その場編集（要フォロー）。超過=next_milestone_date<=当日。

    担当・次回MS種別フィルタ（サーバ側）＋アカウント名検索（クライアント側）＋インライン編集。
    見せ方・編集挙動は「特定日の商談」タブに揃えている。
    exclude_today=Trueで「次回MS日が当日ちょうど」の商談を除外（前日以前の超過のみ）。
    """
    owners = sfa_db.get_master_list(con, "owners")
    stages = sfa_db.get_master_list(con, "deal_stages")
    owner_opts = '<option value="">全担当</option>' + "".join(
        f'<option value="{html.escape(o)}"{" selected" if o == owner else ""}>{html.escape(o)}</option>'
        for o in owners
    )
    # 選択するだけで自動絞り込み（onchangeで即送信・「担当で絞り込み」ボタンは廃止）
    # フィルタUIの並びは全タブ統一（サーバ側select群 → 次回MS種別 → 🔍アカウント検索 → リセット）。
    # 「当日MSを除く」トグル。既定は除外(exclude_today=True)。含める場合は exclude_today=0 を明示。
    _base_q = {"tab": "overdue", "owner": owner or "", "ms_type": ms_type or ""}
    if exclude_today:
        _toggle_href = "/deals?" + urllib.parse.urlencode({**_base_q, "exclude_today": "0"})
        _toggle_btn = (f'<a class="btn sec" href="{_toggle_href}" '
                       f'title="当日ちょうどのMSも含めて表示する">'
                       f'🕑 当日MSを含める</a>')
    else:
        _toggle_href = "/deals?" + urllib.parse.urlencode({**_base_q, "exclude_today": "1"})
        _toggle_btn = (f'<a class="btn" href="{_toggle_href}" '
                       f'style="background:#2563eb" title="次回MS日が当日ちょうどの商談を隠し、前日以前の超過のみ表示">'
                       f'🕑 当日MSを除く</a>')
    form = f"""<form method="get" action="/deals" class="filter-row">
      <input type="hidden" name="tab" value="overdue">
      <input type="hidden" name="exclude_today" value="{'1' if exclude_today else '0'}">
      <select name="owner" onchange="this.form.submit()">{owner_opts}</select>
      <select name="ms_type" onchange="this.form.submit()" title="次回MSの種別で絞り込み">{_ms_type_opts(ms_type)}</select>
      <input type="text" id="accSearchInput" placeholder="🔍 アカウント名で検索..."
        oninput="filterDealsByAccount()" style="max-width:220px">
      {_toggle_btn}
      <a class="btn sec" href="/deals?tab=overdue">リセット</a>
    </form>"""
    _return_qs = urllib.parse.urlencode(
        {**_base_q, "exclude_today": "1" if exclude_today else "0"})
    return_to_url = f"/deals?{_return_qs}"

    deals = [dict(d) for d in sfa_db.list_overdue_deals(
        con, owner=owner, today=_today_jst().isoformat(), exclude_today=exclude_today)]
    deals = _filter_deals_by_ms_type(deals, ms_type)
    _hint = ("次回MS日が<b>前日以前</b>の超過商談のみ（当日ちょうどのMSは除外中）＋次回MS未設定の進行中商談です。"
             if exclude_today else
             "次回MS日が本日以前、または次回MSが未設定の進行中商談です（要フォロー）。")
    return f"""
    <div class="card"><h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
      <span>MS超過の商談 {len(deals)}件</span><a class="btn" href="/deal/new">＋商談追加</a></h2>
    <p class="muted" style="margin:0 0 10px">{_hint}超過分を遅れている順に、次回MS未設定は末尾に並べています。</p>
    {form}
    {unified_deal_table(con, deals, return_to_url=return_to_url, bulk=False)}
    </div>
    """


def _tag_buttons(make_onclick, values: list[str]) -> str:
    """タグ付け用のワンクリック・ボタン群。押すと即保存し「✓値」表示に変わる（ドロップダウン不要）。
    make_onclick(value) が各ボタンの onclick 文字列を返す。onclick属性は二重引用符・JS引数は単引用符で衝突回避。"""
    btns = "".join(
        '<button type="button" class="tagbtn" onclick="' + make_onclick(v) + '">' + html.escape(v) + '</button>'
        for v in values
    )
    return '<div class="tagbtns">' + btns + '</div>'


def data_tagging_page(con) -> str:
    """データ整備タグ付け画面。既存データを遡って素早くタグ付けする（バックフィル）。
      ① 次回MSの種別(アポ/タスク)が未設定の進行中商談 … #28
      ② 終了理由(close_reason)が未設定のクローズ商談 … #26
      ③ 終了理由(lost_reason)が未設定の lost リード … #26
    保存は既存のインライン編集エンドポイント(/deal/{id}/field, /leads/{id}/field)を再利用。
    """
    ms_deals = sfa_db.list_untyped_milestone_deals(con)
    closed_deals = sfa_db.list_unclassified_closed_deals(con)
    lost_leads = sfa_db.list_unclassified_lost_leads(con)

    def _note_preview(v):
        s = (v or "").strip().replace("\n", " ")
        if not s:
            return '<span class="muted">（メモなし）</span>'
        return _esc(s[:140] + ("…" if len(s) > 140 else ""))

    ms_list = []
    for d in ms_deals:
        did = d["id"]
        btns = _tag_buttons(lambda v, i=did: f"tagDeal({i},'next_milestone_type','{v}',this)", sfa_db.NEXT_MS_TYPES)
        ms_list.append(
            f'<tr>'
            f'<td><a href="/deal/{did}?return_to=%2Fdata-tagging">{_esc(d.get("account_name"))}</a></td>'
            f'<td>{_esc(d.get("deal_name"))}</td>'
            f'<td>{btns}</td>'
            f'<td style="white-space:nowrap">{_esc(d.get("next_milestone_date"))}</td>'
            f'<td>{_esc(d.get("next_milestone_label"))}</td></tr>'
        )
    ms_rows = "".join(ms_list) or '<tr><td colspan=5 class=muted>未タグの次回MSはありません。</td></tr>'

    cd_list = []
    for d in closed_deals:
        did = d["id"]
        btns = _tag_buttons(lambda v, i=did: f"tagDeal({i},'close_reason','{v}',this)", sfa_db.CLOSE_REASONS)
        cd_list.append(
            f'<tr>'
            f'<td><a href="/deal/{did}?return_to=%2Fdata-tagging">{_esc(d.get("account_name"))}</a><br>'
            f'<span class="muted" style="font-size:.85em">{_esc(d.get("deal_name"))} / {_esc(d.get("stage"))}</span></td>'
            f'<td>{btns}</td>'
            f'<td style="font-size:.85em;max-width:420px">{_note_preview(d.get("note"))}</td></tr>'
        )
    cd_rows = "".join(cd_list) or '<tr><td colspan=3 class=muted>未分類のクローズ商談はありません。</td></tr>'

    ll_list = []
    for l in lost_leads:
        lid = l["id"]
        btns = _tag_buttons(lambda v, i=lid: f"tagLead({i},'{v}',this)", sfa_db.CLOSE_REASONS)
        ll_list.append(
            f'<tr>'
            f'<td><a href="/leads/{lid}">{_esc(l.get("name"))}</a><br>'
            f'<span class="muted" style="font-size:.85em">{_esc(l.get("company"))}</span></td>'
            f'<td>{btns}</td>'
            f'<td style="font-size:.85em;max-width:420px">{_note_preview(l.get("notes"))}</td></tr>'
        )
    ll_rows = "".join(ll_list) or '<tr><td colspan=3 class=muted>未分類の lost リードはありません。</td></tr>'

    # ④ 日付なしの活動履歴（面談集計に載らない不整合データ）のクリーンアップ
    undated = sfa_db.list_undated_activities(con)
    ud_list = []
    for a in undated:
        aid = a["id"]
        if a.get("deal_id"):
            _dl = (f'<a href="/deal/{a["deal_id"]}?return_to=%2Fdata-tagging">'
                   f'{_esc(a.get("account_name") or a.get("deal_name") or ("商談#" + str(a["deal_id"])))}</a>')
        else:
            _dl = '<span class="muted">（商談なし）</span>'
        ud_list.append(
            f'<tr id="udrow_{aid}">'
            f'<td>{_dl}<br><span class="muted" style="font-size:.85em">{_esc(a.get("deal_name") or "")}</span></td>'
            f'<td>{_esc(a.get("type") or "—")}</td>'
            f'<td>{_esc(a.get("contact_name") or "—")}</td>'
            f'<td style="font-size:.85em;max-width:340px">{_note_preview(a.get("body"))}</td>'
            f'<td style="white-space:nowrap">'
            f'<input type="date" id="ud_date_{aid}" style="font-size:12px;padding:1px 3px"> '
            f'<button class="btn sec" style="font-size:11px;padding:3px 8px" onclick="udSetDate({aid})">日付保存</button> '
            f'<button class="btn" style="font-size:11px;padding:3px 8px;background:#c53030;color:#fff;border-color:#c53030" '
            f'onclick="udDelete({aid})">削除</button></td></tr>')
    ud_rows = "".join(ud_list) or '<tr><td colspan=5 class=muted>日付なしの活動履歴はありません。</td></tr>'

    return f"""
    <style>
      .tagbtns{{display:flex;flex-wrap:wrap;gap:5px}}
      .tagbtn{{font-size:12px;padding:5px 11px;border:1px solid #cbd5e1;border-radius:15px;
        background:#fff;color:#334155;cursor:pointer;white-space:nowrap}}
      .tagbtn:hover{{border-color:#2f6fed;color:#2f6fed;background:#eef4ff}}
      .tagged-ok{{color:#166534;font-weight:700;font-size:13px;white-space:nowrap}}
    </style>
    <div class="card">
      <h2>データ整備 — タグ付け</h2>
      <p class="muted" style="margin:0 0 4px">選択肢を<b>クリックすると即保存</b>され「✓ 値」に変わります（ドロップダウンを開く必要はありません）。</p>
      <p class="muted" style="margin:0">残り：次回MS種別 <b>{len(ms_deals)}</b>件／終了理由(商談) <b>{len(closed_deals)}</b>件／終了理由(リード) <b>{len(lost_leads)}</b>件／日付なし活動 <b>{len(undated)}</b>件</p>
    </div>

    <div class="card">
      <h2>① 次回MSの種別が未設定（{len(ms_deals)}件）</h2>
      <p class="muted" style="margin:0 0 8px">「タスク」にするとSlack翌日アポ通知から除外されます（未設定は投稿されます）。</p>
      <form method="post" action="/data-tagging/bulk-appt" style="margin:0 0 12px"
        onsubmit="return confirm('次回MSが明日以降（7/14以降）で、ラベルに「初回アポ」を含む未タグの商談を、まとめて「アポ」にします。よろしいですか？')">
        <button class="btn sec" type="submit" style="font-size:12px">⚡ 「初回アポ」を含む未タグ（明日以降）を一括でアポにする</button>
      </form>
      <div style="overflow:auto;max-height:60vh">
      <table style="min-width:700px"><tr>
        {_sticky_th('アカウント')}{_sticky_th('案件名')}{_sticky_th('種別')}{_sticky_th('次回MS日')}{_sticky_th('ラベル')}</tr>
      {ms_rows}
      </table></div>
    </div>

    <div class="card">
      <h2>② 終了理由が未設定のクローズ商談（{len(closed_deals)}件）</h2>
      <p class="muted" style="margin:0 0 8px">5区分から選択。展示会ファネルの有効母数（ニーズなし除外）に反映されます。</p>
      <div style="overflow:auto;max-height:60vh">
      <table style="min-width:820px"><tr>
        {_sticky_th('商談')}{_sticky_th('終了理由')}{_sticky_th('メモ（先頭のみ）')}</tr>
      {cd_rows}
      </table></div>
    </div>

    <div class="card">
      <h2>③ 終了理由が未設定の lost リード（{len(lost_leads)}件）</h2>
      <div style="overflow:auto;max-height:60vh">
      <table style="min-width:820px"><tr>
        {_sticky_th('リード')}{_sticky_th('終了理由')}{_sticky_th('メモ（先頭のみ）')}</tr>
      {ll_rows}
      </table></div>
    </div>

    <div class="card">
      <h2>④ 日付なしの活動履歴（{len(undated)}件）</h2>
      <p class="muted" style="margin:0 0 8px">日付が未入力の活動は面談集計に数えません。正しい日付を入れて保存するか、不要なら削除してください。</p>
      <div style="overflow:auto;max-height:60vh">
      <table style="min-width:900px"><tr>
        {_sticky_th('商談')}{_sticky_th('種別')}{_sticky_th('相手')}{_sticky_th('内容（先頭のみ）')}{_sticky_th('日付を入れる / 削除')}</tr>
      {ud_rows}
      </table></div>
    </div>

    <script>
    function _tagPost(url, field, value) {{
      return fetch(url, {{method:'POST',
        headers:{{'Content-Type':'application/x-www-form-urlencoded'}},
        body:'field=' + encodeURIComponent(field) + '&value=' + encodeURIComponent(value)}})
        .then(r => r.json());
    }}
    function _tagOk(btn, value) {{
      var box = btn.parentNode;
      if (box) box.innerHTML = '<span class="tagged-ok">✓ ' + value + '</span>';
    }}
    function tagDeal(id, field, value, btn) {{
      _tagPost('/deal/' + id + '/field', field, value)
        .then(d => {{ if (d.ok) _tagOk(btn, value); else alert('更新エラー'); }})
        .catch(() => alert('通信エラー'));
    }}
    function tagLead(id, value, btn) {{
      _tagPost('/leads/' + id + '/field', 'lost_reason', value)
        .then(d => {{ if (d.ok) _tagOk(btn, value); else alert('更新エラー'); }})
        .catch(() => alert('通信エラー'));
    }}
    function udSetDate(id) {{
      var el = document.getElementById('ud_date_' + id);
      var v = el ? el.value : '';
      if (!v) {{ alert('日付を選択してください'); return; }}
      _tagPost('/activity/' + id + '/field', 'occurred_on', v)
        .then(d => {{ if (d.ok) {{ var r = document.getElementById('udrow_' + id); if (r) r.remove(); }}
                    else alert('更新エラー'); }})
        .catch(() => alert('通信エラー'));
    }}
    function udDelete(id) {{
      if (!confirm('この活動履歴を削除しますか？')) return;
      fetch('/activity/' + id + '/delete', {{method:'POST'}})
        .then(() => {{ var r = document.getElementById('udrow_' + id); if (r) r.remove(); }})
        .catch(() => alert('通信エラー'));
    }}
    </script>
    """


def exhibition_tagging_page(con) -> str:
    """展示会由来商談(lead_pattern='Exh.')に「どの展示会か」をタグ付けする画面。
    各行の入力欄で即保存(ajax)。または複数選択→展示会名を一括設定(POST /exhibition-tagging/bulk)。
    並びは登録が古い順(created_at昇順)。"""
    deals = [dict(r) for r in con.execute(
        "SELECT d.id, d.deal_name, d.exhibition_name, acc.name AS account_name "
        "FROM deals d LEFT JOIN accounts acc ON acc.id=d.account_id "
        "WHERE d.lead_pattern='Exh.' "
        "ORDER BY d.created_at ASC, d.id ASC")]
    names = sfa_db.list_exhibition_names(con)
    _dl = "".join(f'<option value="{_esc(n)}">' for n in names)
    untagged = sum(1 for d in deals if not (d.get("exhibition_name") or "").strip())
    rows = []
    for d in deals:
        did = d["id"]
        rows.append(
            f'<tr id="exrow_{did}">'
            f'<td style="width:28px"><input type="checkbox" class="ex-chk" name="ids" value="{did}"></td>'
            f'<td><a href="/deal/{did}?return_to=%2Fexhibition-tagging">{_esc(d.get("account_name") or "—")}</a>'
            f'<br><span class="muted" style="font-size:.85em">{_esc(d.get("deal_name") or "")}</span></td>'
            f'<td><input type="text" id="exname_{did}" list="exNames" value="{_esc(d.get("exhibition_name") or "")}" '
            f'placeholder="展示会名を入力/選択" style="font-size:12px;padding:2px 6px;width:240px" '
            f'onchange="exSave({did})"> '
            f'<span id="exok_{did}" style="font-size:11px"></span></td></tr>')
    body_rows = "".join(rows) or '<tr><td colspan=3 class=muted>展示会由来(lead_pattern=Exh.)の商談はありません。</td></tr>'
    return f"""
    <div class="card">
      <h2>🎪 展示会名タグ付け（{len(deals)}件／未設定 {untagged}）</h2>
      <p class="muted" style="margin:0 0 8px">展示会由来(lead_pattern='Exh.')の商談に「どの展示会か」を付けます。登録が古い順。
      個別は入力欄で即保存。まとめて付けるときは行を選択→下の「一括設定」。</p>
      <datalist id="exNames">{_dl}</datalist>
      <form method="post" action="/exhibition-tagging/bulk">
      <div class="filter-row" style="margin:0 0 10px">
        <input type="text" name="exhibition_name" list="exNames" placeholder="一括設定する展示会名"
          style="max-width:260px">
        <button class="btn" type="submit"
          onclick="return exBulkConfirm(this)">選択した件に一括設定</button>
        <span class="muted" style="font-size:12px;align-self:center">＝チェックした行すべてにこの展示会名を設定</span>
      </div>
      <div style="overflow:auto;max-height:70vh">
      <table style="min-width:680px"><tr>
        <th class="sticky" style="width:28px"><input type="checkbox" id="exChkAll"
          onchange="document.querySelectorAll('.ex-chk').forEach(function(c){{c.checked=this.checked;}}.bind(this))"></th>
        {_sticky_th('商談')}{_sticky_th('展示会名')}</tr>
      {body_rows}</table></div>
      </form>
    </div>
    <script>
    function exSave(id) {{
      var el = document.getElementById('exname_' + id); var v = el ? el.value : '';
      var s = document.getElementById('exok_' + id); if (s) {{ s.textContent = '…'; s.style.color = '#64748b'; }}
      fetch('/deal/' + id + '/field', {{method:'POST',
        headers:{{'Content-Type':'application/x-www-form-urlencoded'}},
        body:'field=exhibition_name&value=' + encodeURIComponent(v)}})
        .then(r => r.json()).then(d => {{ if (s) {{ s.textContent = d.ok ? '✓ 保存' : 'エラー';
          s.style.color = d.ok ? '#166534' : '#c53030'; }} }})
        .catch(() => {{ if (s) {{ s.textContent = '通信エラー'; s.style.color = '#c53030'; }} }});
    }}
    function exBulkConfirm(btn) {{
      var n = document.querySelectorAll('.ex-chk:checked').length;
      if (!n) {{ alert('対象の行を選択してください'); return false; }}
      var name = btn.form.exhibition_name.value.trim();
      return confirm('選択した ' + n + ' 件の展示会名を「' + (name || '(空=クリア)') + '」に設定します。よろしいですか？');
    }}
    </script>"""


def slack_memo_backfill_page(con, offset: int = 0, limit: int = 25) -> str:
    """過去のSlack確認スレッドから『追記メモ』を拾い直し、1件ずつ確認して現状メモへ転記する画面。
    完了スレッド(slack_threads.state='completed')を新しい順に見て、Slackから本文を再取得→抽出。
    書き込みは行ごとの「追記」ボタン(ajax /slack-memo-backfill/apply)で個別に行う。"""
    from cowork import slack_bot as _sb
    total = con.execute(
        "SELECT COUNT(*) c FROM slack_threads WHERE state='completed' AND deal_id IS NOT NULL"
    ).fetchone()["c"]
    rows = con.execute(
        "SELECT thread_ts, channel_id, deal_id, bot_message_ts FROM slack_threads "
        "WHERE state='completed' AND deal_id IS NOT NULL ORDER BY rowid DESC LIMIT ? OFFSET ?",
        (limit, offset)).fetchall()
    cards = []
    _fetched = 0
    for r in rows:
        did = r["deal_id"]
        deal = sfa_db.get_deal(con, did)
        if not deal:
            continue
        _fetched += 1
        try:
            msgs = _sb.get_thread_messages(r["channel_id"], r["thread_ts"])
            fields = _sb.collect_fields(msgs, r["bot_message_ts"] or "", "")
            memo = (fields.get("追記メモ") or "").strip()
        except Exception as e:  # noqa: BLE001
            cards.append(f'<div class="card"><b>SFA#{did}</b> {_esc(deal.get("deal_name") or "")}'
                         f'<br><span class="muted">Slack取得エラー: {_esc(str(e))}</span></div>')
            continue
        if not memo:
            continue
        _note = deal.get("note") or ""
        _already = memo in _note
        _idx = f"{did}_{r['thread_ts']}"
        _status = ('<span class="tagged-ok">✓ 既に現状メモに反映済み</span>' if _already
                   else '<span class="muted" style="color:#b45309">未反映 — 内容を確認して追記</span>')
        cards.append(f"""
        <div class="card" id="bfcard_{_idx}">
          <div style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:6px">
            <b><a href="/deal/{did}?return_to=%2Fslack-memo-backfill">SFA#{did}</a>　{_esc(deal.get("account_name") or "")} / {_esc(deal.get("deal_name") or "")}</b>
            {_status}
          </div>
          <div class="muted" style="font-size:11px;margin:6px 0 2px">現状メモ（現在）</div>
          <div style="font-size:12px;white-space:pre-wrap;background:#f8fafc;border-radius:6px;padding:6px 8px;max-height:120px;overflow:auto">{_esc(_note) or "（なし）"}</div>
          <div class="muted" style="font-size:11px;margin:8px 0 2px">Slackから復旧した追記メモ（編集可）</div>
          <textarea id="bfmemo_{_idx}" rows="6" style="width:100%;font-size:12px">{_esc(memo)}</textarea>
          <div style="margin-top:6px">
            <button class="btn" onclick="bfApply('{_idx}', {did})">この内容を現状メモに追記</button>
            <span id="bfok_{_idx}" style="font-size:12px;margin-left:8px"></span>
          </div>
        </div>""")
    body = "".join(cards) or '<div class="card muted">このページに転記候補（追記メモあり・未反映）はありません。</div>'
    _next = offset + limit
    nav = ""
    if _next < total:
        nav = (f'<p><a class="btn sec" href="/slack-memo-backfill?offset={_next}">次の{limit}件 →</a> '
               f'<span class="muted">({offset + 1}〜{min(_next, total)} / 完了スレッド{total}件)</span></p>')
    return f"""
    <div class="card">
      <h2>🩹 Slack追記メモの復旧（1件ずつ確認して転記）</h2>
      <p class="muted" style="margin:0">過去のSlack確認スレッドから「追記メモ」を拾い直しました（新しい順・{limit}件ずつ）。
      内容を確認・必要なら編集して、各カードの「追記」で現状メモへ転記します。既に反映済みの分は✓表示。
      Slack履歴が消えているスレッドは復旧できません。</p>
      <p class="muted" style="margin:6px 0 0;font-size:12px">完了スレッド {total}件中 {offset + 1}〜{min(offset + limit, total)} を表示（Slack再取得 {_fetched}件）。</p>
    </div>
    {body}
    {nav}
    <script>
    function bfApply(idx, did) {{
      var ta = document.getElementById('bfmemo_' + idx);
      var memo = ta ? ta.value : '';
      var s = document.getElementById('bfok_' + idx);
      if (!memo.trim()) {{ if (s) {{ s.textContent = '空です'; s.style.color = '#c53030'; }} return; }}
      if (s) {{ s.textContent = '追記中…'; s.style.color = '#64748b'; }}
      fetch('/slack-memo-backfill/apply', {{method:'POST',
        headers:{{'Content-Type':'application/x-www-form-urlencoded'}},
        body:'deal_id=' + did + '&memo=' + encodeURIComponent(memo)}})
        .then(r => r.json()).then(d => {{ if (s) {{ s.textContent = d.ok ? '✓ 現状メモに追記しました' : ('エラー: ' + (d.error || ''));
          s.style.color = d.ok ? '#166534' : '#c53030'; }} }})
        .catch(() => {{ if (s) {{ s.textContent = '通信エラー'; s.style.color = '#c53030'; }} }});
    }}
    </script>"""


def account_form(con, acc=None) -> str:
    acc = acc or {}
    cancel_url = f"/account/{acc['id']}" if acc.get("id") else "/accounts"
    # 業界はマスタ(industries)由来のドロップダウン。マスタ外の既存値は消さないよう先頭に補完。
    _ind_list = sfa_db.get_master_list(con, "industries") or list(sfa_db.INDUSTRIES)
    _cur_ind = acc.get("industry") or ""
    if _cur_ind and _cur_ind not in _ind_list:
        _ind_list = [_cur_ind] + _ind_list
    return f"""
    <div class="card"><h2>{'アカウント編集' if acc.get('id') else '新規アカウント'}</h2>
    {_save_bar('accForm', cancel_url=cancel_url)}
    <form id="accForm" method="post" action="/account/save">
      <input type="hidden" name="id" value="{_esc(acc.get('id'))}">
      <label>企業名 *</label><input name="name" required value="{_esc(acc.get('name'))}">
      <div class="grid">
        <div><label>業界</label>
          <select name="industry">{_opt(_ind_list, _cur_ind)}</select>
        </div>
        <div><label>企業規模</label>
          <select name="company_size">{_opt(sfa_db.COMPANY_SIZES, acc.get('company_size'))}</select>
        </div>
      </div>
      <label>メモ</label><textarea name="note" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="2">{_esc(acc.get('note'))}</textarea>
      <p><button class="btn">保存</button> <a class="btn sec" href="{cancel_url}">キャンセル</a></p>
    </form></div>"""


def account_duplicates_page(con) -> str:
    """同名アカウントのグループを表示し、1つに統合できる画面。"""
    groups = sfa_db.find_duplicate_accounts(con)
    if not groups:
        return """
        <div class="card">
          <p style="margin:0 0 10px"><a class="btn sec" href="/accounts">← アカウント一覧へ</a></p>
          <h2>重複アカウント</h2>
          <p class="muted" style="margin:0">同名の重複アカウントはありません。</p>
        </div>"""
    blocks = ""
    for gi, g in enumerate(groups):
        # 商談・コンタクトが最も多いものを「残す」初期選択にする
        accts = sorted(g["accounts"], key=lambda a: (a["deal_count"] + a["contact_count"]), reverse=True)
        radios = ""
        for i, a in enumerate(accts):
            checked = " checked" if i == 0 else ""
            radios += (
                f'<label style="display:block;padding:6px 0;border-bottom:1px solid #eef1f5">'
                f'<input type="radio" name="keep_id" value="{a["id"]}"{checked} style="width:auto;margin-right:8px">'
                f'#{a["id"]} <b>{_esc(a["name"])}</b>'
                f'<span class="muted" style="margin-left:10px">商談{a["deal_count"]}件 / コンタクト{a["contact_count"]}件'
                f' / 業界:{_esc(a.get("industry") or "—")} / 規模:{_esc(a.get("company_size") or "—")}</span></label>'
            )
        ids_csv = ",".join(str(a["id"]) for a in accts)
        blocks += f"""
        <div class="card">
          <h2 style="margin-bottom:8px">「{_esc(g['name'])}」（{len(accts)}件重複）</h2>
          <form method="post" action="/accounts/merge"
            onsubmit="return confirm('選んだ1件に統合し、他を削除します。\\n商談・コンタクトは残す1件に付け替えられます。\\n（統合前にDBは自動バックアップされます）\\n実行しますか？')">
            <p class="muted" style="margin:0 0 6px">残すアカウントを選択してください（他は削除され、参照は残す側へ移動）:</p>
            {radios}
            <input type="hidden" name="all_ids" value="{ids_csv}">
            <p style="margin-top:10px"><button class="btn" style="background:#c53030">この1件に統合する</button></p>
          </form>
        </div>"""
    return f"""
    <div class="card" style="background:#fff7ed;border:1.5px solid #fed7aa">
      <p style="margin:0 0 10px"><a class="btn sec" href="/accounts">← アカウント一覧へ</a></p>
      <h2>⚠ 重複アカウントの統合（{len(groups)}グループ）</h2>
      <p class="muted" style="margin:0">同じ会社名で複数登録されているアカウントです。1つに統合すると、
      紐づく商談・コンタクトは残す側にまとめられ、他のアカウントは削除されます。
      統合前に自動でDBバックアップが取られます。</p>
    </div>
    {blocks}"""


def accounts_page(con) -> str:
    """アカウント一覧ページ。"""
    accounts = sfa_db.list_accounts(con)
    deal_counts = {
        r["account_id"]: r["cnt"]
        for r in con.execute(
            "SELECT account_id, COUNT(*) as cnt FROM deals WHERE account_id IS NOT NULL GROUP BY account_id"
        )
    }
    rows_html = "".join(
        f'<tr class="acc-row" data-name="{_esc((a.get("name") or "").lower())}" '
        f'data-ind="{_esc(a.get("industry") or "")}" data-size="{_esc(a.get("company_size") or "")}">'
        f'<td style="width:32px"><input type="checkbox" name="ids" value="{a["id"]}"></td>'
        f'<td><a href="/account/{a["id"]}">{_esc(a["name"])}</a></td>'
        f'<td>{_esc(a.get("industry")) or "<span class=muted>―</span>"}</td>'
        f'<td>{_esc(a.get("company_size")) or "<span class=muted>―</span>"}</td>'
        f'<td class="right muted">{deal_counts.get(a["id"], 0)}</td>'
        f'</tr>'
        for a in accounts
    ) or '<tr><td colspan=5 class=muted>アカウントがありません。</td></tr>'
    deal_counts_json = json.dumps(deal_counts, ensure_ascii=False)
    _dup_n = len(sfa_db.find_duplicate_accounts(con))
    dup_link = (f'<a class="btn" style="background:#c53030" href="/accounts/duplicates">'
                f'⚠ 重複を統合（{_dup_n}件）</a>' if _dup_n else "")
    # フィルタ用の選択肢: マスタ ＋ 実データに存在する値（マスタ外の既存値も拾う）。
    _ind_opts_list = list(sfa_db.get_master_list(con, "industries") or sfa_db.INDUSTRIES)
    for a in accounts:
        _iv = a.get("industry")
        if _iv and _iv not in _ind_opts_list:
            _ind_opts_list.append(_iv)
    _ind_options = "".join(f'<option value="{_esc(v)}">{_esc(v)}</option>' for v in _ind_opts_list)
    _size_options = "".join(f'<option value="{_esc(v)}">{_esc(v)}</option>' for v in sfa_db.COMPANY_SIZES)
    filter_row = f"""<div class="filter-row" style="margin-bottom:12px">
      <input type="text" id="accSearch" placeholder="🔍 企業名で検索..."
        oninput="filterAccounts()" style="max-width:220px">
      <select id="accIndFilter" onchange="filterAccounts()">
        <option value="">業界:全て</option>{_ind_options}</select>
      <select id="accSizeFilter" onchange="filterAccounts()">
        <option value="">企業規模:全て</option>{_size_options}</select>
      <a class="btn sec" href="/accounts">リセット</a>
      <span id="accCount" class="muted" style="font-size:12px;align-self:center"></span>
    </div>"""
    return f"""
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>アカウント一覧 ({len(accounts)})</span>
        <span style="display:flex;gap:8px">{dup_link}<a class="btn" href="/account/new">＋手動追加</a></span>
      </h2>
      {filter_row}
      <form id="acc_bulk_form" method="post" action="/accounts/bulk_delete">
      <div style="overflow:auto;max-height:70vh">
      <table>
        <tr><th class="sticky" style="width:32px"><input type="checkbox" id="acc_chk_all" title="全選択"
              onchange="document.querySelectorAll('#acc_bulk_form [name=ids]').forEach(c=>c.checked=this.checked)"></th>
            {_sticky_th('企業名')}{_sticky_th('業界')}{_sticky_th('企業規模')}<th class="right sticky">商談数</th></tr>
        {rows_html}
      </table>
      </div>
      <div style="margin-top:10px">
        <button class="btn" type="button" onclick="accBulkDelete()"
          style="background:#c53030;border-color:#c53030;color:#fff">選択した件を削除</button>
      </div>
      </form>
    </div>
    <script>
    const ACC_DEAL_COUNTS = {deal_counts_json};
    function filterAccounts() {{
      var q = (document.getElementById('accSearch').value || '').toLowerCase();
      var ind = document.getElementById('accIndFilter').value;
      var sz = document.getElementById('accSizeFilter').value;
      var shown = 0;
      document.querySelectorAll('#acc_bulk_form tr.acc-row').forEach(function(r){{
        var okName = !q || (r.getAttribute('data-name') || '').indexOf(q) >= 0;
        var okInd = !ind || (r.getAttribute('data-ind') || '') === ind;
        var okSize = !sz || (r.getAttribute('data-size') || '') === sz;
        var vis = okName && okInd && okSize;
        r.style.display = vis ? '' : 'none';
        if (vis) shown++;
      }});
      var c = document.getElementById('accCount');
      if (c) c.textContent = shown + '件表示';
    }}
    function accBulkDelete() {{
      var ids = Array.from(document.querySelectorAll('#acc_bulk_form [name=ids]:checked')).map(function(c){{return c.value;}});
      if (!ids.length) {{ alert('削除するアカウントを選択してください。'); return; }}
      var dealTotal = ids.reduce(function(sum, id){{ return sum + (ACC_DEAL_COUNTS[id] || 0); }}, 0);
      var msg = ids.length + '件のアカウントを削除します。この操作は取り消せません。';
      if (dealTotal > 0) msg += '\\n※紐づく商談も合計' + dealTotal + '件まとめて削除されます。';
      if (!confirm(msg)) return;
      document.getElementById('acc_bulk_form').submit();
    }}
    </script>"""


def account_detail(con, acc: dict) -> str:
    """アカウント詳細ページ（関連商談含む）。"""
    deals = [dict(r) for r in con.execute(
        "SELECT id, deal_name, stage, owner, sub_owner, status FROM deals WHERE account_id=? ORDER BY id DESC",
        (acc["id"],)
    )]
    deal_rows = "".join(
        f'<tr>'
        f'<td><a href="/deal/{d["id"]}">{_esc(d["deal_name"])}</a></td>'
        f'<td>{_esc(d.get("stage")) or "<span class=muted>―</span>"}</td>'
        f'<td>{_esc(d.get("owner")) or "<span class=muted>―</span>"}</td>'
        f'<td>{_esc(d.get("sub_owner")) or "<span class=muted>―</span>"}</td>'
        f'<td><span class="muted">{_esc(d.get("status") or "open")}</span></td>'
        f'</tr>'
        for d in deals
    ) or '<tr><td colspan=5 class=muted>関連商談がありません</td></tr>'
    note_html = (
        f'<p style="margin-top:8px;white-space:pre-wrap;font-size:13px">{_esc(acc.get("note"))}</p>'
        if acc.get("note") else ""
    )
    _del_msg = f"このアカウントを削除しますか？"
    if deals:
        _del_msg += f" 紐づく商談{len(deals)}件も一緒に削除されます。"
    _del_msg += " この操作は取り消せません。"
    return f"""
    <div class="card">
      <div style="display:flex;justify-content:space-between;align-items:center">
        <h2 style="margin:0">{_esc(acc["name"])}</h2>
        <span style="display:flex;gap:8px">
          <a class="btn sec" href="/account/{acc['id']}/edit" style="font-size:12px;padding:5px 10px">編集</a>
          <form method="post" action="/account/{acc['id']}/delete" style="display:inline;margin:0">
            <button class="btn" style="background:#c53030;border-color:#c53030;color:#fff;font-size:12px;padding:5px 10px"
              onclick="return confirm('{_del_msg}')">削除</button>
          </form>
        </span>
      </div>
      <div class="grid" style="margin-top:12px">
        <div><label>業界</label><p style="margin:2px 0">{_esc(acc.get("industry")) or "―"}</p></div>
        <div><label>企業規模</label><p style="margin:2px 0">{_esc(acc.get("company_size")) or "―"}</p></div>
      </div>
      {note_html}
    </div>
    <div class="card">
      <h2>関連商談 ({len(deals)})</h2>
      <table><tr><th>案件名</th><th>ステージ</th><th>主担当</th><th>サブ担当</th><th>状態</th></tr>
      {deal_rows}
      </table>
      <p style="margin-top:12px">
        <a class="btn" href="/deal/new">＋商談追加</a>
        <a class="btn sec" href="/accounts" style="margin-left:8px">一覧へ</a>
      </p>
    </div>"""


def deal_form(con, deal=None, return_to: str | None = None) -> str:
    deal = deal or {}
    # 特定日/MS超過等から遷移した場合の戻り先。保存後・キャンセル時にここへ戻す（未指定は一覧/詳細）。
    return_to = return_to if (return_to or "").startswith("/") else None
    # 次回MS（複数）。既存行があればそれを、無ければキャッシュから1行を仮生成（保存時に実体化）。
    _ms_list = sfa_db.list_deal_milestones(con, deal["id"]) if deal.get("id") else []
    if not _ms_list:
        if deal.get("next_milestone_date") or deal.get("next_milestone_label"):
            _ms_list = [{"ms_date": deal.get("next_milestone_date"),
                         "ms_label": deal.get("next_milestone_label"),
                         "ms_type": deal.get("next_milestone_type"), "done": 0}]
        else:
            _ms_list = [{}]  # 空行1つ
    ms_editor_html = _ms_editor_html(_ms_list)
    accounts = sfa_db.list_accounts(con)
    acc_opts = ['<option value=""></option>']
    for a in accounts:
        sel = " selected" if a["id"] == deal.get("account_id") else ""
        acc_opts.append(f'<option value="{a["id"]}"{sel}>{html.escape(a["name"])}</option>')

    # 新規作成時のみリード引用セクション
    lead_picker_html = ""
    if not deal.get("id"):
        open_leads = [l for l in sfa_db.list_leads(con)
                      if l.get("lead_status") not in ("converted", "lost")]
        acc_by_name = {a["name"]: a["id"] for a in accounts}
        leads_data = json.dumps([{
            "id": l["id"],
            "account_id": acc_by_name.get(l.get("company", ""), ""),
            "owner": l.get("assigned_to") or "",
            "lead_pattern": _SOURCE_TO_LP.get(l.get("source", "other"), "na"),
            "notes": l.get("notes") or "",
        } for l in open_leads], ensure_ascii=False)
        lead_opts = '<option value="">（リードを引用しない）</option>' + "".join(
            f'<option value="{l["id"]}">{html.escape(l.get("company","?"))} / {html.escape(l.get("name","?"))}</option>'
            for l in open_leads
        )
        lead_picker_html = f"""
        <div style="background:#f0f6ff;border-radius:8px;padding:12px 14px;margin-bottom:14px">
          <label style="color:#2f6fed;font-weight:600;font-size:13px">リードから引用</label>
          <select id="lead_ref" onchange="applyLead()" style="margin-top:6px">{lead_opts}</select>
          <p class="muted" style="margin-top:4px">選ぶとアカウント・担当・経路・メモが自動入力されます</p>
        </div>
        <script>
        const _LEADS = {leads_data};
        function applyLead() {{
          const lid = parseInt(document.getElementById('lead_ref').value);
          if (!lid) return;
          const l = _LEADS.find(x => x.id === lid);
          if (!l) return;
          if (l.account_id) document.querySelector('[name=account_id]').value = l.account_id;
          document.querySelector('[name=owner]').value = l.owner;
          document.querySelector('[name=lead_pattern]').value = l.lead_pattern;
          document.querySelector('[name=note]').value = l.notes;
        }}
        </script>"""

    other_deals_html = ""
    if deal.get("id") and deal.get("account_id"):
        _other_deals = [dict(r) for r in con.execute(
            "SELECT id, deal_name, stage, status, owner, updated_at FROM deals "
            "WHERE account_id=? AND id!=? ORDER BY id DESC",
            (deal["account_id"], deal["id"]),
        )]
        if _other_deals:
            _od_rows = "".join(
                f'<tr><td><a href="/deal/{d["id"]}">{_esc(d.get("deal_name"))}</a></td>'
                f'<td><span class="stage">{_esc(d.get("stage")) or "—"}</span></td>'
                f'<td>{"クローズ" if d.get("status") == "closed" else "進行中"}</td>'
                f'<td>{_esc(d.get("owner")) or "—"}</td>'
                f'<td class="muted">{_esc((d.get("updated_at") or "")[:10])}</td></tr>'
                for d in _other_deals
            )
            other_deals_html = f"""
        <div class="card">
          <h2>この会社の他の商談（{len(_other_deals)}件）</h2>
          <table><tr><th>案件名</th><th>ステージ</th><th>状況</th><th>担当</th><th>更新日</th></tr>{_od_rows}</table>
        </div>"""

    hearing_html = ""
    if deal.get("id"):
        n_hearings = sfa_db.count_hearing_results(con, deal["id"])
        if n_hearings:
            latest = sfa_db.list_hearing_results(con, deal["id"])[0]
            hearing_html = f"""
        <div class="card">
          <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
            <span>初回ヒアリング</span>
            <span style="display:flex;gap:8px">
              <a class="btn" href="/hearing/result/{latest['id']}">📋 初回ヒアリング結果（{n_hearings}件）</a>
              <a class="btn sec" href="/hearing/new?target=deal:{deal['id']}">＋追加ヒアリング</a>
            </span>
          </h2>
          <p class="muted" style="margin:0">最新ヒアリング日: {_esc(latest.get('conducted_on') or '—')}（{_esc(latest.get('template_name') or '')}）</p>
        </div>"""
        else:
            hearing_html = f"""
        <div class="card">
          <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
            <span>初回ヒアリング</span>
            <a class="btn" href="/hearing/new?target=deal:{deal['id']}">ヒアリングを実施</a>
          </h2>
          <p class="muted" style="margin:0">ヒアリング未実施</p>
        </div>"""

    dev_projects_html = ""
    if deal.get("id"):
        dps = sfa_db.list_dev_projects(con, deal_id=deal["id"])
        add_btn = f'<a class="btn sec" href="/dev-projects/new?deal_id={deal["id"]}">＋開発案件を追加</a>'
        if dps:
            dp_rows = "".join(
                f'<tr><td><a href="/dev-project/{p["id"]}/edit">{_esc(p.get("theme"))}</a></td>'
                f'<td><span class="stage">{_esc(p.get("stage"))}</span></td>'
                f'<td>{_esc(p.get("status"))}</td>'
                f'<td>{_esc(p.get("order_potential"))}</td>'
                f'<td>{_esc(p.get("dev_owner"))}</td>'
                f'<td>{_esc(p.get("deadline") or "—")}</td>'
                f'<td>{_tool_link_btn(p.get("tool_url"), tool_id=p.get("tool_login_id"), tool_password=p.get("tool_login_pass")) or "—"}</td></tr>'
                for p in dps
            )
            dev_projects_html = f"""
        <div class="card">
          <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
            <span>開発案件（{len(dps)}件）</span>{add_btn}
          </h2>
          <table><tr><th>テーマ</th><th>ステージ</th><th>状況</th><th>受注余地</th><th>開発担当</th><th>開発期限</th><th>ツール</th></tr>{dp_rows}</table>
        </div>"""
        else:
            dev_projects_html = f"""
        <div class="card">
          <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
            <span>開発案件</span>{add_btn}
          </h2>
          <p class="muted" style="margin:0">開発案件なし</p>
        </div>"""

    # #75: Delivery（受注後アサイン計画）導線。提案以降 or 既存Deliveryがあるとき表示。
    delivery_html = ""
    if deal.get("id"):
        _dvs = sfa_db.list_deliveries(con, deal_id=deal["id"])
        _stg = deal.get("stage") or ""
        if _dvs or _stg in sfa_db.DELIVERY_TRIGGER_STAGES:
            if _dvs:
                _dv_rows = "".join(
                    f'<tr><td><a href="/delivery/{d["id"]}">{_esc(d.get("title") or "(無題)")}</a></td>'
                    f'<td>{_esc(d.get("status") or "")}</td>'
                    f'<td class="muted">{_fmt_week(d.get("start_week"))}〜{_fmt_week(d.get("end_week"))}</td></tr>'
                    for d in _dvs)
                _dv_body = f'<table><tr><th>Delivery</th><th>状態</th><th>期間</th></tr>{_dv_rows}</table>'
            else:
                _dv_body = ('<p class="muted" style="margin:0 0 8px">この商談のDeliveryはまだありません'
                            '（提案到達時に自動起票されます）。</p>')
            delivery_html = f"""
        <div class="card">
          <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
            <span>🚚 Delivery（アサイン計画）</span>
            <form method="post" action="/deliveries/new" style="margin:0">
              <input type="hidden" name="deal_id" value="{deal["id"]}">
              <button class="btn sec" style="font-size:12px">＋Delivery追加</button></form>
          </h2>
          {_dv_body}
        </div>"""

    deal_issues_html = ""
    if deal.get("id"):
        issues = sfa_db.list_deal_issues(con, deal_id=deal["id"])
        add_issue_btn = f'<a class="btn sec" href="/deal-issue/new?deal_id={deal["id"]}">＋論点を追加</a>'
        if issues:
            issue_rows = ""
            for it in issues:
                memos = sfa_db.list_deal_issue_memos(con, it["id"])
                return_to = f'/deal/{deal["id"]}'
                memo_panel = _issue_memo_panel_html(memos, it, return_to=return_to)
                summary_box = _ai_summary_hover_html(it.get('ai_summary'), issue_id=it['id'], return_to=return_to)
                issue_rows += f"""
                <tr>
                  <td>{_esc(it.get('issue'))}</td>
                  <td>{_issue_status_select_html(it['id'], it.get('status'))}</td>
                  <td>{_issue_members_inline_html(it['id'], it.get('members'))}</td>
                  <td>{_issue_due_date_input_html(it['id'], it.get('due_date'))}</td>
                  <td>
                    {summary_box}
                    <details class="di-memo-details">
                      <summary style="cursor:pointer">メモ（{len(memos)}）</summary>
                      {memo_panel}
                    </details>
                  </td>
                  <td><a href="/deal-issue/{it['id']}/edit?return_to={urllib.parse.quote(return_to, safe='')}">編集</a></td>
                </tr>"""
            deal_issues_html = f"""
        <style>
        {AI_SUMMARY_HOVER_CSS}
        </style>
        <div class="card">
          <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
            <span>社内論点（{len(issues)}件）</span>{add_issue_btn}
          </h2>
          <table style="table-layout:fixed;width:100%">
            <tr><th style="width:18%">論点</th><th style="width:9%">ステータス</th><th style="width:18%">議論メンバー</th>
                <th style="width:9%">解消期限</th><th style="width:40%">サマリー・メモ</th><th style="width:6%"></th></tr>
            {issue_rows}
          </table>
        </div>
        <script>
        {DEAL_ISSUE_INLINE_EDIT_JS}
        </script>"""
        else:
            deal_issues_html = f"""
        <div class="card">
          <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
            <span>社内論点</span>{add_issue_btn}
          </h2>
          <p class="muted" style="margin:0">論点なし</p>
        </div>"""

    attachments_widget = ""
    if deal.get("id"):
        attachments = sfa_db.list_deal_attachments(con, deal["id"])
        att_items = "".join(
            f'<div class="attach-item">'
            f'<a href="{_esc(a.get("url"))}" target="_blank" rel="noopener">{_esc(a.get("label"))}</a>'
            f'<form method="post" action="/deal-attachment/{a["id"]}/delete" style="display:inline;margin:0" '
            f'onsubmit="return confirm(\'削除しますか？\')">'
            f'<button type="submit" class="attach-del" title="削除">✕</button></form>'
            f'</div>'
            for a in attachments
        ) or '<div class="muted" style="padding:4px 0">添付ファイルなし</div>'
        attachments_widget = f"""
        <style>
        .attach-wrap {{ position:relative; display:inline-block }}
        .attach-trigger {{ display:inline-block; background:#eff6ff; color:#1d4ed8; border:1px solid #bfdbfe;
          border-radius:6px; padding:4px 10px; font-size:12px; cursor:pointer; white-space:nowrap }}
        /* 固定保存バー(z-index:40)より前面に出す。以前はz-index:30でバーの裏に隠れて操作不能だった */
        .attach-panel {{ display:none; position:absolute; top:100%; right:0; z-index:60; background:#fff;
          border:1px solid #d0e4ff; border-radius:8px; padding:10px 12px; width:260px;
          box-shadow:0 4px 12px rgba(0,0,0,.12); text-align:left }}
        .attach-wrap:hover .attach-panel, .attach-wrap:focus-within .attach-panel {{ display:block }}
        .attach-item {{ display:flex; justify-content:space-between; align-items:center; gap:6px;
          padding:4px 0; border-bottom:1px solid #f1f5f9; font-size:12px }}
        .attach-item a {{ word-break:break-all }}
        .attach-del {{ background:none; border:0; color:#94a3b8; cursor:pointer; font-size:11px; padding:0;
          flex-shrink:0 }}
        .attach-add-form {{ margin-top:8px; display:flex; flex-direction:column; gap:4px }}
        .attach-add-form input {{ font-size:11px; padding:4px 6px; margin:0 }}
        .attach-add-form button {{ font-size:11px; padding:4px 6px; background:#2f6fed; color:#fff;
          border:0; border-radius:5px; cursor:pointer }}
        </style>
        <div class="attach-wrap">
          <span class="attach-trigger">📎 添付ファイル（{len(attachments)}）</span>
          <div class="attach-panel">
            {att_items}
            <form method="post" action="/deal/{deal['id']}/attachment" class="attach-add-form">
              <input name="label" placeholder="ファイル名" required>
              <input type="url" name="url" placeholder="https://..." required>
              <button type="submit">追加</button>
            </form>
          </div>
        </div>"""

    activities_html = ""
    sync_btn = ""
    if deal.get("id"):
        acts = sfa_db.list_activities(con, deal["id"])
        _act_types = sfa_db.get_master_list(con, "activity_types")
        act_rows = "".join(
            f'<tr id="act-{a["id"]}" class="actrow">'
            f'<td style="width:1%;white-space:nowrap">'
            f'<span class="av">{_esc(a.get("occurred_on")) or "—"}</span>'
            f'<input class="ae" type="date" value="{_esc(a.get("occurred_on"))}" '
            f'style="font-size:11px;width:118px;padding:2px 3px" '
            f'onchange="actField({a["id"]},&#39;occurred_on&#39;,this.value)"></td>'
            f'<td style="width:1%;white-space:nowrap">'
            f'<span class="av">{_esc(a.get("type")) or "—"}</span>'
            f'<select class="ae" style="font-size:11px;width:auto;padding:2px 3px" '
            f'onchange="actField({a["id"]},&#39;type&#39;,this.value)">{_opt(_act_types, a.get("type"))}</select></td>'
            f'<td style="width:1%;white-space:nowrap">'
            f'<span class="av">{_esc(a.get("contact_name")) or "—"}</span>'
            f'<input class="ae" value="{_esc(a.get("contact_name"))}" '
            f'style="font-size:11px;width:76px;padding:2px 3px" placeholder="相手" '
            f'onchange="actField({a["id"]},&#39;contact_name&#39;,this.value)"></td>'
            f'<td><span class="av" style="white-space:pre-wrap;font-size:12px">{_esc(a.get("body"))}</span>'
            f'<textarea class="ae ta-expand" rows="2" style="font-size:12px;width:100%" '
            f'onfocus="taExpand(this)" onblur="taShrink(this)" '
            f'onchange="actField({a["id"]},&#39;body&#39;,this.value)">{_esc(a.get("body"))}</textarea></td>'
            f'<td style="width:1%;white-space:nowrap">'
            f'<button type="button" class="btn sec act-editbtn" style="font-size:10px;padding:2px 8px" '
            f'onclick="actEdit({a["id"]})">編集</button> '
            f'<button type="button" class="btn sec ae" style="font-size:10px;padding:2px 6px;color:#c53030" '
            f'onclick="actDelete({a["id"]})">削除</button></td></tr>'
            for a in acts
        ) or '<tr><td colspan=5 class=muted>活動なし</td></tr>'
        activities_html = f"""
        <div class="card" id="activity"><h2>活動履歴 <span class="muted" style="font-size:.5em">（「編集」→入力後、エリア外クリックで自動保存して閉じる）</span></h2>
        <style>
        .actrow .ae{{display:none}}
        .actrow.editing .av{{display:none}}
        .actrow.editing .ae{{display:inline-block}}
        .actrow.editing td{{background:#f5f9ff}}
        </style>
        <table style="width:100%"><tr><th style="width:1%">日付</th><th style="width:1%">種別</th><th style="width:1%">相手</th><th>内容</th><th style="width:1%"></th></tr>{act_rows}</table>
        <script>
        function actEdit(id){{ var r=document.getElementById('act-'+id); if(!r)return;
          if(r.classList.contains('editing')){{ actClose(id); return; }}
          r.classList.add('editing'); var b=r.querySelector('.act-editbtn'); if(b)b.textContent='閉じる';
          var ta=r.querySelector('textarea'); if(ta){{ta.focus();}}
          setTimeout(function(){{ r._out=function(e){{ if(!r.contains(e.target) && !r.contains(document.activeElement)) actClose(id); }};
            document.addEventListener('mousedown', r._out); }},0); }}
        function actClose(id){{ var r=document.getElementById('act-'+id); if(!r||!r.classList.contains('editing'))return;
          if(r._out){{ document.removeEventListener('mousedown', r._out); r._out=null; }}
          var d=r.querySelector('input[type=date]'), s=r.querySelector('select'),
              c=r.querySelector('input:not([type=date])'), b=r.querySelector('textarea'),
              sp=r.querySelectorAll('.av');
          if(d){{ actField(id,'occurred_on',d.value); if(sp[0])sp[0].textContent=d.value||'—'; }}
          if(s){{ actField(id,'type',s.value); if(sp[1])sp[1].textContent=s.value||'—'; }}
          if(c){{ actField(id,'contact_name',c.value); if(sp[2])sp[2].textContent=c.value||'—'; }}
          if(b){{ actField(id,'body',b.value); if(sp[3])sp[3].textContent=b.value; }}
          r.classList.remove('editing'); var eb=r.querySelector('.act-editbtn'); if(eb)eb.textContent='編集'; }}
        function actField(id,f,v){{ return fetch('/activity/'+id+'/field',{{method:'POST',headers:{{'Content-Type':'application/x-www-form-urlencoded'}},body:'field='+encodeURIComponent(f)+'&value='+encodeURIComponent(v)}}).then(function(r){{return r.json();}}).then(function(d){{ if(!d.ok)alert('更新エラー: '+(d.error||'')); }}).catch(function(){{alert('通信エラー');}}); }}
        function actDelete(id){{ if(!confirm('この活動履歴を削除しますか？')) return; var f=document.createElement('form'); f.method='post'; f.action='/activity/'+id+'/delete'; document.body.appendChild(f); f.submit(); }}
        </script>
        <form method="post" action="/activity/add" style="margin-top:16px">
          <input type="hidden" name="deal_id" value="{deal['id']}">
          <div class="grid">
            <div><label>日付</label><input type="date" name="occurred_on"></div>
            <div><label>種別</label><select name="type">{_opt(sfa_db.ACTIVITY_TYPES, '面談')}</select></div>
            <div><label>相手</label><input name="contact_name" placeholder="例：田中部長"></div>
          </div>
          <label>内容・決定事項</label><textarea name="body" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="3"></textarea>
          <div style="margin-top:10px;padding:12px;background:#f8f9fa;border-radius:6px">
            <p style="margin:0 0 8px;font-size:.9em;font-weight:600;color:#555">商談の現状を更新</p>
            <div class="grid">
              <div><label>次回MS日</label><input type="date" name="next_milestone_date" value="{_esc(deal.get('next_milestone_date'))}"></div>
              <div><label>次回MSラベル</label><input name="next_milestone_label" value="{_esc(deal.get('next_milestone_label'))}"></div>
              <div><label>次回MS種別</label><select name="next_milestone_type">{_opt(sfa_db.NEXT_MS_TYPES, deal.get('next_milestone_type') or 'アポ')}</select></div>
            </div>
            <label>現状メモ</label><textarea name="update_note" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="2">{_esc(deal.get('note'))}</textarea>
          </div>
          <p><button class="btn sec">活動を追加して更新</button></p>
        </form></div>"""
        sync_btn = (
            f'<span class="muted" style="font-size:.85em">'
            f'{"🔗 テーマDB連携済 (id="+str(deal.get("theme_id"))+")" if deal.get("theme_id") else "テーマDB未連携（保存時に自動連携）"}'
            f'</span>'
        )
    top_action_buttons = ""
    if deal.get("id"):
        _did = deal["id"]
        top_action_buttons = f"""
        <div style="display:flex;gap:8px;margin:-4px 0 14px;flex-wrap:wrap">
          <a class="btn sec" href="/dev-projects/new?deal_id={_did}">＋新規開発案件</a>
          <a class="btn sec" href="/hearing/new?target=deal:{_did}">＋新規ヒアリング</a>
          <a class="btn sec" href="/deal-issue/new?deal_id={_did}">＋新規論点</a>
        </div>"""
    acc_req = "required" if deal.get("id") else ""
    new_acc_html = ""
    new_acc_js = ""
    # 固定保存バーに表示する「SFA#・アカウント・案件名」（編集時のみ）
    _sb_title = ""
    _sb_extra = ""
    if deal.get("id"):
        _acc_nm = next((a["name"] for a in accounts if a["id"] == deal.get("account_id")), "")
        _sb_title = f"SFA#{deal['id']}　{_acc_nm}／{deal.get('deal_name') or ''}"
        _st_label = "クローズ済" if deal.get("status") == "closed" else "進行中"
        if deal.get("status") == "closed":
            # クローズ済みは「終了理由」をその場で修正できる（リード化時の理由選択ミスを直す導線）。
            _cr_cur = deal.get("close_reason") or ""
            _cr_opts = "".join(
                f'<option value="{_esc(v)}"{" selected" if v == _cr_cur else ""}>{_esc(v)}</option>'
                for v in sfa_db.CLOSE_REASONS)
            _did_js = deal["id"]
            _cr_ctl = (f'<span class="muted" style="font-size:11px">終了理由</span>'
                       f"<select onchange=\"updateDealField({_did_js}, 'close_reason', this.value)\" "
                       f'style="font-size:12px;padding:1px 4px"><option value=""></option>{_cr_opts}</select>')
        else:
            _cr_ctl = '<span class="muted" style="font-size:11px">クローズは画面下部の「クローズ」ボタンから</span>'
        _sb_extra = (f'<span class="muted" style="font-size:11px">ステータス</span>'
                     f'<span class="stage">{_st_label}</span>{_cr_ctl}')
    if not deal.get("id"):
        new_acc_html = (
            '<div style="margin-top:5px;text-align:left">'
            '<label style="font-size:11px;color:#6b7689;cursor:pointer">'
            '<input type="checkbox" id="new_acc_chk" onchange="toggleNewAcc()" style="width:auto;margin-right:4px">'
            '新規アカウントを追加（業界・規模を自動推定）</label>'
            '<div id="new_acc_row" style="display:none;margin-top:4px">'
            '<input name="new_account_name" placeholder="新しい会社名"></div>'
            '</div>'
        )
        new_acc_js = (
            'function toggleNewAcc() {'
            ' var chk=document.getElementById("new_acc_chk");'
            ' document.getElementById("new_acc_row").style.display=chk.checked?"":"none";'
            ' document.getElementById("acc_id_sel").required=!chk.checked;'
            '}'
        )
    revert_btn = ""
    close_btn = ""
    if deal.get("id") and deal.get("status") != "closed":
        # クローズは共通モーダル(理由必須＋詳細任意)に一本化。押すと商談はクローズされリードに戻る。
        revert_btn = (
            f'<button type="button" class="btn" style="background:#c53030;font-size:12px;padding:8px 14px;margin-top:8px"'
            f' onclick="openCloseModal({deal["id"]}, \'/deal/{deal["id"]}\')">'
            '商談をクローズ（リードに戻す）</button>'
        )
    return f"""
    <div class="card">
    <h2 style="display:flex;justify-content:space-between;align-items:flex-start;gap:12px;flex-wrap:wrap">
      <span>{'商談編集' if deal.get('id') else '新規商談'}</span>
      {attachments_widget}
    </h2>
    {_save_bar('dealForm', title=_sb_title, extra=_sb_extra, cancel_url=(return_to or ('/deal/' + str(deal['id']) if deal.get('id') else '/deals')))}
    {top_action_buttons}
    {lead_picker_html}
    <form id="dealForm" method="post" action="/deal/save">
      <input type="hidden" name="id" value="{_esc(deal.get('id'))}">
      {f'<input type="hidden" name="return_to" value="{_esc(return_to)}">' if return_to else ''}
      {ms_editor_html}
      <label>現状メモ</label><textarea name="note" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="2">{_esc(deal.get('note'))}</textarea>
      <hr style="border:none;border-top:1px solid #e6e9f0;margin:16px 0 18px">
      <div class="grid">
        <div><label>アカウント{"" if not deal.get("id") else " *"}</label>
          <select name="account_id" id="acc_id_sel" {acc_req}>{''.join(acc_opts)}</select>
          {new_acc_html}</div>
        <div><label>案件名 *</label>
          <input name="deal_name" required value="{_esc(deal.get('deal_name'))}"></div>
        <div><label>ステージ</label>
          <select name="stage">{_opt(sfa_db.get_master_list(con,'deal_stages'), deal.get('stage'))}</select></div>
        <div><label>主担当</label>
          <select name="owner">{_opt(sfa_db.get_master_list(con,'owners'), deal.get('owner'))}</select></div>
        <div><label>サブ担当</label>
          <select name="sub_owner">{_opt(sfa_db.get_master_list(con,'owners'), deal.get('sub_owner'))}</select></div>
        <div><label>顧客担当者名（先方）</label>
          <input name="client_contact" value="{_esc(deal.get('client_contact'))}" placeholder="例：田中 太郎"></div>
        <div><label>顧客担当 部署（先方）</label>
          <input name="client_dept" value="{_esc(deal.get('client_dept'))}" placeholder="例：購買部"></div>
        <div><label>事業種別L1</label>
          <select name="business_type_l1" id="biz_l1" onchange="updateL2()">{_opt(sfa_db.get_master_list(con,'business_type_l1'), deal.get('business_type_l1'))}</select></div>
        <div><label>事業種別L2</label>
          <select name="business_type_l2" id="biz_l2">{_opt_l2(deal.get('business_type_l1'), deal.get('business_type_l2'))}</select></div>
        <div><label>リード経路</label>
          <select name="lead_pattern">{_opt(sfa_db.get_master_list(con,'lead_patterns'), deal.get('lead_pattern'))}</select></div>
        <div><label>展示会名（展示会由来のとき）</label>
          <input name="exhibition_name" list="dealExhNames" value="{_esc(deal.get('exhibition_name'))}" placeholder="どの展示会か（任意）">
          <datalist id="dealExhNames">{"".join(f'<option value="{_esc(_n)}">' for _n in sfa_db.list_exhibition_names(con))}</datalist></div>
        <div><label>ワンタイム総額（万円）</label>
          <input name="value_lumpsum" value="{_esc(deal.get('value_lumpsum'))}"></div>
        <div><label>ワンタイム月額換算（万円）</label>
          <input name="value_lumpsum_monthly" value="{_esc(deal.get('value_lumpsum_monthly'))}"></div>
        <div><label>継続月額（万円）</label>
          <input name="value_recurring" value="{_esc(deal.get('value_recurring'))}"></div>
        <div><label>クライアント予算</label>
          <input name="client_budget" value="{_esc(deal.get('client_budget'))}"></div>
        <div><label>重要度</label>
          <select name="importance">{_opt(sfa_db.IMPORTANCE_OPTIONS, deal.get('importance'))}</select></div>
      </div>
      <label>ゴール</label><textarea name="goal" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="2">{_esc(deal.get('goal'))}</textarea>
      <div id="cost_section" style="{'display:none' if deal.get('business_type_l1') != 'コスト削減' else ''}">
        <hr style="margin:16px 0">
        <p style="font-weight:600;margin-bottom:8px;color:#555">コスト削減モデル詳細</p>
        <div class="grid">
          <div><label>コスト削減ステージ</label>
            <select name="cost_stage">{_opt(sfa_db.COST_STAGES, deal.get('cost_stage'))}</select></div>
          <div><label>アプローチ額（億円）</label>
            <input name="approach_value" type="number" step="0.01" value="{_esc(deal.get('approach_value'))}"></div>
          <div><label>アプローチ率（%）</label>
            <input name="approach_rate" type="number" step="0.1" value="{_esc(deal.get('approach_rate'))}"></div>
          <div><label>コスト削減率（%）</label>
            <input name="reduction_rate" type="number" step="0.1" value="{_esc(deal.get('reduction_rate'))}"></div>
          <div><label>成果報酬率（%）</label>
            <input name="fee_rate" type="number" step="0.1" value="{_esc(deal.get('fee_rate'))}"></div>
          <div><label>診断原価（万円）</label>
            <input name="diagnosis_cost" type="number" step="1" value="{_esc(deal.get('diagnosis_cost'))}"></div>
        </div>
      </div>
      <p><button class="btn">保存</button> <a class="btn sec" href="/">一覧へ</a> {sync_btn}</p>
    </form>
    {revert_btn}{close_btn}
    <script>
    {new_acc_js}
    const L2_MAP = {json.dumps(sfa_db.BUSINESS_TYPE_L2_BY_L1, ensure_ascii=False)};
    function updateL2() {{
      const l1 = document.getElementById('biz_l1').value;
      const sel = document.getElementById('biz_l2');
      const cur = sel.value;
      sel.innerHTML = '<option value=""></option>' +
        (L2_MAP[l1] || []).map(v => `<option value="${{escH(v)}}"${{v===cur?' selected':''}}>${{escH(v)}}</option>`).join('');
      document.getElementById('cost_section').style.display = l1 === 'コスト削減' ? '' : 'none';
    }}
    </script></div>
    {other_deals_html}
    {hearing_html}
    {dev_projects_html}
    {delivery_html}
    {deal_issues_html}
    {activities_html}"""


# ── 開発案件（商談に紐づく開発テーマ管理）───────────────────────────────────────

def dev_projects_list_page(con, theme_client=None, *, dev_owner: str | None = None, sales_owner: str | None = None,
                            status: str | None = None, stage: str | None = None,
                            order_potential: str | None = None, deadline_week: str | None = None) -> str:
    owners = sfa_db.get_master_list(con, "owners")
    deadline_from = deadline_to = None
    if deadline_week and "-W" in deadline_week:
        try:
            _year, _week = deadline_week.split("-W")
            _mon = date.fromisocalendar(int(_year), int(_week), 1)
            _sun = date.fromisocalendar(int(_year), int(_week), 7)
            deadline_from, deadline_to = _mon.isoformat(), _sun.isoformat()
        except (ValueError, TypeError):
            pass
    projects = sfa_db.list_dev_projects(
        con, dev_owner=dev_owner, sales_owner=sales_owner, status=status,
        stage=stage, order_potential=order_potential,
        deadline_from=deadline_from, deadline_to=deadline_to,
    )

    # 自己修復バックフィル: 分類は入っているのに点数が未設定の行を、開いた時点で自動計算して埋める。
    # （点数付きの行は据え置き＝経験曲線の「過去は据え置き」を尊重。埋めるのは未設定のみ）
    # あわせて、保存済みの開発開始/終了日が現行の逆算式(dev_period_days)とズレている行は再計算する。
    # dev_start/end はSFA側では式由来のみ（ユーザー編集不可）で、人手の期間調整はHisho側の
    # gantt_*_override（devProjectDatesで最優先・当処理では不変）に載るため、再計算は安全。
    # これにより係数変更や分類変更が既存カードへ自動反映される（手動再同期を不要にする）。
    _backfilled = []
    for p in projects:
        if p.get("dev_points") is None and p.get("work_type"):
            _pts = sfa_db.compute_dev_points(
                con, work_type=p.get("work_type"), stage=p.get("stage"),
                difficulty=p.get("difficulty"), has_backend=p.get("has_backend"))
            if _pts is not None:
                con.execute("UPDATE dev_projects SET dev_points=? WHERE id=?", (_pts, p["id"]))
                p["dev_points"] = _pts
                if p["id"] not in _backfilled:
                    _backfilled.append(p["id"])
        if p.get("deadline"):
            _ns, _ne = sfa_db.compute_dev_schedule(
                p.get("deadline"), p.get("stage"), p.get("has_backend"), p.get("difficulty"))
            if _ns and (_ns != p.get("dev_start_date") or _ne != p.get("dev_end_date")):
                con.execute("UPDATE dev_projects SET dev_start_date=?, dev_end_date=? WHERE id=?",
                            (_ns, _ne, p["id"]))
                p["dev_start_date"], p["dev_end_date"] = _ns, _ne
                if p["id"] not in _backfilled:
                    _backfilled.append(p["id"])
    if _backfilled:
        con.commit()
        # 埋めた点数はHishoへ自動同期（手動再同期を不要にする）
        if theme_client is not None:
            for _bid in _backfilled:
                try:
                    dev_project_link.sync_dev_project(theme_client, con, _bid)
                    sfa_db.clear_sync_failure(con, "dev_project", _bid)
                except Exception as _exc:  # noqa: BLE001
                    sfa_db.record_sync_failure(con, "dev_project", _bid, str(_exc))

    def _fopt(values, current):
        return '<option value="">全て</option>' + "".join(
            f'<option value="{html.escape(v)}"{" selected" if v == current else ""}>{html.escape(v)}</option>'
            for v in values
        )

    filter_row = f"""<form method="get" action="/dev-projects" class="filter-row">
      <select name="dev_owner" onchange="this.form.submit()">{_fopt(owners, dev_owner).replace('全て', '開発担当:全て', 1)}</select>
      <select name="sales_owner" onchange="this.form.submit()">{_fopt(owners, sales_owner).replace('全て', '営業担当:全て', 1)}</select>
      <select name="status" onchange="this.form.submit()">{_fopt(sfa_db.DEV_PROJECT_STATUSES, status).replace('全て', '状況:全て', 1)}</select>
      <select name="stage" onchange="this.form.submit()">{_fopt(sfa_db.DEV_PROJECT_STAGES, stage).replace('全て', 'ステージ:全て', 1)}</select>
      <select name="order_potential" onchange="this.form.submit()">{_fopt(sfa_db.DEV_ORDER_POTENTIALS, order_potential).replace('全て', '受注余地:全て', 1)}</select>
      <input type="week" name="deadline_week" value="{_esc(deadline_week)}" title="開発期限（週）で絞り込み" onchange="this.form.submit()">
      <a class="btn sec" href="/dev-projects">リセット</a>
    </form>"""

    def _hearing_link(deal_id):
        n = sfa_db.count_hearing_results(con, deal_id)
        if n:
            latest = sfa_db.list_hearing_results(con, deal_id)[0]
            return f'<a href="/hearing/result/{latest["id"]}">📋ヒアリング（{n}）</a>'
        return f'<a class="muted" href="/hearing/new?target=deal:{deal_id}">ヒアリング未実施</a>'

    _tools_by_dp = sfa_db.list_dev_project_tools_for(con, [p["id"] for p in projects])

    def _tool_cell(p):
        pid = p["id"]
        parts = []
        primary = _tool_link_btn(p.get("tool_url"), tool_id=p.get("tool_login_id"),
                                 tool_password=p.get("tool_login_pass"))
        if primary:
            parts.append(primary)
        for t in _tools_by_dp.get(pid, []):
            parts.append(_tool_link_btn(
                t.get("url"), label="🔧 " + (_esc(t.get("label")) or "リンク"),
                tool_id=t.get("login_id"), tool_password=t.get("login_pass")))
        links = "".join(parts) or '<span class="muted" style="font-size:11px">—</span>'
        add_btn = (f'<button type="button" class="btn sec" style="font-size:10px;padding:2px 7px;margin-top:3px"'
                   f' onclick="openLinkModal({pid})">＋リンク</button>')
        return ('<div style="display:flex;flex-direction:column;gap:3px;align-items:flex-start">'
                + links + add_btn + '</div>')

    _dp_owners = sfa_db.get_master_list(con, "owners")
    _dp_worktypes = sfa_db.dev_work_types(con)
    # 技術シードの複数選択フィルタ（ドロップダウン・L1カテゴリ別チェックボックス, クライアント側で絞り込み）
    _seed_tree_f = sfa_db.get_tech_seed_tree(con)
    _seed_filter_groups = []
    for _l1f, _leavesf in _seed_tree_f.items():
        if not _leavesf:
            continue
        _boxesf = "".join(
            f'<label style="display:flex;align-items:center;gap:5px;font-weight:400;font-size:12px;padding:2px 0;white-space:nowrap">'
            f'<input type="checkbox" class="dp-seedchk" value="{_esc(s)}" onchange="filterDevProjects()" style="width:auto">'
            f'{_esc(s)}</label>' for s in _leavesf)
        _seed_filter_groups.append(
            f'<div style="margin-bottom:6px"><div style="font-size:10px;font-weight:700;color:#4338ca;margin:2px 0">{_esc(_l1f)}</div>{_boxesf}</div>')
    _seed_filter_body = "".join(_seed_filter_groups) or '<span class="muted" style="font-size:11px">シード未登録（管理→🌱技術シードマスタ）</span>'
    _seed_filter_dropdown = (
        '<details class="tb-menu" style="position:relative">'
        '<summary class="btn sec" style="font-size:12px">🏷 技術シードで絞り込み ▾ <span id="dpSeedCount"></span></summary>'
        '<div style="position:absolute;left:0;top:112%;z-index:60;background:#fff;border:1px solid #e6e9f0;'
        'border-radius:8px;box-shadow:0 10px 30px rgba(0,0,0,.18);padding:10px 12px;min-width:240px;max-height:340px;overflow:auto">'
        + _seed_filter_body
        + '<div style="border-top:1px solid #eef1f5;margin-top:6px;padding-top:6px">'
          '<button type="button" class="btn sec" style="font-size:11px" onclick="dpClearSeedFilter()">クリア</button></div>'
        + '</div></details>')

    def _dsel(pid, field, values, current):
        opts = "".join(
            f'<option value="{html.escape(v)}"{" selected" if v == current else ""}>{html.escape(v)}</option>'
            for v in values)
        return (f'<select onchange="updateDevProjectField({pid}, \'{field}\', this.value)"'
                f' style="font-size:12px;padding:3px 4px;min-width:80px;width:100%">'
                f'<option value=""></option>{opts}</select>')

    rows = "".join(
        f'<tr data-search="{_esc(" ".join(str(p.get(k) or "") for k in ("account_name", "deal_name", "theme", "theme_detail", "dev_owner", "sales_owner")).lower())}"'
        f' data-seeds="{_esc("|" + "|".join(s for s in (p.get("tech_seeds") or "").split(",") if s) + "|") if p.get("tech_seeds") else ""}">'
        f'<td class="frz" style="left:0;width:34px;min-width:34px;max-width:34px"><input type="checkbox" name="ids" value="{p["id"]}"></td>'
        f'<td class="frz" style="left:34px;width:240px;min-width:240px;max-width:240px;white-space:normal;word-break:break-word">'
        f'<a href="/dev-project/{p["id"]}/edit">{_esc(p.get("theme"))}</a>'
        f'<div class="muted dp-detail" title="{_esc(p.get("theme_detail") or "")}" '
        f'style="font-size:11px;word-break:break-word">{_esc(p.get("theme_detail") or "")}</div></td>'
        f'<td class="frz" style="left:274px;width:150px;min-width:150px;max-width:150px;white-space:normal;word-break:break-word">'
        f'<a href="/deal/{p["deal_id"]}">{_esc(p.get("account_name"))}</a>'
        f'<div class="muted">{_esc(p.get("deal_name"))}</div>'
        f'<div style="font-size:11px;margin-top:2px">{_hearing_link(p["deal_id"])}</div></td>'
        f'<td>{_dsel(p["id"], "stage", sfa_db.DEV_PROJECT_STAGES, p.get("stage") or "")}</td>'
        f'<td>{_dsel(p["id"], "dev_audience", sfa_db.DEV_AUDIENCES, p.get("dev_audience") or "")}</td>'
        f'<td>{_dsel(p["id"], "work_type", _dp_worktypes, p.get("work_type") or "")}</td>'
        f'<td>{_dsel(p["id"], "difficulty", sfa_db.DEV_DIFFICULTIES, p.get("difficulty") or "")}</td>'
        f'<td>{_dsel(p["id"], "has_backend", sfa_db.DEV_HAS_BACKEND, p.get("has_backend") or "")}</td>'
        f'<td>{_dsel(p["id"], "pricing", sfa_db.DEV_PRICINGS, p.get("pricing") or "")}</td>'
        f'<td>{_dsel(p["id"], "status", sfa_db.DEV_PROJECT_STATUSES, p.get("status") or "")}</td>'
        f'<td>{_dsel(p["id"], "order_potential", sfa_db.DEV_ORDER_POTENTIALS, p.get("order_potential") or "")}</td>'
        f'<td id="dppts{p["id"]}" data-pts="{p.get("dev_points") if p.get("dev_points") is not None else ""}" '
        f'style="text-align:right;font-variant-numeric:tabular-nums" title="自動計算（作業種別×分類×難易度）">'
        f'{(p.get("dev_points") if p.get("dev_points") is not None else "—")}</td>'
        f'<td>{_dsel(p["id"], "dev_owner", _dp_owners, p.get("dev_owner") or "")}</td>'
        f'<td>{_esc(p.get("sales_owner"))}{(" / " + _esc(p["sales_sub_owner"])) if p.get("sales_sub_owner") else ""}</td>'
        f'<td>{_esc(p.get("deadline") or "—")}</td>'
        f'<td style="white-space:normal;min-width:170px;max-width:220px;vertical-align:top">'
        f'<div style="max-height:68px;overflow:auto">{_seed_chips(p.get("tech_seeds"))}</div></td>'
        f'<td>{_tool_cell(p)}</td>'
        f'<td><form method="post" action="/dev-project/{p["id"]}/delete" style="display:inline" '
        f'onsubmit="return confirm(\'削除しますか？\')">'
        f'<button class="btn sec" style="font-size:11px;padding:4px 8px">削除</button></form></td></tr>'
        for p in projects
    ) or '<tr><td colspan=18 class=muted>開発案件がまだありません</td></tr>'
    _pts_total = sum(float(p.get("dev_points") or 0) for p in projects)
    return f"""
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>開発案件一覧（{len(projects)}件）<span class="muted" style="font-size:13px;font-weight:400;margin-left:10px">合計点数 <span id="dpPtsTotal">{_pts_total:g}</span>（≈ <span id="dpPtsFte">{_pts_total/20:.1f}</span> FTE・20点≒1人）</span></span>
        <span style="display:flex;gap:8px">
          <form method="post" action="/dev-projects/resync-hisho" style="margin:0"
            onsubmit="return confirm('全開発案件の点数・分類をHishoダッシュボードへ再同期します。よろしいですか？')">
            <button class="btn sec" type="submit" style="font-size:12px">🔄 Hishoへ再同期</button></form>
          <a class="btn" href="/dev-projects/new">＋新規入力</a>
        </span>
      </h2>
      {filter_row}
      <div style="margin:4px 0 10px;display:flex;gap:10px;align-items:center;flex-wrap:wrap">
        <input type="text" id="dpSearch" placeholder="🔍 アカウント・商談・開発テーマ・担当で検索…"
          oninput="filterDevProjects()" style="flex:1;min-width:240px;max-width:420px;padding:7px 10px">
        {_seed_filter_dropdown}
        <span class="muted" id="dpSearchCount" style="font-size:12px"></span>
      </div>
      <form id="dp_bulk_form" method="post" action="/dev-projects/bulk_delete">
      <div style="overflow:auto;max-height:70vh">
      <table id="dpTable" style="min-width:1920px">
        <tr><th class="sticky frz" style="left:0;width:34px;min-width:34px;max-width:34px"><input type="checkbox" id="dp_chk_all" title="全選択"
              onchange="document.querySelectorAll('#dp_bulk_form [name=ids]').forEach(c=>c.checked=this.checked)"></th>
            <th class="sticky frz" style="left:34px;width:240px;min-width:240px;max-width:240px">開発テーマ</th><th class="sticky frz" style="left:274px;width:150px;min-width:150px;max-width:150px">商談</th>{_sticky_th('ステージ', '84px')}{_sticky_th('提供先', '88px')}{_sticky_th('作業種別', '144px')}{_sticky_th('難易度', '76px')}{_sticky_th('BE有無', '80px')}{_sticky_th('課金', '78px')}{_sticky_th('状況', '84px')}{_sticky_th('受注余地', '76px')}{_sticky_th('点数', '56px')}
            {_sticky_th('開発担当', '88px')}{_sticky_th('営業担当', '88px')}{_sticky_th('開発期限', '88px')}{_sticky_th('技術シード', '180px')}{_sticky_th('ツール')}{_sticky_th('')}</tr>
        {rows}
      </table>
      </div>
      <div style="margin-top:10px">
        <button class="btn" type="button" onclick="dpBulkDelete()"
          style="background:#c53030;border-color:#c53030;color:#fff">選択した件を削除</button>
      </div>
      </form>
    </div>

    <div id="linkModal" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.45);z-index:1000;
      align-items:center;justify-content:center" onclick="if(event.target===this)closeLinkModal()">
      <div style="background:#fff;border-radius:12px;padding:20px 22px;max-width:440px;width:92%;box-shadow:0 10px 40px rgba(0,0,0,.25)">
        <h3 style="margin:0 0 12px">ツールリンクを追加</h3>
        <form id="linkModalForm" method="post">
          <input type="hidden" name="return_to" value="/dev-projects">
          <label style="font-size:12px;color:#6b7689">URL *</label>
          <input type="url" name="url" required placeholder="https://..." style="width:100%;margin:2px 0 8px">
          <label style="font-size:12px;color:#6b7689">表示名（任意）</label>
          <input type="text" name="label" placeholder="例: 設計書、管理画面" style="width:100%;margin:2px 0 8px">
          <div style="display:flex;gap:8px;margin-bottom:14px">
            <div style="flex:1"><label style="font-size:12px;color:#6b7689">ID（任意）</label>
              <input type="text" name="login_id" style="width:100%;margin-top:2px"></div>
            <div style="flex:1"><label style="font-size:12px;color:#6b7689">PASS（任意）</label>
              <input type="text" name="login_pass" style="width:100%;margin-top:2px"></div>
          </div>
          <div style="display:flex;gap:8px;justify-content:flex-end">
            <button type="button" class="btn sec" onclick="closeLinkModal()">キャンセル</button>
            <button type="submit" class="btn">追加</button>
          </div>
        </form>
      </div>
    </div>

    <script>
    function openLinkModal(pid) {{
      var f = document.getElementById('linkModalForm');
      f.reset();
      f.action = '/dev-project/' + pid + '/tools/add';
      document.getElementById('linkModal').style.display = 'flex';
    }}
    function closeLinkModal() {{ document.getElementById('linkModal').style.display = 'none'; }}
    document.addEventListener('keydown', function(e) {{ if (e.key === 'Escape') closeLinkModal(); }});
    function dpBulkDelete() {{
      var ids = Array.from(document.querySelectorAll('#dp_bulk_form [name=ids]:checked')).map(function(c){{return c.value;}});
      if (!ids.length) {{ alert('削除する開発案件を選択してください。'); return; }}
      if (!confirm(ids.length + '件の開発案件を削除します。この操作は取り消せません。')) return;
      document.getElementById('dp_bulk_form').submit();
    }}
    function filterDevProjects() {{
      var q = (document.getElementById('dpSearch').value || '').toLowerCase().trim();
      var checked = Array.prototype.map.call(document.querySelectorAll('.dp-seedchk:checked'),
        function(c) {{ return c.value; }});
      var shown = 0, total = 0;
      document.querySelectorAll('#dpTable tr[data-search]').forEach(function(tr) {{
        total++;
        var textOk = !q || tr.getAttribute('data-search').indexOf(q) >= 0;
        var seeds = tr.getAttribute('data-seeds') || '';
        // 選択シードのいずれかを持つ行を表示（OR条件）
        var seedOk = !checked.length || checked.some(function(s) {{ return seeds.indexOf('|' + s + '|') >= 0; }});
        var vis = textOk && seedOk;
        tr.style.display = vis ? '' : 'none';
        if (vis) shown++;
      }});
      var sc = document.getElementById('dpSeedCount');
      if (sc) sc.textContent = checked.length ? ('(' + checked.length + ')') : '';
      var c = document.getElementById('dpSearchCount');
      if (c) c.textContent = (q || checked.length) ? (shown + ' / ' + total + ' 件') : '';
    }}
    function dpClearSeedFilter() {{
      document.querySelectorAll('.dp-seedchk').forEach(function(c) {{ c.checked = false; }});
      filterDevProjects();
    }}
    function dpRefreshTotal() {{
      var sum = 0;
      document.querySelectorAll('#dpTable td[data-pts]').forEach(function(c) {{
        var v = parseFloat(c.getAttribute('data-pts'));
        if (!isNaN(v)) sum += v;
      }});
      var t = document.getElementById('dpPtsTotal'); if (t) t.textContent = (Math.round(sum * 10) / 10);
      var fe = document.getElementById('dpPtsFte'); if (fe) fe.textContent = (sum / 20).toFixed(1);
    }}
    function updateDevProjectField(id, field, value) {{
      fetch('/dev-project/' + id + '/field', {{
        method: 'POST',
        headers: {{'Content-Type': 'application/x-www-form-urlencoded'}},
        body: 'field=' + encodeURIComponent(field) + '&value=' + encodeURIComponent(value)
      }}).then(r => r.json()).then(d => {{
        if (!d.ok) {{ alert('更新エラー: ' + (d.error || '')); return; }}
        // 再計算された点数をセルに即反映（サーバが自動計算した値）
        var cell = document.getElementById('dppts' + id);
        if (cell && d.dev_points !== undefined) {{
          var v = d.dev_points;
          cell.setAttribute('data-pts', (v === null || v === undefined) ? '' : v);
          cell.textContent = (v === null || v === undefined) ? '—' : v;
          dpRefreshTotal();
        }}
      }})
        .catch(() => alert('通信エラー'));
    }}
    </script>"""


def dev_point_master_page(con) -> str:
    """開発点数マスタ（作業種別→基準点数）＋既存分類/難易度の係数表示＋担当の週次キャパ編集（#41）。"""
    master = sfa_db.list_dev_point_master(con)
    caps = sfa_db.get_owner_capacities(con)
    owners = sfa_db.get_master_list(con, "owners")
    coefs = sfa_db.get_dev_coefs(con)
    stage_coef_inputs = "".join(
        f'<div style="display:inline-block;margin-right:14px">{_esc(k)} ×'
        f'<input type="number" step="0.1" name="cf__stage__{_esc(k)}" '
        f'value="{coefs["stage"].get(k, 1):g}" style="width:64px;margin-left:4px"></div>'
        for k in sfa_db.DEV_PROJECT_STAGES)
    diff_coef_inputs = "".join(
        f'<div style="display:inline-block;margin-right:14px">{_esc(k)} ×'
        f'<input type="number" step="0.1" name="cf__difficulty__{_esc(k)}" '
        f'value="{coefs["difficulty"].get(k, 1):g}" style="width:64px;margin-left:4px"></div>'
        for k in sfa_db.DEV_DIFFICULTIES)
    backend_bonus_inputs = "".join(
        f'<div style="display:inline-block;margin-right:14px">{_esc(k)} ＋'
        f'<input type="number" step="0.1" name="cf__backend__{_esc(k)}" '
        f'value="{coefs["backend"].get(k, 0):g}" style="width:64px;margin-left:4px"> 点</div>'
        for k in sfa_db.DEV_HAS_BACKEND)
    mrows = ""
    for m in master:
        mid = m["id"]
        mrows += (f'<tr><td><input type="text" name="wt__{mid}" value="{_esc(m["work_type"])}" style="width:100%"></td>'
                  f'<td><input type="number" step="0.1" name="bp__{mid}" value="{m["base_points"]:g}" style="width:90px"></td>'
                  f'<td><form method="post" action="/dev-point-master/delete" style="margin:0" '
                  f'onsubmit="return confirm(\'この作業種別を削除しますか？\')">'
                  f'<input type="hidden" name="id" value="{mid}">'
                  f'<button class="btn sec" type="submit" style="font-size:11px;padding:2px 8px;color:#991b1b">削除</button>'
                  f'</form></td></tr>')
    cap_rows = ""
    for o in owners:
        cp = caps.get(o) or {}
        _b = cp.get("base") if cp else ""
        _fr = cp.get("from") or ""
        _p2 = cp.get("base2") if cp.get("base2") is not None else ""
        cap_rows += (
            f'<tr><td>{_esc(o)}</td>'
            f'<td><input type="number" step="0.1" name="cap__{_esc(o)}__base" value="{_b}" placeholder="例: 20" style="width:80px"></td>'
            f'<td><input type="date" name="cap__{_esc(o)}__from" value="{_esc(_fr)}" style="width:150px"></td>'
            f'<td><input type="number" step="0.1" name="cap__{_esc(o)}__p2" value="{_p2}" placeholder="変更後" style="width:80px"></td></tr>')
    return f"""
    <div class="card" style="max-width:760px">
      <p style="margin:0 0 10px"><a href="/dev-projects">← 開発案件一覧</a></p>
      <h2 style="margin:0 0 4px">開発点数マスタ（作業種別）</h2>
      <p class="muted" style="margin:0 0 6px">実際の点数 = <b>作業種別の基準点数</b> × 既存分類係数（プロト/PoC/本番）× 難易度係数。</p>
      <p class="muted" style="margin:0 0 12px">目安: <b>約20点 ≒ 1 FTE（1人分の稼働）</b>。経験曲線＝速くなったら基準点数を下げる（<b>新規案件のみ</b>反映・過去は据え置き）。</p>
      <form method="post" action="/dev-point-master/save">
        <table><tr><th>作業種別</th><th>基準点数</th><th></th></tr>{mrows}
          <tr><td><input type="text" name="new_work_type" placeholder="＋作業種別を追加" style="width:100%"></td>
              <td><input type="number" step="0.1" name="new_base_points" placeholder="点数" style="width:90px"></td><td></td></tr>
        </table>
        <button class="btn" type="submit" style="margin-top:12px">点数マスタを保存</button>
      </form>
      <h2 style="margin:28px 0 4px">加点・係数</h2>
      <p class="muted" style="margin:0 0 12px">点数 = <b>(作業種別の基準点数 ＋ バックエンド加点) × ステージ係数 × 難易度係数</b>。</p>
      <form method="post" action="/dev-point-master/coef">
        <div style="margin-bottom:8px"><b style="font-size:13px">バックエンド加点（基準点数に加算）</b><br>{backend_bonus_inputs}</div>
        <div style="margin-bottom:8px"><b style="font-size:13px">既存分類（ステージ）係数</b><br>{stage_coef_inputs}</div>
        <div><b style="font-size:13px">難易度係数</b><br>{diff_coef_inputs}</div>
        <button class="btn" type="submit" style="margin-top:12px">加点・係数を保存</button>
      </form>
      <h2 style="margin:28px 0 4px">開発担当の週次キャパ</h2>
      <p class="muted" style="margin:0 0 12px">週次上限点数（<b>約20点 ≒ 1 FTE</b>）。負荷率＝その週の割当点数 ÷ この上限（Hishoの点数ダッシュボードで使用）。
        「上限変更週(from)」を入れると、その週以降は「変更後上限」を使います（例: 増員/離任で途中から上限が変わる場合）。</p>
      <form method="post" action="/dev-point-master/capacity">
        <table><tr><th>担当</th><th>週次上限点数</th><th>上限変更週(from)</th><th>変更後上限点数</th></tr>{cap_rows or '<tr><td colspan=4 class=muted>担当マスタ(owners)が空です</td></tr>'}</table>
        <button class="btn" type="submit" style="margin-top:12px">キャパを保存</button>
      </form>
    </div>"""


def dev_project_form(con, project: dict | None = None, deal_id: int | None = None,
                      return_to: str | None = None) -> str:
    """開発案件の新規/編集フォーム。project未指定時は新規入力（商談選択欄あり）。

    return_to: 指定時は保存後・戻る/キャンセル時にこのURLへ戻す（例: 特定日の商談一覧から
    遷移した場合、保存後も同じ一覧に戻れるようにする）。未指定時は従来どおり商談詳細へ戻る。
    """
    is_edit = project is not None
    p = project or {}
    owners = sfa_db.get_master_list(con, "owners")
    # 点数のライブ算出用（作業種別→基準点数＋各係数。係数はマスタ画面で編集された値を反映）
    _pt_base_json = json.dumps({m["work_type"]: m["base_points"] for m in sfa_db.list_dev_point_master(con)},
                               ensure_ascii=False)
    _coefs = sfa_db.get_dev_coefs(con)
    _stage_coef_json = json.dumps(_coefs["stage"], ensure_ascii=False)
    _diff_coef_json = json.dumps(_coefs["difficulty"], ensure_ascii=False)
    _backend_bonus_json = json.dumps(_coefs["backend"], ensure_ascii=False)
    return_to_field = f'<input type="hidden" name="return_to" value="{_esc(return_to)}">' if return_to else ""

    if is_edit:
        deal_label = f'{_esc(p.get("account_name"))} / {_esc(p.get("deal_name"))}'
        deal_field_html = (
            f'<input type="hidden" name="deal_id" value="{p["deal_id"]}">'
            f'<div class="muted" style="margin:4px 0 10px">{deal_label}</div>'
        )
        sales_owner_text = f'{_esc(p.get("sales_owner") or "—")} / {_esc(p.get("sales_sub_owner") or "—")}'
        action = f'/dev-project/{p["id"]}/edit'
        back_href = return_to or f'/deal/{p["deal_id"]}'
        preselect_deal_id = None
    else:
        deals = sfa_db.list_deals(con, status="open")
        opts = "".join(
            f'<option value="{d["id"]}" data-owner="{_esc(d.get("owner"))}" '
            f'data-sub-owner="{_esc(d.get("sub_owner"))}">'
            f'{_esc(d.get("account_name"))} / {_esc(d.get("deal_name"))}</option>'
            for d in deals
        )
        deal_field_html = f"""
          <input type="text" id="dpDealFilter" placeholder="会社名・商談名で絞り込み" oninput="dpFilterDeals()">
          <select name="deal_id" id="dpDealSelect" required size="8" style="height:170px" onchange="dpShowSalesOwner()">
            <option value=""></option>
            {opts}
          </select>
          <p class="muted" id="dpSalesOwnerLine" style="margin-top:6px">営業担当: —</p>"""
        sales_owner_text = None
        action = "/dev-project/new"
        back_href = return_to or (f'/deal/{deal_id}' if deal_id else "/dev-projects")
        preselect_deal_id = deal_id

    sales_owner_block = "" if sales_owner_text is None else (
        f'<label>営業担当（商談の主担当・サブ担当を自動反映）</label>'
        f'<div class="muted" style="margin-bottom:6px">{sales_owner_text}</div>'
    )

    # 追加ツールリンク（主リンク=tool_urlに加えて2つ目以降。編集時のみ登録可能＝dev_project_idが要る）
    if is_edit:
        _extras = sfa_db.list_dev_project_tools(con, p["id"])
        _erows = ""
        for t in _extras:
            _lbl = "🔧 " + (_esc(t.get("label")) or "リンク")
            _erows += (
                '<div style="display:flex;align-items:center;gap:6px;margin-top:4px;justify-content:flex-start">'
                + _tool_link_btn(t.get("url"), label=_lbl, tool_id=t.get("login_id"), tool_password=t.get("login_pass"))
                + f'<form method="post" action="/dev-project/{p["id"]}/tools/{t["id"]}/delete" style="margin:0" '
                  'onsubmit="return confirm(\'この追加リンクを削除しますか？\')">'
                  '<button class="btn sec" type="submit" style="font-size:10px;padding:1px 6px;color:#991b1b;border-color:#fecaca">×</button>'
                  '</form></div>'
            )
        extra_links_html = (
            '<div style="margin:0;padding:8px 10px;background:#f8f9fb;border:1px solid #e6e9f0;border-radius:8px;'
            'width:100%;box-sizing:border-box;text-align:left">'
            '<label style="font-size:11px;color:#6b7689;display:block;margin-bottom:4px">追加のツールリンク（主リンクの他に複数登録できます）</label>'
            + (_erows or '<span class="muted" style="font-size:11px">なし</span>')
            + f'<form method="post" action="/dev-project/{p["id"]}/tools/add" '
              'style="margin-top:8px;display:flex;flex-wrap:wrap;gap:4px;justify-content:flex-start">'
              '<input type="url" name="url" required placeholder="追加リンク https://..." style="width:100%;font-size:12px;padding:4px 6px">'
              '<input type="text" name="label" placeholder="表示名(任意)" style="width:88px;font-size:11px;padding:4px 6px">'
              '<input type="text" name="login_id" placeholder="ID(任意)" style="width:70px;font-size:11px;padding:4px 6px">'
              '<input type="text" name="login_pass" placeholder="PASS(任意)" style="width:70px;font-size:11px;padding:4px 6px">'
              '<button class="btn sec" type="submit" style="font-size:11px">＋追加</button></form>'
            '</div>'
        )
    else:
        extra_links_html = ('<p class="muted" style="font-size:11px;margin:0;text-align:right;width:100%">'
                            '追加リンクは保存後に登録できます</p>')

    # 必要な技術シード（ツリー: L1カテゴリごとにグループ表示・複数選択, #60）
    _seed_tree = sfa_db.get_tech_seed_tree(con)
    _seed_sel = {s for s in (p.get("tech_seeds") or "").split(",") if s}

    def _seed_box(s):
        return (f'<label style="display:flex;align-items:flex-start;gap:6px;font-weight:400;font-size:13px;'
                f'line-height:1.35;margin:0;cursor:pointer">'
                f'<input type="checkbox" name="tech_seeds" value="{_esc(s)}"{" checked" if s in _seed_sel else ""}'
                f' style="margin-top:2px;flex:0 0 auto;width:auto"><span>{_esc(s)}</span></label>')

    _seed_groups = []
    for _l1, _leaves in _seed_tree.items():
        if not _leaves:
            continue
        _boxes = "".join(_seed_box(s) for s in _leaves)
        _seed_groups.append(
            f'<div style="margin-bottom:10px">'
            f'<div style="font-size:12px;font-weight:700;color:#4338ca;margin:0 0 5px">{_esc(_l1)}</div>'
            f'<div style="display:grid;grid-template-columns:repeat(auto-fill,minmax(150px,1fr));gap:8px 14px">{_boxes}</div>'
            f'</div>')
    tech_seeds_block = "".join(_seed_groups) or \
        '<span class="muted">技術シードが未登録です（管理→技術シードマスタで追加できます）</span>'

    return f"""
    <div class="card" style="max-width:900px">
      <div style="display:flex;justify-content:space-between;align-items:flex-start;gap:12px;flex-wrap:wrap">
        <div>
          <p style="margin:0 0 10px"><a class="btn sec" href="{_esc(back_href)}">← 戻る</a></p>
          <h2 style="margin:0">{'開発案件を編集' if is_edit else '開発案件 新規入力'}</h2>
        </div>
        <div style="display:flex;flex-direction:column;align-items:flex-end;gap:8px;width:300px;max-width:100%">
          <div style="text-align:right;width:100%">
            <label style="font-size:11px;color:#6b7689;margin:0 0 2px;display:block">制作したツールのリンク（主）</label>
            <input type="url" name="tool_url" form="dpForm" placeholder="https://..."
              value="{_esc(p.get('tool_url'))}" style="width:220px;font-size:12px;padding:5px 8px">
            <div style="display:flex;gap:4px;margin-top:4px;justify-content:flex-end">
              <input type="text" name="tool_login_id" form="dpForm" placeholder="ID（必要な場合）"
                value="{_esc(p.get('tool_login_id'))}" style="width:106px;font-size:11px;padding:4px 6px">
              <input type="text" name="tool_login_pass" form="dpForm" placeholder="PASS（必要な場合）"
                value="{_esc(p.get('tool_login_pass'))}" style="width:106px;font-size:11px;padding:4px 6px">
            </div>
          </div>
          {extra_links_html}
        </div>
      </div>
      {_save_bar('dpForm', cancel_url=back_href)}
      <form method="post" action="{action}" id="dpForm">
        {return_to_field}
        <label>商談</label>
        {deal_field_html}
        <div style="display:flex;gap:20px;flex-wrap:wrap;align-items:flex-start">
          <div style="flex:1.6;min-width:340px">
            <label>開発テーマ *</label>
            <input name="theme" required value="{_esc(p.get('theme'))}" style="width:100%">
            <label>開発テーマ詳細</label>
            <textarea name="theme_detail" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="3">{_esc(p.get('theme_detail'))}</textarea>
            <div style="display:flex;gap:12px;flex-wrap:wrap;margin-top:4px">
              <div style="flex:1;min-width:130px"><label>開発期限</label>
                <input type="date" name="deadline" value="{_esc(p.get('deadline'))}" style="width:100%"></div>
              <div style="flex:1;min-width:130px"><label>開発MS日</label>
                <input type="date" name="dev_milestone_date" value="{_esc(p.get('dev_milestone_date'))}" style="width:100%"></div>
              <div style="flex:1.5;min-width:170px"><label>開発MS</label>
                <input name="dev_milestone" value="{_esc(p.get('dev_milestone'))}" list="devMsList"
                  placeholder="例: 2次商談（デモあり）" style="width:100%">
                <datalist id="devMsList"><option value="2次商談（デモあり）"></option><option value="初回デモ"></option>
                  <option value="要件定義完了"></option><option value="本番リリース"></option></datalist></div>
            </div>
          </div>
          <div style="flex:1;min-width:280px">
            <label>開発点数（自動・約20点≒1FTE）</label>
            <div id="dpPoints" class="stage" style="font-size:15px;padding:6px 12px;text-align:center">—</div>
            <label>必要な技術シード（カテゴリ別・複数選択可）</label>
            <div style="margin:2px 0 4px;padding:10px 12px;border:1px solid #e6e9f0;border-radius:8px;
              background:#fafbfc;max-height:360px;overflow:auto">
              {tech_seeds_block}
            </div>
          </div>
        </div>
        <div class="grid">
          <div><label>状況</label><select name="status">{_opt(sfa_db.DEV_PROJECT_STATUSES, p.get('status'))}</select></div>
          <div><label>提供先</label><select name="dev_audience">{_opt(sfa_db.DEV_AUDIENCES, p.get('dev_audience'))}</select></div>
          <div><label>ステージ</label><select name="stage" id="dpStage" onchange="dpRecalcPoints()">{_opt(sfa_db.DEV_PROJECT_STAGES, p.get('stage'))}</select></div>
          <div><label>課金</label><select name="pricing">{_opt(sfa_db.DEV_PRICINGS, p.get('pricing'))}</select></div>
          <div><label>作業種別（点数の基準）</label><select name="work_type" id="dpWorkType" onchange="dpRecalcPoints()">{_opt(sfa_db.dev_work_types(con), p.get('work_type'))}</select></div>
          <div><label>バックエンド有無</label><select name="has_backend" id="dpBackend" onchange="dpRecalcPoints()">{_opt(sfa_db.DEV_HAS_BACKEND, p.get('has_backend'))}</select></div>
          <div><label>解像度</label><select name="resolution" id="dpResolution" onchange="dpRecalcPotential()">{_opt(sfa_db.DEV_RESOLUTIONS, p.get('resolution'))}</select></div>
          <div><label>実現難易度</label><select name="difficulty" id="dpDifficulty" onchange="dpRecalcPotential();dpRecalcPoints()">{_opt(sfa_db.DEV_DIFFICULTIES, p.get('difficulty'))}</select></div>
          <div><label>予算確認</label><select name="budget_confirmed" id="dpBudget" onchange="dpRecalcPotential()">{_opt(sfa_db.DEV_BUDGET_CONFIRMED, p.get('budget_confirmed'))}</select></div>
          <div><label>開発担当</label><select name="dev_owner">{_opt(owners, p.get('dev_owner'))}</select></div>
        </div>
        <label>受注余地（自動判定）</label>
        <div><span class="stage" id="dpOrderPotential">{_esc(p.get('order_potential') or '（保存時に判定）')}</span></div>
        <label>開発開始日〜終了日（開発期限から自動計算・読み取り専用。以後の調整はHisho側ダッシュボードで行う）</label>
        <div class="muted">{_esc(p.get('dev_start_date') or '—')} 〜 {_esc(p.get('dev_end_date') or '—')}{'' if is_edit else '（保存時に自動計算）'}</div>
        <label>技術サポート</label>
        <input name="tech_support" value="{_esc(p.get('tech_support'))}">
        {sales_owner_block}
        <label>開発方針</label>
        <textarea name="dev_policy" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="3">{_esc(p.get('dev_policy'))}</textarea>
        <div style="margin-top:16px">
          <button class="btn" type="submit">保存</button>
          <a class="btn sec" href="{_esc(back_href)}">キャンセル</a>
        </div>
      </form>
    </div>
    <script>
    function dpFilterDeals() {{
      const q = document.getElementById('dpDealFilter').value.trim();
      const sel = document.getElementById('dpDealSelect');
      for (const o of sel.options) {{
        if (!o.value) continue;
        o.style.display = (!q || o.text.includes(q)) ? '' : 'none';
      }}
    }}
    function dpShowSalesOwner() {{
      const sel = document.getElementById('dpDealSelect');
      const o = sel.options[sel.selectedIndex];
      const owner = o ? (o.getAttribute('data-owner') || '—') : '—';
      const sub = o ? (o.getAttribute('data-sub-owner') || '—') : '—';
      document.getElementById('dpSalesOwnerLine').textContent = '営業担当: ' + owner + ' / ' + sub;
    }}
    function dpRecalcPotential() {{
      const budget = document.getElementById('dpBudget').value;
      const resolution = document.getElementById('dpResolution').value;
      const difficulty = document.getElementById('dpDifficulty').value;
      let potential = '中';
      if (budget === '×') potential = '低';
      else if (budget === '〇' && resolution === '〇' && (difficulty === '易' || difficulty === '中')) potential = '高';
      document.getElementById('dpOrderPotential').textContent = potential;
    }}
    var DP_BASE = {_pt_base_json};
    var DP_STAGE_COEF = {_stage_coef_json};
    var DP_DIFF_COEF = {_diff_coef_json};
    var DP_BACKEND_BONUS = {_backend_bonus_json};
    function dpRecalcPoints() {{
      var wt = document.getElementById('dpWorkType').value;
      var st = document.getElementById('dpStage').value;
      var df = document.getElementById('dpDifficulty').value;
      var be = document.getElementById('dpBackend').value;
      var el = document.getElementById('dpPoints');
      var base = DP_BASE[wt];
      if (base == null) {{ el.textContent = '—'; return; }}
      var bonus = DP_BACKEND_BONUS[be] || 0;
      var p = Math.round((base + bonus) * (DP_STAGE_COEF[st] || 1) * (DP_DIFF_COEF[df] || 1) * 10) / 10;
      el.textContent = p + '点（≈' + (p / 20).toFixed(1) + 'FTE）';
    }}
    (function() {{
      const sel = document.getElementById('dpDealSelect');
      if (sel) {{
        const pre = {json.dumps(preselect_deal_id)};
        if (pre) sel.value = String(pre);
        dpShowSalesOwner();
      }}
      dpRecalcPoints();
    }})();
    </script>"""


# ── 社内論点（商談に紐づく議論ポイント管理） ─────────────────────────────────────

def _deal_issues_url(*, status=None, member=None, q=None, sort=None, open_issue=None) -> str:
    params = {}
    if status:
        params["status"] = status
    if member:
        params["member"] = member
    if q:
        params["q"] = q
    if sort:
        params["sort"] = sort
    if open_issue:
        params["open_issue"] = open_issue
    qs = urllib.parse.urlencode(params)
    return "/deal-issues" + (f"?{qs}" if qs else "")


def _issue_status_select_html(issue_id: int, current: str | None) -> str:
    opts = "".join(
        f'<option value="{html.escape(s)}"{" selected" if s == current else ""}>{html.escape(s)}</option>'
        for s in sfa_db.DEAL_ISSUE_STATUSES
    )
    return (f'<select onchange="updateDealIssueField({issue_id}, \'status\', this.value, true)" '
            f'style="font-size:11px;padding:2px 4px">{opts}</select>')


def _issue_members_inline_html(issue_id: int, current_members: str | None) -> str:
    selected = set((current_members or "").split(","))
    boxes = "".join(
        f'<label style="display:inline-flex;align-items:center;gap:2px;margin:0 8px 2px 0;'
        f'font-size:11px;font-weight:normal;white-space:nowrap">'
        f'<input type="checkbox" value="{html.escape(m)}"{" checked" if m in selected else ""} '
        f'onchange="diUpdateMembers({issue_id}, this)" style="width:auto">{html.escape(m)}</label>'
        for m in sfa_db.DEAL_ISSUE_MEMBERS
    )
    return f'<div class="di-members-cell">{boxes}</div>'


def _issue_due_date_input_html(issue_id: int, due_date: str | None) -> str:
    is_overdue = bool(due_date) and due_date <= _today_jst().isoformat()
    style = "font-size:11px;padding:2px 4px"
    if is_overdue:
        style += ";color:#dc2626;font-weight:700;border-color:#dc2626"
    return (f'<input type="date" value="{_esc(due_date)}" '
            f'onchange="updateDealIssueField({issue_id}, \'due_date\', this.value, true)" '
            f'style="{style}">')


DEAL_ISSUE_INLINE_EDIT_JS = """
function updateDealIssueField(id, field, value, reload) {
  fetch('/deal-issue/' + id + '/field', {
    method: 'POST',
    headers: {'Content-Type': 'application/x-www-form-urlencoded'},
    body: 'field=' + encodeURIComponent(field) + '&value=' + encodeURIComponent(value)
  }).then(r => r.json()).then(d => {
    if (!d.ok) { alert('更新エラー'); return; }
    // 一覧の絞り込み・並び替えは変更後の値を反映していないため、再読込して整合させる
    // (ステータスは即時、議論メンバーは連続クリックを考慮して少し待って反映)
    if (reload) location.reload();
  }).catch(() => alert('通信エラー'));
}
var _diMembersReloadTimer = null;
function diUpdateMembers(id, checkboxEl) {
  var container = checkboxEl.closest('.di-members-cell');
  var checked = Array.from(container.querySelectorAll('input[type=checkbox]:checked')).map(function(c){ return c.value; });
  updateDealIssueField(id, 'members', checked.join(','), false);
  clearTimeout(_diMembersReloadTimer);
  _diMembersReloadTimer = setTimeout(function() { location.reload(); }, 900);
}
document.addEventListener('click', function(e) {
  document.querySelectorAll('details.di-memo-details[open]').forEach(function(d) {
    if (!d.contains(e.target)) d.removeAttribute('open');
  });
});
"""


AI_SUMMARY_HOVER_CSS = """
.ai-summary-box { margin-bottom: 6px; }
.ai-summary-label { font-size: 10px; color: #8893a8; letter-spacing: .04em; margin-bottom: 2px; }
.ai-summary-text {
  font-size: 12px; color: #3a4760; background: #f8f9fa; border-radius: 6px; padding: 4px 8px;
  display: -webkit-box; -webkit-box-orient: vertical; -webkit-line-clamp: 1; overflow: hidden;
  cursor: default;
}
.ai-summary-box:hover .ai-summary-text {
  display: block; -webkit-line-clamp: unset; max-height: 220px; overflow-y: auto;
}
.ai-summary-title { text-decoration: underline; font-weight: 700; }
"""


def _format_ai_summary_html(summary: str | None) -> str:
    """AIサマリーの各行(・トピック名：内容)を解析し、トピック名を下線付きで表示するHTMLを組み立てる。"""
    if not summary:
        return '<span class="muted">—</span>'
    lines = [l.strip().lstrip("・-").strip() for l in summary.split("\n") if l.strip()]
    parts = []
    for line in lines:
        if "：" in line:
            title, rest = line.split("：", 1)
        elif ":" in line:
            title, rest = line.split(":", 1)
        else:
            title, rest = line, ""
        title_html = f'<span class="ai-summary-title">{_esc(title.strip())}</span>'
        rest_html = f"：{_esc(rest.strip())}" if rest else ""
        parts.append(f"・{title_html}{rest_html}")
    return "<br>".join(parts)


def _ai_summary_hover_html(summary: str | None, *, issue_id: int, return_to: str) -> str:
    """AIサマリー: 通常は1行に折りたたみ、カーソルを合わせると全文を縦に広げて表示する。"""
    text = _format_ai_summary_html(summary)
    return f"""
    <div class="ai-summary-box">
      <div class="ai-summary-label">AIサマリー
        <form method="post" action="/deal-issue/{issue_id}/regenerate_summary" style="display:inline">
          <input type="hidden" name="return_to" value="{_esc(return_to)}">
          <button type="submit" class="muted" style="background:none;border:0;cursor:pointer;padding:0;
            margin-left:6px;font-size:10px;text-decoration:underline" title="メモ履歴から再生成">🔄再生成</button>
        </form>
      </div>
      <div class="ai-summary-text">{text}</div>
    </div>"""


def _issue_memo_panel_html(memos: list[dict], issue: dict, *, return_to: str) -> str:
    """論点1件分の、メモ履歴＋追加フォーム（<details>展開時に表示する内容）。"""
    memo_html = "".join(
        f'<div style="padding:6px 0;border-bottom:1px solid #eef1f5;display:flex;justify-content:space-between;gap:8px">'
        f'<div style="flex:1;min-width:0">'
        f'<div class="muted">{_esc((m.get("created_at") or "")[:16])}</div>'
        f'<div style="white-space:pre-wrap">{_esc(m.get("body"))}</div></div>'
        f'<form method="post" action="/deal-issue-memo/{m["id"]}/delete" style="flex-shrink:0" '
        f'onsubmit="return confirm(\'このメモを削除しますか？\')">'
        f'<input type="hidden" name="return_to" value="{_esc(return_to)}">'
        f'<button type="submit" class="muted" style="background:none;border:0;cursor:pointer;padding:0;'
        f'font-size:12px" title="削除">✕</button></form></div>'
        for m in memos
    ) or '<div class="muted">まだメモがありません</div>'
    return f"""
    <div style="margin-top:8px;width:100%">
      <div style="max-height:220px;overflow-y:auto;margin-bottom:8px">{memo_html}</div>
      <form method="post" action="/deal-issue/{issue['id']}/memo">
        <input type="hidden" name="return_to" value="{_esc(return_to)}">
        <textarea name="body" class="ta-expand" rows="2" placeholder="メモを追加"
          onfocus="taExpand(this)" onblur="taShrink(this)" required
          style="width:100%;box-sizing:border-box"></textarea>
        <p style="margin-top:6px"><button class="btn sec" type="submit">メモを追加</button></p>
      </form>
    </div>"""


def deal_issues_list_page(con, *, status: str | None = None, member: str | None = None,
                           q: str | None = None, sort: str | None = None,
                           open_issue: str | None = None) -> str:
    sort = sort or "due_date"
    issues = sfa_db.list_deal_issues(con, status=status, member=member, q=q, sort=sort)

    def _fopt(values, current):
        return '<option value="">全て</option>' + "".join(
            f'<option value="{html.escape(v)}"{" selected" if v == current else ""}>{html.escape(v)}</option>'
            for v in values
        )

    sort_labels = [("due_date", "解消期限"), ("status", "ステータス"), ("updated_at", "更新日時")]
    sort_opts = "".join(
        f'<option value="{k}"{" selected" if k == sort else ""}>並び替え: {label}</option>'
        for k, label in sort_labels
    )

    filter_row = f"""<form method="get" action="/deal-issues" class="filter-row">
      <input name="q" placeholder="会社名・商談名で検索" value="{_esc(q or '')}">
      <select name="status" onchange="this.form.submit()">{_fopt(sfa_db.DEAL_ISSUE_STATUSES, status).replace('全て', 'ステータス:全て', 1)}</select>
      <select name="member" onchange="this.form.submit()">{_fopt(sfa_db.DEAL_ISSUE_MEMBERS, member).replace('全て', '議論メンバー:全て', 1)}</select>
      <select name="sort" onchange="this.form.submit()">{sort_opts}</select>
      <button class="btn sec" type="submit">検索</button>
      <a class="btn sec" href="/deal-issues">リセット</a>
    </form>"""

    rows = ""
    for it in issues:
        return_to = _deal_issues_url(status=status, member=member, q=q, sort=sort, open_issue=it["id"])
        memos = sfa_db.list_deal_issue_memos(con, it["id"])
        is_open = str(open_issue) == str(it["id"])
        memo_panel = _issue_memo_panel_html(memos, it, return_to=return_to)
        summary_box = _ai_summary_hover_html(it.get('ai_summary'), issue_id=it['id'], return_to=return_to)
        deal_cell = (
            f'<a href="/deal/{it["deal_id"]}">{_esc(it.get("account_name"))}</a>'
            f'<div class="muted">{_esc(it.get("deal_name"))}</div>'
            if it.get('deal_id') else '<span class="muted">商談共通</span>'
        )
        rows += f"""
        <tr>
          <td>{deal_cell}</td>
          <td>{_esc(it.get('issue'))}</td>
          <td>{_issue_status_select_html(it['id'], it.get('status'))}</td>
          <td>{_issue_members_inline_html(it['id'], it.get('members'))}</td>
          <td>{_issue_due_date_input_html(it['id'], it.get('due_date'))}</td>
          <td>
            {summary_box}
            <details class="di-memo-details"{' open' if is_open else ''}>
              <summary style="cursor:pointer">メモ（{len(memos)}）</summary>
              {memo_panel}
            </details>
          </td>
          <td><a href="/deal-issue/{it['id']}/edit">編集</a></td>
        </tr>"""

    return f"""
    <style>
    {AI_SUMMARY_HOVER_CSS}
    </style>
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>社内論点一覧（{len(issues)}件）</span>
        <a class="btn" href="/deal-issue/new">＋新規論点</a>
      </h2>
      {filter_row}
      <div style="overflow:auto;max-height:70vh">
      <table style="table-layout:fixed;width:100%">
        <tr>{_sticky_th('商談', width='16%')}{_sticky_th('論点', width='14%')}{_sticky_th('ステータス', width='9%')}
            {_sticky_th('議論メンバー', width='14%')}{_sticky_th('解消期限', width='9%')}
            {_sticky_th('サマリー・メモ', width='34%')}{_sticky_th('', width='5%')}</tr>
        {rows or '<tr><td colspan=7 class=muted>論点がまだありません</td></tr>'}
      </table>
      </div>
    </div>
    <script>
    {DEAL_ISSUE_INLINE_EDIT_JS}
    </script>"""


def deal_issue_form(con, issue: dict | None = None, deal_id: int | None = None,
                     return_to: str | None = None) -> str:
    """論点の新規/編集フォーム。issue未指定時は新規入力（商談選択欄あり）。"""
    is_edit = issue is not None
    it = issue or {}
    return_to_field = f'<input type="hidden" name="return_to" value="{_esc(return_to)}">' if return_to else ""
    selected_members = set((it.get("members") or "").split(","))
    members_html = "".join(
        f'<label style="display:inline-flex;align-items:center;gap:4px;margin:0 14px 6px 0;font-weight:normal">'
        f'<input type="checkbox" name="members" value="{_esc(m)}"'
        f'{" checked" if m in selected_members else ""} style="width:auto">{_esc(m)}</label>'
        for m in sfa_db.DEAL_ISSUE_MEMBERS
    )
    delete_btn = ""

    if is_edit:
        if it.get("deal_id"):
            deal_label = f'{_esc(it.get("account_name"))} / {_esc(it.get("deal_name"))}'
        else:
            deal_label = '商談共通（特定の商談に紐づかない論点）'
        deal_field_html = (
            f'<input type="hidden" name="deal_id" value="{it.get("deal_id") or ""}">'
            f'<div class="muted" style="margin:4px 0 10px">{deal_label}</div>'
        )
        action = f'/deal-issue/{it["id"]}/edit'
        back_href = return_to or (f'/deal/{it["deal_id"]}' if it.get("deal_id") else "/deal-issues")
        delete_btn = (
            f'<form method="post" action="/deal-issue/{it["id"]}/delete" style="display:inline;margin-left:8px" '
            f'onsubmit="return confirm(\'削除しますか？\')">'
            f'<button class="btn" style="background:#ef4444" type="submit">削除</button></form>'
        )
    else:
        deals = sfa_db.list_deals(con, status="open")
        opts = "".join(
            f'<option value="{d["id"]}"{" selected" if deal_id == d["id"] else ""}>'
            f'{_esc(d.get("account_name"))} / {_esc(d.get("deal_name"))}</option>'
            for d in deals
        )
        deal_field_html = f"""
          <input type="text" id="diDealFilter" placeholder="会社名・商談名で絞り込み" oninput="diFilterDeals()">
          <select name="deal_id" id="diDealSelect" size="8" style="height:170px">
            <option value="">（商談に紐づけない・商談共通）</option>
            {opts}
          </select>"""
        action = "/deal-issue/new"
        back_href = return_to or (f'/deal/{deal_id}' if deal_id else "/deal-issues")

    default_due_date = it.get('due_date') if is_edit else (
        it.get('due_date') or (_today_jst() + timedelta(days=7)).isoformat())

    return f"""
    <div class="card" style="max-width:700px">
      <p style="margin:0 0 10px"><a class="btn sec" href="{_esc(back_href)}">← 戻る</a></p>
      <h2>{'論点を編集' if is_edit else '論点 新規入力'}</h2>
      {_save_bar('issueForm', cancel_url=back_href)}
      <form method="post" action="{action}" id="issueForm">
        {return_to_field}
        <label>商談</label>
        {deal_field_html}
        <label>論点 *</label>
        <input name="issue" required value="{_esc(it.get('issue'))}">
        <label>議論メンバー</label>
        <div style="margin:4px 0 10px">{members_html}</div>
        <div class="grid">
          <div><label>ステータス</label><select name="status">{_opt(sfa_db.DEAL_ISSUE_STATUSES, it.get('status') or '議論中')}</select></div>
          <div><label>解消期限</label><input type="date" name="due_date" value="{_esc(default_due_date)}"></div>
        </div>
        <div style="margin-top:16px">
          <button class="btn" type="submit">保存</button>
          <a class="btn sec" href="{_esc(back_href)}">キャンセル</a>
          {delete_btn}
        </div>
      </form>
    </div>
    <script>
    function diFilterDeals() {{
      const q = document.getElementById('diDealFilter').value.trim();
      const sel = document.getElementById('diDealSelect');
      for (const o of sel.options) {{
        if (!o.value) continue;
        o.style.display = (!q || o.text.includes(q)) ? '' : 'none';
      }}
    }}
    </script>"""


# ── リード / ピッチテーマ ページ（CRM吸収）─────────────────────────────────────

_SOURCE_TO_LP = {"exhibition": "Exh.", "referral": "Connection", "inbound": "HP", "other": "na"}


def convert_lead_to_deal(con, lead: dict) -> int:
    """リードを商談化してdeal_idを返す（アカウント・コンタクト作成、リードをconvertedに）。
    既存のオープン商談がある場合はそのidを返す。クローズ済なら再変換する。"""
    if lead.get("deal_id"):
        _ed = sfa_db.get_deal(con, lead["deal_id"])
        if _ed and _ed.get("status") != "closed":
            return int(lead["deal_id"])
        con.execute("UPDATE leads SET deal_id=NULL WHERE id=?", (lead["id"],))
        con.commit()
    # 1. アカウントを検索または作成
    company_name = (lead.get("company") or "").strip() or "(未設定)"
    existing_acc = con.execute(
        "SELECT id FROM accounts WHERE name=?", (company_name,)
    ).fetchone()
    account_id = (dict(existing_acc)["id"] if existing_acc
                  else sfa_db.upsert_account(
                      con, name=company_name,
                      industry=lead.get("industry"),
                      company_size=lead.get("company_size"),
                  ))
    # 2. コンタクト作成（重複チェック）
    if not con.execute(
        "SELECT id FROM contacts WHERE account_id=? AND name=?",
        (account_id, lead["name"]),
    ).fetchone():
        con.execute(
            "INSERT INTO contacts (account_id,name,title,email,phone) VALUES (?,?,?,?,?)",
            (account_id, lead["name"], lead.get("title"),
             lead.get("email"), lead.get("phone")),
        )
        con.commit()
    # 3. 商談作成
    deal_id = sfa_db.upsert_deal(
        con, account_id=account_id,
        deal_name=company_name, stage="初回アポ実施", status="open",
        lead_pattern=_SOURCE_TO_LP.get(lead.get("source", "other"), "na"),
        owner=lead.get("assigned_to"), note=lead.get("notes"),
    )
    # 4. リードをクローズ（商談化済）してdeal_idをセット
    con.execute(
        "UPDATE leads SET deal_id=?, lead_status='converted', updated_at=datetime('now') WHERE id=?",
        (deal_id, lead["id"]),
    )
    con.commit()
    return int(deal_id)


def find_past_closed_deals(con, company_name: str) -> list[dict]:
    """会社名（アカウント名）が一致する、クローズ済みの過去商談を新しい順で返す。
    リードから商談を復活させる際の選択肢に使う。"""
    company_name = (company_name or "").strip()
    if not company_name:
        return []
    acc = con.execute("SELECT id FROM accounts WHERE name=?", (company_name,)).fetchone()
    if not acc:
        return []
    return [dict(r) for r in con.execute(
        "SELECT id, deal_name, stage, owner, updated_at FROM deals "
        "WHERE account_id=? AND status='closed' ORDER BY updated_at DESC",
        (dict(acc)["id"],),
    )]


def revive_deal_from_lead(con, lead: dict, deal_id: int) -> int:
    """過去にクローズした商談を復活させ、リードをそこに紐づける。"""
    con.execute("UPDATE deals SET status='open', updated_at=datetime('now') WHERE id=?", (deal_id,))
    con.execute(
        "UPDATE leads SET deal_id=?, lead_status='converted', updated_at=datetime('now') WHERE id=?",
        (deal_id, lead["id"]),
    )
    con.execute(
        "INSERT INTO activities (deal_id, type, occurred_on, body) VALUES (?,?,date('now'),?)",
        (deal_id, "メモ", f"リード「{lead.get('name','')}」からアポ再獲得。商談を復活しました。"),
    )
    con.commit()
    return int(deal_id)


def lead_convert_choice_page(con, lead: dict, past_deals: list[dict]) -> str:
    """過去のクローズ済み商談がある場合に、新規作成か復活かを選ばせる確認画面。"""
    opts = "".join(
        f'<label style="display:block;margin:8px 0;font-weight:400">'
        f'<input type="radio" name="mode" value="revive_{d["id"]}" style="width:auto;margin-right:6px">'
        f'過去の商談を復活: <strong>{_esc(d.get("deal_name"))}</strong>'
        f'（ステージ: {_esc(d.get("stage")) or "—"}、担当: {_esc(d.get("owner")) or "—"}、'
        f'最終更新: {_esc((d.get("updated_at") or "")[:10])}）</label>'
        for d in past_deals
    )
    return f"""
    <div class="card" style="max-width:700px">
      <h2>商談化の方法を選択</h2>
      <p class="muted">この会社（{_esc(lead.get('company'))}）には過去にクローズした商談があります。
      過去の商談を復活させて続きから進めますか？それとも新規の商談として開始しますか？</p>
      <form method="post" action="/leads/{lead['id']}/convert">
        <label style="display:block;margin:8px 0;font-weight:400">
          <input type="radio" name="mode" value="new" checked style="width:auto;margin-right:6px">
          新規商談として作成
        </label>
        {opts}
        <div style="margin-top:16px">
          <button class="btn sync" type="submit">確定</button>
          <a class="btn sec" href="/leads/{lead['id']}">キャンセル</a>
        </div>
      </form>
    </div>"""


# ── 初回ヒアリング ───────────────────────────────────────────────────────────────

def hearing_templates_page(con) -> str:
    tmpls = sfa_db.list_hearing_templates(con)
    counts = sfa_db.count_hearing_results_by_template(con)
    rows = ""
    for t in tmpls:
        n_items = len(t.get("items") or [])
        n_results = counts.get(t["id"], 0)
        count_cell = (
            f'<a href="/hearings?template_id={t["id"]}" title="このテンプレートの実施済みヒアリング一覧へ">{n_results}件</a>'
            if n_results else '<span class="muted">0件</span>'
        )
        rows += (
            f'<tr>'
            f'<td><a href="/hearing-templates/{t["id"]}/edit"><strong>{_esc(t["name"])}</strong></a></td>'
            f'<td class="muted">{_esc(t.get("description") or "—")}</td>'
            f'<td style="text-align:center">{n_items}</td>'
            f'<td style="text-align:center">{count_cell}</td>'
            f'<td><form method="post" action="/hearing-templates/{t["id"]}/delete" style="display:inline">'
            f'<button class="btn sec" style="font-size:11px;padding:4px 8px" '
            f'onclick="return confirm(\'削除しますか？（既存のヒアリング結果は残ります）\')">削除</button></form></td>'
            f'</tr>'
        )
    return f"""
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>ヒアリングテンプレート管理</span>
        <span style="display:flex;gap:8px">
          <a class="btn sec" href="/hearings">ヒアリング一覧</a>
          <a class="btn" href="/hearing-templates/new">＋テンプレート追加</a>
        </span>
      </h2>
      <p class="muted" style="margin-bottom:14px">初回商談で使う定型ヒアリング項目を定義します。自由記述／選択肢（単一・複数）を項目ごとに指定できます。</p>
      <table>
        <tr><th>テンプレート名</th><th>説明</th><th>項目数</th><th>実施数</th><th></th></tr>
        {rows or '<tr><td colspan=5 class="muted">テンプレートがありません。</td></tr>'}
      </table>
    </div>"""


def hearing_template_form(con, tmpl=None) -> str:
    tid = tmpl["id"] if tmpl else None
    action = f"/hearing-templates/{tid}/save" if tid else "/hearing-templates/save"
    title = "テンプレート編集" if tmpl else "テンプレート追加"
    items = (tmpl.get("items") if tmpl else None) or []
    items_data = json.dumps(items, ensure_ascii=False)
    return f"""
    <div class="card" style="max-width:1000px">
      <h2>{title}</h2>
      <form method="post" action="{action}" onsubmit="return serializeItems()">
        <label>テンプレート名</label>
        <input name="name" required value="{_esc(tmpl.get('name') if tmpl else '')}">
        <label>説明（任意）</label>
        <input name="description" value="{_esc(tmpl.get('description') if tmpl else '')}">
        <label style="margin-top:14px">ヒアリング項目</label>
        <p class="muted" style="font-size:12px;margin:4px 0 8px">Q&amp;A項目と矢羽セクションを自由に組み合わせられます。</p>
        <div id="items_box"></div>
        <div style="display:flex;gap:8px;margin-top:8px;flex-wrap:wrap">
          <button type="button" class="btn sec" onclick="addItem()">＋Q&amp;A項目を追加</button>
          <button type="button" class="btn sec" style="border-color:#3b82f660;color:#3b82f6"
            onclick="addYabaneItem(null)">＋矢羽セクションを追加</button>
          <button type="button" class="btn sec" style="border-color:#7c3aed60;color:#7c3aed"
            onclick="addRadarItem(null)">＋レーダーチャートを追加</button>
          <button type="button" class="btn sec" style="border-color:#05966960;color:#059669"
            onclick="addTimelineItem(null)">＋タイムラインを追加</button>
          <button type="button" class="btn sec" style="border-color:#d9770660;color:#d97706"
            onclick="addScorecardItem(null)">＋スコアカードを追加</button>
        </div>
        <input type="hidden" name="items_json" id="items_json">
        <div style="margin-top:16px;display:flex;gap:8px">
          <button class="btn" type="submit">保存</button>
          <a class="btn sec" href="/hearing-templates">キャンセル</a>
        </div>
      </form>
    </div>
    <script>
    var _ITEMS = {items_data};

    // ── 矢羽ブロック ──
    function _addYbPairRow(pairsBox, step, dept) {{
      var div = document.createElement('div');
      div.style.cssText = 'display:flex;gap:6px;align-items:center;margin:4px 0';
      var stepInp = document.createElement('input');
      stepInp.type='text'; stepInp.className='yb-pair-step';
      stepInp.value=step||''; stepInp.placeholder='ステップ（例：受注）'; stepInp.style.cssText='flex:1';
      var deptInp = document.createElement('input');
      deptInp.type='text'; deptInp.className='yb-pair-dept';
      deptInp.value=dept||''; deptInp.placeholder='部署（例：営業）'; deptInp.style.cssText='flex:1';
      var btn=document.createElement('button');
      btn.type='button'; btn.className='btn sec';
      btn.style.cssText='font-size:11px;padding:4px 8px;background:#fde8e8;color:#c0392b';
      btn.textContent='削除'; btn.onclick=function(){{div.remove();}};
      div.appendChild(stepInp); div.appendChild(deptInp); div.appendChild(btn); pairsBox.appendChild(div);
    }}

    function addYabaneItem(cfg) {{
      cfg=cfg||{{label:'業務プロセス',rows:[{{step:'ステップ1',dept:'部署1'}},{{step:'ステップ2',dept:'部署2'}}]}};
      // 旧形式（departments/steps）で保存された既存テンプレートを開いた場合の変換
      var rows = cfg.rows;
      if (!rows) {{
        var depts = cfg.departments || [];
        var steps = cfg.steps || [];
        rows = steps.map(function(s, i) {{ return {{step: s.label || '', dept: depts[i] || ''}}; }});
        if (!rows.length) rows = [{{step:'ステップ1',dept:'部署1'}}];
      }}
      var box=document.getElementById('items_box');
      var el=document.createElement('div'); el.className='yb-block';
      el.setAttribute('draggable','true');
      el.style.cssText='border:2px solid #3b82f660;border-radius:8px;padding:12px;margin:8px 0;background:#0a1828';
      // header
      var hdr=document.createElement('div'); hdr.style.cssText='display:flex;align-items:center;gap:8px;margin-bottom:10px';
      var dh=document.createElement('span'); dh.className='drag-handle';
      dh.title='ドラッグで並び替え';
      dh.style.cssText='cursor:grab;color:#555;font-size:18px;user-select:none;line-height:1;flex-shrink:0';
      dh.textContent='⠿';
      var badge=document.createElement('span');
      badge.style.cssText='font-size:10px;font-weight:700;color:#3b82f6;background:#3b82f615;border:1px solid #3b82f640;border-radius:4px;padding:2px 8px;white-space:nowrap;flex-shrink:0';
      badge.textContent='矢羽';
      var lw=document.createElement('div'); lw.style.cssText='flex:1;min-width:0';
      var ll=document.createElement('label'); ll.style.cssText='font-size:12px'; ll.textContent='セクション名';
      var li=document.createElement('input'); li.className='yb-block-label'; li.value=cfg.label||'業務プロセス'; li.placeholder='例：業務プロセス';
      lw.appendChild(ll); lw.appendChild(li);
      var db=document.createElement('button'); db.type='button'; db.className='btn sec';
      db.style.cssText='font-size:11px;padding:4px 8px;background:#fde8e8;color:#c0392b;flex-shrink:0';
      db.textContent='削除'; db.onclick=function(){{this.closest('.yb-block').remove();}};
      hdr.appendChild(dh); hdr.appendChild(badge); hdr.appendChild(lw); hdr.appendChild(db); el.appendChild(hdr);
      // body: ステップ・部署をペアで並べて入力
      var pl=document.createElement('label'); pl.style.cssText='font-size:12px'; pl.textContent='初期のステップ・部署（1行=1ステップ1部署。ヒアリング入力画面はこの並びをそのまま行として表示します）';
      el.appendChild(pl);
      var pairsBox=document.createElement('div'); pairsBox.className='yb-block-pairs'; pairsBox.style.cssText='margin:4px 0';
      el.appendChild(pairsBox);
      var addBtn=document.createElement('button'); addBtn.type='button'; addBtn.className='btn sec';
      addBtn.style.cssText='font-size:11px;padding:4px 8px;margin-top:4px'; addBtn.textContent='＋行追加';
      addBtn.onclick=function(){{_addYbPairRow(pairsBox,'','');}};
      el.appendChild(addBtn);
      box.appendChild(el);
      rows.forEach(function(r){{_addYbPairRow(pairsBox, r.step||'', r.dept||'');}});
    }}

    function _makeBlockHeader(el, blockClass, badgeText, badgeColor, labelClass, defaultLabel) {{
      var hdr=document.createElement('div'); hdr.style.cssText='display:flex;align-items:center;gap:8px;margin-bottom:10px';
      var dh=document.createElement('span'); dh.className='drag-handle'; dh.title='ドラッグで並び替え';
      dh.style.cssText='cursor:grab;color:#555;font-size:18px;user-select:none;line-height:1;flex-shrink:0';
      dh.textContent='⠿';
      var badge=document.createElement('span');
      badge.style.cssText='font-size:10px;font-weight:700;color:'+badgeColor+';background:'+badgeColor+'15;border:1px solid '+badgeColor+'40;border-radius:4px;padding:2px 8px;white-space:nowrap;flex-shrink:0';
      badge.textContent=badgeText;
      var lw=document.createElement('div'); lw.style.cssText='flex:1;min-width:0';
      var ll=document.createElement('label'); ll.style.cssText='font-size:12px'; ll.textContent='セクション名';
      var li=document.createElement('input'); li.className=labelClass; li.value=defaultLabel; li.placeholder='例：'+defaultLabel;
      lw.appendChild(ll); lw.appendChild(li);
      var db=document.createElement('button'); db.type='button'; db.className='btn sec';
      db.style.cssText='font-size:11px;padding:4px 8px;background:#fde8e8;color:#c0392b;flex-shrink:0';
      db.textContent='削除'; db.onclick=function(){{this.closest('.'+blockClass).remove();}};
      hdr.appendChild(dh); hdr.appendChild(badge); hdr.appendChild(lw); hdr.appendChild(db); el.appendChild(hdr);
    }}

    function addRadarItem(cfg) {{
      cfg=cfg||{{label:'DX成熟度評価',axes:['戦略','組織','プロセス','テクノロジー','データ']}};
      var box=document.getElementById('items_box');
      var el=document.createElement('div'); el.className='ra-block'; el.setAttribute('draggable','true');
      el.style.cssText='border:2px solid #7c3aed60;border-radius:8px;padding:12px;margin:8px 0;background:#0a1828';
      _makeBlockHeader(el,'ra-block','レーダー','#7c3aed','ra-block-label',cfg.label||'DX成熟度評価');
      var axDiv=document.createElement('div');
      var axL=document.createElement('label'); axL.style.cssText='font-size:12px'; axL.textContent='評価軸（1行に1つ、3〜8軸推奨）';
      var axTa=document.createElement('textarea'); axTa.className='ra-block-axes'; axTa.rows=5;
      axTa.style.cssText='font-family:inherit'; axTa.placeholder='例：戦略\\n組織\\nプロセス\\nテクノロジー\\nデータ';
      axTa.value=(cfg.axes||[]).join('\\n');
      axDiv.appendChild(axL); axDiv.appendChild(axTa); el.appendChild(axDiv);
      box.appendChild(el);
    }}

    function addTimelineItem(cfg) {{
      cfg=cfg||{{label:'導入スケジュール',milestones:['要件定義','設計・開発','テスト','本稼働']}};
      var box=document.getElementById('items_box');
      var el=document.createElement('div'); el.className='tl-block'; el.setAttribute('draggable','true');
      el.style.cssText='border:2px solid #05966940;border-radius:8px;padding:12px;margin:8px 0;background:#0a1828';
      _makeBlockHeader(el,'tl-block','タイムライン','#059669','tl-block-label',cfg.label||'導入スケジュール');
      var msDiv=document.createElement('div');
      var msL=document.createElement('label'); msL.style.cssText='font-size:12px'; msL.textContent='初期マイルストーン（1行に1つ）';
      var msTa=document.createElement('textarea'); msTa.className='tl-block-milestones'; msTa.rows=4;
      msTa.style.cssText='font-family:inherit'; msTa.placeholder='例：要件定義\\n設計・開発\\nテスト\\n本稼働';
      msTa.value=(cfg.milestones||[]).join('\\n');
      msDiv.appendChild(msL); msDiv.appendChild(msTa); el.appendChild(msDiv);
      box.appendChild(el);
    }}

    function addScorecardItem(cfg) {{
      cfg=cfg||{{label:'競合比較',criteria:['コスト','機能性','サポート','実績'],items:['自社','競合A']}};
      var box=document.getElementById('items_box');
      var el=document.createElement('div'); el.className='sc-block'; el.setAttribute('draggable','true');
      el.style.cssText='border:2px solid #d9770640;border-radius:8px;padding:12px;margin:8px 0;background:#0a1828';
      _makeBlockHeader(el,'sc-block','スコアカード','#d97706','sc-block-label',cfg.label||'競合比較');
      var body=document.createElement('div'); body.style.cssText='display:flex;gap:12px;flex-wrap:wrap';
      var crDiv=document.createElement('div'); crDiv.style.cssText='flex:1;min-width:160px';
      var crL=document.createElement('label'); crL.style.cssText='font-size:12px'; crL.textContent='評価軸（列・1行に1つ）';
      var crTa=document.createElement('textarea'); crTa.className='sc-block-criteria'; crTa.rows=4;
      crTa.style.cssText='font-family:inherit'; crTa.placeholder='例：コスト\\n機能性\\nサポート\\n実績';
      crTa.value=(cfg.criteria||[]).join('\\n');
      crDiv.appendChild(crL); crDiv.appendChild(crTa);
      var itDiv=document.createElement('div'); itDiv.style.cssText='flex:1;min-width:160px';
      var itL=document.createElement('label'); itL.style.cssText='font-size:12px'; itL.textContent='初期対象（行・1行に1つ）';
      var itTa=document.createElement('textarea'); itTa.className='sc-block-items'; itTa.rows=4;
      itTa.style.cssText='font-family:inherit'; itTa.placeholder='例：自社\\n競合A\\n競合B';
      itTa.value=(cfg.items||[]).join('\\n');
      itDiv.appendChild(itL); itDiv.appendChild(itTa);
      body.appendChild(crDiv); body.appendChild(itDiv); el.appendChild(body);
      box.appendChild(el);
    }}

    function rowHtml(it) {{
      it = it || {{label:'',type:'text',multi:false,required:false,options:[],parent_idx:null,parent_value:null}};
      const opts = (it.options || []).join('\\n');
      const hasBranch = it.parent_idx !== null && it.parent_idx !== undefined;
      return `
      <div class="hitem" draggable="true" style="border:1px solid #d4dae4;border-radius:8px;padding:12px;margin:8px 0;background:#fafbfc">
        <div style="display:flex;gap:8px;align-items:flex-end;flex-wrap:wrap">
          <span class="drag-handle" title="ドラッグで並び替え" style="cursor:grab;color:#bbb;font-size:18px;padding-bottom:6px;user-select:none;line-height:1;flex-shrink:0">⠿</span>
          <div style="flex:2;min-width:200px"><label style="font-size:12px">質問ラベル</label>
            <input class="i-label" value="${{(it.label||'').replace(/"/g,'&quot;')}}" placeholder="例：現状の課題"></div>
          <div style="flex:1;min-width:120px"><label style="font-size:12px">回答形式</label>
            <select class="i-type" onchange="syncRow(this)">
              <option value="text"${{it.type==='text'?' selected':''}}>自由記述（テキスト）</option>
              <option value="number"${{it.type==='number'?' selected':''}}>自由記述（数値のみ）</option>
              <option value="choice"${{it.type==='choice'?' selected':''}}>選択肢</option>
            </select></div>
          <div class="i-multi-wrap" style="min-width:120px;${{it.type==='choice'?'':'display:none'}}">
            <label style="font-size:12px">選択方式</label>
            <select class="i-multi">
              <option value="0"${{!it.multi?' selected':''}}>単一選択</option>
              <option value="1"${{it.multi?' selected':''}}>複数選択</option>
            </select></div>
          <div style="display:flex;align-items:center;gap:4px;padding-bottom:6px">
            <input type="checkbox" class="i-required" ${{it.required?'checked':''}} style="width:14px;height:14px">
            <label style="font-size:12px;margin:0">必須</label></div>
          <button type="button" class="btn sec" style="font-size:11px;padding:4px 8px;background:#fde8e8;color:#c0392b"
            onclick="this.closest('.hitem').remove(); refreshBranchSelectors()">削除</button>
        </div>
        <div class="i-opts-wrap" style="margin-top:8px;${{it.type==='choice'?'':'display:none'}}">
          <label style="font-size:12px">選択肢（1行に1つ）</label>
          <textarea class="i-options" rows="3" placeholder="選択肢1&#10;選択肢2"
            oninput="refreshBranchSelectors()">${{opts}}</textarea>
        </div>
        <div style="margin-top:10px;border-top:1px solid #e8ecf0;padding-top:8px">
          <label style="font-size:12px;cursor:pointer;user-select:none">
            <input type="checkbox" class="i-has-branch" ${{hasBranch?'checked':''}}
              onchange="toggleBranchSection(this); refreshBranchSelectors()" style="width:13px;height:13px;margin-right:4px">
            <span style="color:#666">この質問は別の質問の回答を条件に表示する（分岐）</span>
          </label>
          <div class="i-branch-section" style="margin-top:8px;display:${{hasBranch?'flex':'none'}};gap:12px;flex-wrap:wrap;align-items:flex-end">
            <div style="flex:2;min-width:180px">
              <label style="font-size:12px">分岐元の質問（選択肢型のみ）</label>
              <select class="i-parent-idx" data-init="${{hasBranch ? it.parent_idx : ''}}"
                onchange="onParentIdxChange(this)">
                <option value="">— 選択 —</option>
              </select>
            </div>
            <div style="flex:1;min-width:140px">
              <label style="font-size:12px">この回答のときに表示</label>
              <select class="i-parent-value" data-init="${{hasBranch ? (it.parent_value||'') : ''}}">
                <option value="">— 選択 —</option>
              </select>
            </div>
          </div>
        </div>
      </div>`;
    }}

    function toggleBranchSection(cb) {{
      const sec = cb.closest('.hitem').querySelector('.i-branch-section');
      sec.style.display = cb.checked ? 'flex' : 'none';
    }}

    function syncRow(sel) {{
      const row = sel.closest('.hitem');
      const isChoice = sel.value === 'choice';
      row.querySelector('.i-multi-wrap').style.display = isChoice ? '' : 'none';
      row.querySelector('.i-opts-wrap').style.display = isChoice ? '' : 'none';
    }}

    function refreshBranchSelectors() {{
      // 全hitemの情報を収集
      const rows = Array.from(document.querySelectorAll('#items_box .hitem'));
      const choiceItems = rows.map((row, idx) => {{
        const label = row.querySelector('.i-label').value.trim() || `Q${{idx+1}}`;
        const type  = row.querySelector('.i-type').value;
        const opts  = type==='choice'
          ? row.querySelector('.i-options').value.split('\\n').map(s=>s.trim()).filter(Boolean)
          : [];
        return {{idx, label, type, opts}};
      }});

      rows.forEach((row, currentIdx) => {{
        const parentIdxSel  = row.querySelector('.i-parent-idx');
        const parentValSel  = row.querySelector('.i-parent-value');
        if (!parentIdxSel) return;

        const prevIdxVal  = parentIdxSel.value || parentIdxSel.dataset.init || '';
        const prevValVal  = parentValSel.value  || parentValSel.dataset.init  || '';

        // 分岐元ドロップダウンを再構築
        parentIdxSel.innerHTML = '<option value="">— 選択 —</option>';
        choiceItems.forEach(c => {{
          if (c.idx === currentIdx || c.type !== 'choice') return;
          const opt = document.createElement('option');
          opt.value = c.idx;
          opt.textContent = `Q${{c.idx+1}}: ${{c.label}}`;
          if (String(c.idx) === String(prevIdxVal)) opt.selected = true;
          parentIdxSel.appendChild(opt);
        }});
        parentIdxSel.dataset.init = '';

        // 条件値ドロップダウンを再構築
        const selIdx = parseInt(parentIdxSel.value);
        parentValSel.innerHTML = '<option value="">— 選択 —</option>';
        if (!isNaN(selIdx)) {{
          const parent = choiceItems.find(c => c.idx === selIdx);
          if (parent) {{
            parent.opts.forEach(opt => {{
              const o = document.createElement('option');
              o.value = opt;
              o.textContent = opt;
              if (opt === prevValVal) o.selected = true;
              parentValSel.appendChild(o);
            }});
          }}
        }}
        parentValSel.dataset.init = '';
      }});
    }}

    function onParentIdxChange(sel) {{
      const row = sel.closest('.hitem');
      row.querySelector('.i-parent-value').dataset.init = '';
      refreshBranchSelectors();
    }}

    function addItem(it) {{
      const box = document.getElementById('items_box');
      box.insertAdjacentHTML('beforeend', rowHtml(it));
      refreshBranchSelectors();
    }}

    function serializeItems() {{
      var items = [];
      document.querySelectorAll('#items_box > .hitem, #items_box > .yb-block, #items_box > .ra-block, #items_box > .tl-block, #items_box > .sc-block').forEach(function(el) {{
        if (el.classList.contains('yb-block')) {{
          var ybLabel = el.querySelector('.yb-block-label').value.trim() || '業務プロセス';
          var ybRows = [];
          el.querySelectorAll('.yb-block-pairs > div').forEach(function(row) {{
            var step = row.querySelector('.yb-pair-step').value.trim();
            var dept = row.querySelector('.yb-pair-dept').value.trim();
            if (step || dept) ybRows.push({{step:step, dept:dept}});
          }});
          items.push({{label:ybLabel, type:'yabane', rows:ybRows}});
        }} else if (el.classList.contains('ra-block')) {{
          var raLabel=el.querySelector('.ra-block-label').value.trim()||'DX成熟度評価';
          var axes=el.querySelector('.ra-block-axes').value.split('\\n').map(function(s){{return s.trim();}}).filter(Boolean);
          items.push({{label:raLabel, type:'radar', axes:axes}});
        }} else if (el.classList.contains('tl-block')) {{
          var tlLabel=el.querySelector('.tl-block-label').value.trim()||'導入スケジュール';
          var milestones=el.querySelector('.tl-block-milestones').value.split('\\n').map(function(s){{return s.trim();}}).filter(Boolean);
          items.push({{label:tlLabel, type:'timeline', milestones:milestones}});
        }} else if (el.classList.contains('sc-block')) {{
          var scLabel=el.querySelector('.sc-block-label').value.trim()||'競合比較';
          var criteria=el.querySelector('.sc-block-criteria').value.split('\\n').map(function(s){{return s.trim();}}).filter(Boolean);
          var scItems=el.querySelector('.sc-block-items').value.split('\\n').map(function(s){{return s.trim();}}).filter(Boolean);
          items.push({{label:scLabel, type:'scorecard', criteria:criteria, items:scItems}});
        }} else {{
          var stdLabel = el.querySelector('.i-label').value.trim();
          if (!stdLabel) return;
          var type = el.querySelector('.i-type').value;
          var multi = el.querySelector('.i-multi').value === '1';
          var required = el.querySelector('.i-required').checked;
          var options = type==='choice'
            ? el.querySelector('.i-options').value.split('\\n').map(function(s){{return s.trim();}}).filter(Boolean)
            : [];
          var hasBranch = el.querySelector('.i-has-branch').checked;
          var parentIdxRaw = el.querySelector('.i-parent-idx').value;
          var parentIdx = hasBranch && parentIdxRaw !== '' ? parseInt(parentIdxRaw) : null;
          var parentValue = hasBranch ? (el.querySelector('.i-parent-value').value.trim() || null) : null;
          items.push({{label:stdLabel, type:type, multi:type==='choice'?multi:false, required:required,
            options:options, parent_idx:parentIdx, parent_value:parentValue}});
        }}
      }});
      document.getElementById('items_json').value = JSON.stringify(items);
      return true;
    }}

    (function() {{
      if (_ITEMS.length) {{
        _ITEMS.forEach(function(it) {{
          if (it.type === 'yabane') {{ addYabaneItem(it); }}
          else if (it.type === 'radar') {{ addRadarItem(it); }}
          else if (it.type === 'timeline') {{ addTimelineItem(it); }}
          else if (it.type === 'scorecard') {{ addScorecardItem(it); }}
          else {{ addItem(it); }}
        }});
      }} else {{ addItem(); }}
    }})();

    // ── ドラッグ並び替え ──
    (function() {{
      var box = document.getElementById('items_box');
      var _drag = null;
      box.addEventListener('dragstart', function(e) {{
        var t = e.target;
        if (t.tagName==='INPUT'||t.tagName==='TEXTAREA'||t.tagName==='SELECT'||t.tagName==='BUTTON') {{
          e.preventDefault(); return;
        }}
        var item = t.closest('.hitem,.yb-block,.ra-block,.tl-block,.sc-block');
        if (!item) return;
        _drag = item;
        setTimeout(function(){{item.style.opacity='0.4';}}, 0);
      }});
      box.addEventListener('dragend', function() {{
        if (_drag) _drag.style.opacity='';
        _drag = null;
      }});
      box.addEventListener('dragover', function(e) {{
        e.preventDefault();
        if (!_drag) return;
        var over = e.target.closest('.hitem,.yb-block,.ra-block,.tl-block,.sc-block');
        if (!over||over===_drag) return;
        var rect = over.getBoundingClientRect();
        if (e.clientY < rect.top + rect.height/2) box.insertBefore(_drag, over);
        else box.insertBefore(_drag, over.nextSibling);
      }});
    }})();
    </script>"""


def hearing_new_page(con, preselect: str | None = None) -> str:
    """対象（商談 or リード）とテンプレートを選んでヒアリングを開始する画面。"""
    tmpls = sfa_db.list_hearing_templates(con)
    if not tmpls:
        return ('<div class="card"><h2>新規ヒアリング</h2>'
                '<p class="muted">先にテンプレートを作成してください。</p>'
                '<a class="btn" href="/hearing-templates/new">＋テンプレート追加</a></div>')
    tmpl_opts = "".join(
        f'<option value="{t["id"]}">{_esc(t["name"])}</option>' for t in tmpls
    )
    def _sel(v):
        return " selected" if preselect == v else ""
    deals = sfa_db.list_deals(con, status="open")
    deal_opts = ""
    for d in deals:
        val = f"deal:{d['id']}"
        deal_opts += (f'<option value="{val}"{_sel(val)}>商談: '
                      f'{_esc(d.get("account_name") or "")} / {_esc(d.get("deal_name"))}</option>')
    open_leads = [l for l in sfa_db.list_leads(con)
                  if l.get("lead_status") not in ("converted", "lost") and not l.get("deal_id")]
    lead_opts = ""
    for l in open_leads:
        val = f"lead:{l['id']}"
        lead_opts += (f'<option value="{val}"{_sel(val)}>リード: '
                      f'{_esc(l.get("company") or "?")} / {_esc(l.get("name") or "?")}</option>')
    return f"""
    <div class="card" style="max-width:700px">
      <h2>新規ヒアリング</h2>
      <p class="muted" style="margin-bottom:14px">対象とテンプレートを選んでください。リードを選んだ場合は、保存時に自動で商談化されます。</p>
      <form method="get" action="/hearing/start">
        <label>対象（商談 / リード）</label>
        <select name="target" required>
          <option value="">— 選択 —</option>
          <optgroup label="商談">{deal_opts or '<option disabled>なし</option>'}</optgroup>
          <optgroup label="リード（未商談化）">{lead_opts or '<option disabled>なし</option>'}</optgroup>
        </select>
        <label>ヒアリングテンプレート</label>
        <select name="template_id" required>{tmpl_opts}</select>
        <div style="margin-top:16px"><button class="btn" type="submit">ヒアリング入力へ →</button>
        <a class="btn sec" href="/hearings">キャンセル</a></div>
      </form>
    </div>"""


def hearing_input_page(con, *, target_type, target_id, template, target_label,
                       prefill=None, prev_date=None, draft=None,
                       edit_result_id=None, edit_conducted_on=None) -> str:
    """ヒアリング入力画面：ヒアリング項目＋通常の活動履歴入力欄を同一画面に生成。

    draft: get_hearing_draft()の返り値（30秒ごとの自動保存下書き）。存在する場合、
    前回確定結果のprefillよりも優先して復元する（より新しい入力中データのため）。
    edit_result_id: 指定時は既存ヒアリング結果の修正モード。新規の活動履歴は作らず、
    対象のhearing_resultsの内容だけを上書きする。
    """
    prefill = prefill or {}
    draft_data = (draft or {}).get("form_data") or {}
    draft_updated_at = (draft or {}).get("updated_at")

    items = template.get("items") or []
    has_branch = any(
        it.get("parent_idx") is not None and it.get("parent_value") is not None
        for it in items
    )
    fields_html = ""
    for i, it in enumerate(items):
        label = _esc(it.get("label"))
        req = " <span style='color:#c0392b'>*</span>" if it.get("required") else ""
        req_attr = " required" if it.get("required") else ""
        pv = prefill.get(it.get("label"))
        if f"answer_{i}" in draft_data:
            _draft_raw = draft_data[f"answer_{i}"]
            if it.get("type") == "yabane":
                try:
                    pv = json.loads(_draft_raw) if isinstance(_draft_raw, str) else _draft_raw
                except (ValueError, TypeError):
                    pass
            else:
                pv = _draft_raw
        # 分岐設定: data属性でJSに渡す
        parent_idx = it.get("parent_idx")
        parent_value = it.get("parent_value")
        branch_attrs = ""
        branch_class = ""
        if parent_idx is not None and parent_value is not None:
            branch_attrs = f' data-parent-idx="{parent_idx}" data-parent-value="{_esc(parent_value)}"'
            branch_class = " hq-branch"
        if it.get("type") == "radar":
            _ra_axes = it.get("axes") or ["軸1", "軸2", "軸3", "軸4", "軸5"]
            _ra_rows_html = "".join(
                f'<div class="ra-axis-row">'
                f'<input class="ra-axis-label-inp" value="{_esc(_ax)}" placeholder="軸名" onfocus="this.select()">'
                f'<input type="range" class="ra-score-range" min="0" max="5" step="1" value="3"'
                f' oninput="raUpdateScore(this)">'
                f'<span class="ra-score-val">3</span>'
                f'</div>'
                for _ax in _ra_axes
            )
            fields_html += (
                f'<div class="hq-item{branch_class}" style="margin:14px 0"{branch_attrs}>'
                f'<label style="font-weight:700;color:#7c3aed;font-size:13px;margin-bottom:6px;display:block">{label}</label>'
                f'<input type="hidden" name="answer_{i}" id="ra_answer_{i}">'
                f'<div class="ra-wrapper" id="ra_wrapper_{i}" data-ra-idx="{i}">'
                f'<div style="display:flex;gap:20px;flex-wrap:wrap;align-items:flex-start">'
                f'<div class="ra-inputs" style="flex:1;min-width:200px">{_ra_rows_html}</div>'
                f'<div style="flex:0 0 200px;min-width:180px;display:flex;flex-direction:column;align-items:center">'
                f'<svg class="ra-svg" width="200" height="200" style="display:block"></svg>'
                f'</div>'
                f'</div>'
                f'</div>'
                f'</div>'
            )
        elif it.get("type") == "timeline":
            _tl_milestones = it.get("milestones") or []
            _tl_events_html = "".join(
                f'<div class="tl-event">'
                f'<div class="tl-dot"></div>'
                f'<div class="tl-event-body">'
                f'<div class="tl-event-header">'
                f'<input type="month" class="tl-date">'
                f'<input class="tl-event-label" value="{_esc(_ms)}" placeholder="マイルストーン名" onfocus="this.select()">'
                f'<button type="button" class="tl-del-btn" onclick="tlDelEvent(this)">✕</button>'
                f'</div>'
                f'<textarea class="tl-note ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="2" placeholder="詳細・メモ"></textarea>'
                f'</div>'
                f'</div>'
                for _ms in _tl_milestones
            )
            fields_html += (
                f'<div class="hq-item{branch_class}" style="margin:14px 0"{branch_attrs}>'
                f'<label style="font-weight:700;color:#059669;font-size:13px;margin-bottom:6px;display:block">{label}</label>'
                f'<input type="hidden" name="answer_{i}" id="tl_answer_{i}">'
                f'<div class="tl-wrapper" id="tl_wrapper_{i}" data-tl-idx="{i}">'
                f'<div class="tl-track" id="tl_track_{i}">{_tl_events_html}</div>'
                f'<div style="margin-top:10px">'
                f'<button type="button" class="btn sec" style="border-color:#05966960;color:#059669"'
                f' onclick="tlAddEvent({i})">＋マイルストーン追加</button>'
                f'</div>'
                f'</div>'
                f'</div>'
            )
        elif it.get("type") == "scorecard":
            _sc_criteria = it.get("criteria") or []
            _sc_items_list = it.get("items") or []
            _sc_crit_ths = "".join(
                f'<th class="sc-crit-h">'
                f'<input class="sc-crit-name" value="{_esc(_c)}" placeholder="評価軸" onfocus="this.select()">'
                f'<button type="button" class="sc-del-crit-btn" onclick="scDelCrit(this)">✕</button>'
                f'</th>'
                for _c in _sc_criteria
            )
            _sc_item_rows = ""
            for _it_name in _sc_items_list:
                _sc_score_cells = "".join(
                    f'<td class="sc-score-td">'
                    f'<input type="number" class="sc-score-inp" min="1" max="5" placeholder="—"'
                    f' oninput="scUpdateTotal(this)">'
                    f'</td>'
                    for _ in _sc_criteria
                )
                _sc_item_rows += (
                    f'<tr class="sc-item-row">'
                    f'<td class="sc-item-name-td">'
                    f'<input class="sc-item-name" value="{_esc(_it_name)}" placeholder="対象名" onfocus="this.select()">'
                    f'<button type="button" class="sc-del-item-btn" onclick="scDelItem(this)">✕</button>'
                    f'</td>'
                    f'{_sc_score_cells}'
                    f'<td class="sc-total-td">—</td>'
                    f'</tr>'
                )
            fields_html += (
                f'<div class="hq-item{branch_class}" style="margin:14px 0"{branch_attrs}>'
                f'<label style="font-weight:700;color:#d97706;font-size:13px;margin-bottom:6px;display:block">{label}</label>'
                f'<input type="hidden" name="answer_{i}" id="sc_answer_{i}">'
                f'<div class="sc-wrapper" id="sc_wrapper_{i}" data-sc-idx="{i}">'
                f'<div style="overflow-x:auto">'
                f'<table class="sc-table">'
                f'<thead><tr>'
                f'<th class="sc-corner-h">対象 ↓</th>'
                f'{_sc_crit_ths}'
                f'<th class="sc-total-h">合計</th>'
                f'</tr></thead>'
                f'<tbody id="sc_tbody_{i}">{_sc_item_rows}</tbody>'
                f'</table>'
                f'</div>'
                f'<div style="display:flex;gap:8px;margin-top:8px;flex-wrap:wrap">'
                f'<button type="button" class="btn sec" onclick="scAddItem({i})">＋対象追加</button>'
                f'<button type="button" class="btn sec" style="border-color:#d9770660;color:#d97706"'
                f' onclick="scAddCrit({i})">＋評価軸追加</button>'
                f'</div>'
                f'</div>'
                f'</div>'
            )
        elif it.get("type") == "yabane":
            _yb_pairs = it.get("rows")
            if _yb_pairs is None:
                # 旧形式（departments/steps）のテンプレートからの変換: 1ステップ1行、部署は空欄
                _yb_pairs = [{"step": _s.get("label", ""), "dept": ""}
                            for _s in (it.get("steps") or [{"label": "ステップ1"}])]
            _yb_prev = pv if isinstance(pv, dict) else None
            if _yb_prev is not None:
                _yb_rows = _yabane_rows_from_answer(_yb_prev)
            else:
                # テンプレートで定義した「ステップ・部署」の並びをそのまま行にする
                # （掛け算はしない。1テンプレ行 = 1ヒアリング行）。
                _yb_rows = [
                    {"step": _p.get("step", ""), "dept": _p.get("dept", ""), "content": "", "output": "",
                     "issue": "", "target": "", "target_number": ""}
                    for _p in (_yb_pairs or [{"step": "", "dept": ""}])
                ]
            if not _yb_rows:
                _yb_rows = [{"step": "", "dept": "", "content": "", "output": "",
                            "issue": "", "target": "", "target_number": ""}]

            def _yb_row_html(_r):
                return (
                    f'<tr class="yb-row">'
                    f'<td class="yb-step-cell"><input class="yb-row-step" value="{_esc(_r.get("step",""))}"'
                    f' placeholder="ステップ名"></td>'
                    f'<td class="yb-dept-cell"><input class="yb-row-dept" value="{_esc(_r.get("dept",""))}"'
                    f' placeholder="部署名"></td>'
                    f'<td class="yb-wide-cell"><textarea class="yb-row-content ta-expand" rows="2"'
                    f' onfocus="taExpand(this)" onblur="taShrink(this)"'
                    f' placeholder="作業内容">{_esc(_r.get("content",""))}</textarea></td>'
                    f'<td class="yb-wide-cell"><textarea class="yb-row-output ta-expand" rows="2"'
                    f' onfocus="taExpand(this)" onblur="taShrink(this)"'
                    f' placeholder="アウトプット">{_esc(_r.get("output",""))}</textarea></td>'
                    f'<td class="yb-wide-cell"><textarea class="yb-row-issue ta-expand" rows="2"'
                    f' onfocus="taExpand(this)" onblur="taShrink(this)"'
                    f' placeholder="現行課題">{_esc(_r.get("issue",""))}</textarea></td>'
                    f'<td class="yb-wide-cell"><textarea class="yb-row-target ta-expand" rows="2"'
                    f' onfocus="taExpand(this)" onblur="taShrink(this)"'
                    f' placeholder="目指す姿">{_esc(_r.get("target",""))}</textarea></td>'
                    f'<td class="yb-wide-cell"><textarea class="yb-row-target-number ta-expand" rows="2"'
                    f' onfocus="taExpand(this)" onblur="taShrink(this)"'
                    f' placeholder="目標数値（作業時間等）">{_esc(_r.get("target_number",""))}</textarea></td>'
                    f'<td class="yb-del-cell"><button type="button" class="yb-del-row-btn"'
                    f' onclick="ybDelRow(this)">✕</button></td>'
                    f'</tr>'
                )

            _yb_rows_html = "".join(_yb_row_html(_r) for _r in _yb_rows)
            fields_html += (
                f'<div class="hq-item" style="margin:14px 0">'
                f'<label style="font-weight:700;color:#2f6fed;font-size:13px;margin-bottom:6px;display:block">{label}</label>'
                f'<input type="hidden" name="answer_{i}" id="yb_answer_{i}">'
                f'<div class="yb-wrapper" id="yb_wrapper_{i}" data-yb-idx="{i}">'
                f'<table class="yb-table">'
                f'<thead><tr>'
                f'<th class="yb-step-h">ステップ</th><th class="yb-dept-h">部署</th>'
                f'<th class="yb-wide-h">作業内容</th><th class="yb-wide-h">アウトプット</th>'
                f'<th class="yb-wide-h">現行課題</th><th class="yb-wide-h">目指す姿</th>'
                f'<th class="yb-wide-h">目標数値</th><th class="yb-del-h"></th>'
                f'</tr></thead>'
                f'<tbody id="yb_tbody_{i}">{_yb_rows_html}</tbody>'
                f'</table>'
                f'</div>'
                f'<div style="display:flex;gap:8px;margin-top:8px;flex-wrap:wrap">'
                f'<button type="button" class="btn sec" onclick="ybAddRow({i})">＋行追加</button>'
                f'<button type="button" class="btn sec" style="border-color:#3b82f660;color:#3b82f6"'
                f' onclick="ybAddStepBlock({i})">＋ステップ追加（全部署ぶん）</button>'
                f'<button type="button" class="btn sec" style="border-color:#3b82f660;color:#3b82f6"'
                f' onclick="ybAddDeptBlock({i})">＋部署追加（全ステップぶん）</button>'
                f'</div>'
                f'</div>'
            )
        elif it.get("type") == "choice":
            opts = it.get("options") or []
            if it.get("multi"):
                cur = set(pv if isinstance(pv, list) else ([pv] if pv else []))
                boxes = "".join(
                    f'<label style="display:inline-flex;align-items:center;gap:6px;margin:2px 12px 2px 0;font-weight:400">'
                    f'<input type="checkbox" name="answer_{i}" value="{_esc(o)}"'
                    f'{" checked" if o in cur else ""} style="width:14px;height:14px">{_esc(o)}</label>'
                    for o in opts
                )
                fields_html += (f'<div class="hq-item{branch_class}"{branch_attrs} style="margin:10px 0">'
                                f'<label>{label}{req}</label><div>{boxes}</div></div>')
            else:
                radios = "".join(
                    f'<label style="display:inline-flex;align-items:center;gap:6px;margin:2px 12px 2px 0;font-weight:400">'
                    f'<input type="radio" name="answer_{i}" value="{_esc(o)}"'
                    f'{" checked" if pv == o else ""}{req_attr} style="width:14px;height:14px">{_esc(o)}</label>'
                    for o in opts
                )
                fields_html += (f'<div class="hq-item{branch_class}"{branch_attrs} style="margin:10px 0">'
                                f'<label>{label}{req}</label><div>{radios}</div></div>')
        elif it.get("type") == "number":
            val = _esc(str(pv) if pv is not None else "")
            fields_html += (f'<div class="hq-item{branch_class}"{branch_attrs} style="margin:10px 0">'
                            f'<label>{label}{req}</label>'
                            f'<input type="number" name="answer_{i}" value="{val}" step="any"'
                            f' inputmode="numeric" style="max-width:200px"{req_attr}></div>')
        else:
            val = _esc(pv if isinstance(pv, str) else "")
            fields_html += (f'<div class="hq-item{branch_class}"{branch_attrs} style="margin:10px 0">'
                            f'<label>{label}{req}</label>'
                            f'<textarea name="answer_{i}" rows="2" class="ta-expand"'
                            f' onfocus="taExpand(this)" onblur="taShrink(this)"'
                            f'{req_attr}>{val}</textarea></div>')

    prev_note = (f'<p class="muted" style="font-size:12px;margin:0 0 10px">'
                 f'前回ヒアリング（{_esc(prev_date)}）の内容を引用しています。保存すると新しい履歴として追加されます。</p>'
                 if prefill and prev_date and not draft_data else "")
    draft_note = (
        f'<p class="muted" style="font-size:12px;margin:0 0 10px;color:#b45309">'
        f'⚠ 自動保存された下書き（{_esc((draft_updated_at or "")[:16])} 時点）を復元しました。'
        f'内容を確認し、問題なければそのまま保存してください。'
        f'<br>※レーダーチャート／タイムライン／スコアカード形式の項目は自動保存の対象データは残りますが、'
        f'画面への自動復元は未対応です。該当項目があれば再入力してください。</p>'
        if draft_data else ""
    )
    guide_html = (
        '<div style="position:relative;margin-left:auto;font-size:12px"'
        ' onmouseenter="this.querySelector(\'.hq-guide-popup\').style.display=\'block\'"'
        ' onmouseleave="this.querySelector(\'.hq-guide-popup\').style.display=\'none\'">'
        '<div style="cursor:default;color:#2f6fed;font-weight:600;padding:4px 10px;'
        'background:#e8f0fe;border-radius:6px;white-space:nowrap;user-select:none">📋 入力ガイド</div>'
        '<div class="hq-guide-popup" style="display:none;position:absolute;right:0;top:calc(100% + 4px);'
        'z-index:200;background:#fff;border:1px solid #d0e4ff;border-radius:8px;padding:12px 16px;'
        'width:340px;box-shadow:0 4px 16px rgba(0,0,0,.12);line-height:1.75;color:#3a4760">'
        '<p style="margin:0 0 8px;font-weight:700;color:#2f6fed;font-size:13px">このシートの使い方</p>'
        '<ul style="margin:0;padding-left:18px;font-size:12px">'
        '<li><strong>グレーの項目</strong>は<em>分岐（条件付き）質問</em>です。上の質問の回答によって活性化します。</li>'
        '<li>グレーの項目も<strong>直接入力・選択できます</strong>（入力すると分岐元の回答が自動でセットされます）。</li>'
        '<li>分岐元の回答を変えるとグレー項目の入力は<strong>自動でクリア</strong>されます。</li>'
        '<li><span style="color:#c0392b">*</span> 印は必須項目です。</li>'
        '</ul>'
        '</div>'
        '</div>'
    ) if has_branch else ""
    return f"""
    <style>
    .hq-branch {{ transition: opacity .25s, filter .25s; }}
    .hq-branch.hq-inactive {{ opacity: .38; filter: grayscale(.35); }}
    .hq-branch.hq-inactive > label:first-child {{ color: #aaa; }}
    .hq-sticky {{
      position: sticky; top: 0; z-index: 50;
      background: #fff; border-bottom: 1px solid #e2e6ee;
      box-shadow: 0 2px 6px rgba(0,0,0,.06);
      padding: 10px 20px 10px; margin: -20px -16px 16px;
    }}
    .hq-sticky-inner {{
      max-width: 1100px; margin: 0 auto;
      display: flex; align-items: center; gap: 12px; flex-wrap: wrap;
      position: relative;
    }}
    /* ── レーダーチャート ── */
    .ra-wrapper{{margin-top:4px}}
    .ra-axis-row{{display:flex;align-items:center;gap:10px;padding:6px 0;border-bottom:1px solid #e8ecf0}}
    .ra-axis-label-inp{{border:none;background:transparent;color:#1e293b;font-weight:600;font-size:13px;width:120px;padding:2px 4px;outline:none;cursor:text;font-family:inherit}}
    .ra-axis-label-inp:focus{{background:rgba(37,99,235,.07);border-radius:3px}}
    .ra-score-range{{flex:1;accent-color:#2563eb;cursor:pointer;height:4px}}
    .ra-score-val{{font-weight:700;font-size:14px;color:#2563eb;min-width:18px;text-align:right}}
    .ra-svg text{{font-family:inherit}}
    /* ── タイムライン ── */
    .tl-wrapper{{margin-top:4px}}
    .tl-track{{position:relative;padding-left:24px}}
    .tl-track::before{{content:'';position:absolute;left:8px;top:0;bottom:0;width:2px;background:#d1fae5}}
    .tl-event{{position:relative;margin-bottom:12px}}
    .tl-dot{{position:absolute;left:-20px;top:14px;width:10px;height:10px;border-radius:50%;background:#059669;border:2px solid #fff;box-shadow:0 0 0 2px #059669}}
    .tl-event-body{{background:#f0fdf4;border:1px solid #bbf7d0;border-radius:6px;padding:8px 10px}}
    .tl-event-header{{display:flex;align-items:center;gap:8px;margin-bottom:6px;flex-wrap:wrap}}
    .tl-date{{border:1px solid #d4dae4;border-radius:4px;padding:4px 6px;font-size:12px;color:#1e293b;font-family:inherit}}
    .tl-event-label{{flex:1;min-width:120px;border:none;background:transparent;font-weight:600;font-size:13px;color:#065f46;outline:none;cursor:text;padding:2px 4px}}
    .tl-event-label:focus{{background:rgba(5,150,105,.07);border-radius:3px}}
    .tl-del-btn{{font-size:10px;padding:2px 6px;background:#fde8e8;color:#c0392b;border:none;cursor:pointer;border-radius:3px;flex-shrink:0}}
    .tl-note{{width:100%;resize:vertical;font-size:12px;border:1px solid #bbf7d0;border-radius:4px;padding:5px;color:#1e293b;background:#fff;font-family:inherit}}
    /* ── スコアカード ── */
    .sc-wrapper{{overflow-x:auto;margin-top:4px}}
    .sc-table{{border-collapse:collapse;min-width:300px}}
    .sc-corner-h{{background:#fef3c7;color:#92400e;font-weight:700;padding:8px 10px;border:1px solid #fde68a;text-align:center;font-size:11px;white-space:nowrap;min-width:100px}}
    .sc-crit-h{{background:#fef3c7;border:1px solid #fde68a;padding:5px 6px;text-align:center;vertical-align:middle;min-width:90px}}
    .sc-crit-name{{border:none;background:transparent;color:#92400e;font-weight:700;font-size:12px;text-align:center;width:calc(100% - 20px);padding:2px 4px;outline:none;cursor:text;font-family:inherit}}
    .sc-crit-name:focus{{background:rgba(146,64,14,.07);border-radius:3px}}
    .sc-del-crit-btn{{font-size:10px;padding:1px 4px;background:#fde8e8;color:#c0392b;border:none;cursor:pointer;border-radius:2px;margin-left:4px;vertical-align:middle}}
    .sc-total-h{{background:#fef3c7;border:1px solid #fde68a;padding:6px;text-align:center;font-size:11px;color:#92400e;font-weight:700;white-space:nowrap;min-width:44px}}
    .sc-item-name-td{{padding:4px 8px;border:1px solid #fde68a;vertical-align:middle;background:#fffbeb;white-space:nowrap}}
    .sc-item-name{{border:none;background:transparent;color:#d97706;font-weight:600;font-size:13px;padding:2px 4px;outline:none;cursor:text;font-family:inherit;width:calc(100% - 20px)}}
    .sc-item-name:focus{{background:rgba(217,119,6,.07);border-radius:3px}}
    .sc-del-item-btn{{font-size:10px;padding:1px 4px;background:#fde8e8;color:#c0392b;border:none;cursor:pointer;border-radius:2px;margin-left:4px;vertical-align:middle}}
    .sc-score-td{{padding:4px;border:1px solid #fde68a;text-align:center;background:#fff}}
    .sc-score-inp{{width:52px;text-align:center;border:1px solid #d4dae4;border-radius:4px;padding:4px;font-size:13px;font-family:inherit}}
    .sc-total-td{{padding:4px 8px;border:1px solid #fde68a;text-align:center;font-weight:700;font-size:13px;color:#d97706;background:#fffbeb;white-space:nowrap}}
    /* ── 矢羽（1業務1行のフラット表。ステップ列は横スクロール時に固定） ── */
    .yb-wrapper{{overflow-x:auto;margin-top:4px;max-width:100%}}
    .yb-table{{border-collapse:collapse;min-width:900px}}
    .yb-step-h,.yb-dept-h,.yb-wide-h,.yb-del-h{{
      background:#e8eeff;color:#3730a3;font-weight:700;padding:8px 10px;
      border:1px solid #c7d2fe;text-align:center;font-size:11px;white-space:nowrap;
    }}
    .yb-step-h{{position:sticky;left:0;z-index:2;min-width:120px}}
    .yb-dept-h{{min-width:120px}}
    .yb-wide-h{{min-width:240px}}
    .yb-del-h{{min-width:32px}}
    .yb-step-cell{{position:sticky;left:0;z-index:1;background:#eef2ff;border:1px solid #c7d2fe;padding:5px;vertical-align:top}}
    .yb-dept-cell{{border:1px solid #dde4f0;padding:5px;vertical-align:top;background:#fff;min-width:120px}}
    .yb-wide-cell{{border:1px solid #dde4f0;padding:5px;vertical-align:top;background:#fff;min-width:240px}}
    .yb-del-cell{{border:1px solid #dde4f0;padding:0;text-align:center;vertical-align:middle}}
    .yb-row-step,.yb-row-dept{{width:100%;border:1px solid #d4dae4;border-radius:4px;padding:6px;font-size:13px;font-family:inherit}}
    .yb-row-step{{font-weight:700;color:#3730a3}}
    .yb-wide-cell textarea{{width:100%;min-height:54px;resize:vertical;font-size:13px;margin:0;border:1px solid #d4dae4;border-radius:4px;padding:6px;color:#1e293b;font-family:inherit}}
    .yb-del-row-btn{{font-size:11px;padding:6px 8px;background:#fde8e8;color:#c0392b;border:none;cursor:pointer;height:100%}}
    </style>
    <div class="hq-sticky">
      <div class="hq-sticky-inner">
        <div style="flex:1;min-width:0">
          <div style="font-size:15px;font-weight:700;color:#1d2430;margin-bottom:2px">ヒアリング入力</div>
          <div style="font-size:12px;color:#6b7689;white-space:nowrap;overflow:hidden;text-overflow:ellipsis">
            <strong style="color:#3a4760">対象:</strong> {_esc(target_label)}
            <span style="margin:0 6px;color:#d4dae4">|</span>
            <strong style="color:#3a4760">テンプレート:</strong> {_esc(template.get('name', ''))}
          </div>
        </div>
        <button class="btn" type="submit" form="hearing_form" style="white-space:nowrap;align-self:center">💾 保存</button>
        {guide_html}
      </div>
    </div>
    <div class="card" style="max-width:1100px">
      {prev_note}
      {draft_note}
      <p id="hq_autosave_status" class="muted" style="font-size:11px;margin:0 0 8px"></p>
      <form method="post" action="/hearing/submit" id="hearing_form">
        <input type="hidden" name="target_type" value="{_esc(target_type)}">
        <input type="hidden" name="target_id" value="{_esc(target_id)}">
        <input type="hidden" name="template_id" value="{template['id']}">
        {f'<input type="hidden" name="edit_result_id" value="{edit_result_id}">' if edit_result_id else ""}

        <div style="background:#f0f6ff;border-radius:8px;padding:14px 16px;margin-bottom:16px">
          <p style="margin:0 0 6px;font-weight:600;color:#2f6fed">ヒアリング項目</p>
          {fields_html or '<p class="muted">このテンプレートには項目がありません。</p>'}
        </div>

        {f'''<div style="border:1px solid #e2e6ee;border-radius:8px;padding:14px 16px">
          <p style="margin:0 0 8px;font-weight:600;color:#555">ヒアリング実施日</p>
          <input type="date" name="occurred_on" value="{_esc(edit_conducted_on)}" required style="max-width:200px">
        </div>
        <div style="margin-top:16px"><button class="btn" type="submit">更新（この内容でヒアリング結果を修正）</button>
        <a class="btn sec" href="/hearing/result/{edit_result_id}">キャンセル</a></div>''' if edit_result_id else f'''
        <div style="border:1px solid #e2e6ee;border-radius:8px;padding:14px 16px">
          <p style="margin:0 0 8px;font-weight:600;color:#555">活動履歴として記録</p>
          <div class="grid">
            <div><label>ヒアリング日</label><input type="date" name="occurred_on" required></div>
            <div><label>種別</label><select name="type">{_opt(sfa_db.get_master_list(con,'activity_types'), '面談')}</select></div>
            <div><label>相手</label><input name="contact_name" placeholder="例：田中部長"></div>
          </div>
          <label>内容・決定事項</label><textarea name="body" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="3"></textarea>
          <div style="margin-top:10px;padding:12px;background:#f8f9fa;border-radius:6px">
            <p style="margin:0 0 8px;font-size:.9em;font-weight:600;color:#555">商談の現状を更新（任意）</p>
            <div class="grid">
              <div><label>次回MS日</label><input type="date" name="next_milestone_date"></div>
              <div><label>次回MSラベル</label><input name="next_milestone_label"></div>
              <div><label>次回MS種別</label><select name="next_milestone_type">{_opt(sfa_db.NEXT_MS_TYPES, 'アポ')}</select></div>
            </div>
            <label>現状メモ</label><textarea name="update_note" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="2"></textarea>
          </div>
        </div>
        <div style="margin-top:16px"><button class="btn" type="submit">保存（活動履歴＋ヒアリング結果を記録）</button>
        <a class="btn sec" href="/hearings">キャンセル</a></div>'''}
      </form>
    </div>
    <script>
    // ページ共通ヘッダー(position:sticky)の実高さ分だけ、このページ内のstickyバーを下にずらす
    (function() {{
      var pageHeader = document.querySelector('header');
      var stickyBar = document.querySelector('.hq-sticky');
      if (pageHeader && stickyBar) {{
        stickyBar.style.top = pageHeader.getBoundingClientRect().height + 'px';
      }}
    }})();
    // ── レーダーチャート JS ──
    function _svgEsc(s){{return String(s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');}}
    function raRedraw(wrapper) {{
      var rows=Array.from(wrapper.querySelectorAll('.ra-axis-row'));
      var n=rows.length; if(n<3) return;
      var svg=wrapper.querySelector('.ra-svg');
      var cx=100,cy=100,R=72;
      var html='';
      for(var g=1;g<=5;g++){{
        var gr=R*g/5;
        html+='<circle cx="'+cx+'" cy="'+cy+'" r="'+gr+'" fill="none" stroke="#e2e8f0" stroke-width="1"/>';
      }}
      var angles=[],scores=[];
      for(var i=0;i<n;i++){{
        angles.push(-Math.PI/2+i*2*Math.PI/n);
        scores.push(parseFloat(rows[i].querySelector('.ra-score-range').value)||0);
      }}
      for(var i=0;i<n;i++){{
        var ax=cx+R*Math.cos(angles[i]), ay=cy+R*Math.sin(angles[i]);
        html+='<line x1="'+cx+'" y1="'+cy+'" x2="'+ax+'" y2="'+ay+'" stroke="#cbd5e1" stroke-width="1"/>';
        var lx=cx+(R+15)*Math.cos(angles[i]), ly=cy+(R+15)*Math.sin(angles[i]);
        var anc=lx<cx-8?'end':lx>cx+8?'start':'middle';
        var lbl=rows[i].querySelector('.ra-axis-label-inp')?rows[i].querySelector('.ra-axis-label-inp').value:'';
        html+='<text x="'+lx+'" y="'+(ly+4)+'" text-anchor="'+anc+'" font-size="10" fill="#475569">'+_svgEsc(lbl)+'</text>';
      }}
      var poly='';
      for(var i=0;i<n;i++){{
        var pr=R*scores[i]/5;
        poly+=(cx+pr*Math.cos(angles[i]))+','+(cy+pr*Math.sin(angles[i]))+' ';
      }}
      html+='<polygon points="'+poly.trim()+'" fill="rgba(124,58,237,.2)" stroke="#7c3aed" stroke-width="2"/>';
      for(var i=0;i<n;i++){{
        var pr=R*scores[i]/5;
        html+='<circle cx="'+(cx+pr*Math.cos(angles[i]))+'" cy="'+(cy+pr*Math.sin(angles[i]))+'" r="4" fill="#7c3aed"/>';
      }}
      svg.innerHTML=html;
    }}
    function raUpdateScore(rangeEl) {{
      rangeEl.nextElementSibling.textContent=rangeEl.value;
      raRedraw(rangeEl.closest('[data-ra-idx]'));
    }}
    // ── タイムライン JS ──
    function tlDelEvent(btn){{btn.closest('.tl-event').remove();}}
    function tlAddEvent(idx){{
      var track=document.getElementById('tl_track_'+idx);
      var ev=document.createElement('div'); ev.className='tl-event';
      ev.innerHTML='<div class="tl-dot"></div>'
        +'<div class="tl-event-body">'
        +'<div class="tl-event-header">'
        +'<input type="month" class="tl-date">'
        +'<input class="tl-event-label" placeholder="マイルストーン名" onfocus="this.select()">'
        +'<button type="button" class="tl-del-btn" onclick="tlDelEvent(this)">✕</button>'
        +'</div>'
        +'<textarea class="tl-note ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="2" placeholder="詳細・メモ"></textarea>'
        +'</div>';
      track.appendChild(ev);
    }}
    // ── スコアカード JS ──
    function scUpdateTotal(inp){{
      var row=inp.closest('.sc-item-row');
      var sum=0;
      row.querySelectorAll('.sc-score-inp').forEach(function(i){{var v=parseFloat(i.value); if(!isNaN(v))sum+=v;}});
      row.querySelector('.sc-total-td').textContent=sum||'—';
    }}
    function scDelItem(btn){{btn.closest('.sc-item-row').remove();}}
    function scDelCrit(btn){{
      var th=btn.closest('.sc-crit-h');
      var table=th.closest('table');
      var colIdx=Array.from(th.closest('tr').querySelectorAll('.sc-crit-h')).indexOf(th);
      th.remove();
      table.querySelectorAll('.sc-item-row').forEach(function(tr){{
        var cells=tr.querySelectorAll('.sc-score-td');
        if(cells[colIdx]) cells[colIdx].remove();
      }});
    }}
    function scAddItem(idx){{
      var wrapper=document.getElementById('sc_wrapper_'+idx);
      var nCols=wrapper.querySelectorAll('thead .sc-crit-h').length;
      var tr=document.createElement('tr'); tr.className='sc-item-row';
      var ntd=document.createElement('td'); ntd.className='sc-item-name-td';
      var ni=document.createElement('input'); ni.className='sc-item-name'; ni.placeholder='対象名'; ni.onfocus=function(){{this.select();}};
      var db=document.createElement('button'); db.type='button'; db.className='sc-del-item-btn';
      db.textContent='✕'; db.onclick=function(){{scDelItem(this);}};
      ntd.appendChild(ni); ntd.appendChild(db); tr.appendChild(ntd);
      for(var c=0;c<nCols;c++){{
        var td=document.createElement('td'); td.className='sc-score-td';
        var inp=document.createElement('input'); inp.type='number'; inp.className='sc-score-inp';
        inp.min=1; inp.max=5; inp.placeholder='—'; inp.oninput=function(){{scUpdateTotal(this);}};
        td.appendChild(inp); tr.appendChild(td);
      }}
      var ttd=document.createElement('td'); ttd.className='sc-total-td'; ttd.textContent='—';
      tr.appendChild(ttd);
      wrapper.querySelector('tbody').appendChild(tr);
    }}
    function scAddCrit(idx){{
      var wrapper=document.getElementById('sc_wrapper_'+idx);
      var hdrRow=wrapper.querySelector('thead tr');
      var totalTh=hdrRow.querySelector('.sc-total-h');
      var th=document.createElement('th'); th.className='sc-crit-h';
      var inp=document.createElement('input'); inp.className='sc-crit-name'; inp.placeholder='評価軸'; inp.onfocus=function(){{this.select();}};
      var db=document.createElement('button'); db.type='button'; db.className='sc-del-crit-btn';
      db.textContent='✕'; db.onclick=function(){{scDelCrit(this);}};
      th.appendChild(inp); th.appendChild(db); hdrRow.insertBefore(th,totalTh);
      wrapper.querySelectorAll('.sc-item-row').forEach(function(tr){{
        var totalTd=tr.querySelector('.sc-total-td');
        var td=document.createElement('td'); td.className='sc-score-td';
        var sinp=document.createElement('input'); sinp.type='number'; sinp.className='sc-score-inp';
        sinp.min=1; sinp.max=5; sinp.placeholder='—'; sinp.oninput=function(){{scUpdateTotal(this);}};
        td.appendChild(sinp); tr.insertBefore(td,totalTd);
      }});
    }}
    // ── 矢羽入力 JS（1業務1行のフラット表） ──
    function ybDelRow(btn) {{ btn.closest('.yb-row').remove(); }}
    function _ybRowHtml(step, dept) {{
      function esc(s) {{ return (s||'').replace(/"/g,'&quot;'); }}
      return '<tr class="yb-row">' +
        '<td class="yb-step-cell"><input class="yb-row-step" value="'+esc(step)+'" placeholder="ステップ名"></td>' +
        '<td class="yb-dept-cell"><input class="yb-row-dept" value="'+esc(dept)+'" placeholder="部署名"></td>' +
        '<td class="yb-wide-cell"><textarea class="yb-row-content ta-expand" rows="2" onfocus="taExpand(this)" onblur="taShrink(this)" placeholder="作業内容"></textarea></td>' +
        '<td class="yb-wide-cell"><textarea class="yb-row-output ta-expand" rows="2" onfocus="taExpand(this)" onblur="taShrink(this)" placeholder="アウトプット"></textarea></td>' +
        '<td class="yb-wide-cell"><textarea class="yb-row-issue ta-expand" rows="2" onfocus="taExpand(this)" onblur="taShrink(this)" placeholder="現行課題"></textarea></td>' +
        '<td class="yb-wide-cell"><textarea class="yb-row-target ta-expand" rows="2" onfocus="taExpand(this)" onblur="taShrink(this)" placeholder="目指す姿"></textarea></td>' +
        '<td class="yb-wide-cell"><textarea class="yb-row-target-number ta-expand" rows="2" onfocus="taExpand(this)" onblur="taShrink(this)" placeholder="目標数値（作業時間等）"></textarea></td>' +
        '<td class="yb-del-cell"><button type="button" class="yb-del-row-btn" onclick="ybDelRow(this)">✕</button></td>' +
      '</tr>';
    }}
    function ybAddRow(idx, step, dept) {{
      var tbody=document.getElementById('yb_tbody_'+idx);
      tbody.insertAdjacentHTML('beforeend', _ybRowHtml(step||'', dept||''));
    }}
    function ybAddStepBlock(idx) {{
      var tbody=document.getElementById('yb_tbody_'+idx);
      var depts=Array.from(new Set(Array.from(tbody.querySelectorAll('.yb-row-dept'))
        .map(function(i){{return i.value.trim();}}).filter(Boolean)));
      if (!depts.length) depts=[''];
      depts.forEach(function(d){{ ybAddRow(idx, '', d); }});
    }}
    function ybAddDeptBlock(idx) {{
      var tbody=document.getElementById('yb_tbody_'+idx);
      var steps=Array.from(new Set(Array.from(tbody.querySelectorAll('.yb-row-step'))
        .map(function(i){{return i.value.trim();}}).filter(Boolean)));
      if (!steps.length) steps=[''];
      steps.forEach(function(s){{ ybAddRow(idx, s, ''); }});
    }}
    // 各特殊項目のUI状態を隠しinput(answer_N)へ直列化する。
    // submit時・30秒ごとの自動保存時の両方から呼ぶ。
    function serializeAllAnswers() {{
      // レーダーチャート直列化
      document.querySelectorAll('[data-ra-idx]').forEach(function(wrapper) {{
        var idx=wrapper.getAttribute('data-ra-idx');
        var axes=[];
        wrapper.querySelectorAll('.ra-axis-row').forEach(function(row) {{
          axes.push({{
            label:row.querySelector('.ra-axis-label-inp').value.trim(),
            score:parseFloat(row.querySelector('.ra-score-range').value)||0
          }});
        }});
        var h=document.getElementById('ra_answer_'+idx);
        if(h) h.value=JSON.stringify({{axes:axes}});
      }});
      // タイムライン直列化
      document.querySelectorAll('[data-tl-idx]').forEach(function(wrapper) {{
        var idx=wrapper.getAttribute('data-tl-idx');
        var events=[];
        wrapper.querySelectorAll('.tl-event').forEach(function(ev) {{
          events.push({{
            label:ev.querySelector('.tl-event-label').value.trim(),
            date:ev.querySelector('.tl-date').value,
            note:ev.querySelector('.tl-note').value.trim()
          }});
        }});
        var h=document.getElementById('tl_answer_'+idx);
        if(h) h.value=JSON.stringify({{events:events}});
      }});
      // スコアカード直列化
      document.querySelectorAll('[data-sc-idx]').forEach(function(wrapper) {{
        var idx=wrapper.getAttribute('data-sc-idx');
        var crits=Array.from(wrapper.querySelectorAll('thead .sc-crit-name')).map(function(i){{return i.value.trim();}});
        var items=[];
        wrapper.querySelectorAll('.sc-item-row').forEach(function(tr) {{
          var name=tr.querySelector('.sc-item-name').value.trim();
          var scores={{}};
          tr.querySelectorAll('.sc-score-inp').forEach(function(inp,di) {{
            if(crits[di]!==undefined){{var v=parseFloat(inp.value); scores[crits[di]]=isNaN(v)?null:v;}}
          }});
          items.push({{label:name,scores:scores}});
        }});
        var h=document.getElementById('sc_answer_'+idx);
        if(h) h.value=JSON.stringify({{criteria:crits,items:items}});
      }});
      document.querySelectorAll('[data-yb-idx]').forEach(function(wrapper) {{
        var idx=wrapper.getAttribute('data-yb-idx');
        var rows=[];
        wrapper.querySelectorAll('.yb-row').forEach(function(tr) {{
          rows.push({{
            step: tr.querySelector('.yb-row-step').value.trim(),
            dept: tr.querySelector('.yb-row-dept').value.trim(),
            content: tr.querySelector('.yb-row-content').value,
            output: tr.querySelector('.yb-row-output').value,
            issue: tr.querySelector('.yb-row-issue').value,
            target: tr.querySelector('.yb-row-target').value,
            target_number: tr.querySelector('.yb-row-target-number').value
          }});
        }});
        var hidden=document.getElementById('yb_answer_'+idx);
        if(hidden) hidden.value=JSON.stringify({{rows:rows}});
      }});
    }}
    document.getElementById('hearing_form').addEventListener('submit', function() {{
      serializeAllAnswers();
      if (window._hqAutosaveTimer) clearInterval(window._hqAutosaveTimer);
    }});
    // ── 30秒ごとの自動保存（下書き） ──
    (function() {{
      var form = document.getElementById('hearing_form');
      var statusEl = document.getElementById('hq_autosave_status');
      function autoSave() {{
        serializeAllAnswers();
        var data = new URLSearchParams(new FormData(form));
        fetch('/hearing/autosave', {{
          method: 'POST',
          headers: {{'Content-Type': 'application/x-www-form-urlencoded'}},
          body: data.toString()
        }}).then(function(r) {{ return r.json(); }}).then(function(res) {{
          if (statusEl) {{
            var now = new Date();
            var hh = String(now.getHours()).padStart(2,'0');
            var mm = String(now.getMinutes()).padStart(2,'0');
            statusEl.textContent = res.ok ? ('下書きを自動保存しました（' + hh + ':' + mm + '）') : '自動保存に失敗しました';
          }}
        }}).catch(function() {{
          if (statusEl) statusEl.textContent = '自動保存に失敗しました（通信エラー）';
        }});
      }}
      window._hqAutosaveTimer = setInterval(autoSave, 30000);
    }})();
    // レーダーチャート初期描画
    document.querySelectorAll('[data-ra-idx]').forEach(function(w){{raRedraw(w);}});
    // ── 分岐ロジック ──
    (function() {{
      var form = document.getElementById('hearing_form');
      if (!form) return;
      var _initialLoad = true;

      function getParentValues(parentIdx) {{
        // 複数選択（チェックボックス）では複数チェックされ得るため、配列で全件返す
        // （以前はchecked[0]のみ見ており、2つ目以降の親子関係が機能しなかった）。
        var name = 'answer_' + parentIdx;
        var checked = form.querySelectorAll('[name="' + name + '"]:checked');
        if (checked.length) return Array.from(checked).map(function(c) {{ return c.value; }});
        // radio/checkbox は :checked がなければ未回答
        var firstInp = form.querySelector('[name="' + name + '"]');
        if (firstInp && (firstInp.type === 'radio' || firstInp.type === 'checkbox')) return [];
        // textarea / number input
        var inp = form.querySelector('textarea[name="' + name + '"],input[type="number"][name="' + name + '"]');
        return inp ? [inp.value.trim()] : [];
      }}

      function updateBranch() {{
        form.querySelectorAll('.hq-branch').forEach(function(div) {{
          var pIdx = div.dataset.parentIdx;
          var pVal = div.dataset.parentValue;
          if (pIdx === undefined || pVal === undefined) return;
          var shouldInactive = getParentValues(parseInt(pIdx)).indexOf(pVal) === -1;
          div.classList.toggle('hq-inactive', shouldInactive);
        }});
        _initialLoad = false;
      }}

      function autoSetParent(changedEl) {{
        var branchDiv = changedEl.closest('.hq-branch');
        if (!branchDiv) return;
        var pIdx = branchDiv.dataset.parentIdx;
        var pVal = branchDiv.dataset.parentValue;
        if (pIdx === undefined || pVal === undefined) return;
        var parentName = 'answer_' + pIdx;
        var parentInput = form.querySelector('[name="' + parentName + '"][value="' + CSS.escape(pVal) + '"]');
        if (parentInput && (parentInput.type === 'radio' || parentInput.type === 'checkbox') && !parentInput.checked) {{
          parentInput.checked = true;
          updateBranch();
        }}
      }}

      form.addEventListener('change', function(e) {{
        updateBranch();
        if (e.target.type === 'radio' || e.target.type === 'checkbox') autoSetParent(e.target);
      }});
      form.addEventListener('input', function(e) {{
        if (e.target.tagName === 'TEXTAREA' || e.target.type === 'number') {{
          updateBranch();
          autoSetParent(e.target);
        }}
      }});

      updateBranch();
    }})();
    </script>"""


def _format_answer(ans) -> str:
    if isinstance(ans, list):
        return "、".join(str(a) for a in ans)
    return str(ans) if ans is not None else ""


def _yabane_rows_from_answer(ans: dict) -> list:
    """矢羽回答から行リストを取得する。新形式(rows)優先、旧形式(departments/steps[].cells)は変換。"""
    rows = ans.get("rows")
    if rows is not None:
        return rows
    converted = []
    for s in (ans.get("steps") or []):
        for d in (ans.get("departments") or []):
            content = (s.get("cells") or {}).get(d, "")
            if content:
                converted.append({"step": s.get("label", ""), "dept": d, "content": content})
    return converted


def _format_answer_for_export(a: dict) -> str:
    """一覧プレビュー・CSV/xlsx出力向け: 矢羽など特殊タイプも読める文字列に整形する。"""
    ans = a.get("answer")
    if a.get("type") == "yabane" and isinstance(ans, dict):
        parts = []
        for r in _yabane_rows_from_answer(ans):
            frags = [f"{r.get('step','')}/{r.get('dept','')}"]
            for key, fkey in (("content", "作業内容"), ("output", "アウトプット"), ("issue", "現行課題"),
                              ("target", "目指す姿"), ("target_number", "目標数値")):
                v = r.get(key)
                if v:
                    frags.append(f"{fkey}:{v}")
            parts.append(" ".join(frags))
        return "；".join(parts)
    return _format_answer(ans)


def _radar_result_html(ra: dict) -> str:
    axes = ra.get("axes") or []
    if not axes:
        return '<p class="muted">データなし</p>'
    n = len(axes)
    import math
    cx, cy, R = 100, 100, 72
    circles = "".join(
        f'<circle cx="{cx}" cy="{cy}" r="{R*g//5}" fill="none" stroke="#e2e8f0" stroke-width="1"/>'
        for g in range(1, 6)
    )
    axis_lines = ""
    poly_pts = ""
    dots = ""
    for i, ax in enumerate(axes):
        angle = -math.pi / 2 + i * 2 * math.pi / n
        ax_x = cx + R * math.cos(angle)
        ax_y = cy + R * math.sin(angle)
        axis_lines += f'<line x1="{cx}" y1="{cy}" x2="{ax_x:.1f}" y2="{ax_y:.1f}" stroke="#cbd5e1" stroke-width="1"/>'
        lx = cx + (R + 16) * math.cos(angle)
        ly = cy + (R + 16) * math.sin(angle)
        anc = "end" if lx < cx - 8 else ("start" if lx > cx + 8 else "middle")
        axis_lines += (f'<text x="{lx:.1f}" y="{ly+4:.1f}" text-anchor="{anc}"'
                       f' font-size="10" fill="#475569">{_esc(ax.get("label",""))}</text>')
        score = _num(ax.get("score"))
        pr = R * score / 5
        poly_pts += f'{cx+pr*math.cos(angle):.1f},{cy+pr*math.sin(angle):.1f} '
        dots += (f'<circle cx="{cx+pr*math.cos(angle):.1f}" cy="{cy+pr*math.sin(angle):.1f}"'
                 f' r="4" fill="#7c3aed"/>')
    poly = f'<polygon points="{poly_pts.strip()}" fill="rgba(124,58,237,.2)" stroke="#7c3aed" stroke-width="2"/>'
    svg = (f'<svg width="200" height="200" style="display:block;margin:0 auto">'
           f'{circles}{axis_lines}{poly}{dots}</svg>')
    rows = "".join(
        f'<tr>'
        f'<td style="padding:5px 10px;font-weight:600;font-size:13px;border-bottom:1px solid #ede9fe">{_esc(ax.get("label",""))}</td>'
        f'<td style="padding:5px 10px;border-bottom:1px solid #ede9fe">'
        f'<div style="display:flex;align-items:center;gap:8px">'
        f'<div style="flex:1;height:8px;background:#ede9fe;border-radius:4px;overflow:hidden">'
        f'<div style="width:{_num(ax.get("score"))*20}%;height:100%;background:#7c3aed;border-radius:4px"></div>'
        f'</div>'
        f'<span style="font-weight:700;color:#7c3aed;min-width:20px">{_esc(ax.get("score") or 0)}</span>'
        f'</div>'
        f'</td>'
        f'</tr>'
        for ax in axes
    )
    table = (f'<table style="width:100%;border-collapse:collapse;margin-top:8px;max-width:360px">'
             f'<thead><tr>'
             f'<th style="background:#f5f3ff;color:#7c3aed;padding:5px 10px;font-size:11px;text-align:left;border-bottom:1px solid #ede9fe">軸</th>'
             f'<th style="background:#f5f3ff;color:#7c3aed;padding:5px 10px;font-size:11px;text-align:left;border-bottom:1px solid #ede9fe">スコア（/5）</th>'
             f'</tr></thead><tbody>{rows}</tbody></table>')
    return f'<div style="display:flex;gap:20px;flex-wrap:wrap;align-items:flex-start">{svg}{table}</div>'


def _timeline_result_html(tl: dict) -> str:
    events = tl.get("events") or []
    if not events:
        return '<p class="muted">データなし</p>'
    items_html = ""
    for ev in events:
        date_str = ev.get("date") or ""
        ev_label = _esc(ev.get("label") or "")
        ev_note = _esc(ev.get("note") or "")
        _date_badge = (f'<span style="font-size:11px;font-weight:700;color:#059669;background:#d1fae5;'
                       f'padding:1px 8px;border-radius:3px">{_esc(date_str)}</span>') if date_str else ""
        _note_p = (f'<p style="font-size:12px;color:#374151;margin:0;white-space:pre-wrap">{ev_note}</p>'
                   ) if ev_note else ""
        items_html += (
            f'<div style="position:relative;padding-left:24px;margin-bottom:12px">'
            f'<div style="position:absolute;left:4px;top:6px;width:10px;height:10px;border-radius:50%;background:#059669;border:2px solid #fff;box-shadow:0 0 0 2px #059669"></div>'
            f'<div style="background:#f0fdf4;border:1px solid #bbf7d0;border-radius:6px;padding:8px 12px">'
            f'<div style="display:flex;gap:10px;align-items:center;margin-bottom:4px;flex-wrap:wrap">'
            f'{_date_badge}'
            f'<span style="font-weight:700;font-size:13px;color:#065f46">{ev_label}</span>'
            f'</div>'
            f'{_note_p}'
            f'</div>'
            f'</div>'
        )
    return (f'<div style="position:relative;padding-left:8px">'
            f'<div style="position:absolute;left:12px;top:0;bottom:0;width:2px;background:#d1fae5"></div>'
            f'{items_html}</div>')


def _scorecard_result_html(sc: dict) -> str:
    criteria = sc.get("criteria") or []
    items = sc.get("items") or []
    if not criteria or not items:
        return '<p class="muted">データなし</p>'
    crit_ths = "".join(
        f'<th style="background:#fef3c7;color:#92400e;padding:6px 10px;font-size:11px;font-weight:700;border:1px solid #fde68a;text-align:center">{_esc(c)}</th>'
        for c in criteria
    )
    rows_html = ""
    for it in items:
        scores = it.get("scores") or {}
        total = sum(v for v in scores.values() if isinstance(v, (int, float)))
        score_tds = "".join(
            f'<td style="padding:6px 10px;border:1px solid #fde68a;text-align:center;font-size:13px;font-weight:700;color:#d97706">'
            f'{_esc(scores.get(c, "—"))}</td>'
            for c in criteria
        )
        rows_html += (
            f'<tr>'
            f'<td style="padding:6px 10px;border:1px solid #fde68a;font-weight:600;background:#fffbeb;color:#d97706">{_esc(it.get("label",""))}</td>'
            f'{score_tds}'
            f'<td style="padding:6px 10px;border:1px solid #fde68a;text-align:center;font-weight:700;color:#92400e;background:#fef3c7">{total or "—"}</td>'
            f'</tr>'
        )
    return (f'<div style="overflow-x:auto">'
            f'<table style="border-collapse:collapse;min-width:300px">'
            f'<thead><tr>'
            f'<th style="background:#fef3c7;color:#92400e;padding:6px 10px;font-size:11px;font-weight:700;border:1px solid #fde68a">対象</th>'
            f'{crit_ths}'
            f'<th style="background:#fef3c7;color:#92400e;padding:6px 10px;font-size:11px;font-weight:700;border:1px solid #fde68a;text-align:center">合計</th>'
            f'</tr></thead><tbody>{rows_html}</tbody></table></div>')


def _yabane_result_table(yb: dict) -> str:
    """矢羽回答をフラット表（1業務1行）のHTMLに変換。旧形式（departments/steps[].cells）も自動変換して表示する。"""
    rows = _yabane_rows_from_answer(yb)
    if not rows:
        return '<p class="muted">データなし</p>'
    headers = ["ステップ", "部署", "作業内容", "アウトプット", "現行課題", "目指す姿", "目標数値"]
    ths = "".join(
        f'<th style="{"position:sticky;left:0;z-index:1;" if idx == 0 else ""}'
        f'background:#e8eeff;color:#3730a3;font-weight:700;padding:8px 10px;'
        f'border:1px solid #c7d2fe;white-space:nowrap;font-size:12px">{h}</th>'
        for idx, h in enumerate(headers)
    )
    body_rows = ""
    for r in rows:
        tds = (
            f'<td style="position:sticky;left:0;background:#eef2ff;font-weight:700;color:#3730a3;'
            f'padding:8px 10px;border:1px solid #c7d2fe;white-space:nowrap">{_esc(r.get("step",""))}</td>'
            f'<td style="padding:8px 10px;border:1px solid #dde4f0;white-space:nowrap;'
            f'font-weight:600;color:#3730a3">{_esc(r.get("dept",""))}</td>'
        )
        for key in ("content", "output", "issue", "target", "target_number"):
            tds += (
                f'<td style="padding:8px;border:1px solid #dde4f0;vertical-align:top;'
                f'white-space:pre-wrap;font-size:13px;color:#1e293b;min-width:200px">{_esc(r.get(key, ""))}</td>'
            )
        body_rows += f'<tr>{tds}</tr>'
    return (
        f'<div style="overflow-x:auto;margin-top:8px">'
        f'<table style="border-collapse:collapse;width:100%;min-width:900px">'
        f'<thead><tr>{ths}</tr></thead>'
        f'<tbody>{body_rows}</tbody>'
        f'</table></div>'
    )


def hearing_result_page(con, result: dict) -> str:
    """個別ヒアリング結果の表示（Q&A・矢羽混在対応）。"""
    answers = result.get("answers") or []
    result_html_parts = []
    for a in answers:
        _atype = a.get("type")
        if _atype == "yabane":
            result_html_parts.append(
                f'<div style="margin:16px 0">'
                f'<div style="font-size:12px;font-weight:700;color:#3730a3;margin-bottom:4px;'
                f'padding:4px 0;border-bottom:1px solid #e0e7ff">{_esc(a.get("label") or "業務プロセス")}</div>'
                f'{_yabane_result_table(a.get("answer") or {{}})}'
                f'</div>'
            )
        elif _atype == "radar":
            result_html_parts.append(
                f'<div style="margin:16px 0">'
                f'<div style="font-size:12px;font-weight:700;color:#7c3aed;margin-bottom:8px;'
                f'padding:4px 0;border-bottom:1px solid #ede9fe">{_esc(a.get("label") or "レーダーチャート")}</div>'
                f'{_radar_result_html(a.get("answer") or {{}})}'
                f'</div>'
            )
        elif _atype == "timeline":
            result_html_parts.append(
                f'<div style="margin:16px 0">'
                f'<div style="font-size:12px;font-weight:700;color:#059669;margin-bottom:8px;'
                f'padding:4px 0;border-bottom:1px solid #bbf7d0">{_esc(a.get("label") or "タイムライン")}</div>'
                f'{_timeline_result_html(a.get("answer") or {{}})}'
                f'</div>'
            )
        elif _atype == "scorecard":
            result_html_parts.append(
                f'<div style="margin:16px 0">'
                f'<div style="font-size:12px;font-weight:700;color:#d97706;margin-bottom:8px;'
                f'padding:4px 0;border-bottom:1px solid #fde68a">{_esc(a.get("label") or "スコアカード")}</div>'
                f'{_scorecard_result_html(a.get("answer") or {{}})}'
                f'</div>'
            )
        else:
            result_html_parts.append(
                f'<table style="width:100%;border-collapse:collapse;margin:10px 0">'
                f'<tr><td style="white-space:nowrap;font-weight:600;vertical-align:top;'
                f'padding:7px 10px;width:30%;border-bottom:1px solid #f0f4ff">{_esc(a.get("label"))}</td>'
                f'<td style="white-space:pre-wrap;padding:7px 10px;border-bottom:1px solid #f0f4ff">'
                f'{_esc(_format_answer(a.get("answer")))}</td></tr>'
                f'</table>'
            )
    if not result_html_parts:
        result_html_parts = ['<p class="muted">回答なし</p>']
    result_html = "\n".join(result_html_parts)
    other = sfa_db.list_hearing_results(con, result["deal_id"])
    history = ""
    if len(other) > 1:
        links = "".join(
            f'<a href="/hearing/result/{o["id"]}" style="margin-right:10px;font-size:12px'
            f'{";font-weight:700" if o["id"]==result["id"] else ""}">'
            f'{_esc(o.get("conducted_on") or "?")}（{_esc(o.get("template_name") or "")}）</a>'
            for o in other
        )
        history = f'<p class="muted" style="margin-top:12px;font-size:12px">この商談のヒアリング履歴: {links}</p>'
    return f"""
    <div class="card" style="max-width:1100px">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>ヒアリング結果</span>
        <span style="display:flex;gap:8px;flex-wrap:wrap">
          <a class="btn sec" href="/hearing/result/{result['id']}/export.xlsx">📥 xlsxダウンロード</a>
          <a class="btn sec" href="/hearing/result/{result['id']}/export.docx">📥 docxダウンロード</a>
        </span>
      </h2>
      <p style="margin:0 0 4px"><strong>商談:</strong> <a href="/deal/{result['deal_id']}">{_esc(result.get('account_name') or '')} / {_esc(result.get('deal_name') or '')}</a></p>
      <p class="muted" style="margin:0 0 12px"><strong>テンプレート:</strong> {_esc(result.get('template_name') or '')}　<strong>ヒアリング日:</strong> {_esc(result.get('conducted_on') or '—')}</p>
      {result_html}
      {history}
      <div style="margin-top:16px;display:flex;gap:8px;flex-wrap:wrap">
        <a class="btn sec" href="/deal/{result['deal_id']}">商談へ戻る</a>
        <a class="btn sec" href="/hearings">ヒアリング一覧</a>
        {f'<a class="btn sec" href="/hearing/result/{result["id"]}/edit">編集</a>' if result.get('template_id') else ''}
        <form method="post" action="/hearing/result/{result['id']}/delete" style="display:inline;margin:0">
          <button class="btn" style="background:#c53030;border-color:#c53030;color:#fff"
            onclick="return confirm('このヒアリング結果を削除しますか？この操作は取り消せません。')">削除</button>
        </form>
      </div>
    </div>"""


def hearings_page(con, template_id: int | None = None) -> str:
    """ヒアリングタブ：実施済み一覧 + テンプレート絞り込み + xlsx/docxダウンロード。"""
    templates = sfa_db.list_hearing_templates(con)
    results = sfa_db.list_all_hearing_results(con, template_id=template_id)
    tmpl = sfa_db.get_hearing_template(con, template_id) if template_id else None
    tmpl_opts = '<option value="">全テンプレート</option>' + "".join(
        f'<option value="{t["id"]}"{" selected" if template_id == t["id"] else ""}>{_esc(t["name"])}</option>'
        for t in templates
    )
    tmpl_filter_form = f"""<form method="get" action="/hearings" class="filter-row">
      <select name="template_id" onchange="this.form.submit()">{tmpl_opts}</select>
    </form>"""
    _ellip = "overflow:hidden;text-overflow:ellipsis;white-space:nowrap"
    rows = ""
    for r in results:
        preview_full = "　".join(
            f'{a.get("label")}: {_format_answer_for_export(a)}'
            for a in (r.get("answers") or [])[:2]
        )
        preview = preview_full[:160] + ("…" if len(preview_full) > 160 else "")
        _nav = f"location.href='/hearing/result/{r['id']}'"
        rows += (
            f'<tr>'
            f'<td style="width:32px"><input type="checkbox" name="ids" value="{r["id"]}"></td>'
            f'<td style="width:90px;cursor:pointer" onclick="{_nav}">{_esc(r.get("conducted_on") or "—")}</td>'
            f'<td style="width:160px;{_ellip};cursor:pointer" title="{_esc(r.get("account_name") or "")}">'
            f'<a href="/deal/{r["deal_id"]}" onclick="event.stopPropagation()">{_esc(r.get("account_name") or "")}</a></td>'
            f'<td style="width:200px;{_ellip};cursor:pointer" onclick="{_nav}" title="{_esc(r.get("deal_name") or "")}">{_esc(r.get("deal_name") or "")}</td>'
            f'<td style="width:150px;{_ellip};cursor:pointer" onclick="{_nav}" title="{_esc(r.get("template_name") or "")}">{_esc(r.get("template_name") or "")}</td>'
            f'<td class="muted" style="font-size:12px;cursor:pointer" onclick="{_nav}" title="{_esc(preview_full)}">{_esc(preview)}</td>'
            f'</tr>'
        )
    dl_qs = f"?template_id={template_id}" if template_id else ""
    header_actions = []
    if results:
        header_actions.append(f'<a class="btn sec" href="/hearings/export{dl_qs}">📥 xlsxダウンロード</a>')
        header_actions.append(f'<a class="btn sec" href="/hearings/export.docx{dl_qs}">📥 docxダウンロード</a>')
    header_actions.append('<a class="btn sec" href="/hearing-templates">テンプレート管理</a>')
    header_actions.append('<a class="btn" href="/hearing/new">＋新規ヒアリング</a>')
    title = f"ヒアリング（{_esc(tmpl['name'])}）" if tmpl else "ヒアリング"
    desc = ("このテンプレートで実施されたヒアリングのみ表示しています。ダウンロードもこの絞り込み対象のみです。"
            if tmpl else "実施済みのヒアリング結果一覧です。テンプレートを選ぶと絞り込めます（ダウンロードも絞り込み対象のみになります）。")
    return f"""
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>{title}</span>
        <span style="display:flex;gap:8px;flex-wrap:wrap">
          {"".join(header_actions)}
        </span>
      </h2>
      <p class="muted" style="margin-bottom:14px">{desc}</p>
      {tmpl_filter_form}
      <form id="hearing_bulk_form" method="post" action="/hearings/bulk_delete">
      <div style="overflow:auto;max-height:70vh">
      <table style="table-layout:fixed">
        <tr><th class="sticky" style="width:32px"><input type="checkbox" id="hearing_chk_all" title="全選択"
              onchange="document.querySelectorAll('#hearing_bulk_form [name=ids]').forEach(c=>c.checked=this.checked)"></th>
            <th class="sticky" style="width:90px">ヒアリング日</th>
            <th class="sticky" style="width:160px">アカウント</th>
            <th class="sticky" style="width:200px">案件名</th>
            <th class="sticky" style="width:150px">テンプレート</th>
            <th class="sticky">回答プレビュー</th></tr>
        {rows or '<tr><td colspan=6 class="muted">まだヒアリング結果がありません。</td></tr>'}
      </table>
      </div>
      {f'''<div style="margin-top:10px;display:flex;gap:8px;flex-wrap:wrap">
        <button class="btn sec" type="submit" formaction="/hearings/export_selected" name="fmt" value="xlsx">
          📥 選択した件をxlsxダウンロード</button>
        <button class="btn sec" type="submit" formaction="/hearings/export_selected" name="fmt" value="docx">
          📥 選択した件をdocxダウンロード</button>
        <button class="btn" type="button" onclick="hearingBulkDelete()"
          style="background:#c53030;border-color:#c53030;color:#fff">選択した件を削除</button>
      </div>
      <p class="muted" style="font-size:11px;margin-top:6px">ダウンロードは案件（商談）ごとに別ファイルで生成されます。複数の案件をまたいで選択した場合はzipにまとめてダウンロードされます。</p>''' if results else ''}
      </form>
    </div>
    <script>
    function hearingBulkDelete() {{
      var ids = Array.from(document.querySelectorAll('#hearing_bulk_form [name=ids]:checked')).map(function(c){{return c.value;}});
      if (!ids.length) {{ alert('削除するヒアリング結果を選択してください。'); return; }}
      if (!confirm(ids.length + '件のヒアリング結果を削除します。この操作は取り消せません。よろしいですか？')) return;
      var form = document.getElementById('hearing_bulk_form');
      form.action = '/hearings/bulk_delete';
      form.submit();
    }}
    </script>"""


def _docx_add_answer(doc, a: dict) -> None:
    """1つの回答をdocxに追記する。yabane/radar/timeline/scorecardは表（フロー図的な構造）として出力し、
    それ以外は通常のテキストとして出力する。"""
    atype = a.get("type")
    ans = a.get("answer")
    p = doc.add_paragraph()
    p.add_run(str(a.get("label") or "")).bold = True

    def _table(headers, data_rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row_vals in data_rows:
            cells = table.add_row().cells
            for i, v in enumerate(row_vals):
                cells[i].text = "" if v is None else str(v)

    if atype == "yabane" and isinstance(ans, dict):
        rows = _yabane_rows_from_answer(ans)
        if rows:
            _table(
                ["ステップ", "部署", "作業内容", "アウトプット", "現行課題", "目指す姿", "目標数値"],
                [[r.get("step"), r.get("dept"), r.get("content"), r.get("output"),
                  r.get("issue"), r.get("target"), r.get("target_number")] for r in rows],
            )
        else:
            doc.add_paragraph("（データなし）")
    elif atype == "radar" and isinstance(ans, dict):
        axes = ans.get("axes") or []
        if axes:
            _table(["軸", "スコア（/5）"], [[ax.get("label"), ax.get("score")] for ax in axes])
        else:
            doc.add_paragraph("（データなし）")
    elif atype == "timeline" and isinstance(ans, dict):
        events = ans.get("events") or []
        if events:
            _table(["日付", "イベント", "メモ"], [[ev.get("date"), ev.get("label"), ev.get("note")] for ev in events])
        else:
            doc.add_paragraph("（データなし）")
    elif atype == "scorecard" and isinstance(ans, dict):
        criteria = ans.get("criteria") or []
        items = ans.get("items") or []
        if criteria and items:
            data_rows = []
            for it in items:
                scores = it.get("scores") or {}
                total = sum(v for v in scores.values() if isinstance(v, (int, float)))
                data_rows.append([it.get("label")] + [scores.get(c) for c in criteria] + [total or None])
            _table(["対象"] + list(criteria) + ["合計"], data_rows)
        else:
            doc.add_paragraph("（データなし）")
    else:
        doc.add_paragraph(_format_answer(ans))
    doc.add_paragraph("")


def _xlsx_add_answer(ws, row: int, a: dict) -> int:
    """1つの回答をシートに書き込み、次の書き込み開始行を返す。docx版と同様、
    yabane/radar/timeline/scorecardは表として出力する。"""
    from openpyxl.styles import Font
    atype = a.get("type")
    ans = a.get("answer")
    ws.cell(row=row, column=1, value=str(a.get("label") or "")).font = Font(bold=True)
    row += 1

    def _write_table(headers, data_rows):
        nonlocal row
        for i, h in enumerate(headers):
            ws.cell(row=row, column=1 + i, value=h).font = Font(bold=True)
        row += 1
        for row_vals in data_rows:
            for i, v in enumerate(row_vals):
                ws.cell(row=row, column=1 + i, value=v)
            row += 1

    if atype == "yabane" and isinstance(ans, dict):
        rows = _yabane_rows_from_answer(ans)
        if rows:
            _write_table(
                ["ステップ", "部署", "作業内容", "アウトプット", "現行課題", "目指す姿", "目標数値"],
                [[r.get("step"), r.get("dept"), r.get("content"), r.get("output"),
                  r.get("issue"), r.get("target"), r.get("target_number")] for r in rows],
            )
        else:
            ws.cell(row=row, column=1, value="（データなし）")
            row += 1
    elif atype == "radar" and isinstance(ans, dict):
        axes = ans.get("axes") or []
        if axes:
            _write_table(["軸", "スコア（/5）"], [[ax.get("label"), ax.get("score")] for ax in axes])
        else:
            ws.cell(row=row, column=1, value="（データなし）")
            row += 1
    elif atype == "timeline" and isinstance(ans, dict):
        events = ans.get("events") or []
        if events:
            _write_table(["日付", "イベント", "メモ"], [[ev.get("date"), ev.get("label"), ev.get("note")] for ev in events])
        else:
            ws.cell(row=row, column=1, value="（データなし）")
            row += 1
    elif atype == "scorecard" and isinstance(ans, dict):
        criteria = ans.get("criteria") or []
        items = ans.get("items") or []
        if criteria and items:
            data_rows = []
            for it in items:
                scores = it.get("scores") or {}
                total = sum(v for v in scores.values() if isinstance(v, (int, float)))
                data_rows.append([it.get("label")] + [scores.get(c) for c in criteria] + [total or None])
            _write_table(["対象"] + list(criteria) + ["合計"], data_rows)
        else:
            ws.cell(row=row, column=1, value="（データなし）")
            row += 1
    else:
        ws.cell(row=row, column=1, value=_format_answer(ans))
        row += 1
    return row + 1  # 空行を挟む


def build_hearing_result_docx_for_deal(results: list[dict]) -> bytes:
    """1商談分のヒアリング結果（複数可）をdocxにまとめる。"""
    from docx import Document
    from io import BytesIO
    doc = Document()
    first = results[0]
    doc.add_heading(f"{first.get('account_name') or ''} / {first.get('deal_name') or ''}", level=1)
    for r in results:
        doc.add_heading(f"{r.get('template_name') or ''}（{r.get('conducted_on') or '—'}）", level=2)
        for a in (r.get("answers") or []):
            _docx_add_answer(doc, a)
        doc.add_paragraph("")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_hearing_result_xlsx_for_deal(results: list[dict]) -> bytes:
    """1商談分のヒアリング結果（複数可）をxlsxにまとめる。"""
    import openpyxl
    from openpyxl.styles import Font
    from io import BytesIO
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "ヒアリング結果"
    first = results[0]
    ws.cell(row=1, column=1, value=f"{first.get('account_name') or ''} / {first.get('deal_name') or ''}").font = Font(bold=True, size=13)
    row = 3
    for r in results:
        ws.cell(row=row, column=1, value=f"{r.get('template_name') or ''}（{r.get('conducted_on') or '—'}）").font = Font(bold=True, size=12)
        row += 1
        for a in (r.get("answers") or []):
            row = _xlsx_add_answer(ws, row, a)
        row += 1
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_hearings_export_bundle(con, result_ids: list[int], fmt: str) -> tuple[bytes, str, str]:
    """選択されたヒアリング結果を商談（案件）ごとに別ファイルで生成する。
    1商談のみの場合はそのファイルを直接返し、複数商談にまたがる場合はzipにまとめる。
    戻り値: (バイト列, ファイル名, Content-Type)。
    """
    import re as _re
    import zipfile
    from io import BytesIO

    all_results = [r for r in (sfa_db.get_hearing_result(con, rid) for rid in result_ids) if r]
    groups: dict = {}
    for r in all_results:
        groups.setdefault(r["deal_id"], []).append(r)

    ext = "docx" if fmt == "docx" else "xlsx"
    ctype = ("application/vnd.openxmlformats-officedocument.wordprocessingml.document" if ext == "docx"
             else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    builder = build_hearing_result_docx_for_deal if ext == "docx" else build_hearing_result_xlsx_for_deal

    def safe_filename(name: str) -> str:
        name = _re.sub(r'[\\/:*?"<>|]', "_", name or "案件")
        return name[:60] or "案件"

    files = []
    for _deal_id, items in groups.items():
        first = items[0]
        fname = safe_filename(f"{first.get('account_name') or ''}_{first.get('deal_name') or ''}") + f".{ext}"
        files.append((fname, builder(items)))

    if not files:
        raise ValueError("対象のヒアリング結果が見つかりません")
    if len(files) == 1:
        return files[0][1], files[0][0], ctype

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        used_names: set = set()
        for fname, data in files:
            base, _, extn = fname.rpartition(".")
            final_name, n = fname, 1
            while final_name in used_names:
                final_name = f"{base}_{n}.{extn}"
                n += 1
            used_names.add(final_name)
            zf.writestr(final_name, data)
    return buf.getvalue(), "hearings_export.zip", "application/zip"


def build_hearing_results_csv(con, template_id: int) -> bytes:
    """指定テンプレートのヒアリング結果のみをCSVにまとめる（テンプレート管理画面の絞り込みDL用）。"""
    import csv
    import io
    results = sfa_db.list_all_hearing_results(con, template_id=template_id)
    labels: list = []
    for r in results:
        for a in (r.get("answers") or []):
            lbl = a.get("label")
            if lbl and lbl not in labels:
                labels.append(lbl)

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([_csv_safe(h) for h in (["商談ID", "アカウント", "案件名", "ヒアリング日"] + labels)])
    for r in results:
        amap = {a.get("label"): _format_answer_for_export(a) for a in (r.get("answers") or [])}
        writer.writerow([_csv_safe(v) for v in ([
            r.get("deal_id"), r.get("account_name") or "", r.get("deal_name") or "",
            r.get("conducted_on") or "",
        ] + [amap.get(lbl, "") for lbl in labels])])
    return ("﻿" + buf.getvalue()).encode("utf-8")


def leads_page(con, *, status=None, source=None, q=None) -> str:
    # デフォルトは商談化済・見込みなしを除外
    if status is None:
        leads = [l for l in sfa_db.list_leads(con, source=source, q=q)
                 if l.get("lead_status") not in ("converted", "lost")]
    else:
        leads = sfa_db.list_leads(con, status=status, source=source, q=q)

    status_opts = ('<option value="">全ステージ</option>'
                   + "".join(
                       f'<option value="{s}"{" selected" if s == status else ""}>'
                       f'{sfa_db.LEAD_STATUS_LABELS[s]}</option>'
                       for s in sfa_db.LEAD_STATUSES))
    source_opts = ('<option value="">全経路</option>'
                   + "".join(
                       f'<option value="{s}"{" selected" if s == source else ""}>'
                       f'{sfa_db.LEAD_SOURCE_LABELS[s]}</option>'
                       for s in sfa_db.LEAD_SOURCES))

    filter_form = f"""<form method="get" action="/leads" class="filter-row">
      <select name="status" onchange="this.form.submit()">{status_opts}</select>
      <select name="source" onchange="this.form.submit()">{source_opts}</select>
      <a class="btn sec" href="/leads">リセット</a>
    </form>
    <div style="margin:8px 0 10px">
      <input type="text" id="leadSearch" placeholder="🔍 氏名・会社で絞り込み（入力すると即座に絞られます）"
        oninput="filterLeads()" style="max-width:320px">
    </div>"""

    # マスタデータ取得（インライン編集・バルク編集用）
    owners_list = sfa_db.get_master_list(con, "owners")
    industries_list = sfa_db.get_master_list(con, "industries")
    company_sizes_list = sfa_db.get_master_list(con, "company_sizes")
    # バルク編集用JS オブジェクト構築
    bulk_options = {
        "source": [["", "（変更なし）"]] + [[s, sfa_db.LEAD_SOURCE_LABELS[s]] for s in sfa_db.LEAD_SOURCES],
        "assigned_to": [["", "（変更なし）"]] + [[o, o] for o in owners_list],
        "industry": [["", "（変更なし）"]] + [[i, i] for i in industries_list],
        "company_size": [["", "（変更なし）"]] + [[cs, cs] for cs in company_sizes_list],
        "lead_status": [["", "（変更なし）"]] + [[s, sfa_db.LEAD_STATUS_LABELS[s]] for s in sfa_db.LEAD_STATUSES],
    }
    bulk_options_json = json.dumps(bulk_options, ensure_ascii=False)

    def _inline_select_source(lead_id, current):
        opts = "".join(
            f'<option value="{html.escape(s)}"{" selected" if s == current else ""}>{html.escape(sfa_db.LEAD_SOURCE_LABELS[s])}</option>'
            for s in sfa_db.LEAD_SOURCES
        )
        return (f'<select onchange="updateLeadField({lead_id}, \'source\', this.value)"'
                f' style="font-size:12px;padding:2px 4px;max-width:120px">'
                f'<option value=""></option>{opts}</select>')

    def _inline_select_master(lead_id, field, values, current):
        opts = "".join(
            f'<option value="{html.escape(v)}"{" selected" if v == current else ""}>{html.escape(v)}</option>'
            for v in values
        )
        return (f'<select onchange="updateLeadField({lead_id}, \'{field}\', this.value)"'
                f' style="font-size:12px;padding:2px 4px;max-width:120px">'
                f'<option value=""></option>{opts}</select>')

    def _inline_select_status(lead_id, current):
        opts = "".join(
            f'<option value="{html.escape(s)}"{" selected" if s == current else ""}>{html.escape(sfa_db.LEAD_STATUS_LABELS[s])}</option>'
            for s in sfa_db.LEAD_STATUSES
        )
        return (f'<select onchange="updateLeadField({lead_id}, \'lead_status\', this.value)"'
                f' style="font-size:12px;padding:2px 4px;max-width:110px">'
                f'<option value=""></option>{opts}</select>')

    rows = []
    for ld in leads:
        sc = f's-{ld.get("lead_status", "new")}'
        sl = sfa_db.LEAD_STATUS_LABELS.get(ld.get("lead_status", "new"), "")
        deal_badge = (f' <a href="/deal/{ld["deal_id"]}" title="紐付け商談">🔗</a>'
                      if ld.get("deal_id") else "")
        sel_status = _inline_select_status(ld["id"], ld.get("lead_status", "new"))
        sel_source = _inline_select_source(ld["id"], ld.get("source", "other"))
        sel_owner = _inline_select_master(ld["id"], "assigned_to", owners_list, ld.get("assigned_to") or "")
        sel_industry = _inline_select_master(ld["id"], "industry", industries_list, ld.get("industry") or "")
        sel_company_size = _inline_select_master(ld["id"], "company_size", company_sizes_list, ld.get("company_size") or "")
        _search_key = f'{(ld.get("name") or "")} {(ld.get("company") or "")}'.lower()
        rows.append(
            f'<tr data-search="{_esc(_search_key)}">'
            f'<td style="width:32px"><input type="checkbox" name="ids" value="{ld["id"]}"></td>'
            f'<td><a href="/leads/{ld["id"]}">{_esc(ld["name"])}</a>{deal_badge}<br>'
            f'<span class="muted">{_esc(ld.get("company"))}</span></td>'
            f'<td>{sel_status}</td>'
            f'<td class="hide-sm">{sel_source}</td>'
            f'<td class="hide-sm">{sel_owner}</td>'
            f'<td class="hide-sm">{sel_industry}</td>'
            f'<td class="hide-sm">{sel_company_size}</td>'
            f'<td class="muted">{_esc((ld.get("updated_at") or "")[:10])}</td>'
            f'</tr>'
        )

    return f"""
    <div class="card">
      <h2 style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">
        <span>リード一覧 <span class="muted" style="font-weight:normal">({len(leads)}件)</span></span>
        <span style="display:flex;gap:8px">
          <a class="btn sec" href="/email-draft">メールドラフト</a>
          <a class="btn sec" href="/leads/import">CSV取込</a>
          <a class="btn" href="/leads/new">＋新規リード</a>
        </span>
      </h2>
      {filter_form}
      <form id="bulk_form" method="post" action="/leads/bulk_edit">
      <div style="overflow:auto;max-height:70vh">
      <table>
        <tr><th class="sticky" style="width:32px"><input type="checkbox" id="chk_all" title="全選択"
              onchange="document.querySelectorAll('[name=ids]').forEach(c=>c.checked=this.checked)"></th>
            {_sticky_th('氏名 / 会社')}{_sticky_th('ステータス')}
            <th class="hide-sm sticky">経路</th>
            <th class="hide-sm sticky">担当</th>
            <th class="hide-sm sticky">業界</th>
            <th class="hide-sm sticky">企業規模</th>
            {_sticky_th('更新日')}</tr>
        {''.join(rows) or '<tr><td colspan=8 class=muted>リードがありません。「＋新規リード」から追加、またはCSV取込してください。</td></tr>'}
      </table>
      </div>
      <div style="display:flex;align-items:center;gap:8px;margin-top:10px;flex-wrap:wrap">
        <select id="bulk_field" name="field" style="width:auto">
          <option value="lead_status">ステータス</option>
          <option value="source">経路</option>
          <option value="assigned_to">担当</option>
          <option value="industry">業界</option>
          <option value="company_size">企業規模</option>
        </select>
        <select id="bulk_value" name="value" style="width:auto"></select>
        <button class="btn sec" type="submit">選択した件を一括変更</button>
        <button class="btn" type="button" onclick="bulkDelete()"
          style="background:#c53030;border-color:#c53030;color:#fff;margin-left:8px">選択した件を削除</button>
      </div>
      </form>
    </div>
    <script>
    const BULK_OPTIONS = {bulk_options_json};
    function updateLeadField(id, field, value) {{
      fetch('/leads/' + id + '/field', {{
        method: 'POST',
        headers: {{'Content-Type': 'application/x-www-form-urlencoded'}},
        body: 'field=' + encodeURIComponent(field) + '&value=' + encodeURIComponent(value)
      }}).then(r => r.json()).then(d => {{
        if (!d.ok) {{ alert('更新エラー: ' + (d.error || '')); }}
      }}).catch(() => alert('通信エラー'));
    }}
    function repopulateBulkValue() {{
      var field = document.getElementById('bulk_field').value;
      var opts = BULK_OPTIONS[field] || [];
      var sel = document.getElementById('bulk_value');
      sel.innerHTML = opts.map(function(pair) {{
        return '<option value="' + escH(pair[0]) + '">' + escH(pair[1]) + '</option>';
      }}).join('');
    }}
    function bulkDelete() {{
      var ids = Array.from(document.querySelectorAll('[name=ids]:checked')).map(c => c.value);
      if (!ids.length) {{ alert('削除するリードを選択してください。'); return; }}
      if (!confirm(ids.length + '件のリードを削除します。この操作は取り消せません。よろしいですか？')) return;
      var form = document.createElement('form');
      form.method = 'post';
      form.action = '/leads/bulk_delete';
      ids.forEach(function(id) {{
        var inp = document.createElement('input');
        inp.type = 'hidden'; inp.name = 'ids'; inp.value = id;
        form.appendChild(inp);
      }});
      document.body.appendChild(form);
      form.submit();
    }}
    document.getElementById('bulk_field').addEventListener('change', repopulateBulkValue);
    repopulateBulkValue();
    function filterLeads() {{
      var q = (document.getElementById('leadSearch').value || '').toLowerCase();
      document.querySelectorAll('tr[data-search]').forEach(function(tr) {{
        tr.style.display = tr.getAttribute('data-search').indexOf(q) >= 0 ? '' : 'none';
      }});
    }}
    </script>"""


def lead_form(con, lead=None) -> str:
    lead = lead or {}
    accounts = sfa_db.list_accounts(con)
    acc_datalist = "".join(
        f'<option value="{html.escape(a["name"])}"></option>' for a in accounts
    )
    status_items = [(s, sfa_db.LEAD_STATUS_LABELS[s]) for s in sfa_db.LEAD_STATUSES]
    source_items = [(s, sfa_db.LEAD_SOURCE_LABELS[s]) for s in sfa_db.LEAD_SOURCES]

    status_btns = ""
    convert_btn = ""
    deal_link = ""
    activities_html = ""

    if lead.get("id"):
        cur_status = lead.get("lead_status", "new")
        btns = []
        for s in sfa_db.LEAD_STATUSES:
            active_style = ("font-weight:700;box-shadow:inset 0 0 0 2px #2f6fed"
                            if s == cur_status else "opacity:0.55")
            btns.append(
                f'<form method="post" action="/leads/{lead["id"]}/status"'
                f' style="display:inline;margin:0 4px 4px 0">'
                f'<input type="hidden" name="status" value="{s}">'
                f'<button class="btn sec" style="{active_style}">'
                f'{sfa_db.LEAD_STATUS_LABELS[s]}</button></form>'
            )
        status_btns = f'<div style="margin:0 0 14px">{"".join(btns)}</div>'

        acts = sfa_db.list_lead_activities(con, lead["id"])
        act_rows = ""
        for a in acts:
            tl = sfa_db.LEAD_ACTIVITY_LABELS.get(a.get("type", "note"), a.get("type", ""))
            act_rows += (
                f'<tr><td class="muted" style="white-space:nowrap">{_esc((a.get("created_at") or "")[:16])}</td>'
                f'<td>{tl}</td><td>{_esc(a.get("author"))}</td>'
                f'<td style="white-space:pre-wrap">{_esc(a.get("content"))}</td></tr>'
            )
        act_type_opts = "".join(
            f'<option value="{t}">{sfa_db.LEAD_ACTIVITY_LABELS[t]}</option>'
            for t in sfa_db.LEAD_ACTIVITY_TYPES
        )
        activities_html = f"""
        <div class="card"><h2>活動ログ</h2>
        <table><tr><th>日時</th><th>種別</th><th>担当</th><th>内容</th></tr>
        {act_rows or '<tr><td colspan=4 class=muted>活動なし</td></tr>'}
        </table>
        <form method="post" action="/leads/{lead['id']}/activity" style="margin-top:14px">
          <div class="grid">
            <div><label>種別</label><select name="type">{act_type_opts}</select></div>
            <div><label>担当者</label><select name="author">{_opt(sfa_db.get_master_list(con,'owners'), None)}</select></div>
          </div>
          <label>内容 *</label><textarea name="content" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="2" required></textarea>
          <p><button class="btn sec">活動を追加</button></p>
        </form></div>"""

        can_convert = (cur_status not in ("converted", "lost") and not lead.get("deal_id"))
        if can_convert:
            convert_btn = f'<a class="btn sync" href="/leads/{lead["id"]}/convert">アポ獲得 → 商談化</a>'
        if lead.get("deal_id"):
            deal_link = f'<a class="btn sec" href="/deal/{lead["deal_id"]}">紐付け商談を見る 🔗</a>'

    delete_btn = ""
    if lead.get("id"):
        delete_btn = (
            f'<form method="post" action="/leads/{lead["id"]}/delete" style="display:inline;margin:0">'
            '<button class="btn" style="background:#ef4444"'
            ' onclick="return confirm(\'このリードを削除しますか？この操作は元に戻せません。\')">削除</button></form>'
        )

    return f"""
    <div class="card">
      <h2>{'リード編集' if lead.get('id') else '新規リード'}</h2>
      {_save_bar('leadForm', cancel_url=('/leads/' + str(lead['id']) if lead.get('id') else '/leads'))}
      {status_btns}
      <form method="post" action="/leads/save" id="leadForm">
        <input type="hidden" name="id" value="{_esc(lead.get('id'))}">
        <div class="grid">
          <div><label>氏名 *</label>
            <input name="name" required value="{_esc(lead.get('name'))}"></div>
          <div><label>会社名 * <span class="muted">（既存アカウントから選択または新規入力）</span></label>
            <input name="company" required value="{_esc(lead.get('company'))}" list="acc_list" autocomplete="off">
            <datalist id="acc_list">{acc_datalist}</datalist></div>
          <div><label>業界</label>
            <select name="industry">{_opt(sfa_db.get_master_list(con,'industries'), lead.get('industry'))}</select></div>
          <div><label>企業規模</label>
            <select name="company_size">{_opt(sfa_db.get_master_list(con,'company_sizes'), lead.get('company_size'))}</select></div>
          <div><label>役職</label>
            <input name="title" value="{_esc(lead.get('title'))}"></div>
          <div><label>メール</label>
            <input name="email" type="email" value="{_esc(lead.get('email'))}"></div>
          <div><label>電話</label>
            <input name="phone" value="{_esc(lead.get('phone'))}"></div>
          <div><label>担当者</label>
            <select name="assigned_to">{_opt(sfa_db.get_master_list(con,'owners'), lead.get('assigned_to'))}</select></div>
          <div><label>獲得経路</label>
            <select name="source">{_opt_kv(source_items, lead.get('source') or 'other')}</select></div>
          <div><label>ステータス</label>
            <select name="lead_status">{_opt_kv(status_items, lead.get('lead_status') or 'new')}</select></div>
          <div><label>終了理由 <span class="muted" style="font-weight:400;font-size:.8em">（lost時）</span></label>
            <select name="lost_reason"><option value="">（なし）</option>{_opt(sfa_db.CLOSE_REASONS, lead.get('lost_reason'))}</select></div>
        </div>
        <label>メモ</label><textarea name="notes" class="ta-expand" onfocus="taExpand(this)" onblur="taShrink(this)" rows="2">{_esc(lead.get('notes'))}</textarea>
        <p style="display:flex;flex-wrap:wrap;gap:8px">
          <button class="btn">保存</button>
          <a class="btn sec" href="/leads">一覧へ</a>
          {deal_link}
        </p>
      </form>
      <div style="display:flex;flex-wrap:wrap;gap:8px;margin-top:8px">
        {convert_btn}
        {delete_btn}
      </div>
    </div>
    {activities_html}"""


def leads_import_page(con, result: str = "") -> str:
    result_html = f'<div class="flash">{html.escape(result)}</div>' if result else ""
    header_line = ",".join(leads_csv.TEMPLATE_HEADERS)
    example_lines = "\n".join(",".join(row) for row in leads_csv.TEMPLATE_EXAMPLE_ROWS)

    owners_list = sfa_db.get_master_list(con, "owners")
    industries_list = sfa_db.get_master_list(con, "industries")
    sizes_list = sfa_db.get_master_list(con, "company_sizes")
    theme_names = [t["name"] for t in sfa_db.list_pitch_themes(con, active_only=True)]
    source_chips = _chips([f"{s}({sfa_db.LEAD_SOURCE_LABELS[s]})" for s in sfa_db.LEAD_SOURCES])
    status_chips = _chips([f"{s}({sfa_db.LEAD_STATUS_LABELS[s]})" for s in sfa_db.LEAD_STATUSES])

    choices_html = f"""
    <details style="margin:12px 0;background:#f8f9fb;border-radius:6px;padding:10px 14px" open>
      <summary style="cursor:pointer;font-weight:600;color:#3a4760">
        📋 入力選択肢一覧（この中から選んでください。表記が完全一致しないと空欄で取り込まれます）
      </summary>
      <div style="margin-top:10px;font-size:13px;line-height:1.9">
        <div><strong>業界</strong>: {_chips(industries_list)}</div>
        <div><strong>企業規模</strong>: {_chips(sizes_list)}</div>
        <div><strong>獲得経路</strong>（コードで入力。一致しない場合は自動的に other）: {source_chips}</div>
        <div><strong>ステータス</strong>（コードで入力。一致しない場合は自動的に new）: {status_chips}</div>
        <div><strong>ピッチテーマ</strong>: {_chips(theme_names)}</div>
        <div><strong>担当者</strong>: {_chips(owners_list)}</div>
      </div>
    </details>"""

    return f"""
    <div class="card"><h2>リード一括取込</h2>
    {result_html}

    <h3 style="margin:0 0 8px;font-size:14px;color:#3a4760">📇 名刺データ（xlsx）アップロード</h3>
    <p class="muted">名刺管理アプリ（Eight / CAMCARD / Sansan等）からエクスポートしたxlsxを<strong>そのまま</strong>アップロードします（列構成は名刺アプリ側の固定形式）。<br>
    会社名・業界などはAIがWebリサーチで補強します（ANTHROPIC_API_KEY 設定時）。<br>
    <strong>下記テンプレートに沿って作成したCSV/xlsxはこちらではなく、下の「CSVファイルアップロード」欄を使ってください。</strong></p>
    <form method="post" action="/leads/upload_meishi" enctype="multipart/form-data" style="margin-bottom:20px">
      <label>名刺xlsxファイル</label>
      <input type="file" name="meishi_file" accept=".xlsx,.xls,.csv" required style="padding:4px">
      <p style="margin-top:8px"><button class="btn">アップロードして取込</button>
         <a class="btn sec" href="/leads">キャンセル</a></p>
    </form>

    <hr style="margin:20px 0">
    <h3 style="margin:0 0 8px;font-size:14px;color:#3a4760">📋 CSVファイルアップロード／ペースト取込（こちらがメインの取込方法です）</h3>
    <p style="margin:10px 0">
      <a class="btn sec" href="/leads/import/template.csv">📥 テンプレートCSVをダウンロード</a>
    </p>
    <p class="muted">運用ルール: 入力は必ずこのテンプレートに列を揃えて作成してください（列の追加・削除・並び替えは不可、値のみ入力）。</p>

    <form method="post" action="/leads/import" enctype="multipart/form-data"
          style="background:#f0f7ff;border:1px solid #cfe0fb;border-radius:8px;padding:14px;margin:14px 0">
      <label>👇 CSVファイル（テンプレートに沿ったファイルをそのままアップロード。xlsxも可）</label>
      <input type="file" name="csv_file" accept=".csv,.xlsx,.xls" style="padding:4px">
      <label style="margin-top:12px">またはCSVデータを直接ペースト</label>
      <textarea name="csv_text" rows="8"
        style="font-family:monospace;font-size:12px"></textarea>
      <p class="muted" style="margin-top:4px">両方入力した場合はアップロードしたファイルを優先します。</p>
      <p><button class="btn">取込実行</button>
         <a class="btn sec" href="/leads">キャンセル</a></p>
    </form>

    <p class="muted" style="margin-top:12px">CSVフォーマット（1行目はヘッダ行、空行はスキップ）:</p>
    <pre style="background:#f4f6f9;padding:10px;border-radius:6px">{html.escape(header_line)}
{html.escape(example_lines)}</pre>
    <p class="muted" style="margin-top:4px">
      業界 / 企業規模 / 担当者 / ピッチテーマは、下の「入力選択肢一覧」と完全一致した場合のみ取り込まれます
      （一致しない場合は空欄になります）。業界・企業規模が空欄の行はAIが会社名から自動推定します
      （ANTHROPIC_API_KEY設定時）。
    </p>

    {_ai_prompt_block(leads_csv.build_ai_prompt(con), "/leads/import/ai_prompt.md")}

    {choices_html}
    </div>"""


def deals_import_page(con, result: str = "") -> str:
    result_html = f'<div class="flash">{html.escape(result)}</div>' if result else ""
    header_line = ",".join(deals_csv.TEMPLATE_HEADERS)
    example_lines = "\n".join(",".join(row) for row in deals_csv.TEMPLATE_EXAMPLE_ROWS)

    stages = sfa_db.get_master_list(con, "deal_stages")
    biz_l1_list = sfa_db.get_master_list(con, "business_type_l1")
    owners_list = sfa_db.get_master_list(con, "owners")
    patterns_list = sfa_db.get_master_list(con, "lead_patterns")
    sizes_list = sfa_db.get_master_list(con, "company_sizes")
    industries_list = sfa_db.get_master_list(con, "industries")

    biz_l2_html = "".join(
        f'<div style="margin:4px 0 4px 12px">└ <strong>{html.escape(l1)}</strong>: '
        f'{_chips(sfa_db.BUSINESS_TYPE_L2_BY_L1.get(l1, []))}</div>'
        for l1 in biz_l1_list
    )

    choices_html = f"""
    <details style="margin:12px 0;background:#f8f9fb;border-radius:6px;padding:10px 14px" open>
      <summary style="cursor:pointer;font-weight:600;color:#3a4760">
        📋 入力選択肢一覧（この中から選んでください。表記が完全一致しないと空欄で取り込まれます）
      </summary>
      <div style="margin-top:10px;font-size:13px;line-height:1.9">
        <div><strong>種別</strong>: {_chips(["商談", "アカウント"])}</div>
        <div><strong>ステージ</strong>: {_chips(stages)}</div>
        <div><strong>事業種別L1 / L2</strong>: {_chips(biz_l1_list)}
          {biz_l2_html}
        </div>
        <div><strong>リード経路</strong>: {_chips(patterns_list)}</div>
        <div><strong>担当者・サブ担当</strong>: {_chips(owners_list)}</div>
        <div><strong>重要度</strong>: {_chips(sfa_db.IMPORTANCE_OPTIONS)}</div>
        <div><strong>企業規模</strong>: {_chips(sizes_list)}</div>
        <div><strong>業界</strong>（参考。自由入力欄のためこの一覧以外の表記も取り込めます）: {_chips(industries_list)}</div>
      </div>
    </details>"""

    return f"""
    <div class="card"><h2>商談一括取込</h2>
    {result_html}

    <p class="muted">
      1行 = 1商談 または 1アカウント。<strong>種別</strong>列に「商談」「アカウント」のどちらかを指定します
      （省略した場合、商談名があれば商談、無ければアカウントとして扱います）。<br>
      会社名が未登録の場合、商談・アカウントどちらの取込でも<strong>アカウントを自動作成</strong>します
      （既存アカウントは業界・企業規模が空欄の項目のみ補完し、上書きはしません）。
    </p>
    <p style="margin:10px 0">
      <a class="btn sec" href="/deals/import/template.csv">📥 テンプレートCSVをダウンロード</a>
    </p>
    <p class="muted">運用ルール: 入力は必ずこのテンプレートに列を揃えて作成してください（列の追加・削除・並び替えは不可、値のみ入力）。</p>

    <form method="post" action="/deals/import" enctype="multipart/form-data"
          style="background:#f0f7ff;border:1px solid #cfe0fb;border-radius:8px;padding:14px;margin:14px 0">
      <label>👇 CSVファイル（テンプレートに沿ったファイルをそのままアップロード。xlsxも可）</label>
      <input type="file" name="csv_file" accept=".csv,.xlsx,.xls" style="padding:4px">
      <label style="margin-top:12px">またはCSVデータを直接ペースト</label>
      <textarea name="csv_text" rows="8"
        style="font-family:monospace;font-size:12px"></textarea>
      <p class="muted" style="margin-top:4px">両方入力した場合はアップロードしたファイルを優先します。</p>
      <p><button class="btn">取込実行</button>
         <a class="btn sec" href="/deals">キャンセル</a></p>
    </form>

    <p class="muted" style="margin-top:12px">CSVフォーマット（1行目はヘッダ行、空行はスキップ）:</p>
    <pre style="background:#f4f6f9;padding:10px;border-radius:6px">{html.escape(header_line)}
{html.escape(example_lines)}</pre>
    <p class="muted" style="margin-top:4px">
      ステージ / 事業種別L1 / 事業種別L2（L1に対応する値のみ） / リード経路 / 担当者 / 重要度 / 企業規模は、
      下の「入力選択肢一覧」と完全一致した場合のみ取り込まれます（一致しない場合は空欄になります）。<br>
      業界は自由入力欄のためどんな文字列でも取り込めます。事前に業界・企業規模を調査済みであれば、
      その場でCSVに直接入力してください（AI推定より確実です）。
    </p>

    {_ai_prompt_block(deals_csv.build_ai_prompt(con), "/deals/import/ai_prompt.md")}

    {choices_html}
    </div>"""




# ── HTTPハンドラ ───────────────────────────────────────────────────────────────

def _make_handler(db_path: str, theme_client: ThemeDBClient | None):
    class H(BaseHTTPRequestHandler):
        def log_message(self, *a):  # 静音
            pass

        def _send(self, body: bytes, status=200, ctype="text/html; charset=utf-8"):
            self.send_response(status)
            self.send_header("Content-Type", ctype)
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)

        def _send_cors_json(self, body: bytes, status=200):
            self.send_response(status)
            self.send_header("Content-Type", "application/json")
            self.send_header("Content-Length", str(len(body)))
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            self.wfile.write(body)

        def do_OPTIONS(self):
            self.send_response(204)
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
            self.send_header("Access-Control-Allow-Headers", "Content-Type")
            self.end_headers()

        def _redirect(self, location):
            self.send_response(303)
            self.send_header("Location", location)
            self.end_headers()

        def _form(self) -> dict:
            n = int(self.headers.get("Content-Length", 0))
            self._form_raw = self.rfile.read(n).decode("utf-8", errors="replace")
            d = urllib.parse.parse_qs(self._form_raw, keep_blank_values=True)
            return {k: (v[0] if v else "") for k, v in d.items()}

        def _form_list(self, key: str) -> list[str]:
            """_form() 呼び出し後に特定キーの全値リストを返す。"""
            d = urllib.parse.parse_qs(getattr(self, "_form_raw", ""), keep_blank_values=True)
            return d.get(key, [])

        def _form_multi(self) -> dict:
            """multipart/form-data対応。ファイルは(filename, バイト列)のタプルで返す。

            Python 3.13でcgiモジュールが削除された（PEP 594）ため、標準ライブラリのみで
            簡易的なmultipartパーサを自前実装している。
            """
            ctype = self.headers.get("Content-Type", "")
            n = int(self.headers.get("Content-Length", 0))
            body = self.rfile.read(n)

            m = re.search(r'boundary="?([^";]+)"?', ctype)
            if not m:
                return {}
            boundary = ("--" + m.group(1)).encode("utf-8")

            result = {}
            for part in body.split(boundary):
                part = part.strip(b"\r\n")
                if not part or part == b"--":
                    continue
                if b"\r\n\r\n" not in part:
                    continue
                header_blob, content = part.split(b"\r\n\r\n", 1)
                content = content[:-2] if content.endswith(b"\r\n") else content
                headers = {}
                for line in header_blob.decode("utf-8", errors="replace").split("\r\n"):
                    if ":" in line:
                        k, v = line.split(":", 1)
                        headers[k.strip().lower()] = v.strip()
                disposition = headers.get("content-disposition", "")
                name_m = re.search(r'name="([^"]*)"', disposition)
                if not name_m:
                    continue
                field_name = name_m.group(1)
                filename_m = re.search(r'filename="([^"]*)"', disposition)
                if filename_m:
                    result[field_name] = (filename_m.group(1), content)
                else:
                    result[field_name] = content.decode("utf-8", errors="replace")
            return result

        def _qs(self) -> dict:
            qs_raw = self.path.split("?")[1] if "?" in self.path else ""
            return urllib.parse.parse_qs(qs_raw)

        def _check_basic_auth(self) -> bool:
            """ブラウザ向け全ルートの認証（フォームCookieセッション or 従来のBasic認証を許可）。

            除外: /health, /api/*, /slack/*, /login, /logout, /favicon.ico。
            SFA_BASIC_USER/SFA_BASIC_PASS 未設定時はfail-closed（503）。
            未認証: GETは /login へ302誘導（ネイティブBasicダイアログを出さない＝モバイルのループ回避, #54）、
            それ以外は401 JSON。呼び出し側は即returnすること。
            """
            path = self.path.split("?")[0].rstrip("/") or "/"
            if (path in ("/health", "/login", "/logout", "/favicon.ico")
                    or path.startswith("/api/") or path.startswith("/slack/")):
                return True
            if not SFA_BASIC_USER or not SFA_BASIC_PASS:
                body = ("<h1>503</h1><p>SFA_BASIC_USER / SFA_BASIC_PASS が未設定のため"
                        "アクセスを拒否しています（fail-closed）。環境変数を設定してください。</p>").encode("utf-8")
                self._send(body, status=503)
                return False
            # 1) フォームログインの署名Cookieセッション
            try:
                ck = SimpleCookie(self.headers.get("Cookie", ""))
                sess = ck[_SESSION_COOKIE].value if _SESSION_COOKIE in ck else ""
            except Exception:  # noqa: BLE001 — 壊れたCookieヘッダは未認証扱い
                sess = ""
            if sess and _valid_session_token(sess):
                return True
            # 2) 従来のBasic認証（PC等の既存運用を壊さないため併存）
            header = self.headers.get("Authorization", "")
            if header.startswith("Basic "):
                try:
                    userpass = base64.b64decode(header[6:]).decode("utf-8")
                except Exception:  # noqa: BLE001
                    userpass = ""
                if hmac.compare_digest(userpass, f"{SFA_BASIC_USER}:{SFA_BASIC_PASS}"):
                    return True
            # 未認証 → ログイン画面へ（ネイティブダイアログは出さない）
            if self.command == "GET":
                nxt = urllib.parse.quote(self.path, safe="")
                self.send_response(302)
                self.send_header("Location", f"/login?next={nxt}")
                self.send_header("Content-Length", "0")
                self.end_headers()
            else:
                self._send(json.dumps({"error": "unauthorized"}).encode(), status=401,
                           ctype="application/json")
            return False

        def do_GET(self):
            if not self._check_basic_auth():
                return
            path = self.path.split("?")[0].rstrip("/") or "/"
            con = sfa_db.connect(db_path)
            try:
                if path == "/health":
                    self._send(b'{"status":"ok"}', ctype="application/json")
                elif path == "/login":
                    _nxt = self._qs().get("next", ["/"])[0] or "/"
                    self._send(login_page(_nxt), ctype="text/html; charset=utf-8")
                elif path == "/logout":
                    self.send_response(302)
                    self.send_header("Location", "/login")
                    self.send_header("Set-Cookie",
                                     f"{_SESSION_COOKIE}=; Path=/; HttpOnly; SameSite=Lax; Max-Age=0")
                    self.send_header("Content-Length", "0")
                    self.end_headers()
                elif (path.startswith("/deal/") and path.endswith("/milestones")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    # 一覧のMS管理パネル用: 商談の全MSをJSONで返す（レガシーは初回に実体化）
                    _did = int(path.split("/")[2])
                    sfa_db.ensure_milestones_materialized(con, _did)
                    self._send(json.dumps(_ms_panel_json(con, _did), ensure_ascii=False).encode(),
                               ctype="application/json")
                elif (path.startswith("/deal/") and path.endswith("/dev-tools")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    # 一覧のツールリンクパネル用: 商談に紐づく開発案件とそのツールリンクをJSONで返す
                    _did = int(path.split("/")[2])
                    _projs = sfa_db.list_dev_projects(con, deal_id=_did)
                    _out = []
                    for _p in _projs:
                        _extra = sfa_db.list_dev_project_tools(con, _p["id"])
                        _out.append({
                            "id": _p["id"],
                            "theme": _p.get("theme") or f"開発案件#{_p['id']}",
                            "tool_url": _p.get("tool_url") or "",
                            "tool_login_id": _p.get("tool_login_id") or "",
                            "tool_login_pass": _p.get("tool_login_pass") or "",
                            "extra": [{"id": e["id"], "label": e.get("label") or "",
                                       "url": e.get("url") or ""} for e in _extra],
                        })
                    self._send(json.dumps({"ok": True, "deal_id": _did, "projects": _out},
                                          ensure_ascii=False).encode(), ctype="application/json")
                elif path == "/api/deals":
                    qs = self._qs()
                    token = (qs.get("token", [None])[0] or "")
                    if not SFA_API_TOKEN or not hmac.compare_digest(token, SFA_API_TOKEN):
                        self._send(b'{"error":"unauthorized"}', status=401, ctype="application/json")
                    else:
                        status_q = (qs.get("status", ["open"])[0] or "open")
                        effective = None if status_q == "all" else status_q
                        deals = sfa_db.list_deals(con, status=effective)
                        self._send(json.dumps([dict(d) for d in deals], ensure_ascii=False, default=str).encode(), ctype="application/json")
                elif path == "/api/deal_issues":
                    # 論点リマインド(#47)用: 議論中の論点一覧（deal/account名JOIN込み）。トークン認証。
                    qs = self._qs()
                    token = (qs.get("token", [None])[0] or "")
                    if not SFA_API_TOKEN or not hmac.compare_digest(token, SFA_API_TOKEN):
                        self._send(b'{"error":"unauthorized"}', status=401, ctype="application/json")
                    else:
                        _st = (qs.get("status", ["議論中"])[0] or "議論中")
                        issues = sfa_db.list_deal_issues(con, status=(None if _st == "all" else _st))
                        self._send(json.dumps([dict(r) for r in issues], ensure_ascii=False, default=str).encode(),
                                   ctype="application/json")
                elif path == "/api/dev_projects":
                    # スプシ出力用: 開発案件一覧（deals/accounts JOIN込み）
                    qs = self._qs()
                    token = (qs.get("token", [None])[0] or "")
                    if not SFA_API_TOKEN or not hmac.compare_digest(token, SFA_API_TOKEN):
                        self._send(b'{"error":"unauthorized"}', status=401, ctype="application/json")
                    else:
                        status_q = (qs.get("status", ["open"])[0] or "open")
                        projects = sfa_db.list_dev_projects(con)
                        if status_q != "all":
                            projects = [p for p in projects if p.get("status") != "中止"]
                        self._send(json.dumps(projects, ensure_ascii=False, default=str).encode(), ctype="application/json")
                elif path == "/api/memo/list":
                    qs = self._qs()
                    token = (qs.get("token", [None])[0] or "")
                    if not SFA_API_TOKEN or not hmac.compare_digest(token, SFA_API_TOKEN):
                        self._send_cors_json(b'{"error":"unauthorized"}', status=401)
                    else:
                        theme_id_q = qs.get("theme_id", [None])[0]
                        if theme_id_q:
                            notes = con.execute(
                                "SELECT * FROM meeting_notes WHERE theme_id=? ORDER BY note_date ASC, created_at ASC LIMIT 100",
                                (int(theme_id_q),)
                            ).fetchall()
                        else:
                            notes = con.execute(
                                "SELECT * FROM meeting_notes ORDER BY note_date ASC, created_at ASC LIMIT 100"
                            ).fetchall()
                        self._send_cors_json(json.dumps([dict(r) for r in notes], ensure_ascii=False, default=str).encode())
                elif path == "/api/theme_deal_map":
                    # ダッシュボード用: theme_id → SFA deal_id マッピング
                    qs = self._qs()
                    token = (qs.get("token", [None])[0] or "")
                    if not SFA_API_TOKEN or not hmac.compare_digest(token, SFA_API_TOKEN):
                        self._send_cors_json(b'{"error":"unauthorized"}', status=401)
                    else:
                        rows = con.execute(
                            "SELECT id, theme_id FROM deals WHERE theme_id IS NOT NULL"
                        ).fetchall()
                        result = {str(row["theme_id"]): row["id"] for row in rows}
                        self._send_cors_json(json.dumps(result, ensure_ascii=False).encode())
                elif path == "/api/dev_capacity":
                    # ダッシュボード用: 開発担当の週次キャパ {owner: {base, from, base2}}（#42）
                    qs = self._qs()
                    token = (qs.get("token", [None])[0] or "")
                    if not SFA_API_TOKEN or not hmac.compare_digest(token, SFA_API_TOKEN):
                        self._send_cors_json(b'{"error":"unauthorized"}', status=401)
                    else:
                        self._send_cors_json(json.dumps(sfa_db.get_owner_capacities(con), ensure_ascii=False).encode())
                elif path == "/api/delivery_load":
                    # Hishoダッシュボード用: Delivery稼働(FTE%) を owner×week に展開・見込み/確定別（#75）。
                    qs = self._qs()
                    token = (qs.get("token", [None])[0] or "")
                    if not SFA_API_TOKEN or not hmac.compare_digest(token, SFA_API_TOKEN):
                        self._send_cors_json(b'{"error":"unauthorized"}', status=401)
                    else:
                        try:
                            _nw = int(qs.get("weeks", [str(sfa_db.DELIVERY_VIEW_WEEKS)])[0])
                        except (ValueError, TypeError):
                            _nw = sfa_db.DELIVERY_VIEW_WEEKS
                        _sw = (qs.get("start", [None])[0] or None)
                        _load = sfa_db.compute_delivery_load(con, start_week=_sw, n_weeks=_nw)
                        _load["thresholds"] = sfa_db.DELIVERY_HEAT_THRESHOLDS
                        _load["points_per_fte"] = sfa_db.POINTS_PER_FTE
                        self._send_cors_json(json.dumps(_load, ensure_ascii=False).encode())
                elif path == "/api/base_workload":
                    # Hishoダッシュボード用: ベース工数(人×機能×%)。{owner:pct}合算＋明細（#75）。
                    qs = self._qs()
                    token = (qs.get("token", [None])[0] or "")
                    if not SFA_API_TOKEN or not hmac.compare_digest(token, SFA_API_TOKEN):
                        self._send_cors_json(b'{"error":"unauthorized"}', status=401)
                    else:
                        self._send_cors_json(json.dumps({
                            "by_owner": sfa_db.base_workload_by_owner(con),
                            "items": sfa_db.list_base_workload(con),
                        }, ensure_ascii=False).encode())
                elif path == "/api/memo/list_all":
                    # スプシ出力用: 全メモ + deals/accounts JOIN
                    qs = self._qs()
                    token = (qs.get("token", [None])[0] or "")
                    if not SFA_API_TOKEN or not hmac.compare_digest(token, SFA_API_TOKEN):
                        self._send_cors_json(b'{"error":"unauthorized"}', status=401)
                    else:
                        rows = con.execute("""
                            SELECT m.id, m.note_date, m.body, m.task, m.task_owner,
                                   m.task_due, m.task_done, m.created_at,
                                   d.deal_name, a.name AS account_name
                            FROM meeting_notes m
                            LEFT JOIN deals d ON d.theme_id = m.theme_id
                            LEFT JOIN accounts a ON a.id = d.account_id
                            ORDER BY m.note_date DESC, m.created_at DESC
                        """).fetchall()
                        self._send_cors_json(json.dumps([dict(r) for r in rows], ensure_ascii=False, default=str).encode())
                elif path == "/":
                    # ホーム=商談一覧に統合（#38）。ダッシュボードは /dashboard へ。
                    qs = self._qs()
                    def _qs1(k): return (qs.get(k, [None])[0] or None)
                    _d = _qs1("date")
                    if _d:
                        try:
                            date.fromisoformat(_d)
                        except ValueError:
                            _d = None
                    self._send(render(deals_page(
                        con, tab=resolve_default_deals_tab(con, _qs1("tab"), _qs1("owner")),
                        owner=_qs1("owner"),
                        status_filter=_qs1("status"), stage_filter=_qs1("stage"), date=_d,
                        ms_type=_qs1("ms_type"), week=_qs1("week"),
                        exclude_today=(_qs1("exclude_today") != "0"),  # 既定で当日MS除外
                        stages_sel=[s for s in qs.get("stg", []) if s],
                    )))
                elif path == "/dashboard":
                    self._send(render(dashboard_page(con)))
                # ── メールパターン ──
                elif path == "/email-patterns":
                    self._send(render(email_patterns_page(con)))
                elif path == "/email-patterns/new":
                    self._send(render(email_pattern_form(con)))
                elif path == "/email-draft":
                    qs = self._qs()
                    self._send(render(email_draft_page(
                        con,
                        status_filter=(qs.get("status", [None])[0] or None),
                        q=(qs.get("q", [None])[0] or None),
                    )))
                elif path == "/email-draft/eml":
                    qs = self._qs()
                    try:
                        lid = int((qs.get("lead_id", [None])[0]) or 0)
                        pid = int((qs.get("pattern_id", [None])[0]) or 0)
                        lead = sfa_db.get_lead(con, lid)
                        p = sfa_db.get_email_pattern(con, pid)
                        if lead and p:
                            eml = build_eml_bytes(p, lead)
                            self.send_response(200)
                            self.send_header("Content-Type", "message/rfc822")
                            self.send_header("Content-Disposition", 'attachment; filename="draft.eml"')
                            self.send_header("Content-Length", str(len(eml)))
                            self.end_headers()
                            self.wfile.write(eml)
                        else:
                            self._send(b"Not found", 404)
                    except (ValueError, TypeError):
                        self._send(b"Bad request", 400)
                elif path.startswith("/email-patterns/") and path.endswith("/edit"):
                    try:
                        pid = int(path.split("/")[2])
                        p = sfa_db.get_email_pattern(con, pid)
                        self._send(render(email_pattern_form(con, p) if p else "<div class=card>見つかりません</div>"))
                    except (ValueError, IndexError):
                        self._send(render("<div class=card>見つかりません</div>"), 404)
                # ── 初回ヒアリング ──
                elif path == "/hearings/export" or path == "/hearings/export.docx":
                    qs = self._qs()
                    try:
                        tid = int(qs.get("template_id", ["0"])[0] or 0) or None
                    except ValueError:
                        tid = None
                    fmt = "docx" if path.endswith(".docx") else "xlsx"
                    try:
                        result_ids = [r["id"] for r in sfa_db.list_all_hearing_results(con, template_id=tid)]
                        data, fname, ctype = build_hearings_export_bundle(con, result_ids, fmt)
                        self.send_response(200)
                        self.send_header("Content-Type", ctype)
                        self.send_header("Content-Disposition", _content_disposition(fname))
                        self.send_header("Content-Length", str(len(data)))
                        self.end_headers()
                        self.wfile.write(data)
                    except Exception as _ex:
                        print(f"[hearings/export] {_ex}", flush=True)
                        import traceback as _tb; _tb.print_exc()
                        self._send(render("<div class=card>エクスポートに失敗しました</div>"), 500)
                elif path == "/hearings/export.csv":
                    qs = self._qs()
                    try:
                        tid = int(qs.get("template_id", ["0"])[0] or 0)
                    except ValueError:
                        tid = 0
                    if not tid:
                        self._send(render("<div class=card>テンプレートが指定されていません</div>"), 400)
                    else:
                        data = build_hearing_results_csv(con, tid)
                        self.send_response(200)
                        self.send_header("Content-Type", "text/csv; charset=utf-8")
                        self.send_header("Content-Disposition", f'attachment; filename="hearing_results_{tid}.csv"')
                        self.send_header("Content-Length", str(len(data)))
                        self.end_headers()
                        self.wfile.write(data)
                elif path == "/hearings":
                    qs = self._qs()
                    tid_raw = qs.get("template_id", [None])[0]
                    try:
                        tid = int(tid_raw) if tid_raw else None
                    except ValueError:
                        tid = None
                    self._send(render(hearings_page(con, template_id=tid)))
                elif path == "/hearing-templates/new":
                    self._send(render(hearing_template_form(con)))
                elif path == "/hearing-templates":
                    self._send(render(hearing_templates_page(con)))
                elif path.startswith("/hearing-templates/") and path.endswith("/edit"):
                    try:
                        tid = int(path.split("/")[2])
                        t = sfa_db.get_hearing_template(con, tid)
                        self._send(render(hearing_template_form(con, t) if t
                                          else "<div class=card>見つかりません</div>"))
                    except (ValueError, IndexError):
                        self._send(render("<div class=card>見つかりません</div>"), 404)
                elif path == "/hearing/new":
                    qs = self._qs()
                    self._send(render(hearing_new_page(con, preselect=(qs.get("target", [None])[0]))))
                elif path == "/hearing/start":
                    qs = self._qs()
                    target = (qs.get("target", [""])[0] or "")
                    try:
                        tid = int(qs.get("template_id", ["0"])[0] or 0)
                    except ValueError:
                        tid = 0
                    tmpl = sfa_db.get_hearing_template(con, tid) if tid else None
                    if not tmpl or ":" not in target:
                        self._send(render("<div class=card>対象またはテンプレートが不正です。"
                                          "<a href='/hearing/new'>戻る</a></div>"), 400)
                    else:
                        ttype, _, tval = target.partition(":")
                        try:
                            tval_id = int(tval)
                        except ValueError:
                            tval_id = 0
                        prefill, prev_date, label = None, None, ""
                        if ttype == "deal":
                            d = sfa_db.get_deal(con, tval_id)
                            if d:
                                label = f"{d.get('account_name') or ''} / {d.get('deal_name')}"
                                # 2回目以降: 同テンプレの直近結果をプリフィル
                                prev = [r for r in sfa_db.list_hearing_results(con, tval_id)
                                        if r.get("template_id") == tmpl["id"]]
                                if prev:
                                    prefill = {a.get("label"): a.get("answer")
                                               for a in (prev[0].get("answers") or [])}
                                    prev_date = prev[0].get("conducted_on")
                        elif ttype == "lead":
                            ld = sfa_db.get_lead(con, tval_id)
                            if ld:
                                label = f"{ld.get('company') or '?'} / {ld.get('name') or '?'}（リード）"
                        if not label:
                            self._send(render("<div class=card>対象が見つかりません</div>"), 404)
                        else:
                            draft = sfa_db.get_hearing_draft(
                                con, target_type=ttype, target_id=tval_id, template_id=tmpl["id"],
                            )
                            self._send(render(hearing_input_page(
                                con, target_type=ttype, target_id=tval_id, template=tmpl,
                                target_label=label, prefill=prefill, prev_date=prev_date,
                                draft=draft,
                            )))
                elif path.startswith("/hearing/result/") and path.endswith("/edit"):
                    try:
                        rid = int(path.split("/")[3])
                    except (ValueError, IndexError):
                        rid = 0
                    r = sfa_db.get_hearing_result(con, rid) if rid else None
                    tmpl = sfa_db.get_hearing_template(con, r["template_id"]) if r and r.get("template_id") else None
                    if not r or not tmpl:
                        self._send(render("<div class=card>編集対象が見つかりません"
                                          "（テンプレートが削除されている可能性があります）。"
                                          "<a href='/hearings'>一覧へ戻る</a></div>"), 404)
                    else:
                        label = f"{r.get('account_name') or ''} / {r.get('deal_name') or ''}"
                        prefill = {a.get("label"): a.get("answer") for a in (r.get("answers") or [])}
                        draft = sfa_db.get_hearing_draft(
                            con, target_type="hearing_edit", target_id=rid, template_id=tmpl["id"],
                        )
                        self._send(render(hearing_input_page(
                            con, target_type="deal", target_id=r["deal_id"], template=tmpl,
                            target_label=label, prefill=prefill, prev_date=None, draft=draft,
                            edit_result_id=rid, edit_conducted_on=r.get("conducted_on"),
                        )))
                elif path.startswith("/hearing/result/") and (path.endswith("/export.xlsx") or path.endswith("/export.docx")):
                    try:
                        rid = int(path.split("/")[3])
                    except (ValueError, IndexError):
                        rid = 0
                    r = sfa_db.get_hearing_result(con, rid) if rid else None
                    if not r:
                        self._send(render("<div class=card>ヒアリング結果が見つかりません</div>"), 404)
                    else:
                        fmt = "docx" if path.endswith(".docx") else "xlsx"
                        try:
                            data, fname, ctype = build_hearings_export_bundle(con, [rid], fmt)
                            self.send_response(200)
                            self.send_header("Content-Type", ctype)
                            self.send_header("Content-Disposition", _content_disposition(fname))
                            self.send_header("Content-Length", str(len(data)))
                            self.end_headers()
                            self.wfile.write(data)
                        except Exception as _ex:
                            print(f"[hearing/result/export] {_ex}", flush=True)
                            import traceback as _tb; _tb.print_exc()
                            self._send(render("<div class=card>エクスポートに失敗しました</div>"), 500)
                elif path.startswith("/hearing/result/"):
                    try:
                        rid = int(path.split("/")[3])
                        r = sfa_db.get_hearing_result(con, rid)
                        self._send(render(hearing_result_page(con, r) if r
                                          else "<div class=card>ヒアリング結果が見つかりません</div>"),
                                   200 if r else 404)
                    except (ValueError, IndexError):
                        self._send(render("<div class=card>ページが見つかりません</div>"), 404)
                elif path == "/deals":
                    qs = self._qs()
                    def qs1(k): return (qs.get(k, [None])[0] or None)
                    _date_q = qs1("date")
                    if _date_q:
                        try:
                            date.fromisoformat(_date_q)
                        except ValueError:
                            _date_q = None
                    self._send(render(deals_page(
                        con, tab=resolve_default_deals_tab(con, qs1("tab"), qs1("owner")),
                        owner=qs1("owner"),
                        status_filter=qs1("status"), stage_filter=qs1("stage"), date=_date_q,
                        ms_type=qs1("ms_type"), week=qs1("week"),
                        exclude_today=(qs1("exclude_today") != "0"),  # 既定で当日MS除外
                        stages_sel=[s for s in qs.get("stg", []) if s],
                    )))
                elif path == "/deals/import":
                    self._send(render(deals_import_page(con)))
                elif path == "/deals/import/template.csv":
                    body = deals_csv.build_template_csv()
                    self.send_response(200)
                    self.send_header("Content-Type", "text/csv; charset=utf-8")
                    self.send_header("Content-Disposition", 'attachment; filename="deals_import_template.csv"')
                    self.send_header("Content-Length", str(len(body)))
                    self.end_headers()
                    self.wfile.write(body)
                elif path == "/deals/import/ai_prompt.md":
                    body = deals_csv.build_ai_prompt(con).encode("utf-8")
                    self.send_response(200)
                    self.send_header("Content-Type", "text/markdown; charset=utf-8")
                    self.send_header("Content-Disposition", 'attachment; filename="deals_import_ai_prompt.md"')
                    self.send_header("Content-Length", str(len(body)))
                    self.end_headers()
                    self.wfile.write(body)
                elif path == "/masters":
                    self._send(render(masters_page(con)))
                elif path == "/dev-point-master":
                    self._send(render(dev_point_master_page(con)))
                elif path == "/tech-seed-master":
                    self._send(render(tech_seed_master_page(con)))
                elif path == "/tech-seed-tagging":
                    self._send(render(tech_seed_tagging_page(con)))
                elif path == "/tasks":
                    _tq = self._qs()
                    self._send(render(tasks_page(
                        con, assignee=(_tq.get("assignee", [""])[0] or None),
                        category=(_tq.get("category", [""])[0] or None),
                        project=(_tq.get("project", [""])[0] or None),
                        urgency=(_tq.get("urgency", [""])[0] or None))))
                elif path == "/tasks/digest":
                    self._send(render(tasks_digest_page(con)))
                elif path == "/task-projects":
                    self._send(render(task_projects_page(con)))
                elif path == "/tasks/new":
                    self._send(render(task_form(con)))
                elif (path.startswith("/tasks/") and path.endswith("/edit")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    _tk = sfa_db.get_task(con, int(path.split("/")[2]))
                    if _tk:
                        self._send(render(task_form(con, _tk)))
                    else:
                        self._redirect("/tasks")
                elif (path.startswith("/task/") and path.endswith("/notes")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    _tid = int(path.split("/")[2])
                    _kind = (self._qs().get("kind", [""])[0] or None)
                    _notes = sfa_db.list_task_notes(con, _tid, kind=_kind)
                    _tk = sfa_db.get_task(con, _tid)
                    self._send(json.dumps({"ok": True,
                        "summary": (_tk.get("summary") if _tk else None), "notes": [
                        {"body": n.get("body"), "author": n.get("author"),
                         "created_at": n.get("created_at")} for n in _notes]},
                        ensure_ascii=False).encode(), ctype="application/json")
                elif path == "/sync-health":
                    self._send(render(sync_health_page(con, theme_client)))
                elif path == "/weekly-numbers":
                    self._send(render(weekly_numbers_page(con)))
                elif path == "/weekly-numbers/audit":
                    _aq = (self._qs().get("as_of", [""])[0] or "").strip()
                    _aod = None
                    if _aq:
                        try:
                            _aod = date.fromisoformat(_aq)
                        except ValueError:
                            _aod = None
                    _exf = (self._qs().get("exh", [""])[0] or "").strip() or None
                    self._send(render(weekly_numbers_audit_page(con, as_of=_aod, exh_filter=_exf)))
                elif path == "/exhibition-tagging":
                    self._send(render(exhibition_tagging_page(con)))
                elif path == "/slack-memo-backfill":
                    try:
                        _off = max(0, int(self._qs().get("offset", ["0"])[0]))
                    except (ValueError, TypeError):
                        _off = 0
                    self._send(render(slack_memo_backfill_page(con, offset=_off)))
                elif path == "/data-tagging":
                    self._send(render(data_tagging_page(con)))
                elif path == "/deliveries":
                    self._send(render(deliveries_page(con)))
                elif path == "/base-workload":
                    self._send(render(base_workload_page(con)))
                elif (path.startswith("/delivery/") and len(path.split("/")) == 3
                      and path.split("/")[2].isdigit()):
                    self._send(render(delivery_form(con, int(path.split("/")[2]))))
                elif path == "/reports":
                    self._send(reports_index_page(con).encode("utf-8"))
                elif path == "/reports/manage":
                    _slug = (self._qs().get("slug", [""])[0] or "")
                    self._send(render(reports_manage_page(con, _slug)))
                elif path == "/reports/rail-preview":
                    # #39: 編集画面の数字プレビュー。指定週(ws=月曜)のレールHTML断片を返す。
                    import cowork.weekly_report as weekly_report
                    _ws = (self._qs().get("ws", [""])[0] or "").strip()
                    try:
                        _as_of = date.fromisoformat(_ws) if _ws else None
                        _nums = weekly_report.compute_weekly_numbers(con, as_of=_as_of)
                        _html = weekly_report.render_number_rail(_nums)
                    except Exception as _exc:  # noqa: BLE001
                        _html = f'<span class="muted" style="color:#b91c1c">集計エラー: {html.escape(str(_exc))}</span>'
                    self._send(_html.encode("utf-8"), ctype="text/html; charset=utf-8")
                elif path.startswith("/reports/"):
                    import cowork.weekly_report as weekly_report
                    _slug = path[len("/reports/"):].strip("/")
                    _rep = sfa_db.get_weekly_report(con, _slug) if _SLUG_RE.fullmatch(_slug or "") else None
                    if _rep:
                        # 数字レール自動注入: 対象週(week_start)の数字を集計してレール化し <!--NUMBERS--> と差替え
                        _rail = ""
                        if (_rep.get("html_body") or "").find("<!--NUMBERS-->") >= 0:
                            _ws = (_rep.get("week_start") or "").strip()
                            try:
                                _as_of = date.fromisoformat(_ws) if _ws else None
                                _nums = weekly_report.compute_weekly_numbers(con, as_of=_as_of)
                                _rail = weekly_report.render_number_rail(_nums)
                            except Exception:
                                _rail = ""
                        self._send(report_article_html(_rep, _rail).encode("utf-8"))
                    else:
                        self._redirect("/reports")
                elif path == "/backups":
                    self._send(render(backups_page(db_path)))
                elif path == "/backups/download":
                    name = (self._qs().get("name", [""])[0] or "")
                    # list_backupsに載る名前のみ許可（パストラバーサル防止）
                    valid = {b["name"] for b in sfa_db.list_backups(db_path)}
                    if name not in valid:
                        self._send(render("<div class=card>バックアップが見つかりません</div>"), 404)
                    else:
                        fpath = sfa_db._backup_dir(db_path) / name
                        data = fpath.read_bytes()
                        self.send_response(200)
                        self.send_header("Content-Type", "application/octet-stream")
                        self.send_header("Content-Disposition", _content_disposition(name))
                        self.send_header("Content-Length", str(len(data)))
                        self.end_headers()
                        self.wfile.write(data)
                elif path == "/activity/new":
                    self._send(render(activity_deal_picker(con)))
                elif path == "/deal/new":
                    self._send(render(deal_form(con)))
                # ── 開発案件 ──
                elif path == "/dev-projects":
                    qs = self._qs()
                    def qs1(k): return (qs.get(k, [None])[0] or None)
                    self._send(render(dev_projects_list_page(
                        con, dev_owner=qs1("dev_owner"), sales_owner=qs1("sales_owner"),
                        status=qs1("status"), stage=qs1("stage"),
                        order_potential=qs1("order_potential"), deadline_week=qs1("deadline_week"),
                        theme_client=theme_client,
                    )))
                elif path == "/dev-projects/new":
                    qs = self._qs()
                    did_raw = qs.get("deal_id", [None])[0]
                    try:
                        did = int(did_raw) if did_raw else None
                    except ValueError:
                        did = None
                    return_to = qs.get("return_to", [None])[0]
                    self._send(render(dev_project_form(con, deal_id=did, return_to=return_to)))
                elif path.startswith("/dev-project/") and path.endswith("/edit"):
                    try:
                        pid = int(path.split("/")[2])
                        proj = sfa_db.get_dev_project(con, pid)
                        return_to = self._qs().get("return_to", [None])[0]
                        self._send(
                            render(dev_project_form(con, proj, return_to=return_to) if proj
                                   else "<div class=card>開発案件が見つかりません</div>"),
                            200 if proj else 404,
                        )
                    except (ValueError, IndexError):
                        self._send(render("<div class=card>ページが見つかりません</div>"), 404)
                # ── 社内論点 ──
                elif path == "/deal-issues":
                    qs = self._qs()
                    def qs1(k): return (qs.get(k, [None])[0] or None)
                    # フィルタフォーム未送信(初回訪問・リセットリンク)はデフォルトで議論中のみ表示。
                    # フォーム送信済み(sortが必ず送られる)なら、status=""(「全て」明示選択)を
                    # そのまま絞り込みなし(None)として扱う。
                    # 注: _qs()の parse_qs は空値パラメータを保持しないため、statusキーの有無では
                    # 「フォーム未送信」と「全て選択」を区別できない。sortの有無で判定する。
                    _status_default = "議論中" if "sort" not in qs else qs1("status")
                    self._send(render(deal_issues_list_page(
                        con, status=_status_default, member=qs1("member"),
                        q=qs1("q"), sort=qs1("sort"), open_issue=qs1("open_issue"),
                    )))
                elif path == "/deal-issue/new":
                    qs = self._qs()
                    did_raw = qs.get("deal_id", [None])[0]
                    try:
                        did = int(did_raw) if did_raw else None
                    except ValueError:
                        did = None
                    return_to = qs.get("return_to", [None])[0]
                    self._send(render(deal_issue_form(con, deal_id=did, return_to=return_to)))
                elif path.startswith("/deal-issue/") and path.endswith("/edit"):
                    try:
                        iid = int(path.split("/")[2])
                        iss = sfa_db.get_deal_issue(con, iid)
                        return_to = self._qs().get("return_to", [None])[0]
                        self._send(
                            render(deal_issue_form(con, iss, return_to=return_to) if iss
                                   else "<div class=card>論点が見つかりません</div>"),
                            200 if iss else 404,
                        )
                    except (ValueError, IndexError):
                        self._send(render("<div class=card>ページが見つかりません</div>"), 404)
                elif path == "/accounts":
                    self._send(render(accounts_page(con)))
                elif path == "/accounts/duplicates":
                    self._send(render(account_duplicates_page(con)))
                elif path == "/account/new":
                    self._send(render(account_form(con)))
                # ── リード ──
                elif path == "/leads":
                    qs = self._qs()
                    def qs1(k): return (qs.get(k, [None])[0] or None)
                    self._send(render(leads_page(
                        con, status=qs1("status"), source=qs1("source"), q=qs1("q"),
                    )))
                elif path == "/leads/new":
                    self._send(render(lead_form(con)))
                elif path == "/leads/import":
                    self._send(render(leads_import_page(con)))
                elif path == "/leads/import/template.csv":
                    body = leads_csv.build_template_csv()
                    self.send_response(200)
                    self.send_header("Content-Type", "text/csv; charset=utf-8")
                    self.send_header("Content-Disposition", 'attachment; filename="leads_import_template.csv"')
                    self.send_header("Content-Length", str(len(body)))
                    self.end_headers()
                    self.wfile.write(body)
                elif path == "/leads/import/ai_prompt.md":
                    body = leads_csv.build_ai_prompt(con).encode("utf-8")
                    self.send_response(200)
                    self.send_header("Content-Type", "text/markdown; charset=utf-8")
                    self.send_header("Content-Disposition", 'attachment; filename="leads_import_ai_prompt.md"')
                    self.send_header("Content-Length", str(len(body)))
                    self.end_headers()
                    self.wfile.write(body)
                elif path.startswith("/leads/") and path.endswith("/convert"):
                    try:
                        lid = int(path.split("/")[2])
                    except (ValueError, IndexError):
                        self._send(render("<div class=card>ページが見つかりません</div>"), 404)
                    else:
                        lead = sfa_db.get_lead(con, lid)
                        if not lead:
                            self._send(render("<div class=card>リードが見つかりません</div>"), 404)
                        else:
                            # GETではDBに書き込まない（リンクプリフェッチ等での誤変換防止）。
                            # 過去商談があれば選択ページ、無ければ確認ページを表示しPOSTで実行する。
                            past_deals = find_past_closed_deals(con, lead.get("company"))
                            if past_deals:
                                self._send(render(lead_convert_choice_page(con, lead, past_deals)))
                            else:
                                confirm_html = f"""
                                <div class="card" style="max-width:560px">
                                  <h2>リードを商談化</h2>
                                  <p>{_esc(lead.get('company'))} / {_esc(lead.get('name'))} を商談化します。</p>
                                  <form method="post" action="/leads/{lid}/convert">
                                    <input type="hidden" name="mode" value="new">
                                    <button class="btn" type="submit">商談化する</button>
                                    <a class="btn sec" href="/leads/{lid}">キャンセル</a>
                                  </form>
                                </div>"""
                                self._send(render(confirm_html))
                elif path.startswith("/leads/"):
                    try:
                        lid = int(path.split("/")[2])
                        lead = sfa_db.get_lead(con, lid)
                        if lead:
                            self._send(render(lead_form(con, lead)))
                        else:
                            self._send(render("<div class=card>リードが見つかりません</div>"), 404)
                    except (ValueError, IndexError):
                        self._send(render("<div class=card>ページが見つかりません</div>"), 404)
                # ── 商談・アカウント ──
                elif path.startswith("/deal/"):
                    try:
                        did = int(path.split("/")[2])
                    except (ValueError, IndexError):
                        self._send(render("<div class=card>ページが見つかりません</div>"), 404)
                    else:
                        deal = sfa_db.get_deal(con, did)
                        _rt = self._qs().get("return_to", [None])[0]
                        self._send(
                            render(deal_form(con, deal, return_to=_rt)) if deal
                            else render("<div class=card>商談が見つかりません</div>"),
                            200 if deal else 404,
                        )
                elif path.startswith("/account/"):
                    parts = path.split("/")
                    try:
                        aid = int(parts[2])
                    except (ValueError, IndexError):
                        self._send(render("<div class=card>ページが見つかりません</div>"), 404)
                    else:
                        acc = con.execute("SELECT * FROM accounts WHERE id=?", (aid,)).fetchone()
                        if len(parts) >= 4 and parts[3] == "edit":
                            self._send(render(account_form(con, dict(acc) if acc else None)))
                        else:
                            self._send(
                                render(account_detail(con, dict(acc))) if acc
                                else render("<div class=card>アカウントが見つかりません</div>"),
                                200 if acc else 404,
                            )
                else:
                    self._send(render("<div class=card>ページが見つかりません</div>"), 404)
            finally:
                con.close()

        def do_POST(self):
            if not self._check_basic_auth():
                return
            path = self.path.split("?")[0].rstrip("/")
            con = sfa_db.connect(db_path)
            ctype = self.headers.get("Content-Type", "")
            try:
                if "multipart/form-data" in ctype:
                    f = self._form_multi()
                    f_list = {}  # multipart returns single values; list values handled separately
                else:
                    n = int(self.headers.get("Content-Length", 0))
                    raw = self.rfile.read(n).decode("utf-8")
                    import urllib.parse as _up
                    d = _up.parse_qs(raw, keep_blank_values=True)
                    f_list = {k: v for k, v in d.items()}
                    f = {k: (v[0] if v else "") for k, v in d.items()}

                # ── ログイン（フォーム認証・Cookieセッション付与, #54） ──
                if path == "/login":
                    _u = f.get("username", "")
                    _p = f.get("password", "")
                    _nxt = f.get("next", "/") or "/"
                    if not _nxt.startswith("/"):
                        _nxt = "/"
                    if (SFA_BASIC_USER and SFA_BASIC_PASS
                            and hmac.compare_digest(_u, SFA_BASIC_USER)
                            and hmac.compare_digest(_p, SFA_BASIC_PASS)):
                        _secure = "; Secure" if self.headers.get("X-Forwarded-Proto", "") == "https" else ""
                        self.send_response(303)
                        self.send_header("Location", _nxt)
                        self.send_header("Set-Cookie",
                                         f"{_SESSION_COOKIE}={_make_session_token()}; Path=/; HttpOnly; "
                                         f"SameSite=Lax; Max-Age={_SESSION_MAX_AGE}{_secure}")
                        self.send_header("Content-Length", "0")
                        self.end_headers()
                    else:
                        self._send(login_page(_nxt, error="ユーザー名またはパスワードが違います。"),
                                   status=401, ctype="text/html; charset=utf-8")

                # ── マスタ ──
                elif path == "/masters/save":
                    for key in sfa_db.MASTER_KEYS:
                        values = f_list.get(f"{key}[]", [])
                        values = [v.strip() for v in values if v.strip()]
                        sfa_db.set_master_list(con, key, values)
                    self._redirect("/")

                # ── 技術シード マスタ（ツリー L1→L2, #60） ──
                elif path == "/tech-seed-master/save":
                    _names = f_list.get("l1_name[]", [])
                    _lines = f_list.get("l2_lines[]", [])
                    _tree = {}
                    for _i, _nm in enumerate(_names):
                        _nm = (_nm or "").strip()
                        if not _nm:
                            continue
                        _raw = _lines[_i] if _i < len(_lines) else ""
                        _leaves = [ln.strip() for ln in _raw.replace("\r", "").split("\n") if ln.strip()]
                        _tree[_nm] = _leaves
                    sfa_db.set_tech_seed_tree(con, _tree)
                    self._redirect("/tech-seed-master")

                # ── 技術シード 一括タグ付け（#58）: 表示中の全開発案件のシードをまとめて保存 ──
                elif path == "/tech-seed-tagging/save":
                    for _pid_s in f_list.get("pids[]", []):
                        try:
                            _pid = int(_pid_s)
                        except (ValueError, TypeError):
                            continue
                        _seeds = ",".join(s for s in f_list.get(f"seeds__{_pid}", []) if s) or None
                        con.execute("UPDATE dev_projects SET tech_seeds=?, updated_at=datetime('now') WHERE id=?",
                                    (_seeds, _pid))
                    con.commit()
                    self._redirect("/tech-seed-tagging")

                # ── タスク管理（#30） ──
                elif path == "/tasks/save":
                    try:
                        _tid = int(f["id"]) if f.get("id") else None
                    except ValueError:
                        _tid = None
                    _link = f.get("link", "") or ""
                    _lt, _li = None, None
                    if _link.startswith("dev_project:"):
                        _lt = "dev_project"
                        try:
                            _li = int(_link.split(":", 1)[1])
                        except ValueError:
                            _li = None
                    _cat = f.get("category") or None
                    if not _cat:  # 種類が空ならAI(Haiku)で自動判定（間違えたら手修正）
                        _cat = _ai_guess_task_category(f.get("title") or "", f.get("detail") or "")
                    _saved_id = sfa_db.upsert_task(
                        con, id=_tid,
                        title=f.get("title") or "(無題)",
                        detail=f.get("detail") or None,
                        project=f.get("project") or None,
                        next_action=f.get("next_action") or None,
                        assignee=f.get("assignee") or None,
                        due_date=f.get("due_date") or None,
                        status=f.get("status") or "受信箱",
                        category=_cat,
                        link_type=_lt, link_id=_li, source="web",
                    )
                    # 担当＋期限が揃っていれば受信箱→未着手へ自動整理
                    _task_auto_triage(con, _saved_id)
                    self._redirect("/tasks")

                elif path == "/task-projects/save":
                    _pid = None
                    if f.get("id"):
                        try:
                            _pid = int(f["id"])
                        except ValueError:
                            _pid = None
                    _pname = (f.get("name") or "").strip()
                    if _pname:
                        sfa_db.upsert_task_project(
                            con, id=_pid, name=_pname,
                            deadline=(f.get("deadline") or None),
                            status=(f.get("status") or "進行中"))
                    self._redirect("/task-projects")
                elif (path.startswith("/task-project/") and path.endswith("/delete")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    sfa_db.delete_task_project(con, int(path.split("/")[2]))
                    self._redirect("/task-projects")
                elif (path.startswith("/task-project/") and path.endswith("/summary")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    _pid = int(path.split("/")[2])
                    _p = next((x for x in sfa_db.list_task_projects(con) if x["id"] == _pid), None)
                    _summ = None
                    if _p:
                        _tks = [t for t in sfa_db.list_tasks(con) if (t.get("project") or "") == _p["name"]]
                        _lines = []
                        for t in _tks:
                            _s = (t.get("summary") or t.get("next_action") or "").strip()
                            _lines.append(f"[{t.get('status')}] {t.get('title')}" + (f" — {_s}" if _s else ""))
                        _meta = (f"期限{_p.get('deadline') or '—'} / 状態{_p.get('status') or '進行中'} "
                                 f"/ タスク{len(_tks)}件")
                        _summ = _ai_summarize_project(_p["name"], _meta, _lines)
                        if _summ:
                            sfa_db.set_project_summary(con, _p["name"], _summ)
                    self._send(json.dumps({"ok": bool(_summ), "summary": _summ},
                                          ensure_ascii=False).encode(), ctype="application/json")

                elif path == "/tasks/seed-test":
                    _n = sfa_db.seed_sample_tasks(con)
                    self._send(render(tasks_page(con),
                                      flash=f"🧪 テストデータを {_n} 件投入しました（【テスト】で始まるタスク）。"))
                elif path == "/tasks/delete-test":
                    _n = sfa_db.delete_test_tasks(con)
                    self._send(render(tasks_page(con), flash=f"🧪 テストデータを {_n} 件削除しました。"))

                elif path == "/tasks/digest/send":
                    _only = (f.get("only") or "").strip() or None
                    _res = send_task_digests(con, only=_only)
                    self._send(render(tasks_digest_page(con, result=_res)))

                elif (path.startswith("/task/") and path.endswith("/note")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    _tid = int(path.split("/")[2])
                    _body = (f.get("body") or "").strip()
                    _kind = "discussion" if f.get("kind") == "discussion" else "progress"
                    _new_status = None
                    _summary = None
                    if _body:
                        sfa_db.add_task_note(con, _tid, _body, kind=_kind)
                        if _kind == "progress":
                            _new_status = _task_auto_start(con, _tid)  # 進捗追記＝着手→対応中へ
                        else:
                            # 議論メモ→タスクのAIサマリを自動再生成
                            _memos = [n["body"] for n in sfa_db.list_task_notes(con, _tid, kind="discussion")]
                            _tk = sfa_db.get_task(con, _tid)
                            _summary = _ai_summarize_task(_tk.get("title", "") if _tk else "", _memos)
                            if _summary:
                                sfa_db.set_task_summary(con, _tid, _summary)
                    if f.get("ajax") == "1":
                        _notes = sfa_db.list_task_notes(con, _tid, kind=_kind)
                        self._send(json.dumps({"ok": True, "status": _new_status, "summary": _summary, "notes": [
                            {"body": n.get("body"), "author": n.get("author"),
                             "created_at": n.get("created_at")} for n in _notes]},
                            ensure_ascii=False).encode(), ctype="application/json")
                    else:
                        self._redirect(f"/tasks/{_tid}/edit")

                elif (path.startswith("/task/") and path.endswith("/field")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    _tid = int(path.split("/")[2])
                    _field = f.get("field", "")
                    _value = f.get("value", "")
                    _allowed = {"status", "assignee", "due_date", "category", "priority",
                                "title", "project", "next_action", "pinned"}
                    if _field not in _allowed:
                        self._send(json.dumps({"ok": False, "error": "不正なフィールド"}).encode(), ctype="application/json")
                    elif _field == "status" and _value not in sfa_db.TASK_STATUSES:
                        self._send(json.dumps({"ok": False, "error": "不正な状態"}).encode(), ctype="application/json")
                    elif _field == "priority" and _value and _value not in sfa_db.TASK_PRIORITIES:
                        self._send(json.dumps({"ok": False, "error": "不正な優先度"}).encode(), ctype="application/json")
                    elif _field == "status":
                        sfa_db.set_task_status(con, _tid, _value)
                        self._send(json.dumps({"ok": True, "status": _value}, ensure_ascii=False).encode(),
                                   ctype="application/json")
                    elif _field == "pinned":
                        con.execute("UPDATE tasks SET pinned=?, updated_at=datetime('now') WHERE id=?",
                                    (1 if _value in ("1", "true", "on") else 0, _tid))
                        con.commit()
                        self._send(json.dumps({"ok": True}).encode(), ctype="application/json")
                    else:
                        con.execute(f"UPDATE tasks SET {_field}=?, updated_at=datetime('now') WHERE id=?",
                                    (_value or None, _tid))
                        con.commit()
                        # 担当＋期限が揃ったら受信箱→未着手へ自動整理。現在のstatusを返す。
                        _st = _task_auto_triage(con, _tid)
                        self._send(json.dumps({"ok": True, "status": _st}, ensure_ascii=False).encode(),
                                   ctype="application/json")

                elif (path.startswith("/task/") and path.endswith("/ai-category")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    _tid = int(path.split("/")[2])
                    _tk = sfa_db.get_task(con, _tid)
                    _cat = _ai_guess_task_category(_tk.get("title", ""), _tk.get("detail") or "") if _tk else None
                    if _cat:
                        con.execute("UPDATE tasks SET category=?, updated_at=datetime('now') WHERE id=?",
                                    (_cat, _tid))
                        con.commit()
                    self._send(json.dumps({"ok": True, "category": _cat}, ensure_ascii=False).encode(),
                               ctype="application/json")

                elif (path.startswith("/task/") and path.endswith("/delete")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    sfa_db.delete_task(con, int(path.split("/")[2]))
                    if f.get("ajax") == "1":
                        self._send(json.dumps({"ok": True}).encode(), ctype="application/json")
                    else:
                        self._redirect("/tasks")

                # ── 開発点数マスタ / 担当キャパ（#41） ──
                elif path == "/dev-point-master/save":
                    # 既存行の更新（作業種別名も編集可・id基準）
                    for k, v in f.items():
                        if not k.startswith("wt__") or not k[4:].isdigit():
                            continue
                        _mid = k[4:]
                        _name = (v or "").strip()
                        _bp = (f.get(f"bp__{_mid}") or "").strip()
                        if _name and _bp:
                            try:
                                sfa_db.update_dev_point_master(con, int(_mid), _name, float(_bp))
                            except ValueError:
                                pass
                    # 追加行
                    _nwt = (f.get("new_work_type") or "").strip()
                    _nbp = (f.get("new_base_points") or "").strip()
                    if _nwt and _nbp:
                        try:
                            sfa_db.upsert_dev_point_master(con, _nwt, float(_nbp))
                        except ValueError:
                            pass
                    self._redirect("/dev-point-master")
                elif path == "/dev-point-master/delete":
                    _mid = (f.get("id") or "").strip()
                    if _mid.isdigit():
                        sfa_db.delete_dev_point_master_by_id(con, int(_mid))
                    self._redirect("/dev-point-master")
                elif path == "/dev-point-master/coef":
                    for k, v in f.items():
                        if not k.startswith("cf__") or not str(v).strip():
                            continue
                        _p = k[4:].split("__", 1)
                        if len(_p) != 2 or _p[0] not in ("stage", "difficulty", "backend"):
                            continue
                        try:
                            sfa_db.set_dev_coef(con, _p[0], _p[1], float(v))
                        except ValueError:
                            continue
                    self._redirect("/dev-point-master")
                elif path == "/dev-point-master/capacity":
                    _capin: dict = {}
                    for k, v in f.items():
                        if not k.startswith("cap__") or "__" not in k[5:]:
                            continue
                        _owner, _sub = k[5:].rsplit("__", 1)
                        _capin.setdefault(_owner, {})[_sub] = v
                    for _owner, _d in _capin.items():
                        _base = (_d.get("base") or "").strip()
                        if not _base:
                            continue
                        try:
                            _bv = float(_base)
                        except ValueError:
                            continue
                        _from = (_d.get("from") or "").strip() or None
                        _p2s = (_d.get("p2") or "").strip()
                        _p2v = None
                        if _p2s:
                            try:
                                _p2v = float(_p2s)
                            except ValueError:
                                _p2v = None
                        sfa_db.set_owner_capacity(con, _owner, _bv, _from, _p2v)
                    self._redirect("/dev-point-master")

                # ── Delivery（受注後・アサイン計画。#75） ──
                elif path == "/deliveries/bulk_new":
                    # 商談一覧でチェックした商談にDeliveryを登録（既に有ればスキップ＝重複防止）
                    _ids = [x for x in f_list.get("ids", []) if str(x).isdigit()]
                    _created = []
                    for _idr in _ids:
                        _dealid = int(_idr)
                        _deal = sfa_db.get_deal(con, _dealid)
                        if not _deal:
                            continue
                        if sfa_db.list_deliveries(con, deal_id=_dealid):
                            continue  # 既存があれば作らない
                        _created.append(sfa_db.create_delivery(
                            con, deal_id=_dealid, title=_deal.get("deal_name") or ""))
                    if len(_created) == 1:
                        self._redirect(f"/delivery/{_created[0]}")
                    else:
                        self._redirect("/deliveries")
                elif path == "/deliveries/new":
                    _did = (f.get("deal_id", "") or "").strip()
                    if _did.isdigit() and sfa_db.get_deal(con, int(_did)):
                        _nid = sfa_db.create_delivery(
                            con, deal_id=int(_did),
                            title=(sfa_db.get_deal(con, int(_did)) or {}).get("deal_name") or "")
                        self._redirect(f"/delivery/{_nid}")
                    else:
                        self._redirect("/deliveries")
                elif (path.startswith("/delivery/") and path.endswith("/save")
                      and path.split("/")[2].isdigit()):
                    _dvid = int(path.split("/")[2])
                    sfa_db.update_delivery(
                        con, _dvid,
                        title=(f.get("title", "") or "").strip(),
                        start_week=_snap_monday(f.get("start_week", "")),
                        end_week=_snap_monday(f.get("end_week", "")),
                        status=(f.get("status", "") or "進行中").strip(),
                        overview=(f.get("overview", "") or "").strip())
                    self._redirect(f"/delivery/{_dvid}")
                elif (path.startswith("/delivery/") and path.endswith("/assignment/add")
                      and path.split("/")[2].isdigit()):
                    _dvid = int(path.split("/")[2])
                    _ow = (f.get("owner", "") or "").strip()
                    _fw = _snap_monday(f.get("from_week", ""))
                    _tw = _snap_monday(f.get("to_week", ""))
                    if _ow and _fw and _tw:
                        if _tw < _fw:
                            _fw, _tw = _tw, _fw   # 逆順は入れ替え
                        try:
                            _fte = float(f.get("fte_pct", "0") or 0)
                        except ValueError:
                            _fte = 0.0
                        sfa_db.add_delivery_assignment(
                            con, delivery_id=_dvid, owner=_ow, from_week=_fw, to_week=_tw,
                            fte_pct=_fte, note=(f.get("note", "") or "").strip())
                    self._redirect(f"/delivery/{_dvid}")
                elif (path.startswith("/delivery/") and "/assignment/" in path
                      and path.endswith("/delete") and path.split("/")[2].isdigit()):
                    _dvid = int(path.split("/")[2])
                    _aid = path.split("/")[4]
                    if _aid.isdigit():
                        sfa_db.delete_delivery_assignment(con, int(_aid))
                    self._redirect(f"/delivery/{_dvid}")
                elif (path.startswith("/delivery/") and path.endswith("/delete")
                      and path.split("/")[2].isdigit() and len(path.split("/")) == 4):
                    _dvid = int(path.split("/")[2])
                    sfa_db.delete_delivery(con, _dvid)
                    self._redirect("/deliveries")
                elif path == "/base-workload/save":
                    _ow = (f.get("owner", "") or "").strip()
                    _fn = (f.get("function", "") or "").strip()
                    if _ow and _fn:
                        try:
                            _p = float(f.get("pct", "0") or 0)
                        except ValueError:
                            _p = 0.0
                        sfa_db.upsert_base_workload(con, _ow, _fn, _p)
                    self._redirect("/base-workload")
                elif (path.startswith("/base-workload/") and path.endswith("/delete")
                      and path.split("/")[2].isdigit()):
                    sfa_db.delete_base_workload(con, int(path.split("/")[2]))
                    self._redirect("/base-workload")

                # ── 週次レポート（本文はDBのみ・Git非保存） ──
                elif path == "/reports/manage/save":
                    _slug = (f.get("slug", "") or "").strip()
                    _body = f.get("html_body", "") or ""
                    if not _SLUG_RE.fullmatch(_slug) or not _body.strip():
                        self._redirect("/reports/manage")
                    else:
                        _cover = f.get("cover_image", "") or ""
                        if not _cover.startswith("data:image/"):
                            _cover = ""
                        sfa_db.upsert_weekly_report(
                            con, _slug,
                            (f.get("report_date", "") or "").strip(),
                            (f.get("title", "") or "").strip(),
                            (f.get("lead", "") or "").strip(),
                            _body,
                            _cover,
                            week_start=(f.get("week_start", "") or "").strip(),
                        )
                        self._redirect(f"/reports/{_slug}")
                elif path == "/reports/manage/delete":
                    _slug = (f.get("slug", "") or "").strip()
                    if _SLUG_RE.fullmatch(_slug):
                        sfa_db.delete_weekly_report(con, _slug)
                    self._redirect("/reports")

                # ── バックアップ ──
                elif path == "/backups/create":
                    try:
                        path_out = sfa_db.backup_now(db_path, tag="manual")
                        import os as _os
                        _name = _os.path.basename(path_out) if path_out else "?"
                        self._send(render(backups_page(db_path),
                                          flash=f"バックアップを作成しました: {_name}"))
                    except Exception as _exc:  # noqa: BLE001
                        self._send(render(backups_page(db_path), flash=f"バックアップ作成に失敗: {_exc}"))
                elif path == "/backups/restore":
                    name = f.get("name") or ""
                    try:
                        pre = sfa_db.restore_backup(db_path, name)
                        import os as _os
                        self._send(render(backups_page(db_path),
                                          flash=f"「{name}」から復元しました。復元前の状態は "
                                                f"{_os.path.basename(pre)} として退避済みです。"))
                    except Exception as _exc:  # noqa: BLE001
                        self._send(render(backups_page(db_path), flash=f"復元に失敗: {_exc}"))

                # ── アカウント ──
                elif path == "/account/save":
                    try:
                        _acc_id = int(f["id"]) if f.get("id") else None
                    except ValueError:
                        _acc_id = None
                    saved_acc_id = sfa_db.upsert_account(
                        con, id=_acc_id,
                        name=f.get("name") or "(無名)",
                        industry=f.get("industry") or None,
                        company_size=f.get("company_size") or None,
                        note=f.get("note") or None,
                    )
                    self._redirect(f"/account/{saved_acc_id}")

                elif path.startswith("/account/") and path.endswith("/delete"):
                    parts = path.split("/")
                    if len(parts) == 4 and parts[3] == "delete" and parts[2].isdigit():
                        con.execute("DELETE FROM accounts WHERE id=?", (int(parts[2]),))
                        con.commit()
                    self._redirect("/accounts")

                elif path == "/accounts/bulk_delete":
                    ids = f_list.get("ids", [])
                    for acc_id in ids:
                        if str(acc_id).isdigit():
                            con.execute("DELETE FROM accounts WHERE id=?", (int(acc_id),))
                    if ids:
                        con.commit()
                    self._redirect("/accounts")

                elif path == "/slack-memo-backfill/apply":
                    # 過去のSlack追記メモを、確認済みの内容で現状メモへ追記（1件ずつ）。
                    try:
                        _did = int(f["deal_id"])
                    except (ValueError, KeyError):
                        self._send(json.dumps({"ok": False, "error": "不正なID"}).encode(),
                                   ctype="application/json")
                    else:
                        _memo = (f.get("memo") or "").strip()
                        _deal = sfa_db.get_deal(con, _did)
                        if not _memo or not _deal:
                            self._send(json.dumps({"ok": False, "error": "メモ空 or 商談なし"}).encode(),
                                       ctype="application/json")
                        else:
                            _existing = _deal.get("note") or ""
                            _marker = f"[Slack復旧 {_today_jst().isoformat()}] "
                            _new = (_existing + "\n" + _marker + _memo).strip() if _existing else (_marker + _memo)
                            con.execute("UPDATE deals SET note=?, updated_at=datetime('now') WHERE id=?",
                                        (_new, _did))
                            con.commit()
                            if theme_client is not None:
                                try:
                                    theme_link.sync_deal(theme_client, con, _did)
                                except Exception as _exc:  # noqa: BLE001
                                    print(f"[theme_link] sync_deal failed (memo backfill): {_exc}")
                            self._send(json.dumps({"ok": True}).encode(), ctype="application/json")

                elif path == "/exhibition-tagging/bulk":
                    # 選択した展示会由来商談に展示会名を一括設定（空なら未設定へクリア）。
                    _ids = [int(i) for i in f_list.get("ids", []) if str(i).isdigit()]
                    _exname = (f.get("exhibition_name") or "").strip() or None
                    _n = 0
                    for _did in _ids:
                        con.execute("UPDATE deals SET exhibition_name=? WHERE id=?", (_exname, _did))
                        _n += 1
                        if theme_client is not None:
                            try:
                                theme_link.sync_deal(theme_client, con, _did)
                            except Exception as _exc:  # noqa: BLE001
                                print(f"[theme_link] sync_deal failed (exhibition bulk): {_exc}")
                    if _ids:
                        con.commit()
                    self._send(render(exhibition_tagging_page(con),
                                      flash=f"{_n}件の展示会名を「{_exname or '(未設定)'}」に設定しました。"))

                elif path == "/accounts/merge":
                    try:
                        keep_id = int(f.get("keep_id"))
                        all_ids = [int(x) for x in (f.get("all_ids") or "").split(",") if x.strip().isdigit()]
                    except (ValueError, TypeError):
                        self._send(render(account_duplicates_page(con), flash="統合対象の指定が不正です。"))
                    else:
                        drop_ids = [i for i in all_ids if i != keep_id]
                        if not drop_ids:
                            self._send(render(account_duplicates_page(con), flash="統合対象がありません。"))
                        else:
                            # 破壊的操作の前に必ずバックアップ
                            try:
                                sfa_db.backup_now(db_path, tag="premerge")
                            except Exception as _bexc:  # noqa: BLE001
                                print(f"[merge] pre-merge backup failed: {_bexc}", flush=True)
                            res = sfa_db.merge_accounts(con, keep_id=keep_id, drop_ids=drop_ids)
                            # 統合で商談のアカウントが変わったのでテーマDBへ再同期（該当商談のみ）
                            if theme_client is not None:
                                for drow in con.execute(
                                        "SELECT id FROM deals WHERE account_id=?", (keep_id,)):
                                    try:
                                        theme_link.sync_deal(theme_client, con, drow["id"])
                                    except Exception as _sexc:  # noqa: BLE001
                                        sfa_db.record_sync_failure(con, "deal", drow["id"], str(_sexc))
                            self._send(render(account_duplicates_page(con),
                                              flash=f"統合しました（商談{res['moved_deals']}件・"
                                                    f"コンタクト{res['moved_contacts']}件を移動、"
                                                    f"{res['dropped']}件のアカウントを削除）。"))

                # ── 商談一括編集 ──
                elif path == "/deals/bulk_edit":
                    _DEAL_ALLOWED = {"stage", "owner", "sub_owner", "business_type_l1"}
                    ids = f_list.get("ids", [])
                    field = f.get("field", "")
                    value = f.get("value", "")
                    if field in _DEAL_ALLOWED and ids:
                        if field == "stage":
                            valid = sfa_db.get_master_list(con, "deal_stages")
                            if value and value not in valid:
                                self._redirect("/deals")
                                return
                        for did in ids:
                            if str(did).isdigit():
                                con.execute(
                                    f"UPDATE deals SET {field}=?, updated_at=datetime('now') WHERE id=?",
                                    (value or None, int(did)),
                                )
                        con.commit()
                        if field == "stage" and theme_client is not None:
                            for did in ids:
                                if str(did).isdigit():
                                    try:
                                        theme_link.sync_deal(theme_client, con, int(did))
                                    except Exception:
                                        pass
                    self._redirect("/deals")

                # ── 商談 ──
                elif path == "/deal/save":
                    def num(k):
                        v = f.get(k, "").strip()
                        try:
                            return float(v) if v else None
                        except ValueError:
                            return None
                    # 新規アカウント自動作成（deal/new フォームで「新規アカウントを追加」チェック時）
                    try:
                        deal_account_id = int(f["account_id"]) if f.get("account_id") else None
                    except ValueError:
                        deal_account_id = None
                    new_acc_name = (f.get("new_account_name") or "").strip()
                    if new_acc_name and not deal_account_id:
                        existing_acc = con.execute(
                            "SELECT id FROM accounts WHERE name=?", (new_acc_name,)
                        ).fetchone()
                        if existing_acc:
                            deal_account_id = existing_acc["id"]
                        else:
                            industries_m = sfa_db.get_master_list(con, "industries")
                            sizes_m = sfa_db.get_master_list(con, "company_sizes")
                            try:
                                est = leads_csv.estimate_companies([new_acc_name], industries_m, sizes_m)
                                est1 = est.get(new_acc_name, {})
                            except Exception:
                                est1 = {}
                            deal_account_id = sfa_db.upsert_account(
                                con, name=new_acc_name,
                                industry=est1.get("industry"),
                                company_size=est1.get("company_size"),
                            )
                    try:
                        _deal_id_in = int(f["id"]) if f.get("id") else None
                    except ValueError:
                        _deal_id_in = None
                    # クローズはモーダル経由に一本化したため、編集フォームはstatusを送らない。
                    # 既存商談のstatus（open/closed）は保持し、編集で誤って再オープンしない。
                    _keep_status = "open"
                    if _deal_id_in:
                        _ex_deal = sfa_db.get_deal(con, _deal_id_in)
                        if _ex_deal:
                            _keep_status = _ex_deal.get("status") or "open"
                    did = sfa_db.upsert_deal(
                        con, id=_deal_id_in,
                        account_id=deal_account_id,
                        deal_name=f.get("deal_name") or "(無題)",
                        stage=f.get("stage") or None,
                        business_type_l1=f.get("business_type_l1") or None,
                        business_type_l2=f.get("business_type_l2") or None,
                        lead_pattern=f.get("lead_pattern") or None,
                        owner=f.get("owner") or None,
                        sub_owner=f.get("sub_owner") or None,
                        client_contact=f.get("client_contact") or None,
                        client_dept=f.get("client_dept") or None,
                        value_lumpsum=num("value_lumpsum"),
                        value_lumpsum_monthly=num("value_lumpsum_monthly"),
                        value_recurring=num("value_recurring"),
                        client_budget=f.get("client_budget") or None,
                        next_milestone_date=f.get("next_milestone_date") or None,
                        next_milestone_label=f.get("next_milestone_label") or None,
                        next_milestone_type=f.get("next_milestone_type") or None,
                        note=f.get("note") or None,
                        goal=f.get("goal") or None,
                        importance=f.get("importance") or None,
                        status=_keep_status,
                        cost_stage=f.get("cost_stage") or None,
                        approach_value=num("approach_value"),
                        approach_rate=num("approach_rate"),
                        reduction_rate=num("reduction_rate"),
                        fee_rate=num("fee_rate"),
                        diagnosis_cost=num("diagnosis_cost"),
                    )
                    # #75: 提案以降ステージで保存されたらDeliveryを自動起票（未作成時のみ）
                    try:
                        sfa_db.ensure_delivery_on_stage(con, did, f.get("stage") or "")
                    except Exception as _exc:  # noqa: BLE001
                        print(f"[delivery] ensure_delivery_on_stage failed: {_exc}")
                    # 次回MS（複数, #48）: フォームからのMS行で置き換え→キャッシュ(next_milestone_*)再計算。
                    # deal_formは ms_*[] 配列を、quick-add等は単一 next_milestone_* を送る。両対応。
                    _ms_dates = f_list.get("ms_date[]", [])
                    if _ms_dates or ("ms_label[]" in f_list):
                        _labels = f_list.get("ms_label[]", [])
                        _types = f_list.get("ms_type[]", [])
                        _dones = f_list.get("ms_done[]", [])
                        _n = max(len(_ms_dates), len(_labels), len(_types))
                        _items = [{
                            "date": _ms_dates[i] if i < len(_ms_dates) else "",
                            "label": _labels[i] if i < len(_labels) else "",
                            "type": _types[i] if i < len(_types) else "",
                            "done": (_dones[i] if i < len(_dones) else "0") == "1",
                        } for i in range(_n)]
                        sfa_db.set_deal_milestones(con, did, _items)
                    elif f.get("next_milestone_date") or f.get("next_milestone_label"):
                        sfa_db.set_deal_milestones(con, did, [{
                            "date": f.get("next_milestone_date") or "",
                            "label": f.get("next_milestone_label") or "",
                            "type": f.get("next_milestone_type") or "",
                            "done": False}])
                    else:
                        # MS欄を持たないフォーム: 既存MS行からキャッシュを復元（upsertのNULL上書きを打ち消す）
                        sfa_db.recompute_deal_next_milestone(con, did)
                    # 終了理由はDEAL_FIELDS外（部分更新でのNULL上書き事故を避けるため個別UPDATE）。
                    # 送信された時のみ更新（空送信で既存値を消さない）。
                    _cr = f.get("close_reason")
                    if _cr is not None and "close_reason" in f:
                        con.execute("UPDATE deals SET close_reason=? WHERE id=?",
                                    (_cr or None, did))
                        con.commit()
                    # 展示会名もDEAL_FIELDS外。フォームから送られた時のみ更新。
                    if "exhibition_name" in f:
                        con.execute("UPDATE deals SET exhibition_name=? WHERE id=?",
                                    (f.get("exhibition_name") or None, did))
                        con.commit()
                    if theme_client is not None:
                        try:
                            theme_link.sync_deal(theme_client, con, did)
                            sfa_db.clear_sync_failure(con, "deal", did)
                        except Exception as exc:  # noqa: BLE001
                            sfa_db.record_sync_failure(con, "deal", did, str(exc))
                            print(f"[theme_link] sync_deal failed: {exc}")
                    # 保存後は商談一覧へ（連打・多重保存の防止）。特定画面から来た場合は元の画面へ戻る。
                    _rt = f.get("return_to") or ""
                    self._redirect(_rt if _rt.startswith("/") else "/deals")

                # ── 開発案件 ──
                elif path == "/dev-project/new":
                    try:
                        deal_id_val = int(f["deal_id"]) if f.get("deal_id") else None
                    except ValueError:
                        deal_id_val = None
                    if not deal_id_val:
                        self._redirect("/dev-projects")
                        return
                    _dp_deadline = f.get("deadline") or None
                    _dp_stage = f.get("stage") or None
                    _dp_backend = f.get("has_backend") or None
                    _dp_difficulty = f.get("difficulty") or None
                    _dp_start, _dp_end = sfa_db.compute_dev_schedule(
                        _dp_deadline, _dp_stage, _dp_backend, _dp_difficulty)
                    pid = sfa_db.upsert_dev_project(
                        con, id=None, deal_id=deal_id_val,
                        theme=f.get("theme") or "(無題)",
                        theme_detail=f.get("theme_detail") or None,
                        status=f.get("status") or None,
                        stage=_dp_stage,
                        resolution=f.get("resolution") or None,
                        budget_confirmed=f.get("budget_confirmed") or None,
                        difficulty=_dp_difficulty,
                        has_backend=_dp_backend,
                        dev_audience=f.get("dev_audience") or None,
                        work_type=f.get("work_type") or None,
                        pricing=f.get("pricing") or None,
                        dev_points=f.get("dev_points") or None,
                        dev_owner=f.get("dev_owner") or None,
                        tech_support=f.get("tech_support") or None,
                        dev_milestone=f.get("dev_milestone") or None,
                        dev_milestone_date=f.get("dev_milestone_date") or None,
                        deadline=_dp_deadline,
                        dev_start_date=_dp_start,
                        dev_end_date=_dp_end,
                        dev_policy=f.get("dev_policy") or None,
                        tech_seeds=",".join(s for s in f_list.get("tech_seeds", []) if s) or None,
                        tool_url=f.get("tool_url") or None,
                        tool_login_id=f.get("tool_login_id") or None,
                        tool_login_pass=f.get("tool_login_pass") or None,
                    )
                    if theme_client is not None:
                        try:
                            dev_project_link.sync_dev_project(theme_client, con, pid)
                            sfa_db.clear_sync_failure(con, "dev_project", pid)
                        except Exception as exc:  # noqa: BLE001
                            sfa_db.record_sync_failure(con, "dev_project", pid, str(exc))
                            print(f"[dev_project_link] sync_dev_project failed: {exc}")
                    # 新規作成後は開発案件一覧へ遷移（保存後に画面が変わらず二重登録される問題の対策）
                    self._redirect("/dev-projects")

                elif path.startswith("/dev-project/") and path.endswith("/edit"):
                    try:
                        pid = int(path.split("/")[2])
                    except (ValueError, IndexError):
                        self._redirect("/dev-projects")
                        return
                    existing = sfa_db.get_dev_project(con, pid)
                    if not existing:
                        self._redirect("/dev-projects")
                        return
                    _dp_deadline = f.get("deadline") or None
                    _dp_stage = f.get("stage") or None
                    _dp_backend = f.get("has_backend") or None
                    _dp_difficulty = f.get("difficulty") or None
                    _dp_start = existing.get("dev_start_date")
                    _dp_end = existing.get("dev_end_date")
                    if not _dp_start or not _dp_end:
                        _dp_start, _dp_end = sfa_db.compute_dev_schedule(
                            _dp_deadline, _dp_stage, _dp_backend, _dp_difficulty)
                    sfa_db.upsert_dev_project(
                        con, id=pid, deal_id=existing["deal_id"],
                        theme=f.get("theme") or "(無題)",
                        theme_detail=f.get("theme_detail") or None,
                        status=f.get("status") or None,
                        stage=_dp_stage,
                        resolution=f.get("resolution") or None,
                        budget_confirmed=f.get("budget_confirmed") or None,
                        difficulty=_dp_difficulty,
                        has_backend=_dp_backend,
                        dev_audience=f.get("dev_audience") or None,
                        work_type=f.get("work_type") or None,
                        pricing=f.get("pricing") or None,
                        dev_points=f.get("dev_points") or None,
                        dev_owner=f.get("dev_owner") or None,
                        tech_support=f.get("tech_support") or None,
                        dev_milestone=f.get("dev_milestone") or None,
                        dev_milestone_date=f.get("dev_milestone_date") or None,
                        deadline=_dp_deadline,
                        dev_start_date=_dp_start,
                        dev_end_date=_dp_end,
                        dev_policy=f.get("dev_policy") or None,
                        tech_seeds=",".join(s for s in f_list.get("tech_seeds", []) if s) or None,
                        tool_url=f.get("tool_url") or None,
                        tool_login_id=f.get("tool_login_id") or None,
                        tool_login_pass=f.get("tool_login_pass") or None,
                    )
                    if theme_client is not None:
                        try:
                            dev_project_link.sync_dev_project(theme_client, con, pid)
                            sfa_db.clear_sync_failure(con, "dev_project", pid)
                        except Exception as exc:  # noqa: BLE001
                            sfa_db.record_sync_failure(con, "dev_project", pid, str(exc))
                            print(f"[dev_project_link] sync_dev_project failed: {exc}")
                    _return_to = f.get("return_to") or ""
                    self._redirect(_return_to if _return_to.startswith("/") else f"/deal/{existing['deal_id']}")

                # 追加ツールリンクの削除（汎用 /delete より先に置く＝先勝ち）
                elif path.startswith("/dev-project/") and "/tools/" in path and path.endswith("/delete"):
                    parts = path.split("/")  # ['', 'dev-project', '{id}', 'tools', '{tool_id}', 'delete']
                    if len(parts) == 6 and parts[2].isdigit() and parts[4].isdigit():
                        sfa_db.delete_dev_project_tool(con, int(parts[4]))
                        self._redirect(f"/dev-project/{parts[2]}/edit")
                    else:
                        self._redirect("/dev-projects")

                # 追加ツールリンクの追加
                elif path.startswith("/dev-project/") and path.endswith("/tools/add"):
                    parts = path.split("/")  # ['', 'dev-project', '{id}', 'tools', 'add']
                    if len(parts) == 5 and parts[2].isdigit():
                        pid = int(parts[2])
                        url = (f.get("url") or "").strip()
                        if url and url.lower().startswith(("http://", "https://")):
                            sfa_db.add_dev_project_tool(
                                con, dev_project_id=pid, url=url,
                                label=f.get("label") or None,
                                login_id=f.get("login_id") or None,
                                login_pass=f.get("login_pass") or None)
                        _rt = f.get("return_to") or ""
                        self._redirect(_rt if _rt.startswith("/") else f"/dev-project/{pid}/edit")
                    else:
                        self._redirect("/dev-projects")

                elif (path.startswith("/dev-project/") and path.endswith("/tool-main")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    # 一覧のツールリンクパネルから「制作したツールのリンク（主）」を保存（ajax）。
                    pid = int(path.split("/")[2])
                    con.execute(
                        "UPDATE dev_projects SET tool_url=?, tool_login_id=?, tool_login_pass=?, "
                        "updated_at=datetime('now') WHERE id=?",
                        (f.get("url") or None, f.get("login_id") or None, f.get("login_pass") or None, pid))
                    con.commit()
                    if theme_client is not None:
                        try:
                            dev_project_link.sync_dev_project(theme_client, con, pid)
                        except Exception as _exc:  # noqa: BLE001
                            print(f"[dev_project_link] sync_dev_project failed (tool-main): {_exc}")
                    self._send(json.dumps({"ok": True}).encode(), ctype="application/json")

                elif (path.startswith("/dev-project/") and path.endswith("/tool-add")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    # 一覧のツールリンクパネルから「追加リンク」を保存（ajax）。
                    pid = int(path.split("/")[2])
                    _url = (f.get("url") or "").strip()
                    if _url and _url.lower().startswith(("http://", "https://")):
                        sfa_db.add_dev_project_tool(
                            con, dev_project_id=pid, url=_url,
                            label=f.get("label") or None, login_id=f.get("login_id") or None,
                            login_pass=f.get("login_pass") or None)
                        self._send(json.dumps({"ok": True}).encode(), ctype="application/json")
                    else:
                        self._send(json.dumps({"ok": False, "error": "URLはhttp(s)で入力してください"}).encode(),
                                   ctype="application/json")

                elif path.startswith("/dev-project/") and path.endswith("/delete"):
                    try:
                        pid = int(path.split("/")[2])
                    except (ValueError, IndexError):
                        self._redirect("/dev-projects")
                        return
                    existing = sfa_db.get_dev_project(con, pid)
                    if existing:
                        sfa_db.delete_dev_project(con, pid)
                        if theme_client is not None and existing.get("hisho_id"):
                            try:
                                dev_project_link.delete_dev_project_remote(theme_client, existing["hisho_id"])
                            except Exception as exc:  # noqa: BLE001
                                print(f"[dev_project_link] delete_dev_project_remote failed: {exc}")
                    self._redirect(f"/deal/{existing['deal_id']}" if existing else "/dev-projects")

                elif path == "/dev-projects/bulk_delete":
                    ids = f_list.get("ids", [])
                    for pid_raw in ids:
                        if not str(pid_raw).isdigit():
                            continue
                        pid = int(pid_raw)
                        existing = sfa_db.get_dev_project(con, pid)
                        if not existing:
                            continue
                        sfa_db.delete_dev_project(con, pid)
                        if theme_client is not None and existing.get("hisho_id"):
                            try:
                                dev_project_link.delete_dev_project_remote(theme_client, existing["hisho_id"])
                            except Exception as exc:  # noqa: BLE001
                                print(f"[dev_project_link] delete_dev_project_remote failed: {exc}")
                    self._redirect("/dev-projects")

                elif path == "/dev-projects/resync-hisho":
                    # 全開発案件をHishoへ再同期（点数・分類を反映。既存はsync漏れがあるため）
                    if theme_client is None:
                        self._send(render(dev_projects_list_page(con),
                                          flash="テーマDB連携が無効です（THEME_API_TOKEN未設定）。"))
                    else:
                        _ok = _ng = 0
                        for _p in sfa_db.list_dev_projects(con):
                            try:
                                dev_project_link.sync_dev_project(theme_client, con, _p["id"])
                                sfa_db.clear_sync_failure(con, "dev_project", _p["id"])
                                _ok += 1
                            except Exception as _exc:  # noqa: BLE001
                                sfa_db.record_sync_failure(con, "dev_project", _p["id"], str(_exc))
                                _ng += 1
                        self._send(render(dev_projects_list_page(con),
                                          flash=f"Hishoへ再同期しました（成功{_ok}件 / 失敗{_ng}件）。"))

                # ── 社内論点 ──
                elif path == "/deal-issue/new":
                    try:
                        deal_id_val = int(f["deal_id"]) if f.get("deal_id") else None
                    except ValueError:
                        deal_id_val = None
                    iid = sfa_db.upsert_deal_issue(
                        con, id=None, deal_id=deal_id_val,
                        issue=f.get("issue") or "(無題)",
                        members=",".join(f_list.get("members", [])),
                        status=f.get("status") or "議論中",
                        due_date=f.get("due_date") or None,
                    )
                    _return_to = f.get("return_to") or ""
                    self._redirect(_return_to if _return_to.startswith("/")
                                   else (f"/deal/{deal_id_val}" if deal_id_val else "/deal-issues"))

                elif path.startswith("/deal-issue/") and path.endswith("/edit"):
                    try:
                        iid = int(path.split("/")[2])
                    except (ValueError, IndexError):
                        self._redirect("/deal-issues")
                        return
                    existing = sfa_db.get_deal_issue(con, iid)
                    if not existing:
                        self._redirect("/deal-issues")
                        return
                    sfa_db.upsert_deal_issue(
                        con, id=iid, deal_id=existing["deal_id"],
                        issue=f.get("issue") or "(無題)",
                        members=",".join(f_list.get("members", [])),
                        status=f.get("status") or "議論中",
                        due_date=f.get("due_date") or None,
                    )
                    _return_to = f.get("return_to") or ""
                    self._redirect(_return_to if _return_to.startswith("/") else f"/deal/{existing['deal_id']}")

                elif path.startswith("/deal-issue/") and path.endswith("/delete"):
                    try:
                        iid = int(path.split("/")[2])
                    except (ValueError, IndexError):
                        self._redirect("/deal-issues")
                        return
                    existing = sfa_db.get_deal_issue(con, iid)
                    if existing:
                        sfa_db.delete_deal_issue(con, iid)
                    self._redirect(f"/deal/{existing['deal_id']}" if existing else "/deal-issues")

                elif path.startswith("/deal-issue/") and path.endswith("/memo"):
                    try:
                        iid = int(path.split("/")[2])
                    except (ValueError, IndexError):
                        self._redirect("/deal-issues")
                        return
                    existing = sfa_db.get_deal_issue(con, iid)
                    body = f.get("body") or ""
                    if existing and body.strip():
                        sfa_db.add_deal_issue_memo(con, issue_id=iid, body=body)
                        memos = sfa_db.list_deal_issue_memos(con, iid)
                        summary = _generate_issue_ai_summary(existing, memos)
                        if summary:
                            sfa_db.set_deal_issue_ai_summary(con, iid, summary)
                    _return_to = f.get("return_to") or ""
                    self._redirect(_return_to if _return_to.startswith("/")
                                   else (f"/deal/{existing['deal_id']}" if existing else "/deal-issues"))

                elif path.startswith("/deal-issue-memo/") and path.endswith("/delete"):
                    try:
                        mid = int(path.split("/")[2])
                    except (ValueError, IndexError):
                        self._redirect("/deal-issues")
                        return
                    memo = sfa_db.get_deal_issue_memo(con, mid)
                    if memo:
                        sfa_db.delete_deal_issue_memo(con, mid)
                    _return_to = f.get("return_to") or ""
                    self._redirect(_return_to if _return_to.startswith("/") else "/deal-issues")

                elif path.startswith("/deal-issue/") and path.endswith("/regenerate_summary"):
                    try:
                        iid = int(path.split("/")[2])
                    except (ValueError, IndexError):
                        self._redirect("/deal-issues")
                        return
                    existing = sfa_db.get_deal_issue(con, iid)
                    if existing:
                        memos = sfa_db.list_deal_issue_memos(con, iid)
                        summary = _generate_issue_ai_summary(existing, memos)
                        if summary:
                            sfa_db.set_deal_issue_ai_summary(con, iid, summary)
                    _return_to = f.get("return_to") or ""
                    self._redirect(_return_to if _return_to.startswith("/")
                                   else (f"/deal/{existing['deal_id']}" if existing else "/deal-issues"))

                elif path.startswith("/deal-issue/") and path.endswith("/field"):
                    _DEAL_ISSUE_ALLOWED_FIELDS = {"status", "members", "due_date"}
                    parts = path.split("/")
                    _ok = False
                    _err = ""
                    if len(parts) == 4 and parts[3] == "field" and parts[2].isdigit():
                        iid = int(parts[2])
                        field = f.get("field", "")
                        value = f.get("value", "")
                        if field not in _DEAL_ISSUE_ALLOWED_FIELDS:
                            _err = "不正なフィールド"
                        elif field == "status" and value and value not in sfa_db.DEAL_ISSUE_STATUSES:
                            _err = "不正なステータス値"
                        else:
                            con.execute(
                                f"UPDATE deal_issues SET {field}=?, updated_at=datetime('now') WHERE id=?",
                                (value or None, iid),
                            )
                            con.commit()
                            _ok = True
                    else:
                        _err = "不正なリクエスト"
                    _resp = json.dumps({"ok": _ok} if _ok else {"ok": False, "error": _err}).encode("utf-8")
                    self._send(_resp, ctype="application/json")

                elif path.startswith("/deal/") and path.endswith("/attachment"):
                    try:
                        did = int(path.split("/")[2])
                    except (ValueError, IndexError):
                        self._redirect("/deals")
                        return
                    label = (f.get("label") or "").strip()
                    url = (f.get("url") or "").strip()
                    if label and url:
                        sfa_db.add_deal_attachment(con, deal_id=did, label=label, url=url)
                    self._redirect(f"/deal/{did}")

                elif path.startswith("/deal-attachment/") and path.endswith("/delete"):
                    try:
                        aid = int(path.split("/")[2])
                    except (ValueError, IndexError):
                        self._redirect("/deals")
                        return
                    att = sfa_db.get_deal_attachment(con, aid)
                    if att:
                        sfa_db.delete_deal_attachment(con, aid)
                    self._redirect(f"/deal/{att['deal_id']}" if att else "/deals")

                elif (path.startswith("/activity/") and path.endswith("/field")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    _aid = int(path.split("/")[2])
                    _field = f.get("field", "")
                    _value = f.get("value", "")
                    if _field not in sfa_db.ACTIVITY_EDIT_FIELDS:
                        self._send(json.dumps({"ok": False, "error": "不正なフィールド"}).encode(),
                                   ctype="application/json")
                    else:
                        sfa_db.update_activity_field(con, _aid, _field, _value)
                        self._send(json.dumps({"ok": True}).encode(), ctype="application/json")

                elif (path.startswith("/activity/") and path.endswith("/delete")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    _aid = int(path.split("/")[2])
                    _act = sfa_db.get_activity(con, _aid)
                    _did = _act.get("deal_id") if _act else None
                    sfa_db.delete_activity(con, _aid)
                    self._redirect(f"/deal/{_did}#activity" if _did else "/deals")

                elif path == "/activity/add":
                    try:
                        did = int(f["deal_id"])
                    except (ValueError, KeyError):
                        self._redirect("/deals")
                        return
                    # 活動履歴は「日付＋種別」が必須。相手/内容/日付のいずれかが入っていれば
                    # 活動登録の意図とみなし、日付か種別が欠けていればエラーで差し戻す。
                    # 何も入っていなければ活動は作らず、状況メモ/次回MSの更新だけ行う
                    # （このフォームはメモ/MS更新も兼ねるため。日付なし空「面談」の混入を防ぐ）。
                    _occ = (f.get("occurred_on") or "").strip()
                    _typ = (f.get("type") or "").strip()
                    _contact = (f.get("contact_name") or "").strip()
                    _body = (f.get("body") or "").strip()
                    _wants_activity = bool(_occ or _contact or _body)
                    if _wants_activity and (not _occ or not _typ):
                        _deal = sfa_db.get_deal(con, did)
                        _rt = self._qs().get("return_to", [None])[0]
                        self._send(render(
                            deal_form(con, _deal, return_to=_rt) if _deal
                            else "<div class=card>商談が見つかりません</div>",
                            flash="❌ 活動履歴の登録には「日付」と「種別」が必須です。"
                                  "（状況メモ・次回MSだけを更新したい場合は、日付・相手・内容を空のままにしてください）"))
                        return
                    if _wants_activity:
                        sfa_db.add_activity(
                            con, deal_id=did, type=_typ, occurred_on=_occ,
                            contact_name=_contact or None, body=_body or None,
                        )
                    # 商談の現状メモ・次回MSを同時更新（入力があった場合のみ）
                    update_note = f.get("update_note", "").strip()
                    ms_date = f.get("next_milestone_date", "").strip()
                    ms_label = f.get("next_milestone_label", "").strip()
                    if update_note or ms_date or ms_label:
                        deal = sfa_db.get_deal(con, did)
                        if deal:
                            sfa_db.upsert_deal(
                                con, id=did,
                                account_id=deal["account_id"],
                                theme_id=deal.get("theme_id"),
                                deal_name=deal["deal_name"],
                                stage=deal.get("stage"),
                                business_type_l1=deal.get("business_type_l1"),
                                business_type_l2=deal.get("business_type_l2"),
                                lead_pattern=deal.get("lead_pattern"),
                                owner=deal.get("owner"),
                                sub_owner=deal.get("sub_owner"),
                                client_contact=deal.get("client_contact"),
                                client_dept=deal.get("client_dept"),
                                value_lumpsum=deal.get("value_lumpsum"),
                                value_lumpsum_monthly=deal.get("value_lumpsum_monthly"),
                                value_recurring=deal.get("value_recurring"),
                                client_budget=deal.get("client_budget"),
                                next_milestone_date=ms_date or deal.get("next_milestone_date"),
                                next_milestone_label=ms_label or deal.get("next_milestone_label"),
                                next_milestone_type=f.get("next_milestone_type") or deal.get("next_milestone_type"),
                                note=update_note or deal.get("note"),
                                goal=deal.get("goal"),
                                status=deal.get("status"),
                            )
                            # 次回MS(複数, #48): 現状更新で入った値を最古MSへ反映（行とキャッシュを整合）
                            if ms_date or ms_label or f.get("next_milestone_type"):
                                sfa_db.upsert_earliest_milestone(
                                    con, did,
                                    date=ms_date or deal.get("next_milestone_date"),
                                    label=ms_label or deal.get("next_milestone_label"),
                                    ms_type=f.get("next_milestone_type") or deal.get("next_milestone_type"))
                    self._redirect(f"/deal/{did}")

                # ── 商談インライン編集 ──
                elif path.startswith("/deal/") and path.endswith("/field"):
                    _DEAL_ALLOWED_FIELDS = {"stage", "owner", "sub_owner", "business_type_l1", "business_type_l2",
                                             "client_budget", "value_lumpsum", "deal_name",
                                             "next_milestone_date", "next_milestone_label", "next_milestone_type",
                                             "close_reason", "exhibition_name"}
                    parts = path.split("/")
                    _ok = False
                    _err = ""
                    if len(parts) == 4 and parts[3] == "field" and parts[2].isdigit():
                        deal_id = int(parts[2])
                        field = f.get("field", "")
                        value = f.get("value", "")
                        if field not in _DEAL_ALLOWED_FIELDS:
                            _err = "不正なフィールド"
                        elif field == "next_milestone_type" and value and value not in sfa_db.NEXT_MS_TYPES:
                            _err = "不正な次回MS種別"
                        elif field == "close_reason" and value and value not in sfa_db.CLOSE_REASONS:
                            _err = "不正な終了理由"
                        elif field == "stage":
                            valid_stages = sfa_db.get_master_list(con, "deal_stages")
                            if value and value not in valid_stages:
                                _err = "不正なステージ値"
                            else:
                                con.execute(
                                    "UPDATE deals SET stage=?, updated_at=datetime('now') WHERE id=?",
                                    (value or None, deal_id),
                                )
                                con.commit()
                                _ok = True
                                # #75: 提案以降に到達したらDeliveryを自動起票（未作成時のみ）
                                try:
                                    sfa_db.ensure_delivery_on_stage(con, deal_id, value)
                                except Exception as _exc:  # noqa: BLE001
                                    print(f"[delivery] ensure_delivery_on_stage failed: {_exc}")
                        elif field in ("next_milestone_date", "next_milestone_label", "next_milestone_type"):
                            # 次回MSは複数対応（#48）: 未完了で最古のMSを更新（無ければ作成）→キャッシュ再計算
                            _mf = {"next_milestone_date": "date", "next_milestone_label": "label",
                                   "next_milestone_type": "type"}[field]
                            sfa_db.set_earliest_milestone_field(con, deal_id, _mf, value)
                            _ok = True
                        else:
                            con.execute(
                                f"UPDATE deals SET {field}=?, updated_at=datetime('now') WHERE id=?",
                                (value or None, deal_id),
                            )
                            con.commit()
                            _ok = True
                        if _ok and theme_client is not None:
                            try:
                                theme_link.sync_deal(theme_client, con, deal_id)
                            except Exception as _exc:
                                print(f"[theme_link] sync_deal failed: {_exc}")
                    else:
                        _err = "不正なリクエスト"
                    _resp = json.dumps({"ok": _ok} if _ok else {"ok": False, "error": _err}).encode("utf-8")
                    self._send(_resp, ctype="application/json")

                # ── 次回MS（複数）一覧パネル操作（#48）: 追加/項目更新/削除。応答は最新の全MS ──
                elif (path.startswith("/deal/") and path.endswith("/milestones/add")
                      and len(path.split("/")) == 5 and path.split("/")[2].isdigit()):
                    _did = int(path.split("/")[2])
                    _mt = f.get("ms_type") or ""
                    if _mt and _mt not in sfa_db.NEXT_MS_TYPES:
                        self._send(json.dumps({"ok": False, "error": "不正な種別"}).encode(),
                                   ctype="application/json")
                    else:
                        sfa_db.add_deal_milestone(con, _did, date=f.get("ms_date"),
                                                  label=f.get("ms_label"), ms_type=_mt)
                        if theme_client is not None:
                            try:
                                theme_link.sync_deal(theme_client, con, _did)
                            except Exception as _exc:  # noqa: BLE001
                                print(f"[theme_link] sync_deal failed: {_exc}")
                        self._send(json.dumps(_ms_panel_json(con, _did), ensure_ascii=False).encode(),
                                   ctype="application/json")

                elif (path.startswith("/milestone/") and path.endswith("/field")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    _mid = int(path.split("/")[2])
                    _field = f.get("field", "")
                    _value = f.get("value", "")
                    if _field == "type" and _value and _value not in sfa_db.NEXT_MS_TYPES:
                        self._send(json.dumps({"ok": False, "error": "不正な種別"}).encode(),
                                   ctype="application/json")
                    else:
                        _did = sfa_db.update_deal_milestone(con, _mid, _field, _value)
                        if _did and theme_client is not None:
                            try:
                                theme_link.sync_deal(theme_client, con, _did)
                            except Exception as _exc:  # noqa: BLE001
                                print(f"[theme_link] sync_deal failed: {_exc}")
                        if _did:
                            self._send(json.dumps(_ms_panel_json(con, _did), ensure_ascii=False).encode(),
                                       ctype="application/json")
                        else:
                            self._send(json.dumps({"ok": False, "error": "見つかりません"}).encode(),
                                       status=404, ctype="application/json")

                elif (path.startswith("/milestone/") and path.endswith("/delete")
                      and len(path.split("/")) == 4 and path.split("/")[2].isdigit()):
                    _mid = int(path.split("/")[2])
                    _did = sfa_db.delete_deal_milestone(con, _mid)
                    if _did and theme_client is not None:
                        try:
                            theme_link.sync_deal(theme_client, con, _did)
                        except Exception as _exc:  # noqa: BLE001
                            print(f"[theme_link] sync_deal failed: {_exc}")
                    if _did:
                        self._send(json.dumps(_ms_panel_json(con, _did), ensure_ascii=False).encode(),
                                   ctype="application/json")
                    else:
                        self._send(json.dumps({"ok": False, "error": "見つかりません"}).encode(),
                                   status=404, ctype="application/json")

                elif path.startswith("/dev-project/") and path.endswith("/field"):
                    _DEV_PROJECT_ALLOWED_FIELDS = {
                        "dev_owner", "stage", "status", "order_potential", "dev_audience",
                        "work_type", "pricing", "difficulty", "has_backend", "dev_points",
                    }
                    # 変わったら点数を再計算する対象（(基準+BE加点)×分類係数×難易度係数）
                    _RECALC_FIELDS = {"work_type", "stage", "difficulty", "has_backend"}
                    parts = path.split("/")
                    _ok = False
                    _err = ""
                    if len(parts) == 4 and parts[3] == "field" and parts[2].isdigit():
                        pid = int(parts[2])
                        field = f.get("field", "")
                        value = f.get("value", "")
                        if field not in _DEV_PROJECT_ALLOWED_FIELDS:
                            _err = "不正なフィールド"
                        elif field == "stage" and value and value not in sfa_db.DEV_PROJECT_STAGES:
                            _err = "不正なステージ値"
                        elif field == "status" and value and value not in sfa_db.DEV_PROJECT_STATUSES:
                            _err = "不正な状況値"
                        elif field == "order_potential" and value and value not in sfa_db.DEV_ORDER_POTENTIALS:
                            _err = "不正な受注余地値"
                        elif field == "dev_audience" and value and value not in sfa_db.DEV_AUDIENCES:
                            _err = "不正な提供先値"
                        elif field == "pricing" and value and value not in sfa_db.DEV_PRICINGS:
                            _err = "不正な課金値"
                        elif field == "difficulty" and value and value not in sfa_db.DEV_DIFFICULTIES:
                            _err = "不正な難易度値"
                        elif field == "has_backend" and value and value not in sfa_db.DEV_HAS_BACKEND:
                            _err = "不正なバックエンド値"
                        if not _err and field == "dev_points":
                            # 点数の手動上書き（数値のみ）
                            try:
                                _pv = float(value) if value not in (None, "") else None
                            except ValueError:
                                _err = "不正な点数"
                            else:
                                con.execute("UPDATE dev_projects SET dev_points=?, updated_at=datetime('now') WHERE id=?",
                                            (_pv, pid))
                                con.commit()
                                _ok = True
                        elif not _err:
                            con.execute(
                                f"UPDATE dev_projects SET {field}=?, updated_at=datetime('now') WHERE id=?",
                                (value or None, pid),
                            )
                            # 分類変更時は点数をマスタ×難易度で再計算（スナップショット更新）
                            if field in _RECALC_FIELDS:
                                _row = sfa_db.get_dev_project(con, pid)
                                if _row:
                                    _pts = sfa_db.compute_dev_points(
                                        con, work_type=_row.get("work_type"), stage=_row.get("stage"),
                                        difficulty=_row.get("difficulty"), has_backend=_row.get("has_backend"))
                                    if _pts is not None:
                                        con.execute("UPDATE dev_projects SET dev_points=? WHERE id=?", (_pts, pid))
                            con.commit()
                            _ok = True
                        if _ok and theme_client is not None:
                            try:
                                dev_project_link.sync_dev_project(theme_client, con, pid)
                            except Exception as _exc:
                                print(f"[dev_project_link] sync_dev_project failed: {_exc}")
                    else:
                        _err = "不正なリクエスト"
                    if _ok:
                        # 再計算後の点数を返し、一覧の点数セルを即更新できるようにする
                        _r2 = con.execute("SELECT dev_points FROM dev_projects WHERE id=?", (pid,)).fetchone()
                        _dp_val = _r2["dev_points"] if _r2 else None
                        _resp = json.dumps({"ok": True, "dev_points": _dp_val}).encode("utf-8")
                    else:
                        _resp = json.dumps({"ok": False, "error": _err}).encode("utf-8")
                    self._send(_resp, ctype="application/json")

                # ── メールパターン ──
                elif path == "/email-patterns/save":
                    cc_list = f_list.get("cc", [])
                    sfa_db.save_email_pattern(
                        con,
                        name=f.get("name", ""),
                        subject=f.get("subject", ""),
                        body=f.get("body", ""),
                        from_address=f.get("from_address") or None,
                        cc_addresses=",".join(cc_list) if cc_list else None,
                    )
                    self._redirect("/email-patterns")
                elif path.startswith("/email-patterns/") and path.endswith("/save"):
                    try:
                        pid = int(path.split("/")[2])
                        cc_list = f_list.get("cc", [])
                        sfa_db.save_email_pattern(
                            con, id=pid,
                            name=f.get("name", ""),
                            subject=f.get("subject", ""),
                            body=f.get("body", ""),
                            from_address=f.get("from_address") or None,
                            cc_addresses=",".join(cc_list) if cc_list else None,
                        )
                        self._redirect("/email-patterns")
                    except (ValueError, IndexError):
                        self._send(render("<div class=card>不正なリクエスト</div>"), 400)
                elif path.startswith("/email-patterns/") and path.endswith("/delete"):
                    try:
                        pid = int(path.split("/")[2])
                        sfa_db.delete_email_pattern(con, pid)
                        self._redirect("/email-patterns")
                    except (ValueError, IndexError):
                        self._send(render("<div class=card>不正なリクエスト</div>"), 400)

                # ── 初回ヒアリング ──
                elif path == "/hearing-templates/save" or (
                        path.startswith("/hearing-templates/") and path.endswith("/save")):
                    try:
                        items = json.loads(f.get("items_json") or "[]")
                        if not isinstance(items, (list, dict)):
                            items = []
                    except (ValueError, TypeError):
                        items = []
                    tid = None
                    if path != "/hearing-templates/save":
                        try:
                            tid = int(path.split("/")[2])
                        except (ValueError, IndexError):
                            self._send(render("<div class=card>不正なリクエスト</div>"), 400)
                            return
                    sfa_db.save_hearing_template(
                        con, id=tid,
                        name=f.get("name", "") or "(無題)",
                        description=f.get("description") or None,
                        items=items,
                    )
                    self._redirect("/hearing-templates")
                elif path.startswith("/hearing-templates/") and path.endswith("/delete"):
                    try:
                        tid = int(path.split("/")[2])
                        sfa_db.delete_hearing_template(con, tid)
                        self._redirect("/hearing-templates")
                    except (ValueError, IndexError):
                        self._send(render("<div class=card>不正なリクエスト</div>"), 400)
                elif path.startswith("/hearing/result/") and path.endswith("/delete"):
                    parts = path.split("/")
                    if len(parts) == 5 and parts[4] == "delete" and parts[3].isdigit():
                        rid = int(parts[3])
                        r = sfa_db.get_hearing_result(con, rid)
                        sfa_db.delete_hearing_result(con, rid)
                        self._redirect(f"/deal/{r['deal_id']}" if r else "/hearings")
                    else:
                        self._redirect("/hearings")

                elif path == "/hearings/bulk_delete":
                    ids = f_list.get("ids", [])
                    for rid in ids:
                        if str(rid).isdigit():
                            sfa_db.delete_hearing_result(con, int(rid))
                    self._redirect("/hearings")

                elif path == "/hearings/export_selected":
                    result_ids = [int(rid) for rid in f_list.get("ids", []) if str(rid).isdigit()]
                    fmt = f.get("fmt") or "xlsx"
                    if not result_ids:
                        self._redirect("/hearings")
                    else:
                        try:
                            data, fname, ctype = build_hearings_export_bundle(con, result_ids, fmt)
                            self.send_response(200)
                            self.send_header("Content-Type", ctype)
                            self.send_header("Content-Disposition", _content_disposition(fname))
                            self.send_header("Content-Length", str(len(data)))
                            self.end_headers()
                            self.wfile.write(data)
                        except Exception as _ex:
                            print(f"[hearings/export_selected] {_ex}", flush=True)
                            import traceback as _tb; _tb.print_exc()
                            self._send(render("<div class=card>エクスポートに失敗しました</div>"), 500)

                elif path == "/hearing/autosave":
                    try:
                        ttype = f.get("target_type", "")
                        tval_id = int(f.get("target_id") or 0)
                        tmpl_id = int(f.get("template_id") or 0)
                        edit_rid = f.get("edit_result_id", "")
                        if edit_rid:
                            # 修正モードの下書きは、新規ヒアリングの下書きと衝突しないよう別名前空間に保存
                            ttype, tval_id = "hearing_edit", int(edit_rid)
                        if not (ttype and tval_id and tmpl_id):
                            raise ValueError("missing target/template")
                        form_data = {k: (v[0] if len(v) == 1 else v) for k, v in f_list.items()}
                        sfa_db.save_hearing_draft(
                            con, target_type=ttype, target_id=tval_id, template_id=tmpl_id,
                            form_data=form_data,
                        )
                        self._send_cors_json(json.dumps({"ok": True}).encode())
                    except Exception as exc:  # noqa: BLE001
                        self._send_cors_json(
                            json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False).encode(),
                            status=400,
                        )

                elif path == "/hearing/submit":
                    try:
                        ttype = f.get("target_type", "")
                        tval_id = int(f.get("target_id") or 0)
                        tmpl_id = int(f.get("template_id") or 0)
                    except ValueError:
                        self._send(render("<div class=card>不正なリクエスト</div>"), 400)
                        return
                    tmpl = sfa_db.get_hearing_template(con, tmpl_id) if tmpl_id else None
                    # 対象 deal_id を確定（リードは商談化）
                    deal_id = None
                    if ttype == "lead":
                        lead = sfa_db.get_lead(con, tval_id)
                        if lead:
                            try:
                                deal_id = convert_lead_to_deal(con, lead)
                            except Exception as _e:
                                print(f"[hearing/submit] convert failed: {_e}", flush=True)
                    elif ttype == "deal":
                        d = sfa_db.get_deal(con, tval_id)
                        deal_id = d["id"] if d else None
                    if not deal_id or not tmpl:
                        self._send(render("<div class=card>対象またはテンプレートが見つかりません</div>"), 404)
                        return
                    # 回答を組み立て（Q&A・矢羽混在対応）
                    answers = []
                    for i, it in enumerate(tmpl.get("items") or []):
                        if it.get("type") == "yabane":
                            try:
                                yb_ans = json.loads(f.get(f"answer_{i}") or "{}")
                            except (ValueError, TypeError):
                                yb_ans = {}
                            answers.append({"label": it.get("label") or "業務プロセス",
                                            "type": "yabane", "answer": yb_ans})
                        elif it.get("type") == "choice" and it.get("multi"):
                            ans = [v for v in f_list.get(f"answer_{i}", []) if v]
                            answers.append({"label": it.get("label"),
                                            "type": it.get("type"), "answer": ans})
                        else:
                            ans = (f.get(f"answer_{i}", "") or "").strip()
                            answers.append({"label": it.get("label"),
                                            "type": it.get("type"), "answer": ans})
                    conducted_on = f.get("occurred_on") or None

                    # 修正モード: 新規の活動履歴は作らず、既存hearing_resultsを上書きするだけ
                    edit_rid_raw = f.get("edit_result_id", "")
                    if edit_rid_raw:
                        try:
                            edit_rid = int(edit_rid_raw)
                        except ValueError:
                            self._send(render("<div class=card>不正なリクエスト</div>"), 400)
                            return
                        sfa_db.update_hearing_result(con, edit_rid, conducted_on=conducted_on, answers=answers)
                        sfa_db.delete_hearing_draft(
                            con, target_type="hearing_edit", target_id=edit_rid, template_id=tmpl_id,
                        )
                        self._redirect(f"/hearing/result/{edit_rid}")
                        return

                    # 活動履歴を1件追加
                    act_id = sfa_db.add_activity(
                        con, deal_id=deal_id,
                        type=f.get("type") or None,
                        occurred_on=conducted_on,
                        contact_name=f.get("contact_name") or None,
                        body=f.get("body") or None,
                    )
                    # ヒアリング結果を保存（活動履歴と相互リンク）
                    sfa_db.add_hearing_result(
                        con, deal_id=deal_id, template_id=tmpl["id"],
                        template_name=tmpl.get("name"), conducted_on=conducted_on,
                        answers=answers, activity_id=act_id,
                    )
                    # 確定保存できたので自動保存の下書きは不要（元のtarget_type/idで保存されているため）
                    sfa_db.delete_hearing_draft(con, target_type=ttype, target_id=tval_id, template_id=tmpl_id)
                    # 商談の現状メモ・次回MSを更新（入力があった場合のみ）
                    update_note = (f.get("update_note") or "").strip()
                    ms_date = (f.get("next_milestone_date") or "").strip()
                    ms_label = (f.get("next_milestone_label") or "").strip()
                    if update_note or ms_date or ms_label:
                        deal = sfa_db.get_deal(con, deal_id)
                        if deal:
                            sfa_db.upsert_deal(
                                con, id=deal_id,
                                account_id=deal["account_id"], theme_id=deal.get("theme_id"),
                                deal_name=deal["deal_name"], stage=deal.get("stage"),
                                business_type_l1=deal.get("business_type_l1"),
                                business_type_l2=deal.get("business_type_l2"),
                                lead_pattern=deal.get("lead_pattern"), owner=deal.get("owner"),
                                sub_owner=deal.get("sub_owner"),
                                client_contact=deal.get("client_contact"),
                                client_dept=deal.get("client_dept"),
                                value_lumpsum=deal.get("value_lumpsum"),
                                value_lumpsum_monthly=deal.get("value_lumpsum_monthly"),
                                value_recurring=deal.get("value_recurring"),
                                client_budget=deal.get("client_budget"),
                                next_milestone_date=ms_date or deal.get("next_milestone_date"),
                                next_milestone_label=ms_label or deal.get("next_milestone_label"),
                                next_milestone_type=f.get("next_milestone_type") or deal.get("next_milestone_type"),
                                note=update_note or deal.get("note"),
                                goal=deal.get("goal"), status=deal.get("status"),
                            )
                            # 次回MS(複数, #48): 現状更新値を最古MSへ反映（行とキャッシュを整合）
                            if ms_date or ms_label or f.get("next_milestone_type"):
                                sfa_db.upsert_earliest_milestone(
                                    con, deal_id,
                                    date=ms_date or deal.get("next_milestone_date"),
                                    label=ms_label or deal.get("next_milestone_label"),
                                    ms_type=f.get("next_milestone_type") or deal.get("next_milestone_type"))
                    if theme_client is not None:
                        try:
                            theme_link.sync_deal(theme_client, con, deal_id)
                        except Exception as exc:  # noqa: BLE001
                            print(f"[theme_link] sync_deal failed: {exc}")
                    self._redirect(f"/deal/{deal_id}")

                # ── リード ──
                elif path == "/leads/save":
                    try:
                        existing_id = int(f["id"]) if f.get("id") else None
                    except ValueError:
                        existing_id = None
                    existing_deal_id = None
                    existing_pitch_theme_id = None
                    if existing_id:
                        existing = sfa_db.get_lead(con, existing_id)
                        existing_deal_id = existing.get("deal_id") if existing else None
                        existing_pitch_theme_id = existing.get("pitch_theme_id") if existing else None
                    company_name = f.get("company") or "(未設定)"
                    industry = f.get("industry") or None
                    company_size = f.get("company_size") or None
                    lid = sfa_db.upsert_lead(
                        con, id=existing_id,
                        name=f.get("name") or "(無名)",
                        company=company_name,
                        industry=industry,
                        company_size=company_size,
                        title=f.get("title") or None,
                        email=f.get("email") or None,
                        phone=f.get("phone") or None,
                        source=f.get("source") or "other",
                        pitch_theme_id=existing_pitch_theme_id,
                        lead_status=f.get("lead_status") or "new",
                        notes=f.get("notes") or None,
                        assigned_to=f.get("assigned_to") or None,
                        deal_id=existing_deal_id,
                    )
                    # 終了理由(lost_reason)はLEAD_FIELDS外なので個別UPDATE（送信時のみ）
                    if "lost_reason" in f:
                        _lr = (f.get("lost_reason") or "").strip()
                        con.execute("UPDATE leads SET lost_reason=? WHERE id=?",
                                    (_lr if _lr in sfa_db.CLOSE_REASONS else None, lid))
                        con.commit()
                    # アカウント自動追加・補完
                    existing_acc = con.execute(
                        "SELECT id, industry, company_size FROM accounts WHERE name=?",
                        (company_name,)
                    ).fetchone()
                    if existing_acc is None:
                        sfa_db.upsert_account(
                            con, name=company_name,
                            industry=industry,
                            company_size=company_size,
                        )
                    else:
                        acc = dict(existing_acc)
                        updates = {}
                        if industry and not acc.get("industry"):
                            updates["industry"] = industry
                        if company_size and not acc.get("company_size"):
                            updates["company_size"] = company_size
                        if updates:
                            set_clause = ", ".join(f"{k}=?" for k in updates)
                            con.execute(
                                f"UPDATE accounts SET {set_clause}, updated_at=datetime('now') WHERE id=?",
                                (*updates.values(), acc["id"]),
                            )
                            con.commit()
                    self._redirect(f"/leads/{lid}")

                elif path == "/leads/upload_meishi":
                    file_item = f.get("meishi_file")
                    if not file_item or not isinstance(file_item, tuple):
                        self._send(render(leads_import_page(con), flash="ファイルが選択されていません。"))
                    else:
                        filename, data = file_item
                        try:
                            from . import meishi_import
                            added, skipped, errors = meishi_import.import_meishi_file(con, data, filename)
                            msg = f"取込完了: {added}件追加、{skipped}件スキップ。"
                            if errors:
                                msg += " エラー: " + "; ".join(errors[:3])
                            self._send(render(leads_import_page(con, result=msg)))
                        except ImportError:
                            self._send(render(leads_import_page(con), flash="meishi_importモジュールが見つかりません。"))
                        except Exception as exc:
                            self._send(render(leads_import_page(con), flash=f"取込エラー: {exc}"))

                elif path == "/leads/import":
                    try:
                        csv_text = _decode_uploaded_csv(f.get("csv_file")) or f.get("csv_text", "")
                        ok, skip = leads_csv.import_leads(
                            con, csv_text,
                            industries=sfa_db.get_master_list(con, "industries"),
                            company_sizes=sfa_db.get_master_list(con, "company_sizes"),
                        )
                        self._send(render(
                            leads_import_page(con),
                            flash=f"取込完了: {ok}件追加。" + (f"スキップ {skip}件。" if skip else ""),
                        ))
                    except Exception as exc:
                        self._send(render(leads_import_page(con), flash=f"取込エラー: {exc}"))

                elif path == "/deals/import":
                    try:
                        csv_text = _decode_uploaded_csv(f.get("csv_file")) or f.get("csv_text", "")
                        ok_deals, ok_accounts, dup_skip, skip = deals_csv.import_deals(
                            con, csv_text,
                            industries=sfa_db.get_master_list(con, "industries"),
                            company_sizes=sfa_db.get_master_list(con, "company_sizes"),
                        )
                        msg = f"取込完了: 商談{ok_deals}件・アカウント{ok_accounts}件を追加。"
                        if dup_skip:
                            msg += f" 重複スキップ {dup_skip}件（同一内容の商談が既存）。"
                        if skip:
                            msg += f" エラースキップ {skip}件。"
                        if ok_deals and theme_client is not None:
                            import threading as _threading
                            _threading.Thread(
                                target=theme_link.sync_all_pending, args=(theme_client, db_path),
                                daemon=True,
                            ).start()
                            msg += " テーマDBへの同期をバックグラウンドで開始しました。"
                        self._send(render(deals_import_page(con), flash=msg))
                    except Exception as exc:
                        self._send(render(deals_import_page(con), flash=f"取込エラー: {exc}"))

                elif path == "/deals/sync_pending":
                    if theme_client is None:
                        self._send(render(home_page(con), flash="テーマDB連携が無効です（THEME_API_TOKEN未設定）。"))
                    else:
                        pending = con.execute("SELECT COUNT(*) c FROM deals WHERE theme_id IS NULL").fetchone()["c"]
                        import threading as _threading
                        _threading.Thread(
                            target=theme_link.sync_all_pending, args=(theme_client, db_path),
                            daemon=True,
                        ).start()
                        self._send(render(
                            home_page(con),
                            flash=f"テーマDB未同期の商談{pending}件の同期をバックグラウンドで開始しました。"
                                  f"数分後に商談一覧の「連携」列で確認できます。",
                        ))

                elif path == "/deals/migrate_lost_stage":
                    # #67: ステージ='失注'の商談を「クローズ＋理由=失注」にし、フォロー中リードとして
                    #      作成/再活性化した上でHishoへ再同期。
                    res = sfa_db.migrate_lost_stage_to_closed(con)
                    _ids = res["deal_ids"]
                    synced = 0
                    if theme_client is not None:
                        for _did in _ids:
                            try:
                                theme_link.sync_deal(theme_client, con, _did)
                                synced += 1
                            except Exception as exc:  # noqa: BLE001
                                print(f"[theme_link] sync_deal failed (migrate_lost_stage): {exc}")
                    if _ids:
                        _lead_bits = []
                        if res["leads_created"]:
                            _lead_bits.append(f"リード新規作成 {res['leads_created']}件")
                        if res["leads_reactivated"]:
                            _lead_bits.append(f"既存リード再活性化 {res['leads_reactivated']}件")
                        _lead_txt = ("、" + "・".join(_lead_bits)) if _lead_bits else "（リードは既存を維持）"
                        _msg = (f"失注商談 {len(_ids)}件を処理しました"
                                f"（うち新規クローズ {res['newly_closed']}件{_lead_txt}"
                                + (f"／Hisho再同期 {synced}件" if theme_client is not None else "")
                                + "）。リード一覧で「フォロー中」として確認できます。")
                    else:
                        _msg = "ステージ='失注'の商談はありませんでした（移行不要）。"
                    self._send(render(sync_health_page(con, theme_client), flash=_msg))

                elif path == "/data-tagging/bulk-appt":
                    # 明日以降・ラベルに「初回アポ」を含む未タグ次回MSを一括で「アポ」にする
                    n = sfa_db.bulk_tag_appt_by_label(con, after_date=_today_jst().isoformat())
                    self._send(render(data_tagging_page(con),
                                      flash=f"「初回アポ」を含む未タグ（明日以降）{n}件を「アポ」にしました。"))

                elif path == "/sync-failures/retry":
                    # 記録済みのHisho同期失敗を1件ずつ再同期し、成功したものは記録から消す
                    if theme_client is None:
                        self._send(render(home_page(con), flash="テーマDB連携が無効です（THEME_API_TOKEN未設定）。"))
                    else:
                        failures = sfa_db.list_sync_failures(con)
                        ok = 0
                        still_failing = 0
                        for frec in failures:
                            kind, ref_id = frec["kind"], frec["ref_id"]
                            try:
                                if kind == "deal":
                                    theme_link.sync_deal(theme_client, con, ref_id)
                                elif kind == "dev_project":
                                    dev_project_link.sync_dev_project(theme_client, con, ref_id)
                                elif kind == "dev_project_delete":
                                    dev_project_link.delete_dev_project_remote(theme_client, ref_id)
                                else:
                                    continue
                                sfa_db.clear_sync_failure(con, kind, ref_id)
                                ok += 1
                            except Exception as exc:  # noqa: BLE001
                                sfa_db.record_sync_failure(con, kind, ref_id, str(exc))
                                still_failing += 1
                        self._send(render(
                            home_page(con),
                            flash=f"再同期: 成功{ok}件 / 失敗{still_failing}件。"
                                  + ("失敗分は記録に残しています。" if still_failing else ""),
                        ))

                elif path == "/leads/bulk_source":
                    ids = f_list.get("ids", [])
                    source = f.get("source", "")
                    if source in sfa_db.LEAD_SOURCES and ids:
                        for lead_id in ids:
                            if lead_id.isdigit():
                                con.execute(
                                    "UPDATE leads SET source=?, updated_at=datetime('now') WHERE id=?",
                                    (source, int(lead_id)),
                                )
                        con.commit()
                    self._redirect("/leads")

                elif path == "/leads/bulk_edit":
                    _LEAD_ALLOWED_FIELDS = {"source", "assigned_to", "industry", "company_size", "lead_status"}
                    ids = f_list.get("ids", [])
                    field = f.get("field", "")
                    value = f.get("value", "")
                    if field in _LEAD_ALLOWED_FIELDS and ids:
                        if field == "source" and value and value not in sfa_db.LEAD_SOURCES:
                            pass  # invalid, skip
                        elif field == "lead_status" and value and value not in sfa_db.LEAD_STATUSES:
                            pass  # invalid, skip
                        else:
                            for lead_id in ids:
                                if str(lead_id).isdigit():
                                    con.execute(
                                        f"UPDATE leads SET {field}=?, updated_at=datetime('now') WHERE id=?",
                                        (value or None, int(lead_id)),
                                    )
                            con.commit()
                    self._redirect("/leads")

                elif path == "/leads/bulk_delete":
                    ids = f_list.get("ids", [])
                    for lead_id in ids:
                        if str(lead_id).isdigit():
                            con.execute("DELETE FROM leads WHERE id=?", (int(lead_id),))
                    if ids:
                        con.commit()
                    self._redirect("/leads")

                elif path.startswith("/leads/") and path.endswith("/set_pattern"):
                    parts = path.split("/")
                    _ok = False
                    _err = ""
                    if len(parts) == 4 and parts[2].isdigit():
                        lid = int(parts[2])
                        pid_str = f.get("pattern_id", "")
                        pid = int(pid_str) if pid_str and pid_str.isdigit() else None
                        sfa_db.set_lead_email_pattern(con, lid, pid)
                        _ok = True
                    else:
                        _err = "不正なリクエスト"
                    self._send(json.dumps({"ok": _ok} if _ok else {"ok": False, "error": _err}).encode(), ctype="application/json")
                elif path.startswith("/leads/") and path.endswith("/field"):
                    _LEAD_ALLOWED_FIELDS = {"source", "assigned_to", "industry", "company_size", "lead_status",
                                            "lost_reason"}
                    parts = path.split("/")
                    _ok = False
                    _err = ""
                    if len(parts) == 4 and parts[3] == "field" and parts[2].isdigit():
                        lid = int(parts[2])
                        field = f.get("field", "")
                        value = f.get("value", "")
                        if field not in _LEAD_ALLOWED_FIELDS:
                            _err = "不正なフィールド"
                        elif field == "source" and value and value not in sfa_db.LEAD_SOURCES:
                            _err = "不正な経路値"
                        elif field == "lead_status" and value and value not in sfa_db.LEAD_STATUSES:
                            _err = "不正なステータス値"
                        elif field == "lost_reason" and value and value not in sfa_db.CLOSE_REASONS:
                            _err = "不正な終了理由"
                        else:
                            con.execute(
                                f"UPDATE leads SET {field}=?, updated_at=datetime('now') WHERE id=?",
                                (value or None, lid),
                            )
                            con.commit()
                            _ok = True
                    else:
                        _err = "不正なリクエスト"
                    _resp = json.dumps({"ok": _ok} if _ok else {"ok": False, "error": _err}).encode("utf-8")
                    self._send(_resp, ctype="application/json")

                elif path.startswith("/leads/") and path.endswith("/delete"):
                    parts = path.split("/")
                    if len(parts) == 4 and parts[3] == "delete" and parts[2].isdigit():
                        lid = int(parts[2])
                        con.execute("DELETE FROM leads WHERE id=?", (lid,))
                        con.commit()
                    self._redirect("/leads")

                elif path.startswith("/leads/") and path.endswith("/activity"):
                    parts = path.split("/")
                    if len(parts) != 4 or not parts[2].isdigit():
                        self._redirect("/leads")
                        return
                    lid = int(parts[2])
                    sfa_db.create_lead_activity(
                        con, lead_id=lid,
                        type=f.get("type") or "note",
                        content=f.get("content") or "(内容なし)",
                        author=f.get("author") or None,
                    )
                    self._redirect(f"/leads/{lid}")

                elif path.startswith("/leads/") and path.endswith("/status"):
                    parts = path.split("/")
                    if len(parts) != 4 or not parts[2].isdigit():
                        self._redirect("/leads")
                        return
                    lid = int(parts[2])
                    new_status = f.get("status", "")
                    if new_status in sfa_db.LEAD_STATUSES:
                        con.execute(
                            "UPDATE leads SET lead_status=?, updated_at=datetime('now') WHERE id=?",
                            (new_status, lid),
                        )
                        con.commit()
                    self._redirect(f"/leads/{lid}")

                elif path.startswith("/leads/") and path.endswith("/convert"):
                    parts = path.split("/")
                    if len(parts) != 4 or not parts[2].isdigit():
                        self._redirect("/leads")
                        return
                    lid = int(parts[2])
                    lead = sfa_db.get_lead(con, lid)
                    if not lead:
                        self._redirect("/leads")
                    else:
                        mode = f.get("mode") or "new"
                        try:
                            if mode.startswith("revive_") and mode[7:].isdigit():
                                deal_id = revive_deal_from_lead(con, lead, int(mode[7:]))
                            else:
                                deal_id = convert_lead_to_deal(con, lead)
                            if theme_client is not None:
                                try:
                                    theme_link.sync_deal(theme_client, con, deal_id)
                                except Exception as _exc:
                                    print(f"[theme_link] sync_deal failed: {_exc}")
                            self._redirect(f"/deal/{deal_id}")
                        except Exception as _conv_e:
                            print(f"[convert] error lid={lid}: {_conv_e}", flush=True)
                            import traceback as _tb; _tb.print_exc()
                            self._redirect(f"/leads/{lid}")

                # ── 商談 → リード戻し ──
                elif path.endswith("/revert_to_lead") and "/deal/" in path:
                    deal_id_str = path.split("/deal/")[1].split("/")[0]
                    _redirect_to = "/deals"
                    if deal_id_str.isdigit():
                        _did = int(deal_id_str)
                        _deal = sfa_db.get_deal(con, _did)
                        if _deal and _deal.get("status") != "closed":
                            _memo = (f.get("memo") or "").strip()
                            _cr = (f.get("close_reason") or "").strip()
                            _cr = _cr if _cr in sfa_db.CLOSE_REASONS else None
                            if _cr is None:
                                # 終了理由は必須。未選択ならクローズせずに戻す（直POST等への防御）
                                _rt0 = f.get("return_to") or ""
                                self._redirect(_rt0 if _rt0.startswith("/") else f"/deal/{_did}")
                                return
                            _acct_row = con.execute(
                                "SELECT * FROM accounts WHERE id=?", (_deal.get("account_id"),)
                            ).fetchone()
                            _acct = dict(_acct_row) if _acct_row else {}

                            # 現状メモ: リードに戻す時のメモ（あれば）を一番上に追記した上でクローズ理由も付与
                            _close_line = "アポ未獲得のためクローズ（リードに戻す）"
                            if _cr:
                                _close_line += f"／終了理由: {_cr}"
                            _existing_note = _deal.get("note") or ""
                            _body = f"{_existing_note}\n{_close_line}" if _existing_note else _close_line
                            _new_note = f"[リードに戻す時のメモ] {_memo}\n{_body}" if _memo else _body
                            con.execute(
                                "UPDATE deals SET status='closed', note=?, "
                                "close_reason=COALESCE(?, close_reason), updated_at=datetime('now') WHERE id=?",
                                (_new_note, _cr, _did),
                            )

                            _lid = None
                            # 既存リード検索（deal_id が紐付いているもの）
                            _lead_row = con.execute(
                                "SELECT * FROM leads WHERE deal_id=? LIMIT 1", (_did,)
                            ).fetchone()
                            if _lead_row:
                                _lead = dict(_lead_row)
                                _lid = _lead["id"]
                                _lead_notes = (_lead.get("notes") or "")
                                _lead_new_notes = f"{_new_note}\n{_lead_notes}" if _lead_notes else _new_note
                                con.execute(
                                    "UPDATE leads SET lead_status='following', deal_id=NULL, "
                                    "industry=COALESCE(?, industry), company_size=COALESCE(?, company_size), "
                                    "notes=?, updated_at=datetime('now') WHERE id=?",
                                    (_acct.get("industry"), _acct.get("company_size"), _lead_new_notes, _lid),
                                )
                                _activity_note = "アポ未獲得のため商談からリードへ戻す（フォロー中に変更）。"
                                if _memo:
                                    _activity_note += f" メモ: {_memo}"
                                con.execute(
                                    "INSERT INTO lead_activities (lead_id,type,content,author) VALUES (?,?,?,?)",
                                    (_lid, "note", _activity_note, "システム"),
                                )
                            else:
                                # 既存リードがなければアカウントから新規作成（業界・企業規模・現状メモを連携）
                                _origin_line = f"商談 #{_did}（{_deal.get('deal_name','')}）からリードに戻す"
                                _lid = sfa_db.upsert_lead(
                                    con, name=_acct.get("name", "（不明）"),
                                    company=_acct.get("name", "（不明）"),
                                    industry=_acct.get("industry"),
                                    company_size=_acct.get("company_size"),
                                    lead_status="following",
                                    notes=f"{_origin_line}\n{_new_note}",
                                    assigned_to=_deal.get("owner"),
                                )
                            con.commit()
                            if theme_client is not None:
                                try:
                                    theme_link.sync_deal(theme_client, con, _did)
                                except Exception as _exc:
                                    print(f"[theme_link] sync_deal failed: {_exc}")
                            _return_to = f.get("return_to") or ""
                            if _return_to.startswith("/"):
                                _redirect_to = _return_to
                            elif _lid:
                                _redirect_to = f"/leads/{_lid}"
                    self._redirect(_redirect_to)

                # 旧「商談クローズ（リードに戻さない）」は廃止。クローズは /deal/{id}/revert_to_lead に一本化。

                # ── メモ保存 ──
                elif path == "/api/memo/save":
                    qs = self._qs()
                    token = (qs.get("token", [None])[0] or "")
                    if not SFA_API_TOKEN or not hmac.compare_digest(token, SFA_API_TOKEN):
                        self._send_cors_json(b'{"error":"unauthorized"}', status=401)
                    else:
                        try:
                            data = json.loads(raw)
                        except Exception:
                            data = f
                        tid = data.get("theme_id")
                        note_id = con.execute(
                            "INSERT INTO meeting_notes(theme_id,note_date,body,task,task_owner,task_due) VALUES(?,?,?,?,?,?)",
                            (int(tid) if tid else None, data.get("note_date") or None,
                             data.get("body") or None, data.get("task") or None,
                             data.get("task_owner") or None, data.get("task_due") or None),
                        ).lastrowid
                        con.commit()
                        self._send_cors_json(json.dumps({"ok": True, "id": note_id}, ensure_ascii=False).encode())

                # ── メモ削除 ──
                elif path == "/api/memo/delete":
                    qs = self._qs()
                    token = (qs.get("token", [None])[0] or "")
                    if not SFA_API_TOKEN or not hmac.compare_digest(token, SFA_API_TOKEN):
                        self._send_cors_json(b'{"error":"unauthorized"}', status=401)
                    else:
                        try:
                            data = json.loads(raw)
                        except Exception:
                            data = f
                        note_id = data.get("id")
                        con.execute("DELETE FROM meeting_notes WHERE id=?", (int(note_id),))
                        con.commit()
                        self._send_cors_json(json.dumps({"ok": True}, ensure_ascii=False).encode())

                # ── タスク完了トグル ──
                elif path == "/api/memo/toggle_task":
                    qs = self._qs()
                    token = (qs.get("token", [None])[0] or "")
                    if not SFA_API_TOKEN or not hmac.compare_digest(token, SFA_API_TOKEN):
                        self._send_cors_json(b'{"error":"unauthorized"}', status=401)
                    else:
                        try:
                            data = json.loads(raw)
                        except Exception:
                            data = f
                        note_id = data.get("id")
                        done = 1 if data.get("done") else 0
                        con.execute("UPDATE meeting_notes SET task_done=? WHERE id=?", (done, int(note_id)))
                        con.commit()
                        self._send_cors_json(json.dumps({"ok": True}, ensure_ascii=False).encode())

                # ── 翌日アポSlack通知の重複防止マーカー ──
                elif path == "/api/deals/mark_notified":
                    qs = self._qs()
                    token = (qs.get("token", [None])[0] or "")
                    if not SFA_API_TOKEN or not hmac.compare_digest(token, SFA_API_TOKEN):
                        self._send_cors_json(b'{"error":"unauthorized"}', status=401)
                    else:
                        try:
                            data = json.loads(raw)
                        except Exception:
                            data = f
                        deal_id = data.get("deal_id")
                        notified_date = data.get("date")
                        if deal_id and notified_date:
                            con.execute(
                                "UPDATE deals SET slack_notified_date=? WHERE id=?",
                                (notified_date, int(deal_id)),
                            )
                            con.commit()
                            self._send_cors_json(json.dumps({"ok": True}, ensure_ascii=False).encode())
                        else:
                            self._send_cors_json(b'{"error":"deal_id and date required"}', status=400)

                # ── Slack スラッシュコマンド /task（起票モーダルを開く）──
                elif path == "/slack/commands":
                    import threading as _threading
                    from cowork import slack_bot as _sb
                    if not _sb.verify_signature(
                        raw.encode("utf-8"),
                        self.headers.get("X-Slack-Request-Timestamp", ""),
                        self.headers.get("X-Slack-Signature", ""),
                    ):
                        self._send(b'{"error":"invalid signature"}', 401, ctype="application/json")
                        return
                    import urllib.parse as _up2
                    _form = {k: (v[0] if v else "")
                             for k, v in _up2.parse_qs(raw, keep_blank_values=True).items()}
                    self.send_response(200)
                    self.send_header("Content-Type", "text/plain")
                    self.end_headers()
                    self.wfile.write(b"")

                    def _proc_cmd():
                        _con = sfa_db.connect(db_path)
                        try:
                            from cowork import slack_tasks as _st
                            _st.handle_slash(_con, _form)
                        except Exception as _e:  # noqa: BLE001
                            print(f"[slack_commands] error: {_e}", flush=True)
                        finally:
                            _con.close()
                    _threading.Thread(target=_proc_cmd, daemon=True).start()
                    return

                # ── Slack インタラクティブ（ボタン・モーダル送信）──
                elif path == "/slack/interactive":
                    import threading as _threading
                    from cowork import slack_bot as _sb
                    if not _sb.verify_signature(
                        raw.encode("utf-8"),
                        self.headers.get("X-Slack-Request-Timestamp", ""),
                        self.headers.get("X-Slack-Signature", ""),
                    ):
                        self._send(b'{"error":"invalid signature"}', 401, ctype="application/json")
                        return
                    import urllib.parse as _up2
                    _payload_raw = _up2.parse_qs(raw, keep_blank_values=True).get("payload", [""])[0]
                    try:
                        _payload = json.loads(_payload_raw)
                    except Exception:
                        self._send(b"{}", 400, ctype="application/json")
                        return
                    from cowork import slack_tasks as _st
                    if _payload.get("type") == "view_submission":
                        # モーダル送信: 3秒以内にresponse_actionを同期で返す。AIは背景へ退避。
                        _con2 = sfa_db.connect(db_path)
                        try:
                            _resp = _st.handle_interactive(_con2, _payload) or {}
                        finally:
                            _con2.close()
                        _defer_cat = _resp.pop("_defer_category", None)
                        _defer_sum = _resp.pop("_defer_summary", None)
                        _rb = json.dumps(_resp, ensure_ascii=False).encode() if _resp else b""
                        self.send_response(200)
                        self.send_header("Content-Type", "application/json")
                        self.send_header("Content-Length", str(len(_rb)))
                        self.end_headers()
                        self.wfile.write(_rb)
                        if _defer_cat or _defer_sum:
                            def _bg():
                                _c3 = sfa_db.connect(db_path)
                                try:
                                    if _defer_cat:
                                        _st.regenerate_task_category(_c3, _defer_cat)
                                    if _defer_sum:
                                        _st.regenerate_task_summary(_c3, _defer_sum)
                                finally:
                                    _c3.close()
                            _threading.Thread(target=_bg, daemon=True).start()
                        return
                    else:
                        # ボタン(block_actions): 即200→背景処理（モーダルを開く操作もここ）
                        self.send_response(200)
                        self.send_header("Content-Type", "text/plain")
                        self.end_headers()
                        self.wfile.write(b"")

                        def _proc_iv():
                            _con = sfa_db.connect(db_path)
                            try:
                                _st.handle_interactive(_con, _payload)
                            except Exception as _e:  # noqa: BLE001
                                print(f"[slack_interactive] error: {_e}", flush=True)
                            finally:
                                _con.close()
                        _threading.Thread(target=_proc_iv, daemon=True).start()
                        return

                # ── Slack Events API ──
                elif path == "/slack/events":
                    import threading as _threading
                    from cowork import slack_bot as _sb
                    # body は do_POST 先頭の raw 変数で読み込み済み（rfile は再読不可）
                    # Slack署名検証（fail-closed）: 偽イベントによるDB書き込み・API消費を防ぐ
                    if not _sb.verify_signature(
                        raw.encode("utf-8"),
                        self.headers.get("X-Slack-Request-Timestamp", ""),
                        self.headers.get("X-Slack-Signature", ""),
                    ):
                        self._send(b'{"error":"invalid signature"}', 401, ctype="application/json")
                        return
                    try:
                        data = json.loads(raw)
                    except Exception:
                        self._send(b"<error/>", 400)
                        return

                    # URL検証チャレンジ
                    if data.get("type") == "url_verification":
                        self.send_response(200)
                        self.send_header("Content-Type", "application/json")
                        self.end_headers()
                        self.wfile.write(json.dumps({"challenge": data["challenge"]}).encode())
                        return

                    # Slackに即時200を返してからバックグラウンド処理
                    self.send_response(200)
                    self.send_header("Content-Type", "text/plain")
                    self.end_headers()
                    self.wfile.write(b"ok")

                    # イベントをバックグラウンドで処理（conはスレッドセーフのため再接続）
                    def _process():
                        _con = sfa_db.connect(db_path)
                        try:
                            from cowork import slack_bot
                            _inner = data.get("event", {}) or {}
                            if _inner.get("type") == "reaction_added":
                                # 🎯等のリアクションでタスク化。Slack再送対策に冪等化。
                                _eid = data.get("event_id")
                                if _eid and not slack_bot._mark_event_processed(_con, _eid):
                                    return
                                from cowork import slack_tasks as _st
                                _st.handle_reaction(_con, _inner)
                            else:
                                slack_bot.handle_event(data, _con, theme_client)
                        except Exception as _e:
                            print(f"[slack_events] error: {_e}")
                        finally:
                            _con.close()
                    _threading.Thread(target=_process, daemon=True).start()
                    return

                else:
                    self._send(render("<div class=card>不明な操作</div>"), 404)
            finally:
                con.close()

    return H


def start(db_path: str = sfa_db.DEFAULT_DB_PATH, port: int = 8787,
          theme_client: ThemeDBClient | None = None) -> None:
    sfa_db.init_db(db_path)
    handler = _make_handler(db_path, theme_client)
    srv = ThreadingHTTPServer(("0.0.0.0", port), handler)
    print(f"Inproc Salesforce: http://localhost:{port}  (DB={db_path})")
    srv.serve_forever()
